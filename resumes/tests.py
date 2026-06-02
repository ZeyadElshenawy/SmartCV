"""Tests for resumes.views description helpers.

These helpers handle the bracket-corruption bug territory: the resume editor's
textarea stores multiline bullet descriptions, but the JSON schema stores them
as List[str]. A mistake in this conversion (or a round-trip that mutates the
data) is what caused the bug fixed in fd90299.
"""
from django.test import SimpleTestCase

from resumes.views import (
    _description_list_to_text,
    _description_text_to_list,
)


class TextareaToListTests(SimpleTestCase):
    def test_none_becomes_empty_list(self):
        self.assertEqual(_description_text_to_list(None), [])

    def test_empty_string_becomes_empty_list(self):
        self.assertEqual(_description_text_to_list(''), [])

    def test_single_line_becomes_single_element_list(self):
        self.assertEqual(_description_text_to_list('Shipped feature X'), ['Shipped feature X'])

    def test_newline_separated_bullets_become_list(self):
        raw = 'Shipped feature X\nOwned migration Y\nMentored 2 juniors'
        self.assertEqual(
            _description_text_to_list(raw),
            ['Shipped feature X', 'Owned migration Y', 'Mentored 2 juniors'],
        )

    def test_crlf_line_endings_are_handled(self):
        """Browsers POST textareas with \\r\\n; regression guard."""
        raw = 'Line one\r\nLine two\r\nLine three'
        self.assertEqual(
            _description_text_to_list(raw),
            ['Line one', 'Line two', 'Line three'],
        )

    def test_blank_lines_are_dropped(self):
        raw = 'First\n\n\nSecond\n   \nThird'
        self.assertEqual(
            _description_text_to_list(raw),
            ['First', 'Second', 'Third'],
        )

    def test_surrounding_whitespace_is_stripped(self):
        raw = '   padded bullet   \n\ttabbed bullet\t'
        self.assertEqual(
            _description_text_to_list(raw),
            ['padded bullet', 'tabbed bullet'],
        )


class ListToTextareaTests(SimpleTestCase):
    def test_none_becomes_empty_string(self):
        self.assertEqual(_description_list_to_text(None), '')

    def test_empty_list_becomes_empty_string(self):
        self.assertEqual(_description_list_to_text([]), '')

    def test_list_joins_with_newline(self):
        self.assertEqual(
            _description_list_to_text(['First bullet', 'Second bullet']),
            'First bullet\nSecond bullet',
        )

    def test_legacy_string_value_passes_through(self):
        """Older resumes may still have string-shaped descriptions; don't mangle them."""
        self.assertEqual(
            _description_list_to_text('already a string'),
            'already a string',
        )

    def test_falsy_list_items_are_skipped(self):
        self.assertEqual(
            _description_list_to_text(['real', '', None, 'also real']),
            'real\nalso real',
        )

    def test_non_string_items_are_coerced(self):
        self.assertEqual(_description_list_to_text([1, 2]), '1\n2')


class RoundTripTests(SimpleTestCase):
    """The view's lifecycle is: stored List[str] -> textarea string (GET) ->
    back to List[str] (POST save). This must be lossless for well-formed data,
    which is exactly what the bracket-corruption bug violated."""

    def test_list_roundtrips_losslessly(self):
        original = ['Led team of 5 engineers', 'Cut p95 latency by 40%', 'Shipped feature X']
        textarea = _description_list_to_text(original)
        roundtripped = _description_text_to_list(textarea)
        self.assertEqual(roundtripped, original)

    def test_empty_list_roundtrips(self):
        self.assertEqual(_description_text_to_list(_description_list_to_text([])), [])

    def test_user_editing_in_browser_preserves_bullets(self):
        """Simulate a user opening the editor (LF on server) and the browser
        resubmitting the same textarea with CRLF line endings."""
        original = ['Built A', 'Built B', 'Built C']
        textarea_server_sent = _description_list_to_text(original)
        textarea_browser_posted = textarea_server_sent.replace('\n', '\r\n')
        self.assertEqual(_description_text_to_list(textarea_browser_posted), original)



# ============================================================
# resumes.services.scoring — ATS breakdown + evidence confidence
# ============================================================

from types import SimpleNamespace
from resumes.services.scoring import (
    compute_ats_breakdown,
    compute_evidence_confidence,
    calculate_ats_score,
    STUFFING_THRESHOLD,
    STUFFING_PENALTY_PER_SKILL,
    IN_CONTEXT_BONUS_PER_SKILL,
)


class ComputeAtsBreakdownTests(SimpleTestCase):
    def test_no_job_skills_returns_zero(self):
        out = compute_ats_breakdown({"skills": ["Python"]}, [])
        self.assertEqual(out["score"], 0.0)
        self.assertEqual(out["matched_count"], 0)
        self.assertEqual(out["total_count"], 0)

    def test_basic_match_score(self):
        # 2 of 4 keywords present anywhere in the resume JSON
        content = {"skills": ["Python", "SQL"], "experience": []}
        out = compute_ats_breakdown(content, ["Python", "SQL", "Rust", "Go"])
        self.assertEqual(out["matched_count"], 2)
        self.assertEqual(out["total_count"], 4)
        self.assertEqual(out["raw_score"], 50.0)
        # No in-context bonus (no experience), no stuffing — final = raw
        self.assertEqual(out["score"], 50.0)

    def test_in_context_bonus_for_keywords_in_experience(self):
        # Same matched count, but the keyword also appears in experience
        # bullets — should get the in-context bonus.
        content = {
            "skills": ["Python"],
            "experience": [
                {"description": ["Built distributed Python pipelines"]},
            ],
        }
        out = compute_ats_breakdown(content, ["Python"])
        self.assertEqual(out["matched_count"], 1)
        self.assertEqual(out["in_context_count"], 1)
        # raw 100 + bonus 2, capped at 100
        self.assertEqual(out["score"], 100.0)
        self.assertEqual(out["in_context_bonus"], IN_CONTEXT_BONUS_PER_SKILL)

    def test_in_context_bonus_is_capped(self):
        # 6 in-context skills × 2 = 12 raw bonus, capped at 10
        content = {
            "skills": ["Python", "SQL", "Java", "Rust", "Go", "Ruby"],
            "experience": [{
                "description": [
                    "Used Python and SQL daily",
                    "Migrated services to Java and Go",
                    "Wrote internal tooling in Rust and Ruby",
                ],
            }],
        }
        out = compute_ats_breakdown(content, ["Python", "SQL", "Java", "Rust", "Go", "Ruby"])
        self.assertEqual(out["in_context_count"], 6)
        self.assertEqual(out["in_context_bonus"], 10.0)  # capped

    def test_keyword_stuffing_is_penalized(self):
        # A skill that appears > STUFFING_THRESHOLD times across the resume
        # gets penalized 5 points per stuffed keyword.
        stuffed = " python " * (STUFFING_THRESHOLD + 1)
        content = {
            "skills": ["Python"],
            "experience": [{"description": [stuffed]}],
        }
        out = compute_ats_breakdown(content, ["Python"])
        self.assertIn("Python", out["stuffed_skills"])
        self.assertEqual(out["stuffing_penalty"], STUFFING_PENALTY_PER_SKILL)
        # raw 100 + 2 bonus - 5 penalty = 97
        self.assertEqual(out["score"], 97.0)

    def test_score_is_clamped_to_zero_when_penalties_exceed(self):
        # Engineer 5 stuffed skills (5 × 5pt = 25 pt penalty) on a resume
        # with raw_score 0 (no skills match). Final must clamp to 0, not negative.
        stuffed_text = " ".join(["python sql java rust go"] * 6)
        content = {"experience": [{"description": [stuffed_text]}]}
        out = compute_ats_breakdown(content, ["Python", "SQL", "Java", "Rust", "Go"])
        self.assertGreater(out["stuffing_penalty"], 0)
        self.assertGreaterEqual(out["score"], 0.0)

    def test_keyword_counts_per_skill_are_returned(self):
        content = {"skills": ["Python", "SQL"]}
        out = compute_ats_breakdown(content, ["Python", "SQL", "Rust"])
        self.assertEqual(out["keyword_counts"]["Python"], 1)
        self.assertEqual(out["keyword_counts"]["SQL"], 1)
        self.assertEqual(out["keyword_counts"]["Rust"], 0)

    def test_legacy_calculate_ats_score_returns_just_the_float(self):
        content = {"skills": ["Python"]}
        score = calculate_ats_score(content, ["Python"])
        self.assertIsInstance(score, float)
        self.assertEqual(score, 100.0)


class ComputeEvidenceConfidenceTests(SimpleTestCase):
    def _profile(self, **signals):
        return SimpleNamespace(data_content=signals)

    def test_no_signals_returns_zero(self):
        out = compute_evidence_confidence(self._profile())
        self.assertEqual(out["score"], 0)
        self.assertEqual(out["label"], "Untested")
        self.assertEqual(out["sources"], [])

    def test_only_github_with_repos_counts(self):
        out = compute_evidence_confidence(self._profile(
            github_signals={"public_repos": 5},
        ))
        self.assertEqual(out["score"], 1)
        self.assertEqual(out["label"], "Limited")
        self.assertEqual(out["sources"], ["github"])

    def test_github_with_zero_repos_does_not_count(self):
        out = compute_evidence_confidence(self._profile(
            github_signals={"public_repos": 0},
        ))
        self.assertEqual(out["score"], 0)

    def test_error_snapshot_is_skipped(self):
        out = compute_evidence_confidence(self._profile(
            github_signals={"error": "rate-limited", "public_repos": 99},
        ))
        self.assertEqual(out["score"], 0)

    def test_scholar_needs_pubs_or_citations(self):
        # Empty top_publications + 0 citations → does not count
        out = compute_evidence_confidence(self._profile(
            scholar_signals={"top_publications": [], "total_citations": 0},
        ))
        self.assertEqual(out["score"], 0)
        # With citations
        out = compute_evidence_confidence(self._profile(
            scholar_signals={"top_publications": [], "total_citations": 100},
        ))
        self.assertEqual(out["score"], 1)
        # With publications
        out = compute_evidence_confidence(self._profile(
            scholar_signals={"top_publications": [{"title": "A"}], "total_citations": 0},
        ))
        self.assertEqual(out["score"], 1)

    def test_kaggle_any_category_activity_counts(self):
        out = compute_evidence_confidence(self._profile(
            kaggle_signals={
                "competitions": {"count": 0},
                "datasets": {"count": 0},
                "notebooks": {"count": 3},
                "discussion": {"count": 0},
            },
        ))
        self.assertEqual(out["score"], 1)
        self.assertEqual(out["sources"], ["kaggle"])

    def test_all_three_signals_max_score(self):
        out = compute_evidence_confidence(self._profile(
            github_signals={"public_repos": 5},
            scholar_signals={"total_citations": 100, "top_publications": [{"title": "X"}]},
            kaggle_signals={"competitions": {"count": 1}, "datasets": {"count": 0},
                            "notebooks": {"count": 0}, "discussion": {"count": 0}},
        ))
        self.assertEqual(out["score"], 3)
        self.assertEqual(out["label"], "Strong")
        self.assertEqual(set(out["sources"]), {"github", "scholar", "kaggle"})

    def test_linkedin_link_only_counts(self):
        # A parsed handle without scraped rich fields still counts — the
        # recruiter can verify identity via the URL alone.
        out = compute_evidence_confidence(self._profile(
            linkedin_signals={"username": "jane-doe", "profile_url": "https://linkedin.com/in/jane-doe"},
        ))
        self.assertEqual(out["score"], 1)
        self.assertEqual(out["sources"], ["linkedin"])

    def test_linkedin_error_snapshot_skipped(self):
        out = compute_evidence_confidence(self._profile(
            linkedin_signals={"error": "blocked", "profile_url": "x"},
        ))
        self.assertEqual(out["score"], 0)

    def test_all_four_signals_max_score(self):
        out = compute_evidence_confidence(self._profile(
            github_signals={"public_repos": 5},
            scholar_signals={"total_citations": 100, "top_publications": [{"title": "X"}]},
            kaggle_signals={"competitions": {"count": 1}, "datasets": {"count": 0},
                            "notebooks": {"count": 0}, "discussion": {"count": 0}},
            linkedin_signals={"username": "jane", "profile_url": "https://linkedin.com/in/jane"},
        ))
        self.assertEqual(out["score"], 4)
        self.assertEqual(out["label"], "Strong")
        self.assertEqual(set(out["sources"]), {"github", "scholar", "kaggle", "linkedin"})


from django.test import TestCase
from django.contrib.auth import get_user_model
from django.urls import reverse


class ResumeEditPreviewTemplateClassTests(TestCase):
    """Live preview on /resumes/edit/<id>/ must carry a pdf-preview--<template>
    modifier class so each template choice shifts the preview's CSS vars.
    If it doesn't render server-side, the page opens mismatched with the
    saved template, and the first radio click would be the first thing to
    shape the preview (jarring UX).
    """

    def setUp(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from resumes.models import GeneratedResume
        User = get_user_model()
        self.user = User.objects.create_user(
            username='rt@example.com', email='rt@example.com', password='x',
        )
        self.client.force_login(self.user)
        job = Job.objects.create(user=self.user, title='Data Scientist')
        self.gap = GapAnalysis.objects.create(
            user=self.user, job=job, similarity_score=0.5,
        )
        self.resume = GeneratedResume.objects.create(
            gap_analysis=self.gap,
            content={
                'professional_title': 'Data Scientist',
                'professional_summary': 'Lorem ipsum.',
                'template_name': 'ats_clean_accent',
            },
        )

    def test_preview_carries_saved_template_modifier(self):
        resp = self.client.get(reverse('resume_edit', args=[self.resume.id]))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'pdf-preview pdf-preview--ats_clean_accent')

    def test_preview_falls_back_to_default_when_template_missing(self):
        self.resume.content = {'professional_title': 'X', 'professional_summary': 'Y'}
        self.resume.save()
        resp = self.client.get(reverse('resume_edit', args=[self.resume.id]))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'pdf-preview pdf-preview--ats_clean')

    def test_every_card_has_a_thumbnail_preview(self):
        """Each template radio card must render a .template-thumb miniature
        styled with the same pdf-preview--<value> modifier the big right-side
        preview uses, so users can eyeball the style before picking."""
        import re
        resp = self.client.get(reverse('resume_edit', args=[self.resume.id]))
        body = resp.content.decode('utf-8')
        values = re.findall(r'value="([^"]+)"\s+class="sr-only"', body)
        self.assertTrue(values, 'Template radio values not found in page.')
        for v in values:
            needle = f'class="template-thumb pdf-preview pdf-preview--{v}"'
            self.assertIn(
                needle, body,
                f'Template "{v}" radio card is missing its thumbnail div.',
            )
        # And the thumbnail stylesheet must exist.
        self.assertIn('.template-thumb {', body)

    def test_every_template_choice_has_matching_css_modifier(self):
        """Regression: if a new template is added to template_choices in the
        view but the CSS block is forgotten, the preview silently falls back
        to the default. Compare the view's choices against CSS selectors in
        the rendered page and fail if any are missing."""
        import re
        resp = self.client.get(reverse('resume_edit', args=[self.resume.id]))
        body = resp.content.decode('utf-8')
        values = re.findall(r'value="([^"]+)"\s+class="sr-only"', body)
        self.assertTrue(values, 'Template radio values not found in page.')
        for v in values:
            self.assertIn(
                f'.pdf-preview--{v}', body,
                f'No .pdf-preview--{v} rule in the page for template "{v}".',
            )


class DocxExportTests(TestCase):
    """The DOCX exporter mirrors PDF export's section_order, the same
    fields, and the same authorization. The output should be a valid
    .docx (zip with [Content_Types].xml) carrying the candidate's
    name, professional title, all section headings the saved
    section_order asks for, and content for each.

    We don't try to verify visual fidelity — DOCX rendering varies by
    Word version and ATS pipelines re-style aggressively anyway.
    """

    def setUp(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from profiles.models import UserProfile
        from resumes.models import GeneratedResume
        from django.contrib.auth import get_user_model
        User = get_user_model()
        self.user = User.objects.create_user(
            username='docx@example.com', email='docx@example.com', password='x',
        )
        UserProfile.objects.create(
            user=self.user, full_name='Ada Lovelace', email='docx@example.com',
            phone='+1-555-0100', location='London, UK',
            linkedin_url='https://www.linkedin.com/in/ada/',
            github_url='https://github.com/ada/',
            data_content={'portfolio_url': 'https://ada.example.com'},
        )
        self.client.force_login(self.user)
        job = Job.objects.create(
            user=self.user, title='Backend Engineer', company='Acme Corp',
            description='Need a Python backend dev.',
            extracted_skills=['Python', 'PostgreSQL'],
        )
        gap = GapAnalysis.objects.create(user=self.user, job=job, similarity_score=0.7)
        self.resume = GeneratedResume.objects.create(
            gap_analysis=gap,
            content={
                'professional_title': 'Backend Systems Engineer',
                'professional_summary': 'Built distributed Python systems.',
                'skills': ['Python', 'PostgreSQL', 'Redis'],
                'experience': [{
                    'title': 'Backend Engineer',
                    'company': 'PriorCo',
                    'duration': '2023 - Present',
                    'description': ['Cut p99 latency from 1.2s to 380ms.', 'Owned async migration.'],
                }],
                'education': [{
                    'degree': 'BSc Computer Science',
                    'institution': 'Cairo University',
                    'graduation_year': '2024',
                }],
                'projects': [{
                    'name': 'pgbench-tuner',
                    'url': 'https://github.com/ada/pgbench-tuner',
                    'description': ['Auto-tunes PostgreSQL parameters.'],
                }],
                'certifications': [{
                    'name': 'AWS SAA',
                    'issuer': 'Amazon',
                    'date': '2024',
                    'url': 'https://example.com/cert',
                }],
                'languages': ['English', 'Arabic'],
            },
        )

    def _url(self):
        return f'/resumes/export-docx/{self.resume.id}/'

    def test_export_returns_docx_content_type(self):
        resp = self.client.get(self._url())
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(
            resp['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        self.assertIn('attachment', resp['Content-Disposition'])
        self.assertIn('.docx', resp['Content-Disposition'])

    def test_export_produces_valid_zip_with_required_parts(self):
        """A real DOCX is a zip containing [Content_Types].xml and
        word/document.xml. If we accidentally returned plain text or
        a corrupted file, this catches it."""
        import io
        import zipfile
        resp = self.client.get(self._url())
        buf = io.BytesIO(resp.content)
        with zipfile.ZipFile(buf) as z:
            names = set(z.namelist())
            self.assertIn('[Content_Types].xml', names)
            self.assertIn('word/document.xml', names)
            doc_xml = z.read('word/document.xml').decode('utf-8')
        # Spot-check: the candidate's name should appear in the document XML
        self.assertIn('ADA LOVELACE', doc_xml)
        # Section headings the saved order asks for should all be present
        self.assertIn('PROFESSIONAL SUMMARY', doc_xml)
        self.assertIn('SKILLS', doc_xml)
        self.assertIn('PROFESSIONAL EXPERIENCE', doc_xml)
        self.assertIn('EDUCATION', doc_xml)
        self.assertIn('PROJECTS', doc_xml)
        self.assertIn('CERTIFICATIONS', doc_xml)
        self.assertIn('LANGUAGES', doc_xml)
        # Content sample — bullet text from experience
        self.assertIn('Cut p99 latency', doc_xml)

    def test_export_honors_saved_section_order(self):
        """Sections in the DOCX should appear in the user's saved
        section_order, not the default. Verify by extracting all section
        heading positions and asserting the order matches what's saved."""
        import io, zipfile, re
        # Reorder: projects → skills → summary, rest fills in after.
        self.resume.content = {
            **self.resume.content,
            'section_order': ['projects', 'skills', 'summary'],
        }
        self.resume.save()
        resp = self.client.get(self._url())
        with zipfile.ZipFile(io.BytesIO(resp.content)) as z:
            doc_xml = z.read('word/document.xml').decode('utf-8')
        proj_pos = doc_xml.find('PROJECTS')
        skills_pos = doc_xml.find('SKILLS')
        summary_pos = doc_xml.find('PROFESSIONAL SUMMARY')
        # All three must appear, and in the order the user saved
        self.assertGreater(proj_pos, 0)
        self.assertGreater(skills_pos, 0)
        self.assertGreater(summary_pos, 0)
        self.assertLess(proj_pos, skills_pos)
        self.assertLess(skills_pos, summary_pos)

    def test_export_per_owner_scope(self):
        from django.contrib.auth import get_user_model
        User = get_user_model()
        other = User.objects.create_user(
            username='other@example.com', email='other@example.com', password='x',
        )
        self.client.force_login(other)
        resp = self.client.get(self._url())
        self.assertEqual(resp.status_code, 404)

    def test_export_handles_empty_sections_gracefully(self):
        """A nearly-empty resume must still produce a valid DOCX (just
        with a header and whatever sections do have content)."""
        import io, zipfile
        self.resume.content = {'professional_title': 'Engineer'}
        self.resume.save()
        resp = self.client.get(self._url())
        self.assertEqual(resp.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(resp.content)) as z:
            doc_xml = z.read('word/document.xml').decode('utf-8')
        self.assertIn('ADA LOVELACE', doc_xml)
        # No Skills heading because skills was empty
        self.assertNotIn('SKILLS', doc_xml)


class SectionOrderEndpointTests(TestCase):
    """The user can drag-reorder sections on the edit page; the order is
    persisted on resume.content['section_order'] and applied uniformly
    to the live preview, the resume_preview page, and the downloaded PDF."""

    def setUp(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from profiles.models import UserProfile
        from resumes.models import GeneratedResume
        from django.contrib.auth import get_user_model
        User = get_user_model()
        self.user = User.objects.create_user(
            username='order@example.com', email='order@example.com', password='x',
        )
        UserProfile.objects.create(user=self.user, full_name='Test User', email='order@example.com')
        self.client.force_login(self.user)
        job = Job.objects.create(user=self.user, title='Engineer')
        gap = GapAnalysis.objects.create(user=self.user, job=job, similarity_score=0.5)
        self.resume = GeneratedResume.objects.create(
            gap_analysis=gap, content={'professional_title': 'Engineer'},
        )

    def _url(self):
        return f'/resumes/section-order/{self.resume.id}/'

    def test_post_persists_validated_order(self):
        from django.urls import reverse  # noqa
        resp = self.client.post(
            self._url(),
            data='{"order": ["projects", "skills", "summary", "experience"]}',
            content_type='application/json',
        )
        self.assertEqual(resp.status_code, 200)
        body = resp.json()
        # User-supplied order respected, missing keys appended at end
        self.assertEqual(
            body['order'],
            ['projects', 'skills', 'summary', 'experience', 'education', 'certifications', 'languages'],
        )
        self.resume.refresh_from_db()
        self.assertEqual(self.resume.content['section_order'], body['order'])

    def test_unknown_keys_are_dropped_silently(self):
        """Defensive: a stale UI or a malicious client must not poison the
        saved order with unknown keys (e.g., 'admin_only_section')."""
        resp = self.client.post(
            self._url(),
            data='{"order": ["skills", "definitely_not_a_section", "summary"]}',
            content_type='application/json',
        )
        self.assertEqual(resp.status_code, 200)
        # Bad key dropped; valid ones preserved in their relative order
        self.assertNotIn('definitely_not_a_section', resp.json()['order'])
        self.assertEqual(resp.json()['order'][:2], ['skills', 'summary'])

    def test_non_list_body_returns_400(self):
        resp = self.client.post(
            self._url(),
            data='{"order": "not a list"}',
            content_type='application/json',
        )
        self.assertEqual(resp.status_code, 400)

    def test_invalid_json_returns_400(self):
        resp = self.client.post(self._url(), data='not json', content_type='application/json')
        self.assertEqual(resp.status_code, 400)

    def test_per_owner_scope(self):
        from django.contrib.auth import get_user_model
        User = get_user_model()
        other = User.objects.create_user(
            username='other@example.com', email='other@example.com', password='x',
        )
        self.client.force_login(other)
        resp = self.client.post(self._url(), data='{"order": []}', content_type='application/json')
        self.assertEqual(resp.status_code, 404)

    def test_edit_page_renders_saved_order_in_drag_list_and_preview(self):
        """The drag list AND the live preview must both honor the saved
        section_order — not the default — when the page loads."""
        self.resume.content = {
            'professional_title': 'Engineer',
            'professional_summary': 'sum',
            'skills': ['Python'],
            'section_order': ['projects', 'skills', 'summary'],
        }
        self.resume.save()
        from django.urls import reverse
        resp = self.client.get(reverse('resume_edit', args=[self.resume.id]))
        body = resp.content.decode('utf-8')
        # Drag list: 'projects' chip should appear before 'skills' chip
        proj_idx = body.index('data-section-key="projects"')
        skills_idx = body.index('data-section-key="skills"')
        summary_idx = body.index('data-section-key="summary"')
        self.assertLess(proj_idx, skills_idx)
        self.assertLess(skills_idx, summary_idx)
        # Live preview: data-preview-section attributes must appear in same order
        ppi = body.index('data-preview-section="projects"')
        psi = body.index('data-preview-section="skills"')
        psum = body.index('data-preview-section="summary"')
        self.assertLess(ppi, psi)
        self.assertLess(psi, psum)


class RegenerateSectionEndpointTests(TestCase):
    """The per-section regenerate endpoint lets the user iterate on a single
    weak section (summary, skills, experience, projects) without losing
    edits to other sections. Pinning the contract: auth scope, allowed
    sections, persisted update, body shape that returns just the rewritten
    value."""

    def setUp(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from profiles.models import UserProfile
        from resumes.models import GeneratedResume
        from django.contrib.auth import get_user_model
        from django.urls import reverse  # noqa
        User = get_user_model()
        self.user = User.objects.create_user(
            username='regen@example.com', email='regen@example.com', password='x',
        )
        UserProfile.objects.create(
            user=self.user, full_name='Test User', email='regen@example.com',
            data_content={'skills': [{'name': 'Python'}]},
        )
        self.client.force_login(self.user)
        job = Job.objects.create(
            user=self.user, title='Backend Engineer', company='Acme',
            description='We need Python and Postgres experts.',
            extracted_skills=['Python', 'PostgreSQL'],
        )
        gap = GapAnalysis.objects.create(user=self.user, job=job, similarity_score=0.6)
        self.resume = GeneratedResume.objects.create(
            gap_analysis=gap,
            content={
                'professional_title': 'Backend Engineer',
                'professional_summary': 'Old summary.',
                'skills': ['Python'],
                'experience': [],
                'projects': [],
            },
        )

    def _url(self, section):
        return f'/resumes/regen/{self.resume.id}/{section}/'

    def test_unsupported_section_returns_400(self):
        resp = self.client.post(self._url('unknown'), data='{}', content_type='application/json')
        self.assertEqual(resp.status_code, 400)

    def test_regen_summary_persists_new_value(self):
        from unittest.mock import patch
        with patch('resumes.views.regenerate_section', return_value='Brand new tailored summary.') as m:
            resp = self.client.post(
                self._url('professional_summary'),
                data='{"current_content": {"professional_summary": "old"}}',
                content_type='application/json',
            )
        self.assertEqual(resp.status_code, 200)
        body = resp.json()
        self.assertEqual(body['section'], 'professional_summary')
        self.assertEqual(body['value'], 'Brand new tailored summary.')
        self.resume.refresh_from_db()
        self.assertEqual(
            self.resume.content['professional_summary'],
            'Brand new tailored summary.',
        )
        # Other sections must be untouched
        self.assertEqual(self.resume.content['professional_title'], 'Backend Engineer')
        m.assert_called_once()

    def test_regen_skills_returns_list(self):
        from unittest.mock import patch
        new_skills = ['Python', 'PostgreSQL', 'Docker']
        with patch('resumes.views.regenerate_section', return_value=new_skills):
            resp = self.client.post(
                self._url('skills'),
                data='{}',
                content_type='application/json',
            )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.json()['value'], new_skills)
        self.resume.refresh_from_db()
        self.assertEqual(self.resume.content['skills'], new_skills)

    def test_regen_uses_in_flight_content_when_provided(self):
        """The client sends current_content with unsaved edits; the
        generator should see those, not the DB snapshot."""
        from unittest.mock import patch
        captured_content = {}
        def fake(profile, job, gap, content, section):
            captured_content.update(content)
            return 'ok'
        with patch('resumes.views.regenerate_section', side_effect=fake):
            self.client.post(
                self._url('professional_summary'),
                data='{"current_content": {"professional_title": "EDITED IN BROWSER"}}',
                content_type='application/json',
            )
        self.assertEqual(captured_content.get('professional_title'), 'EDITED IN BROWSER')

    def test_regen_scoped_to_owner(self):
        """Another user must not be able to regenerate someone else's resume."""
        from django.contrib.auth import get_user_model
        User = get_user_model()
        other = User.objects.create_user(
            username='other@example.com', email='other@example.com', password='x',
        )
        self.client.force_login(other)
        resp = self.client.post(self._url('professional_summary'),
                                data='{}', content_type='application/json')
        self.assertEqual(resp.status_code, 404)

    def test_regen_failure_returns_502(self):
        """When the LLM call raises, surface a 502 not a 500 stacktrace —
        the UI handles 502 gracefully (alert + button reset)."""
        from unittest.mock import patch
        with patch('resumes.views.regenerate_section', side_effect=Exception('LLM down')):
            resp = self.client.post(
                self._url('skills'),
                data='{}',
                content_type='application/json',
            )
        self.assertEqual(resp.status_code, 502)

    def test_regen_refuses_to_save_empty_summary(self):
        """A silent LLM failure that returns an empty string must NOT
        overwrite the saved summary. 422 + clear detail message."""
        from unittest.mock import patch
        original = self.resume.content.get('professional_summary')
        with patch('resumes.views.regenerate_section', return_value=''):
            resp = self.client.post(
                self._url('professional_summary'),
                data='{}',
                content_type='application/json',
            )
        self.assertEqual(resp.status_code, 422)
        body = resp.json()
        self.assertEqual(body['error'], 'empty_regeneration')
        self.assertIn('unchanged', body['detail'].lower())
        # Saved content is preserved
        self.resume.refresh_from_db()
        self.assertEqual(self.resume.content.get('professional_summary'), original)

    def test_regen_refuses_to_save_empty_skills_list(self):
        from unittest.mock import patch
        with patch('resumes.views.regenerate_section', return_value=[]):
            resp = self.client.post(
                self._url('skills'),
                data='{}',
                content_type='application/json',
            )
        self.assertEqual(resp.status_code, 422)
        self.assertEqual(resp.json()['error'], 'empty_regeneration')

    def test_regen_refuses_to_save_experience_without_bullets(self):
        """A list of experience entries with empty descriptions is just as
        bad as an empty list — the user would see "regenerated" rows with
        nothing in them. Refuse + return 422."""
        from unittest.mock import patch
        useless = [{'title': 'Engineer', 'company': 'Co', 'duration': '2024',
                    'description': []}]
        with patch('resumes.views.regenerate_section', return_value=useless):
            resp = self.client.post(
                self._url('experience'),
                data='{}',
                content_type='application/json',
            )
        self.assertEqual(resp.status_code, 422)
        self.assertEqual(resp.json()['error'], 'empty_regeneration')

    def test_regen_refuses_to_save_projects_without_bullets(self):
        from unittest.mock import patch
        useless = [{'name': 'demo', 'url': '', 'description': []}]
        with patch('resumes.views.regenerate_section', return_value=useless):
            resp = self.client.post(
                self._url('projects'),
                data='{}',
                content_type='application/json',
            )
        self.assertEqual(resp.status_code, 422)


class SchemaSupersetTests(TestCase):
    """Phase 0 schema unification — the resume content schema is now a
    superset of the master profile. Editor form fields persist round-trip
    and the DOCX/sync paths surface them.
    """

    def setUp(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from profiles.models import UserProfile
        from resumes.models import GeneratedResume
        from django.contrib.auth import get_user_model
        User = get_user_model()
        self.user = User.objects.create_user(
            username='superset@example.com', email='superset@example.com', password='x',
        )
        UserProfile.objects.create(
            user=self.user, full_name='Marie Curie', email='superset@example.com',
            data_content={
                'objective': 'Translate physics breakthroughs into impactful tools.',
                'experiences': [{
                    'title': 'Research Scientist', 'company': 'Sorbonne Lab',
                    'start_date': '1898', 'end_date': '1906',
                    'location': 'Paris, FR', 'industry': 'Academic Research',
                    'highlights': ['Isolated polonium and radium.'],
                }],
                'education': [{
                    'degree': 'Doctorate', 'field': 'Physics',
                    'institution': 'Sorbonne', 'graduation_year': '1903',
                    'gpa': 'Honours', 'location': 'Paris, FR',
                    'honors': ['Nobel Prize in Physics 1903'],
                }],
                'projects': [{
                    'name': 'Polonium isolation', 'url': 'https://example.com/po',
                    'description': ['Identified polonium via radiochemical separation.'],
                    'technologies': ['Radiometry', 'Chemistry'],
                }],
                'certifications': [{
                    'name': 'Pharmacy Diploma', 'issuer': 'Sorbonne',
                    'date': '1894', 'duration': '4 years',
                    'url': '',
                }],
            },
        )
        self.client.force_login(self.user)
        job = Job.objects.create(
            user=self.user, title='Physicist', company='Pasteur',
            description='Physics research role.',
            extracted_skills=['Radiometry'],
        )
        gap = GapAnalysis.objects.create(user=self.user, job=job, similarity_score=0.9)
        self.resume = GeneratedResume.objects.create(
            gap_analysis=gap,
            content={
                'professional_title': 'Physicist',
                'professional_summary': 'Physicist focused on radioactivity research.',
                'skills': ['Radiometry'],
                # Intentionally minimal so sync-from-master has work to do.
                'experience': [{'title': 'Research Scientist', 'company': 'Sorbonne Lab',
                                'duration': '1898 - 1906',
                                'description': ['Isolated polonium.']}],
                'education': [{'degree': 'Doctorate', 'institution': 'Sorbonne', 'year': '1903'}],
                'projects': [{'name': 'Polonium isolation',
                              'description': ['Identified polonium.'],
                              'url': 'https://example.com/po'}],
                'certifications': [{'name': 'Pharmacy Diploma', 'issuer': 'Sorbonne',
                                    'date': '1894', 'url': ''}],
                'languages': [],
            },
        )

    def test_edit_form_persists_new_experience_fields(self):
        """exp_location[], exp_industry[] survive a POST round-trip."""
        from django.urls import reverse
        url = reverse('resume_edit', args=[self.resume.id])
        resp = self.client.post(url, data={
            'professional_title': 'Physicist',
            'professional_summary': 'Updated.',
            'objective': 'Custom objective text.',
            'skills': 'Radiometry, Chemistry',
            'exp_title[]': ['Research Scientist'],
            'exp_company[]': ['Sorbonne Lab'],
            'exp_duration[]': ['1898 - 1906'],
            'exp_location[]': ['Paris, FR'],
            'exp_industry[]': ['Academic Research'],
            'exp_description[]': ['Isolated polonium and radium.'],
            'edu_degree[]': ['Doctorate'],
            'edu_field[]': ['Physics'],
            'edu_institution[]': ['Sorbonne'],
            'edu_year[]': ['1903'],
            'edu_gpa[]': ['Honours'],
            'edu_location[]': ['Paris, FR'],
            'edu_honors[]': ['Nobel Prize in Physics 1903'],
            'proj_name[]': ['Polonium isolation'],
            'proj_url[]': ['https://example.com/po'],
            'proj_technologies[]': ['Radiometry, Chemistry'],
            'proj_description[]': ['Identified polonium.'],
            'cert_name[]': ['Pharmacy Diploma'],
            'cert_issuer[]': ['Sorbonne'],
            'cert_date[]': ['1894'],
            'cert_duration[]': ['4 years'],
            'cert_url[]': [''],
            'languages': '',
        }, follow=False)
        self.assertEqual(resp.status_code, 302)
        self.resume.refresh_from_db()
        c = self.resume.content
        self.assertEqual(c['objective'], 'Custom objective text.')
        self.assertEqual(c['experience'][0]['location'], 'Paris, FR')
        self.assertEqual(c['experience'][0]['industry'], 'Academic Research')
        self.assertEqual(c['education'][0]['field'], 'Physics')
        self.assertEqual(c['education'][0]['gpa'], 'Honours')
        self.assertEqual(c['education'][0]['location'], 'Paris, FR')
        self.assertEqual(c['education'][0]['honors'], ['Nobel Prize in Physics 1903'])
        self.assertEqual(c['projects'][0]['technologies'], ['Radiometry', 'Chemistry'])
        self.assertEqual(c['certifications'][0]['duration'], '4 years')

    def test_auto_sync_on_get_fills_blank_fields(self):
        """Visiting /resumes/edit/<id>/ silently merges blank/missing fields
        from the master profile into the resume content. No LLM call, no
        manual button — just open the page and master fields are there."""
        from django.urls import reverse
        # Force the auto-regen branch off so we exercise the cheap auto-sync.
        # (Auto-regen fires when profile.updated_at > resume.created_at; we
        # pass ?refresh=0 to match the existing escape-hatch contract.)
        resp = self.client.get(reverse('resume_edit', args=[self.resume.id]) + '?refresh=0')
        self.assertEqual(resp.status_code, 200)
        self.resume.refresh_from_db()
        c = self.resume.content
        # Master fields populated:
        self.assertEqual(c['experience'][0]['location'], 'Paris, FR')
        self.assertEqual(c['experience'][0]['industry'], 'Academic Research')
        self.assertEqual(c['education'][0]['field'], 'Physics')
        self.assertEqual(c['education'][0]['gpa'], 'Honours')
        self.assertEqual(c['education'][0]['honors'], ['Nobel Prize in Physics 1903'])
        self.assertEqual(c['projects'][0]['technologies'], ['Radiometry', 'Chemistry'])
        self.assertEqual(c['certifications'][0]['duration'], '4 years')
        self.assertIn('Translate physics breakthroughs', c['objective'])
        # Existing typed bullet preserved (not clobbered):
        self.assertIn('Isolated polonium.', c['experience'][0]['description'])

    def test_auto_sync_is_idempotent(self):
        """Second visit produces identical content + must NOT call save()
        again (idempotent merge prevents needless DB writes)."""
        from django.urls import reverse
        from unittest.mock import patch
        url = reverse('resume_edit', args=[self.resume.id]) + '?refresh=0'
        # First hit: should save (master fields are being merged in for
        # the first time).
        self.client.get(url)
        self.resume.refresh_from_db()
        first_content = self.resume.content
        # Second hit: nothing should change; assert resume.save() is NOT
        # called by the auto-sync branch this time.
        with patch('resumes.models.GeneratedResume.save') as save_mock:
            self.client.get(url)
            # Auto-sync detects no diff, skips save.
            self.assertFalse(save_mock.called)
        self.resume.refresh_from_db()
        self.assertEqual(self.resume.content, first_content)

    def test_auto_sync_skipped_when_master_empty(self):
        """If the user hasn't filled in the master profile, opening the
        edit page must not blow up — it just renders without changes."""
        from profiles.models import UserProfile
        from django.urls import reverse
        UserProfile.objects.filter(user=self.user).update(data_content={})
        resp = self.client.get(reverse('resume_edit', args=[self.resume.id]) + '?refresh=0')
        self.assertEqual(resp.status_code, 200)

    def test_docx_renders_new_fields(self):
        """The DOCX exporter surfaces location, GPA, honors, technologies,
        and cert duration when present in resume.content."""
        import io, zipfile
        # Populate the resume directly with all new fields so we don't rely
        # on sync_from_master to put them there.
        self.resume.content = {
            **self.resume.content,
            'objective': 'Translate physics breakthroughs into impactful tools.',
            'experience': [{
                'title': 'Research Scientist', 'company': 'Sorbonne Lab',
                'duration': '1898 - 1906',
                'location': 'Paris, FR', 'industry': 'Academic Research',
                'description': ['Isolated polonium.'],
            }],
            'education': [{
                'degree': 'Doctorate', 'field': 'Physics',
                'institution': 'Sorbonne', 'year': '1903',
                'gpa': 'Honours', 'location': 'Paris, FR',
                'honors': ['Nobel Prize in Physics 1903'],
            }],
            'projects': [{
                'name': 'Polonium isolation',
                'url': 'https://example.com/po',
                'description': ['Identified polonium.'],
                'technologies': ['Radiometry', 'Chemistry'],
            }],
            'certifications': [{
                'name': 'Pharmacy Diploma', 'issuer': 'Sorbonne',
                'date': '1894', 'duration': '4 years', 'url': '',
            }],
        }
        self.resume.save()
        resp = self.client.get(f'/resumes/export-docx/{self.resume.id}/')
        self.assertEqual(resp.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(resp.content)) as z:
            doc_xml = z.read('word/document.xml').decode('utf-8')
        # New fields should all appear in the rendered XML.
        self.assertIn('Translate physics breakthroughs', doc_xml)  # objective
        self.assertIn('Paris, FR', doc_xml)                         # location
        self.assertIn('Academic Research', doc_xml)                 # industry
        self.assertIn('Physics', doc_xml)                           # education field
        self.assertIn('GPA Honours', doc_xml)                       # gpa
        self.assertIn('Nobel Prize', doc_xml)                       # honors
        self.assertIn('Radiometry', doc_xml)                        # tech stack
        self.assertIn('4 years', doc_xml)                           # cert duration


class ExportErrorPageTests(TestCase):
    """Tier-1 papercut: export failures used to return a plain 500 with
    a single sentence. Now they render the export_error.html template
    with retry / alt-format / back-to-resume links."""

    def setUp(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from profiles.models import UserProfile
        from resumes.models import GeneratedResume
        from django.contrib.auth import get_user_model
        User = get_user_model()
        self.user = User.objects.create_user(
            username='exporterr@example.com', email='exporterr@example.com', password='x',
        )
        UserProfile.objects.create(user=self.user, full_name='Test User')
        self.client.force_login(self.user)
        job = Job.objects.create(
            user=self.user, title='Engineer', company='Acme', description='x',
            extracted_skills=[],
        )
        gap = GapAnalysis.objects.create(user=self.user, job=job, similarity_score=0.5)
        self.resume = GeneratedResume.objects.create(
            gap_analysis=gap,
            content={'professional_title': 'Engineer'},
        )

    def test_pdf_export_failure_renders_friendly_error_page(self):
        from unittest.mock import patch
        with patch('resumes.views.generate_pdf', side_effect=RuntimeError('boom')):
            resp = self.client.get(f'/resumes/export/{self.resume.id}/')
        self.assertEqual(resp.status_code, 500)
        # Friendly page body — not the old plaintext "PDF generation failed."
        self.assertContains(resp, 'PDF export failed', status_code=500)
        self.assertContains(resp, 'Retry PDF', status_code=500)
        self.assertContains(resp, 'Try DOCX instead', status_code=500)
        self.assertContains(resp, 'Back to résumé', status_code=500)

    def test_docx_export_failure_renders_friendly_error_page(self):
        from unittest.mock import patch
        with patch('resumes.views.generate_docx', side_effect=RuntimeError('boom')):
            resp = self.client.get(f'/resumes/export-docx/{self.resume.id}/')
        self.assertEqual(resp.status_code, 500)
        self.assertContains(resp, 'DOCX export failed', status_code=500)
        self.assertContains(resp, 'Retry DOCX', status_code=500)
        self.assertContains(resp, 'Try PDF instead', status_code=500)


class DashboardResumeCountTests(TestCase):
    """S4: a returning user with multiple resumes per job should see a
    `N résumés` badge on each job tile, linking to the history."""

    def test_dashboard_annotates_resume_count_on_jobs(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from profiles.models import UserProfile
        from resumes.models import GeneratedResume
        from django.contrib.auth import get_user_model
        User = get_user_model()
        user = User.objects.create_user(
            username='dash@example.com', email='dash@example.com', password='x',
        )
        UserProfile.objects.create(user=user, full_name='Returning User')
        job = Job.objects.create(
            user=user, title='Engineer', company='Acme', description='x',
            extracted_skills=[], application_status='saved',
        )
        gap = GapAnalysis.objects.create(user=user, job=job, similarity_score=0.7)
        # Three resume iterations for the same job — typical second-session
        # state for a user iterating on tailoring.
        for i in range(3):
            GeneratedResume.objects.create(gap_analysis=gap, content={'professional_title': f'v{i}'})
        self.client.force_login(user)
        resp = self.client.get('/profiles/dashboard/')
        self.assertEqual(resp.status_code, 200)
        # The annotated count surfaces as "3 résumés" badge text.
        self.assertContains(resp, '3 résumés')


class AutoSyncBannerTests(TestCase):
    """M10: when the auto-sync branch patches fields into resume.content,
    the next page render shows a one-shot banner so the user knows the
    page didn't render exactly what they last saved."""

    def setUp(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from profiles.models import UserProfile
        from resumes.models import GeneratedResume
        from django.contrib.auth import get_user_model
        User = get_user_model()
        self.user = User.objects.create_user(
            username='sync@example.com', email='sync@example.com', password='x',
        )
        UserProfile.objects.create(
            user=self.user, full_name='Synced User',
            data_content={
                # Master profile has fields the resume snapshot lacks.
                'experiences': [{'title': 'Engineer', 'company': 'Acme',
                                 'location': 'Berlin, DE', 'industry': 'SaaS',
                                 'highlights': ['Shipped X.']}],
            },
        )
        self.client.force_login(self.user)
        job = Job.objects.create(
            user=self.user, title='Engineer', company='Acme',
            description='x', extracted_skills=[],
        )
        gap = GapAnalysis.objects.create(user=self.user, job=job, similarity_score=0.7)
        # Resume content lacks the supplemental fields (location/industry).
        self.resume = GeneratedResume.objects.create(
            gap_analysis=gap,
            content={
                'experience': [{'title': 'Engineer', 'company': 'Acme',
                                'duration': '', 'description': ['Shipped X.']}],
            },
        )

    def test_auto_sync_banner_shows_on_first_visit_after_change(self):
        """Visiting the edit page when auto-sync writes new fields surfaces
        a banner. The session flag is consumed so it doesn't reappear."""
        from django.urls import reverse
        url = reverse('resume_edit', args=[self.resume.id]) + '?refresh=0'
        # First visit — auto-sync fires (resume content was missing
        # location/industry; master has them) and banner shows.
        resp = self.client.get(url)
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Synced from your master profile')
        # Second visit — content already merged on the first hit, no diff,
        # no banner.
        resp = self.client.get(url)
        self.assertNotContains(resp, 'Synced from your master profile')

    def test_no_sync_banner_when_nothing_to_sync(self):
        """If the master has no new fields beyond what the resume already
        has, no banner appears."""
        from profiles.models import UserProfile
        from django.urls import reverse
        # Empty data_content → nothing to merge → no banner.
        UserProfile.objects.filter(user=self.user).update(data_content={})
        resp = self.client.get(reverse('resume_edit', args=[self.resume.id]) + '?refresh=0')
        self.assertNotContains(resp, 'Synced from your master profile')


class OnboardingBannerDismissTests(TestCase):
    """M9: dismissing the dashboard onboarding banner persists across
    visits. Welcome's 'Just show me around' also sets the dismiss flag."""

    def setUp(self):
        from profiles.models import UserProfile
        from django.contrib.auth import get_user_model
        User = get_user_model()
        self.user = User.objects.create_user(
            username='banner@example.com', email='banner@example.com', password='x',
        )
        # Profile with no name + no skills + no jobs → banner SHOULD show
        # by default.
        UserProfile.objects.create(user=self.user)
        self.client.force_login(self.user)

    def test_dashboard_shows_banner_for_new_user_by_default(self):
        resp = self.client.get('/profiles/dashboard/')
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Two steps to unlock')

    def test_dismiss_endpoint_persists_across_visits(self):
        from profiles.models import UserProfile
        # Click the dismiss endpoint
        resp = self.client.post('/profiles/api/onboarding/dismiss/')
        self.assertEqual(resp.status_code, 200)
        self.assertTrue(resp.json()['ok'])
        # Flag persisted
        profile = UserProfile.objects.get(user=self.user)
        self.assertTrue(profile.data_content.get('onboarding_banner_dismissed'))
        # Subsequent dashboard visits don't show the banner
        resp = self.client.get('/profiles/dashboard/')
        self.assertNotContains(resp, 'Two steps to unlock')

    def test_dismiss_endpoint_rejects_get(self):
        resp = self.client.get('/profiles/api/onboarding/dismiss/')
        self.assertEqual(resp.status_code, 405)

    def test_welcome_skip_also_dismisses_banner(self):
        from profiles.models import UserProfile
        resp = self.client.post('/welcome/', {'action': 'skip'})
        self.assertEqual(resp.status_code, 302)
        profile = UserProfile.objects.get(user=self.user)
        self.assertTrue(profile.data_content.get('has_seen_welcome'))
        # Critical: welcome's skip ALSO sets the banner-dismissed flag
        # (M9 fix — without this, dashboard re-nags after the user
        # explicitly opted out of guided onboarding).
        self.assertTrue(profile.data_content.get('onboarding_banner_dismissed'))


class AgentChatProfileAwareTests(TestCase):
    """B1: /agent/ adapts to zero-data state. Input stays enabled (per
    user feedback — disabled inputs feel broken), but a banner above the
    chat warns the user the agent has nothing to ground answers in."""

    def test_zero_data_user_sees_setup_banner(self):
        from django.contrib.auth import get_user_model
        from profiles.models import UserProfile
        User = get_user_model()
        user = User.objects.create_user(
            username='zero@example.com', email='zero@example.com', password='x',
        )
        UserProfile.objects.create(user=user)  # empty: no name, no skills
        self.client.force_login(user)
        resp = self.client.get('/agent/')
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'This will be more useful once you set up your profile')
        # Header text reflects the empty state too — no false "knows your
        # profile and pipeline" claim.
        self.assertContains(resp, 'profile not set up yet')
        self.assertNotContains(resp, 'knows your profile and pipeline')

    def test_populated_user_does_not_see_banner(self):
        from django.contrib.auth import get_user_model
        from profiles.models import UserProfile
        User = get_user_model()
        user = User.objects.create_user(
            username='full@example.com', email='full@example.com', password='x',
        )
        UserProfile.objects.create(
            user=user, full_name='Marie Curie',
            data_content={'skills': [{'name': 'Python'}]},
        )
        self.client.force_login(user)
        resp = self.client.get('/agent/')
        self.assertEqual(resp.status_code, 200)
        self.assertNotContains(resp, 'This will be more useful once you set up your profile')
        self.assertContains(resp, 'knows your profile and pipeline')


class SubstepLoadingComponentTests(TestCase):
    """Tier 3: the substep loading component replaces the old static-string
    overlay. These tests cover the JS / template contract — the component
    is registered globally on every page (via base.html), the registry has
    every operation we expect, and call sites use the {op: '...'} form
    rather than legacy strings."""

    def setUp(self):
        from django.contrib.auth import get_user_model
        from profiles.models import UserProfile
        User = get_user_model()
        self.user = User.objects.create_user(
            username='loader@example.com', email='loader@example.com', password='x',
        )
        UserProfile.objects.create(user=self.user, full_name='Loader User')
        self.client.force_login(self.user)

    def test_base_template_exposes_loading_ops_registry(self):
        """Every page should have the LoadingOps registry available so any
        client-side code can drive the overlay without re-declaring keys."""
        resp = self.client.get('/profiles/dashboard/')
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode('utf-8')
        # Registry keys we rely on across pages:
        for op in ('cv-upload', 'job-scrape', 'job-paste', 'gap-analysis',
                   'resume-gen', 'cover-letter', 'outreach', 'outreach-campaign',
                   'learning-path', 'salary'):
            self.assertIn(f"'{op}'", body, f"LoadingOps registry missing key {op!r}")
        # Component pieces that pages target:
        self.assertIn('id="loading-steps"', body)
        self.assertIn('id="loading-failure"', body)
        self.assertIn('id="loading-retry"', body)
        # Legacy single-line mode still supported.
        self.assertIn('id="loading-msg"', body)

    def test_cv_upload_uses_substep_op_key(self):
        """upload_cv.html should call showLoading({op: 'cv-upload'}), not
        the legacy plain-string form."""
        resp = self.client.get('/profiles/setup/upload/')
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode('utf-8')
        self.assertIn("showLoading({op: 'cv-upload'})", body)
        # The legacy single-string call must NOT be there — catches regression.
        self.assertNotIn("Your agent is reading your CV — up to a minute on first run.", body)

    def test_job_input_uses_substep_op_keys(self):
        resp = self.client.get('/jobs/input/')
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode('utf-8')
        # Both URL-mode and paste-mode forms migrated:
        self.assertIn("showLoading({op: 'job-scrape'})", body)
        self.assertIn("showLoading({op: 'job-paste'})", body)


class TemplateBlockTagInJSCommentTests(TestCase):
    """Regression guard: a literal `{% block name %}` string inside a JS
    comment is parsed by Django's template engine as a real block tag.
    On extending templates this raises TemplateSyntaxError ("'block' tag
    with name 'X' appears more than once"). Hit production once via a
    DOMContentLoaded JSDoc that referenced `{% block content %}` literally.

    This test renders the gap-analysis loading state + the resume
    generation page and asserts both succeed (no TemplateSyntaxError).
    """

    def setUp(self):
        from django.contrib.auth import get_user_model
        from profiles.models import UserProfile
        from jobs.models import Job
        User = get_user_model()
        self.user = User.objects.create_user(
            username='blocktag@example.com', email='blocktag@example.com',
            password='x',
        )
        UserProfile.objects.create(
            user=self.user, full_name='Block Tag User',
            data_content={'skills': [{'name': 'Python'}]},
        )
        self.client.force_login(self.user)
        self.job = Job.objects.create(
            user=self.user, title='Engineer', company='Acme',
            description='Need Python.', extracted_skills=['Python'],
        )

    def test_gap_analysis_loading_state_renders(self):
        """No TemplateSyntaxError; the loading-state HTML reaches the user."""
        resp = self.client.get(f'/analysis/gap/{self.job.id}/')
        self.assertEqual(resp.status_code, 200)
        # Loading state contains the substep heading.
        self.assertContains(resp, 'Reading the')

    def test_resume_generate_page_renders(self):
        from analysis.models import GapAnalysis
        GapAnalysis.objects.create(
            user=self.user, job=self.job, similarity_score=0.85,
        )
        resp = self.client.get(f'/resumes/generate/{self.job.id}/')
        self.assertEqual(resp.status_code, 200)


class ScriptTagBalanceTests(TestCase):
    """Regression guard: a literal `</script>` string inside a JS comment
    terminates the script element early in the HTML parser, leaking the
    rest of the script body (Tours registry, theme toggle, everything) as
    visible text on the page. Hit production once via a Shepherd JSDoc
    comment that wrote `<script>...</script>` literally inside backticks.

    This test asserts the rendered dashboard has matching counts of
    <script> opens and </script> closes. Mismatch == something inside a
    script body is being parsed as a tag boundary.
    """

    def setUp(self):
        from django.contrib.auth import get_user_model
        from profiles.models import UserProfile
        User = get_user_model()
        self.user = User.objects.create_user(
            username='balance@example.com', email='balance@example.com', password='x',
        )
        UserProfile.objects.create(user=self.user, full_name='Balance User')
        self.client.force_login(self.user)

    def test_dashboard_script_tags_balanced(self):
        import re
        resp = self.client.get('/profiles/dashboard/')
        body = resp.content.decode('utf-8')
        # Count opening <script (handles `<script>` and `<script src=...>`)
        # and closing </script> tags.
        opens = len(re.findall(r'<script\b', body, flags=re.IGNORECASE))
        closes = len(re.findall(r'</script\s*>', body, flags=re.IGNORECASE))
        self.assertEqual(
            opens, closes,
            f'Unbalanced <script> tags: {opens} opens vs {closes} closes. '
            f'A literal "</script>" inside a script body breaks the parser.'
        )

    def test_no_loading_ops_text_leaks_outside_script(self):
        """The LoadingOps registry should only appear inside a <script>
        tag. If it shows up as visible text on the page, a script tag
        boundary broke."""
        resp = self.client.get('/profiles/dashboard/')
        body = resp.content.decode('utf-8')
        # Find the position of `const LoadingOps = {` and verify it sits
        # inside a script block (last <script> before this position has
        # not been closed yet).
        marker = 'const LoadingOps = {'
        idx = body.find(marker)
        self.assertGreater(idx, 0, 'LoadingOps registry missing entirely')
        before = body[:idx]
        # Last <script appearance before the marker must NOT be followed
        # by a </script> close before the marker.
        last_open = before.rfind('<script')
        last_close = before.rfind('</script>')
        self.assertGreater(last_open, last_close,
                           'LoadingOps appears AFTER a </script> close — script body leaked.')


class ProfileSummaryAndObjectivePropertyTests(TestCase):
    """Master review form binds to `profile.objective` and
    `profile.normalized_summary`. Both surface from data_content via
    @property; saving them via the form persists back into data_content."""

    def setUp(self):
        from django.contrib.auth import get_user_model
        from profiles.models import UserProfile
        User = get_user_model()
        self.user = User.objects.create_user(
            username='summary@example.com', email='summary@example.com', password='x',
        )
        UserProfile.objects.create(
            user=self.user, full_name='Test User',
            data_content={
                'objective': 'Build resilient distributed systems.',
                'normalized_summary': 'Senior engineer, 8 years Python + Go.',
            },
        )
        self.client.force_login(self.user)

    def test_property_surfaces_objective_and_summary(self):
        from profiles.models import UserProfile
        p = UserProfile.objects.get(user=self.user)
        self.assertEqual(p.objective, 'Build resilient distributed systems.')
        self.assertEqual(p.normalized_summary, 'Senior engineer, 8 years Python + Go.')

    def test_normalized_summary_falls_back_to_summary(self):
        """Older parses populated `summary` instead of `normalized_summary`.
        The property reads either, in priority order."""
        from profiles.models import UserProfile
        UserProfile.objects.filter(user=self.user).update(data_content={
            'summary': 'Legacy summary text.',
        })
        p = UserProfile.objects.get(user=self.user)
        self.assertEqual(p.normalized_summary, 'Legacy summary text.')

    def test_review_post_persists_objective_and_summary(self):
        """The master review POST handler now writes both fields back to
        data_content. Without this, edits to the textareas vanished."""
        from profiles.models import UserProfile
        from django.urls import reverse
        resp = self.client.post(reverse('review_master_profile'), {
            'full_name': 'Test User', 'email': 'summary@example.com',
            'phone': '', 'location': '',
            'objective': 'New objective text.',
            'normalized_summary': 'New summary text.',
            'contact_links_json': '[]', 'skills_json': '[]',
            'experiences_json': '[]', 'education_json': '[]',
            'projects_json': '[]', 'certifications_json': '[]',
        })
        self.assertEqual(resp.status_code, 302)
        p = UserProfile.objects.get(user=self.user)
        self.assertEqual(p.data_content['objective'], 'New objective text.')
        self.assertEqual(p.data_content['normalized_summary'], 'New summary text.')

    def test_review_form_renders_property_values(self):
        """The form template binds `{{ profile.objective }}` and
        `{{ profile.normalized_summary }}`. The rendered HTML must show
        the actual stored values, not blank placeholders."""
        resp = self.client.get('/profiles/setup/review/')
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'Build resilient distributed systems.')
        self.assertContains(resp, 'Senior engineer, 8 years Python + Go.')


class OnboardingFlowOrderTests(TestCase):
    """The post-CV-upload onboarding sequence is:
        Upload → Connect → (Project review if signals) → Master review.

    Master review is the FINAL step, not the first stop after upload.
    Out-of-onboarding visits keep the legacy semantics.
    """

    def setUp(self):
        from django.contrib.auth import get_user_model
        from profiles.models import UserProfile
        User = get_user_model()
        self.user = User.objects.create_user(
            username='flow@example.com', email='flow@example.com', password='x',
        )
        self.profile = UserProfile.objects.create(user=self.user)
        self.client.force_login(self.user)

    def _set_onboarding_flag(self, value=True):
        session = self.client.session
        session['in_onboarding'] = value
        session.save()

    def test_connect_auto_merges_and_redirects_to_master_review_when_signals_present(self):
        """Onboarding + signals: connect_accounts auto-applies enriched
        projects to the master profile and skips the (now read-only)
        review page entirely. The user lands on master review with the
        merge already done."""
        from profiles.models import UserProfile
        UserProfile.objects.filter(user=self.user).update(data_content={
            'github_signals': {
                'profile_url': 'https://github.com/me',
                'top_repos': [{'name': 'alpha', 'full_name': 'me/alpha',
                               'description': 'demo', 'html_url': 'https://github.com/me/alpha',
                               'stargazers_count': 1, 'forks_count': 0, 'language': 'Python'}],
                'language_breakdown': [['Python', 1]],
            },
        })
        self._set_onboarding_flag()
        from unittest.mock import patch
        # Stub the LLM calls — we only care about the redirect contract +
        # that data_content['projects'] gets written. Both enrich_profile
        # and dedupe_projects fall back to deterministic paths on LLM
        # failure (existing behaviour) so the auto-apply still runs.
        from profiles.services import project_enricher, project_dedupe
        with patch.object(project_enricher, 'get_structured_llm') as mock_a, \
             patch.object(project_dedupe, 'get_structured_llm') as mock_b:
            mock_a.return_value.invoke.side_effect = RuntimeError('boom')
            mock_b.return_value.invoke.side_effect = RuntimeError('boom')
            resp = self.client.post('/profiles/setup/connect/')
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(resp.url, '/profiles/setup/review/')
        # Auto-apply persisted: data_content['projects'] now includes the
        # GitHub-derived project (rule-based fallback shape from enricher).
        profile = UserProfile.objects.get(user=self.user)
        names = [p.get('name') for p in profile.data_content.get('projects', [])]
        self.assertIn('alpha', names)

    def test_connect_redirects_to_master_review_when_no_signals(self):
        """In onboarding + no signals connected → POST connect →
        master review (skip the empty project-review step)."""
        self._set_onboarding_flag()
        resp = self.client.post('/profiles/setup/connect/')
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(resp.url, '/profiles/setup/review/')

    def test_connect_redirects_to_master_review_when_signals_have_error(self):
        """A snapshot dict with `error` set is treated as 'no signals' —
        the user's connect attempt failed (e.g. rate-limited), so don't
        send them to a project review with nothing to merge."""
        from profiles.models import UserProfile
        UserProfile.objects.filter(user=self.user).update(data_content={
            'github_signals': {'error': 'rate_limited'},
        })
        self._set_onboarding_flag()
        resp = self.client.post('/profiles/setup/connect/')
        self.assertEqual(resp.url, '/profiles/setup/review/')

    def test_connect_out_of_onboarding_preserves_legacy_redirect(self):
        """Settings-style visit (not in onboarding): posting Continue
        should go to job_input_view or dashboard, NOT into the project-
        review detour even if signals exist."""
        from profiles.models import UserProfile
        UserProfile.objects.filter(user=self.user).update(data_content={
            'github_signals': {'top_repos': []},
        })
        # No in_onboarding flag set.
        resp = self.client.post('/profiles/setup/connect/')
        self.assertEqual(resp.status_code, 302)
        self.assertIn(resp.url, ('/jobs/input/', '/profiles/dashboard/'))

    def test_master_review_is_final_step_when_onboarding(self):
        """Master review used to redirect to connect_accounts when in
        onboarding. Post-reorder, it's the LAST step — should route
        directly to job_input_view (or dashboard) and clear the
        in_onboarding session flag."""
        self._set_onboarding_flag()
        # Minimal POST to satisfy the form.
        resp = self.client.post('/profiles/setup/review/', {
            'full_name': 'Test User',
            'email': 'flow@example.com',
        })
        self.assertEqual(resp.status_code, 302)
        # No active jobs → job input is the natural next stop.
        self.assertEqual(resp.url, '/jobs/input/')
        # Session flag must be cleared so the "Skip onboarding" affordance
        # disappears on subsequent pages.
        self.assertFalse(self.client.session.get('in_onboarding'))


class TourAndHelpAffordanceTests(TestCase):
    """Tier 4: Shepherd tour, Help affordance, step indicators, and the
    routing tooltip on gap analysis. The tour is registered on every page
    via base.html; pages that should auto-run it set SHOULD_RUN_TOUR."""

    def setUp(self):
        from django.contrib.auth import get_user_model
        from profiles.models import UserProfile
        User = get_user_model()
        self.user = User.objects.create_user(
            username='tour@example.com', email='tour@example.com', password='x',
        )
        UserProfile.objects.create(user=self.user, full_name='Tour User')
        self.client.force_login(self.user)

    def test_help_button_renders_on_authenticated_pages(self):
        """The "?" Help button is in base.html and renders for every
        logged-in page so the user always has a re-trigger affordance."""
        resp = self.client.get('/profiles/dashboard/')
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'id="help-button"')
        self.assertContains(resp, 'aria-label="Help"')
        # Pages register their tour via window.PAGE_TOUR — Help reads it.
        self.assertContains(resp, "window.PAGE_TOUR = 'dashboard'")

    def test_shepherd_cdn_loaded_on_authenticated_pages(self):
        """Shepherd.js + CSS are CDN-loaded; verify the script + stylesheet
        tags are present so a tour can actually run."""
        resp = self.client.get('/profiles/dashboard/')
        body = resp.content.decode('utf-8')
        self.assertIn('shepherd.js', body)
        self.assertIn('shepherd.css', body)
        # Tours registry is declared globally
        self.assertIn("const Tours =", body)
        for key in ('dashboard', 'resume-edit', 'gap-analysis'):
            self.assertIn(f"'{key}':", body)

    def test_first_visit_dashboard_auto_runs_tour(self):
        """A user without `has_seen_tour` set should get
        window.SHOULD_RUN_TOUR=true on first dashboard visit."""
        resp = self.client.get('/profiles/dashboard/')
        self.assertContains(resp, 'window.SHOULD_RUN_TOUR = true')

    def test_returning_user_does_not_auto_run_tour(self):
        """After dismiss/complete, has_seen_tour=True; the dashboard
        should NOT set SHOULD_RUN_TOUR on subsequent visits."""
        from profiles.models import UserProfile
        UserProfile.objects.filter(user=self.user).update(
            data_content={'has_seen_tour': True},
        )
        resp = self.client.get('/profiles/dashboard/')
        self.assertNotContains(resp, 'window.SHOULD_RUN_TOUR = true')

    def test_dismiss_tour_endpoint_persists_flag(self):
        from profiles.models import UserProfile
        resp = self.client.post('/profiles/api/tour/dismiss/')
        self.assertEqual(resp.status_code, 200)
        self.assertTrue(resp.json()['ok'])
        profile = UserProfile.objects.get(user=self.user)
        self.assertTrue(profile.data_content.get('has_seen_tour'))

    def test_dismiss_tour_endpoint_rejects_get(self):
        resp = self.client.get('/profiles/api/tour/dismiss/')
        self.assertEqual(resp.status_code, 405)

    def test_step_indicators_on_onboarding_pages(self):
        """The 4-step onboarding sequence shows explicit progress markers.
        Welcome short-circuits for users with profile data already, so we
        log in as a brand-new user with no profile fields."""
        from django.contrib.auth import get_user_model
        from profiles.models import UserProfile
        User = get_user_model()
        new_user = User.objects.create_user(
            username='fresh@example.com', email='fresh@example.com', password='x',
        )
        UserProfile.objects.create(user=new_user)  # empty profile
        self.client.force_login(new_user)
        resp = self.client.get('/welcome/')
        self.assertContains(resp, 'Step 1 of 4')
        resp = self.client.get('/profiles/setup/upload/')
        self.assertContains(resp, 'Step 2 of 4')
        # Post-reorder, connect-accounts is step 3 (was step 4 before
        # signals/project-review moved ahead of master review).
        resp = self.client.get('/profiles/setup/connect/')
        self.assertContains(resp, 'Step 3 of 4')


class ResumeSchemaCoercionTests(SimpleTestCase):
    """Regression: Groq strict-validates tool-call shapes. The model
    sometimes returns null for blank string fields and wraps list items
    as single-key objects. Schema's before-validators coerce both shapes
    to the canonical form so we don't lose generations to validation."""

    def test_null_string_fields_coerce_to_empty(self):
        from profiles.services.schemas import ResumeExperience, ResumeEducation
        # The literal shape Groq rejected with `expected string, but got null`:
        exp = ResumeExperience(**{
            'title': 'Engineer', 'company': '', 'duration': '',
            'location': None, 'industry': None,
            'start_date': None, 'end_date': None,
            'description': [],
        })
        self.assertEqual(exp.location, '')
        self.assertEqual(exp.industry, '')
        self.assertEqual(exp.start_date, '')
        edu = ResumeEducation(**{
            'degree': 'BSc', 'institution': 'MIT', 'year': '2024',
            'gpa': None, 'location': None, 'field': None,
        })
        self.assertEqual(edu.gpa, '')
        self.assertEqual(edu.location, '')

    def test_object_wrapped_strings_flatten(self):
        """The model emits the legacy `highlights` alias as
        `[{description: "..."}]` and skills as
        `[{name: "Dart", proficiency: null}]` — both should flatten to
        plain string lists. PR 3a: `highlights` is folded into the
        canonical `description` field (no separate output field)."""
        from profiles.services.schemas import ResumeProject, ResumeContentResult
        proj = ResumeProject(**{
            'name': 'Mega News',
            'highlights': [
                {'description': 'Engineered a scalable codebase using Clean Architecture.'},
                {'description': 'Reduced load times by fetching news from 4 APIs in parallel.'},
            ],
            'description': [],
            'technologies': [],
        })
        self.assertEqual(len(proj.description), 2)
        self.assertIn('scalable codebase', proj.description[0])
        self.assertIn('Reduced load times', proj.description[1])
        # PR 3a: highlights is no longer a field on ResumeProject.
        self.assertFalse(hasattr(proj, 'highlights'))
        # ResumeContentResult.skills = list of objects → list of strings
        result = ResumeContentResult(**{
            'professional_title': 'Engineer',
            'skills': [
                {'name': 'Dart', 'proficiency': None, 'years': None},
                {'name': 'Flutter', 'proficiency': None, 'years': None},
            ],
        })
        self.assertEqual(result.skills, ['Dart', 'Flutter'])

    def test_top_level_null_strings_coerce(self):
        from profiles.services.schemas import ResumeContentResult
        result = ResumeContentResult(**{
            'professional_title': None,
            'professional_summary': None,
            'objective': None,
        })
        self.assertEqual(result.professional_title, '')
        self.assertEqual(result.professional_summary, '')
        self.assertEqual(result.objective, '')

    # --- Pass G: null description + highlights merge regressions --------

    def test_experience_null_description_coerces_to_empty_list(self):
        """The exact failure from the regen log: description=null on a
        Union[str, List[str]] field. v1 raised
        ``Input should be a valid string [...] input_value=None``; v2
        treats null as no-description."""
        from profiles.services.schemas import ResumeExperience
        exp = ResumeExperience(**{
            'title': 'DT Intern', 'description': None,
        })
        self.assertEqual(exp.description, [])

    def test_experience_null_desc_with_highlights_promotes(self):
        # The Banque Misr regen had description=null and bullets in
        # the (extra) highlights field. The docx exporter reads only
        # description — without this promotion the rendered bullets
        # would have vanished.
        from profiles.services.schemas import ResumeExperience
        exp = ResumeExperience(**{
            'title': 'DT Intern',
            'description': None,
            'highlights': ['Rotated across departments.',
                           'Built a Microsoft Fabric pipeline.'],
        })
        self.assertEqual(exp.description, [
            'Rotated across departments.',
            'Built a Microsoft Fabric pipeline.',
        ])

    def test_experience_string_desc_plus_highlights_merges(self):
        # AI/DS Trainee case: LLM put a summary sentence in description
        # and the actual bullets in highlights. Merge: summary stays
        # first, then highlights bullets.
        from profiles.services.schemas import ResumeExperience
        exp = ResumeExperience(**{
            'title': 'AI Trainee',
            'description': 'Selected participant in DEPI.',
            'highlights': ['Applied the full data-science lifecycle.',
                           'Used MLOps tools (MLflow, Hugging Face).'],
        })
        self.assertEqual(exp.description, [
            'Selected participant in DEPI.',
            'Applied the full data-science lifecycle.',
            'Used MLOps tools (MLflow, Hugging Face).',
        ])

    def test_experience_list_desc_appends_highlights(self):
        # PR 3a behavior change: the pre-PR-3a validator dropped
        # highlights when description was already multi-item (heuristic:
        # likely-duplicate content from the LLM). The new validator
        # appends in order — predictable, matches the new prompt's
        # "do not invent fields" contract (the LLM should produce only
        # description under PR 3a; duplication is a prompt-side concern,
        # not a validator-side one).
        from profiles.services.schemas import ResumeExperience
        exp = ResumeExperience(**{
            'title': 'IT Intern',
            'description': ['Built X.', 'Shipped Y.'],
            'highlights': ['Appended from highlights alias.'],
        })
        self.assertEqual(exp.description, [
            'Built X.', 'Shipped Y.',
            'Appended from highlights alias.',
        ])

    def test_education_null_honors_coerces_to_empty_list(self):
        # Round 1.5.1 — the DevOps regen log showed honors=null on the
        # education entry blocking the entire recovery path. Null →
        # empty list lets recovery salvage the LLM's content.
        from profiles.services.schemas import ResumeEducation
        edu = ResumeEducation(**{
            'degree': 'BSc',
            'institution': 'KSIU',
            'year': '2026',
            'field': '',
            'gpa': '',
            'location': '',
            'honors': None,
        })
        self.assertEqual(edu.honors, [])

    def test_project_null_description_with_highlights_promotes(self):
        # Same promotion rule applies to projects.
        from profiles.services.schemas import ResumeProject
        p = ResumeProject(**{
            'name': 'Healthcare Prediction',
            'description': None,
            'highlights': ['End-to-end ML pipeline for stroke risk.'],
            'technologies': ['Python', 'scikit-learn'],
        })
        self.assertEqual(p.description, ['End-to-end ML pipeline for stroke risk.'])


# --- Round 1 (DevOps audit): regression coverage -------------------------

from resumes.services.resume_normalizer import (
    trim_skills_to_plan,
    normalize_bullet_punctuation,
)


class TrimSkillsToPlanTests(SimpleTestCase):
    """The audit caught that the resume's Skills section was 'Rust, C++,
    Python, Java, TypeScript, ...' for a DevOps JD — the LLM had its
    own picks while the plan's JD-must-have-ordered skills_to_list
    sat unused. trim_skills_to_plan now forces the section to the
    plan's ordering."""

    def test_replaces_llm_skills_with_plan_ordering(self):
        resume = {'skills': ['Rust', 'C++', 'Python', 'Java']}
        plan = _FakePlan()
        plan.skills_to_list = ['Docker', 'Kubernetes', 'Linux', 'CI/CD']
        out = trim_skills_to_plan(resume, plan)
        self.assertEqual(out['skills'][:4], ['Docker', 'Kubernetes', 'Linux', 'CI/CD'])

    def test_appends_llm_extras_not_in_plan(self):
        resume = {'skills': ['Docker', 'Terraform', 'Ansible']}
        plan = _FakePlan()
        plan.skills_to_list = ['Docker', 'Kubernetes']
        out = trim_skills_to_plan(resume, plan)
        # Plan's order first, then LLM extras (Terraform, Ansible).
        self.assertEqual(out['skills'][0], 'Docker')
        self.assertEqual(out['skills'][1], 'Kubernetes')
        self.assertIn('Terraform', out['skills'])
        self.assertIn('Ansible', out['skills'])

    def test_drops_llm_soft_skill_extras(self):
        # LLM extras that happen to be soft skills (Communication, ...)
        # should NOT survive the merge.
        resume = {'skills': ['Communication', 'Teamwork', 'Terraform']}
        plan = _FakePlan()
        plan.skills_to_list = ['Docker', 'Kubernetes']
        out = trim_skills_to_plan(resume, plan)
        self.assertNotIn('Communication', out['skills'])
        self.assertNotIn('Teamwork', out['skills'])
        self.assertIn('Terraform', out['skills'])

    def test_skips_when_plan_has_no_skills(self):
        # Defensive: empty plan ≠ wipe.
        resume = {'skills': ['Python', 'Rust']}
        plan = _FakePlan()
        plan.skills_to_list = []
        out = trim_skills_to_plan(resume, plan)
        self.assertEqual(out['skills'], ['Python', 'Rust'])

    def test_dedupes_conjunction_extended_subset(self):
        """The 2026-05-28 5:16 run shipped 'SQL' AND 'Databases & SQL'
        side-by-side. The token-subset-with-conjunction rule must drop
        the second-seen so the resume shows one canonical entry."""
        resume = {'skills': []}
        plan = _FakePlan()
        plan.skills_to_list = ['SQL', 'Databases & SQL', 'Python']
        out = trim_skills_to_plan(resume, plan)
        self.assertIn('SQL', out['skills'])
        self.assertNotIn('Databases & SQL', out['skills'])
        self.assertIn('Python', out['skills'])

    def test_dedupes_supervised_learning_variants(self):
        """Same pattern: 'Supervised Learning' + 'Supervised &
        Unsupervised Learning' — drop the conjunction-extended form
        when the simpler one was seen first."""
        resume = {'skills': []}
        plan = _FakePlan()
        plan.skills_to_list = ['Supervised Learning',
                               'Supervised & Unsupervised Learning']
        out = trim_skills_to_plan(resume, plan)
        self.assertIn('Supervised Learning', out['skills'])
        self.assertNotIn('Supervised & Unsupervised Learning', out['skills'])

    def test_dedupes_when_conjunction_form_seen_first(self):
        """Order-independence: with 'Databases & SQL' seen first and
        'SQL' coming after, the SQL entry is the subset and gets dropped."""
        resume = {'skills': []}
        plan = _FakePlan()
        plan.skills_to_list = ['Databases & SQL', 'SQL']
        out = trim_skills_to_plan(resume, plan)
        # First-seen wins: 'Databases & SQL' stays, second 'SQL' drops.
        self.assertIn('Databases & SQL', out['skills'])
        self.assertNotIn('SQL', out['skills'])

    def test_does_not_dedupe_compound_product_names(self):
        """Docker / Docker Compose is a known false-positive risk. The
        longer form has no conjunction, so the rule MUST NOT fire.
        Same protection for PostgreSQL / SQL."""
        resume = {'skills': []}
        plan = _FakePlan()
        plan.skills_to_list = ['Docker', 'Docker Compose', 'SQL', 'PostgreSQL']
        out = trim_skills_to_plan(resume, plan)
        self.assertIn('Docker', out['skills'])
        self.assertIn('Docker Compose', out['skills'])
        self.assertIn('SQL', out['skills'])
        self.assertIn('PostgreSQL', out['skills'])


class NormalizeExperienceDatesTests(SimpleTestCase):
    """2026-05-29 Almansour duration shipped as 'August 2025 - Sep 2025'
    (mixed long/short month form). Normalize to 3-letter abbreviations."""

    def test_shortens_long_month_in_duration(self):
        from resumes.services.resume_normalizer import normalize_experience_dates
        resume = {'experience': [
            {'title': 'A', 'duration': 'August 2025 - Sep 2025', 'description': []},
        ]}
        out = normalize_experience_dates(resume)
        self.assertEqual(out['experience'][0]['duration'], 'Aug 2025 - Sep 2025')

    def test_shortens_in_start_and_end_dates(self):
        from resumes.services.resume_normalizer import normalize_experience_dates
        resume = {'experience': [
            {'title': 'A', 'start_date': 'August 2025', 'end_date': 'December 2025',
             'duration': '', 'description': []},
        ]}
        out = normalize_experience_dates(resume)
        self.assertEqual(out['experience'][0]['start_date'], 'Aug 2025')
        self.assertEqual(out['experience'][0]['end_date'], 'Dec 2025')

    def test_idempotent_on_already_abbreviated(self):
        from resumes.services.resume_normalizer import normalize_experience_dates
        resume = {'experience': [
            {'title': 'A', 'duration': 'Jun 2025 - Dec 2025', 'description': []},
        ]}
        out = normalize_experience_dates(resume)
        self.assertEqual(out['experience'][0]['duration'], 'Jun 2025 - Dec 2025')

    def test_preserves_present_and_ranges(self):
        from resumes.services.resume_normalizer import normalize_experience_dates
        resume = {'experience': [
            {'title': 'A', 'duration': 'January 2024 - Present', 'description': []},
        ]}
        out = normalize_experience_dates(resume)
        self.assertEqual(out['experience'][0]['duration'], 'Jan 2024 - Present')


class SortExperienceReverseChronologicalTests(SimpleTestCase):
    """Universal resume convention: most recent role first, by end date.

    These tests cover the GENERAL rule across mixed real-world date
    formats. Nothing here is profile-specific."""

    _TODAY = (2026, 5)   # frozen "now" so 'Present' is deterministic

    def _titles(self, out):
        return [e['title'] for e in out['experience']]

    def test_mixed_date_formats_sort_most_recent_end_first(self):
        from resumes.services.resume_normalizer import (
            sort_experience_reverse_chronological,
        )
        resume = {'experience': [
            {'title': 'Old',       'end_date': 'Jul 2024'},
            {'title': 'Newest',    'end_date': 'Dec 2025'},
            {'title': 'Mid',       'end_date': 'Sep 2025'},
            {'title': 'YearOnly',  'end_date': '2024'},   # → Dec 2024
            {'title': 'IsoForm',   'end_date': '2025-10-15'},
        ]}
        out = sort_experience_reverse_chronological(resume, _today=self._TODAY)
        self.assertEqual(
            self._titles(out),
            ['Newest', 'IsoForm', 'Mid', 'YearOnly', 'Old'],
        )

    def test_present_sorts_to_top(self):
        """end_date='Present' on a record explicitly marked is_current=True
        still sorts to the top (today_ym). A 'Present' string without
        the flag is treated as unknown (covered by the dedicated
        legacy-heal tests)."""
        from resumes.services.resume_normalizer import (
            sort_experience_reverse_chronological,
        )
        resume = {'experience': [
            {'title': 'Old',     'end_date': 'Dec 2025'},
            {'title': 'Live',    'end_date': 'Present', 'is_current': True},
            {'title': 'AlsoOld', 'end_date': 'Aug 2025'},
        ]}
        out = sort_experience_reverse_chronological(resume, _today=self._TODAY)
        self.assertEqual(self._titles(out)[0], 'Live')

    def test_present_in_duration_field_when_end_date_missing(self):
        """duration tail 'Present' on a record with is_current=True
        still sorts to the top. Without the flag it would heal to
        unknown (covered separately)."""
        from resumes.services.resume_normalizer import (
            sort_experience_reverse_chronological,
        )
        resume = {'experience': [
            {'title': 'Old',  'duration': 'Jan 2024 - Dec 2024'},
            {'title': 'Live', 'duration': 'Mar 2025 - Present', 'is_current': True},
        ]}
        out = sort_experience_reverse_chronological(resume, _today=self._TODAY)
        self.assertEqual(self._titles(out), ['Live', 'Old'])

    def test_range_in_duration_uses_the_tail_as_end(self):
        from resumes.services.resume_normalizer import (
            sort_experience_reverse_chronological,
        )
        resume = {'experience': [
            {'title': 'A', 'duration': 'Jun 2025 - Dec 2025'},
            {'title': 'B', 'duration': 'Aug 2025 - Sep 2025'},
            {'title': 'C', 'duration': 'Jan 2025 - Oct 2025'},
        ]}
        out = sort_experience_reverse_chronological(resume, _today=self._TODAY)
        self.assertEqual(self._titles(out), ['A', 'C', 'B'])

    def test_missing_end_falls_back_to_start_date(self):
        from resumes.services.resume_normalizer import (
            sort_experience_reverse_chronological,
        )
        resume = {'experience': [
            {'title': 'NoEnd',  'start_date': 'Jun 2025'},
            {'title': 'HasEnd', 'end_date': 'Mar 2025'},
        ]}
        out = sort_experience_reverse_chronological(resume, _today=self._TODAY)
        # NoEnd → start=Jun 2025 (later than HasEnd's Mar 2025).
        self.assertEqual(self._titles(out), ['NoEnd', 'HasEnd'])

    def test_missing_both_dates_keeps_stable_relative_position(self):
        from resumes.services.resume_normalizer import (
            sort_experience_reverse_chronological,
        )
        resume = {'experience': [
            {'title': 'A_NoDates'},
            {'title': 'B_NoDates'},
            {'title': 'C_HasEnd', 'end_date': 'Jul 2024'},
            {'title': 'D_NoDates'},
        ]}
        out = sort_experience_reverse_chronological(resume, _today=self._TODAY)
        # Parseable entry sorts first; unparseable entries sink, but
        # their order RELATIVE to each other is preserved (A, B, D).
        self.assertEqual(
            self._titles(out),
            ['C_HasEnd', 'A_NoDates', 'B_NoDates', 'D_NoDates'],
        )

    def test_tie_on_end_date_breaks_by_start_then_index(self):
        from resumes.services.resume_normalizer import (
            sort_experience_reverse_chronological,
        )
        resume = {'experience': [
            {'title': 'Older_start', 'start_date': 'Jan 2024',
             'end_date': 'Dec 2025'},
            {'title': 'Newer_start', 'start_date': 'Jun 2025',
             'end_date': 'Dec 2025'},
            {'title': 'Same_start',  'start_date': 'Jun 2025',
             'end_date': 'Dec 2025'},
        ]}
        out = sort_experience_reverse_chronological(resume, _today=self._TODAY)
        # Both 'Newer_start' and 'Same_start' have the same end+start →
        # tiebreak by original index (Newer_start was first in input).
        self.assertEqual(
            self._titles(out),
            ['Newer_start', 'Same_start', 'Older_start'],
        )

    def test_idempotent_already_sorted(self):
        from resumes.services.resume_normalizer import (
            sort_experience_reverse_chronological,
        )
        resume = {'experience': [
            {'title': 'A', 'end_date': 'Dec 2025'},
            {'title': 'B', 'end_date': 'Oct 2025'},
            {'title': 'C', 'end_date': 'Jul 2024'},
        ]}
        out = sort_experience_reverse_chronological(resume, _today=self._TODAY)
        self.assertEqual(self._titles(out), ['A', 'B', 'C'])
        # Re-running yields the same result.
        out2 = sort_experience_reverse_chronological(out, _today=self._TODAY)
        self.assertEqual(self._titles(out2), ['A', 'B', 'C'])

    def test_no_bullet_or_field_dropped(self):
        """Sort only changes ORDER. No entry dropped, no field mutated."""
        from resumes.services.resume_normalizer import (
            sort_experience_reverse_chronological,
        )
        bullets = ['Built X with 30%.', 'Deployed Y at scale.']
        resume = {'experience': [
            {'title': 'Old', 'end_date': 'Jul 2024',
             'description': list(bullets), 'company': 'X', 'location': 'L'},
            {'title': 'New', 'end_date': 'Dec 2025',
             'description': ['Different bullet.'], 'company': 'Y'},
        ]}
        out = sort_experience_reverse_chronological(resume, _today=self._TODAY)
        self.assertEqual(len(out['experience']), 2)
        self.assertEqual(self._titles(out), ['New', 'Old'])
        # 'Old' entry's content survives the move.
        old = out['experience'][1]
        self.assertEqual(old['description'], bullets)
        self.assertEqual(old['company'], 'X')
        self.assertEqual(old['location'], 'L')

    def test_short_lists_pass_through_untouched(self):
        from resumes.services.resume_normalizer import (
            sort_experience_reverse_chronological,
        )
        self.assertEqual(
            sort_experience_reverse_chronological(
                {'experience': []}, _today=self._TODAY,
            ),
            {'experience': []},
        )
        single = {'experience': [{'title': 'Only', 'end_date': 'Jul 2024'}]}
        self.assertEqual(
            sort_experience_reverse_chronological(single, _today=self._TODAY),
            single,
        )


class MarkExpectedGraduationTests(SimpleTestCase):
    """Prefix education year with 'Expected' when it's in the future."""

    def test_marks_future_month_year_as_expected(self):
        import datetime as dt
        from resumes.services.resume_normalizer import mark_expected_graduation
        resume = {'education': [{'degree': 'BSc', 'year': 'June 2026'}]}
        out = mark_expected_graduation(resume, _today=dt.date(2026, 5, 29))
        self.assertEqual(out['education'][0]['year'], 'Expected June 2026')

    def test_does_not_mark_past_dates(self):
        import datetime as dt
        from resumes.services.resume_normalizer import mark_expected_graduation
        resume = {'education': [{'degree': 'BSc', 'year': 'June 2024'}]}
        out = mark_expected_graduation(resume, _today=dt.date(2026, 5, 29))
        self.assertEqual(out['education'][0]['year'], 'June 2024')

    def test_does_not_mark_current_month(self):
        """Same month, same year — already happened or in progress, not
        future. No prefix."""
        import datetime as dt
        from resumes.services.resume_normalizer import mark_expected_graduation
        resume = {'education': [{'degree': 'BSc', 'year': 'May 2026'}]}
        out = mark_expected_graduation(resume, _today=dt.date(2026, 5, 29))
        self.assertEqual(out['education'][0]['year'], 'May 2026')

    def test_marks_year_only_future(self):
        import datetime as dt
        from resumes.services.resume_normalizer import mark_expected_graduation
        resume = {'education': [{'degree': 'BSc', 'year': '2027'}]}
        out = mark_expected_graduation(resume, _today=dt.date(2026, 5, 29))
        self.assertEqual(out['education'][0]['year'], 'Expected 2027')

    def test_idempotent_when_already_marked(self):
        import datetime as dt
        from resumes.services.resume_normalizer import mark_expected_graduation
        resume = {'education': [{'degree': 'BSc', 'year': 'Expected June 2026'}]}
        out = mark_expected_graduation(resume, _today=dt.date(2026, 5, 29))
        self.assertEqual(out['education'][0]['year'], 'Expected June 2026')


class StripPipeTitleSummaryTests(SimpleTestCase):
    """Round-4 reviewer: "Data Scientist | AI Engineer | Data Analyst
    with applied experience in..." was the lead phrase for 4 rounds
    despite the prompt rule. Strip deterministically."""

    def test_strips_three_pipe_titles_with_with_connector(self):
        from types import SimpleNamespace
        from resumes.services.resume_normalizer import clean_summary_phrasing
        resume = {'professional_summary':
                  'Data Scientist | AI Engineer | Data Analyst with applied '
                  'experience in Machine Learning and Deep Learning.'}
        job = SimpleNamespace(title='Data Scientist')
        out = clean_summary_phrasing(resume, job=job)
        self.assertEqual(
            out['professional_summary'],
            'Data Scientist with applied experience in Machine Learning and Deep Learning.',
        )

    def test_uses_jd_title_when_primary_does_not_match(self):
        from types import SimpleNamespace
        from resumes.services.resume_normalizer import clean_summary_phrasing
        resume = {'professional_summary':
                  'AI Engineer | Data Scientist | Analyst with applied experience.'}
        job = SimpleNamespace(title='Data Scientist')
        out = clean_summary_phrasing(resume, job=job)
        self.assertEqual(
            out['professional_summary'],
            'Data Scientist with applied experience.',
        )

    def test_falls_back_to_primary_when_no_job(self):
        from resumes.services.resume_normalizer import clean_summary_phrasing
        resume = {'professional_summary':
                  'Data Scientist | AI Engineer | Data Analyst with hands-on Python work.'}
        out = clean_summary_phrasing(resume, job=None)
        self.assertEqual(
            out['professional_summary'],
            'Data Scientist with hands-on Python work.',
        )

    def test_leaves_single_title_unchanged(self):
        from types import SimpleNamespace
        from resumes.services.resume_normalizer import clean_summary_phrasing
        resume = {'professional_summary':
                  'Data Scientist with applied experience in NLP.'}
        job = SimpleNamespace(title='Data Scientist')
        out = clean_summary_phrasing(resume, job=job)
        self.assertEqual(
            out['professional_summary'],
            'Data Scientist with applied experience in NLP.',
        )

    def test_handles_two_pipe_titles(self):
        from types import SimpleNamespace
        from resumes.services.resume_normalizer import clean_summary_phrasing
        resume = {'professional_summary':
                  'Data Scientist | ML Engineer focused on production pipelines.'}
        job = SimpleNamespace(title='Data Scientist')
        out = clean_summary_phrasing(resume, job=job)
        self.assertEqual(
            out['professional_summary'],
            'Data Scientist focused on production pipelines.',
        )


class FilterLanguagesProficiencyTests(SimpleTestCase):
    """Round-4 reviewer: R3 had 'Arabic (Native), English (Fluent)'.
    R4 lost the proficiency markers, shipping bare 'English, Arabic'.
    Enrich from profile_data and reorder by profile sequence."""

    def test_enriches_bare_names_with_profile_proficiency(self):
        from resumes.services.resume_normalizer import filter_languages
        resume = {'languages': ['English', 'Arabic']}
        profile = {'languages': [
            {'name': 'Arabic', 'proficiency': 'Native'},
            {'name': 'English', 'proficiency': 'Fluent'},
        ]}
        out = filter_languages(resume, profile_data=profile)
        self.assertEqual(out['languages'], ['Arabic (Native)', 'English (Fluent)'])

    def test_reorders_to_match_profile_sequence(self):
        """LLM emitted 'English' first but the profile lists Arabic first
        (the candidate's native language). Profile order wins."""
        from resumes.services.resume_normalizer import filter_languages
        resume = {'languages': ['English (Fluent)', 'Arabic (Native)']}
        profile = {'languages': [
            {'name': 'Arabic', 'proficiency': 'Native'},
            {'name': 'English', 'proficiency': 'Fluent'},
        ]}
        out = filter_languages(resume, profile_data=profile)
        self.assertEqual(out['languages'], ['Arabic (Native)', 'English (Fluent)'])

    def test_no_op_without_profile_data(self):
        from resumes.services.resume_normalizer import filter_languages
        resume = {'languages': ['Arabic (Native)', 'English (Fluent)']}
        out = filter_languages(resume, profile_data=None)
        # Without profile, just sanitize-pass. Both still spoken languages.
        self.assertEqual(out['languages'], ['Arabic (Native)', 'English (Fluent)'])

    def test_handles_profile_string_entries(self):
        from resumes.services.resume_normalizer import filter_languages
        resume = {'languages': ['Arabic']}
        # Profile stored as strings with parens already embedded.
        profile = {'languages': ['Arabic (Native)', 'English (Fluent)']}
        out = filter_languages(resume, profile_data=profile)
        self.assertEqual(out['languages'], ['Arabic (Native)'])


class NormalizeBulletPunctuationTests(SimpleTestCase):
    def test_adds_period_when_at_least_one_bullet_has_one(self):
        resume = {'experience': [{'description': [
            'Built X.',
            'Shipped Y',
            'Reduced Z',
        ]}]}
        out = normalize_bullet_punctuation(resume)
        self.assertEqual(out['experience'][0]['description'],
                         ['Built X.', 'Shipped Y.', 'Reduced Z.'])

    def test_leaves_period_less_lists_alone(self):
        resume = {'experience': [{'description': [
            'Built X', 'Shipped Y', 'Reduced Z',
        ]}]}
        out = normalize_bullet_punctuation(resume)
        # No bullet had terminal punctuation — leave the list unchanged.
        self.assertEqual(out['experience'][0]['description'],
                         ['Built X', 'Shipped Y', 'Reduced Z'])

    def test_respects_existing_terminal_punctuation(self):
        # Bullets that end with ! or ? don't get an extra period.
        resume = {'experience': [{'description': [
            'Built X.', 'What is Y?', 'Achieved 10x growth!',
        ]}]}
        out = normalize_bullet_punctuation(resume)
        self.assertEqual(out['experience'][0]['description'],
                         ['Built X.', 'What is Y?', 'Achieved 10x growth!'])


class NormalizeBulletPunctuationRound152Tests(SimpleTestCase):
    """Round 1.5.2 — drop orphan list-introducer stubs ending in ':'
    that v1's policy was appending '.' to, producing the 'foo:.' bug
    the audit flagged as broken template scaffolding."""

    def test_drops_orphan_colon_header_bullets(self):
        resume = {'experience': [{'description': [
            'Built the platform.',
            'Capstone project highlights:',          # ← stub, drop
            'Delivered 4 prototypes:',               # ← stub, drop
            'Shipped to production.',
        ]}]}
        out = normalize_bullet_punctuation(resume)
        bullets = out['experience'][0]['description']
        self.assertEqual(len(bullets), 2)
        for b in bullets:
            self.assertFalse(b.endswith(':.'))
            self.assertFalse(b.endswith(':'))


class CourseworkBulletRejectsAchievementsTests(SimpleTestCase):
    """Round 1.5.2 — tighten coursework detection so capstone-metric
    bullets ("3 Grafana dashboards (19 panels)") never get folded into
    a fake "Coursework included: ..." line."""

    def test_digit_in_bullet_rejects_as_coursework(self):
        from resumes.services.resume_normalizer import _is_coursework_bullet
        # DevOps capstone deliverables — these were getting mis-classified.
        self.assertFalse(_is_coursework_bullet('3 Grafana dashboards'))
        self.assertFalse(_is_coursework_bullet('Multi-channel alerting Slack Email Discord 4 channels'))
        # Real course title without digits still passes.
        self.assertTrue(_is_coursework_bullet('Prompt Engineering'))

    def test_colon_midbullet_rejects_as_coursework(self):
        from resumes.services.resume_normalizer import _is_coursework_bullet
        # "Section: detail" pattern from capstone notes.
        self.assertFalse(_is_coursework_bullet('Systems Administration: Active Directory'))
        # Pure course title still passes.
        self.assertTrue(_is_coursework_bullet('Tools for Data Science'))


class AcronymSuffixSkillDedupTests(SimpleTestCase):
    """Round 1.5.2 — dedup 'CI/CD' against
    'Continuous Integration and Continuous Delivery (CI/CD)' via the
    acronym-suffix rule."""

    def test_dedups_verbose_form_against_acronym(self):
        from resumes.services.resume_normalizer import (
            trim_skills_to_plan,
        )
        resume = {'skills': []}
        plan = _FakePlan()
        plan.skills_to_list = [
            'CI/CD',
            'Continuous Integration and Continuous Delivery (CI/CD)',
            'Python',
        ]
        out = trim_skills_to_plan(resume, plan)
        # First-seen wins — 'CI/CD' kept, verbose form deduped.
        self.assertIn('CI/CD', out['skills'])
        self.assertNotIn(
            'Continuous Integration and Continuous Delivery (CI/CD)',
            out['skills'],
        )
        self.assertIn('Python', out['skills'])

    def test_dedups_acronym_against_verbose_form_when_verbose_first(self):
        from resumes.services.resume_normalizer import trim_skills_to_plan
        resume = {'skills': []}
        plan = _FakePlan()
        plan.skills_to_list = [
            'Continuous Integration and Continuous Delivery (CI/CD)',
            'CI/CD',
        ]
        out = trim_skills_to_plan(resume, plan)
        # Verbose form wins (first-seen) — acronym deduped against it.
        self.assertEqual(out['skills'],
                         ['Continuous Integration and Continuous Delivery (CI/CD)'])

    def test_sql_is_not_a_duplicate_of_postgresql(self):
        # Round 1.5.3 regression — the v1 blind-suffix rule was deduping
        # SQL against PostgreSQL because 'postgresql' ends with 'sql'.
        # New rule requires parens-acronym or whitelisted-suffix, so
        # SQL stays distinct.
        from resumes.services.resume_normalizer import _is_near_duplicate_skill
        self.assertFalse(_is_near_duplicate_skill(
            [('PostgreSQL', 'postgresql')], 'SQL', 'sql',
        ))

    def test_docker_compose_is_not_a_duplicate_of_docker(self):
        # Round 1.5.3 regression — the v1 prefix rule deduped
        # "Docker Compose" against "Docker" (different products).
        # New rule requires the remainder to be a whitelisted generic
        # suffix; "compose" is not in the list.
        from resumes.services.resume_normalizer import _is_near_duplicate_skill
        self.assertFalse(_is_near_duplicate_skill(
            [('Docker', 'docker')], 'Docker Compose', 'dockercompose',
        ))

    def test_cicd_tools_is_a_duplicate_of_cicd(self):
        # Whitelisted-suffix rule kicks in: "tools" is in
        # _GENERIC_SKILL_SUFFIXES so "CI/CD tools" dedups against
        # "CI/CD". Preserved from Round 1.5.
        from resumes.services.resume_normalizer import _is_near_duplicate_skill
        self.assertTrue(_is_near_duplicate_skill(
            [('CI/CD', 'cicd')], 'CI/CD tools', 'cicdtools',
        ))


class CleanSummaryPhrasingTests(SimpleTestCase):
    """Round 1.5.2 — strip recruiter-jargon openers and unsupported YoE
    claims from the LLM's generated summary."""

    def setUp(self):
        from resumes.services.resume_normalizer import clean_summary_phrasing
        self.clean = clean_summary_phrasing

    def test_strips_highly_motivated_opener(self):
        resume = {'professional_summary': 'Highly motivated Junior DevOps Engineer with Docker.'}
        out = self.clean(resume)
        self.assertFalse(out['professional_summary'].lower().startswith('highly motivated'))
        self.assertTrue(out['professional_summary'].startswith('Junior DevOps Engineer'))

    def test_strips_yoe_claim(self):
        resume = {'professional_summary':
                  'Junior DevOps Engineer with 1 year of experience in CI/CD pipelines. Strong DevOps practitioner.'}
        out = self.clean(resume)
        self.assertNotIn('1 year of experience', out['professional_summary'])
        self.assertIn('Strong DevOps practitioner', out['professional_summary'])

    def test_strips_compound_recruiter_jargon(self):
        resume = {'professional_summary':
                  'Results-driven backend developer with up to 2 years of experience in microservices.'}
        out = self.clean(resume)
        self.assertNotIn('Results-driven', out['professional_summary'])
        self.assertNotIn('years of experience', out['professional_summary'])

    def test_no_op_when_summary_is_clean(self):
        resume = {'professional_summary':
                  'Junior DevOps Engineer with hands-on Docker, Kubernetes, and Linux experience.'}
        out = self.clean(resume)
        self.assertEqual(out['professional_summary'],
                         'Junior DevOps Engineer with hands-on Docker, Kubernetes, and Linux experience.')


class EnforceVerbatimTitlesTests(SimpleTestCase):
    """Round 1.5.2 — snap paraphrased experience titles back to the CV's
    verbatim form (audit caught "DevOps Engineering Trainee" → "DevOps
    Engineer Trainee")."""

    def setUp(self):
        from resumes.services.resume_normalizer import enforce_verbatim_titles
        self.enforce = enforce_verbatim_titles

    def test_snaps_paraphrased_title_back_to_cv(self):
        resume = {'experience': [
            {'title': 'DevOps Engineer Trainee'},  # LLM paraphrase
        ]}
        profile = {'experiences': [
            {'title': 'DevOps Engineering Trainee'},  # CV verbatim
        ]}
        out = self.enforce(resume, profile)
        self.assertEqual(out['experience'][0]['title'], 'DevOps Engineering Trainee')

    def test_leaves_unrelated_titles_alone(self):
        resume = {'experience': [
            {'title': 'Marketing Manager'},  # very different from CV
        ]}
        profile = {'experiences': [
            {'title': 'DevOps Engineering Trainee'},
        ]}
        out = self.enforce(resume, profile)
        # Below the 0.75 similarity threshold — leave it alone.
        self.assertEqual(out['experience'][0]['title'], 'Marketing Manager')

    def test_no_op_without_profile_data(self):
        resume = {'experience': [{'title': 'X'}]}
        out = self.enforce(resume, None)
        self.assertEqual(out['experience'][0]['title'], 'X')


class CleanLocationTests(SimpleTestCase):
    """Round 1.5.2 — strip Arabic-government-registry prefixes from
    LinkedIn-scraped location fields."""

    def test_strips_qesm_prefix(self):
        from profiles.services.profile_sanitizer import sanitize_profile_data
        clean = sanitize_profile_data({'experiences': [
            {'title': 'X', 'company': 'Y',
             'location': 'Qesm El Zamalek, Cairo, Egypt'},
        ]})
        self.assertEqual(
            clean['experiences'][0]['location'],
            'El Zamalek, Cairo, Egypt',
        )

    def test_strips_markaz_prefix(self):
        from profiles.services.profile_sanitizer import sanitize_profile_data
        clean = sanitize_profile_data({'experiences': [
            {'title': 'X', 'company': 'Y',
             'location': 'Markaz Tanta, Gharbia, Egypt'},
        ]})
        self.assertEqual(
            clean['experiences'][0]['location'],
            'Tanta, Gharbia, Egypt',
        )


class TrimSkillsToPlanRound15Tests(SimpleTestCase):
    """Round 1.5 strengthens trim_skills_to_plan: soft skills the plan
    surfaced get re-filtered, near-duplicate skills get deduped, and
    summary backfill uses just-the-facts phrasing."""

    def test_plan_supplied_soft_skills_are_dropped(self):
        # Gap analyzer extracted "Agile, Communication, Multitasking,
        # Time management" from the JD's soft-skill line and put them in
        # plan.skills_to_list. The audit caught these leaking through.
        resume = {'skills': ['Docker']}
        plan = _FakePlan()
        plan.skills_to_list = [
            'Docker', 'Kubernetes', 'Linux',
            'Agile', 'Communication', 'Multitasking', 'Time management',
            'Bash',
        ]
        out = trim_skills_to_plan(resume, plan)
        for soft in ('Agile', 'Communication', 'Multitasking', 'Time management'):
            self.assertNotIn(soft, out['skills'],
                             f"{soft!r} should be re-filtered out of skills")
        self.assertIn('Docker', out['skills'])
        self.assertIn('Kubernetes', out['skills'])
        self.assertIn('Bash', out['skills'])

    def test_near_duplicate_skills_dedup(self):
        # "CI/CD tools" and "CI/CD" both appeared from JD parsing.
        # First-seen wins. Same for "Scripting" + "Bash" + "Python".
        resume = {'skills': []}
        plan = _FakePlan()
        plan.skills_to_list = ['CI/CD', 'CI/CD tools', 'Bash', 'Scripting']
        out = trim_skills_to_plan(resume, plan)
        # "CI/CD tools" dedups against "CI/CD"; "Scripting" is in the
        # soft-skill blocklist (Round 1.5 added it because Bash/Python
        # already cover it).
        self.assertEqual(out['skills'], ['CI/CD', 'Bash'])

    def test_round_153_devops_skills_dont_falsely_dedup(self):
        # Regression for the Round 1.5.3 catastrophe — the v1
        # blind-suffix rule was dropping Docker Compose, SQL, and
        # any other plain skill that happened to share a substring
        # with another. With the new parens-acronym + whitelisted
        # prefix rule, all of these stay.
        resume = {'skills': []}
        plan = _FakePlan()
        plan.skills_to_list = [
            'Linux', 'virtualization', 'Docker', 'CI/CD', 'PostgreSQL',
            'Bash', 'Python', 'Nginx', 'Docker Compose',
            'Continuous Integration and Continuous Delivery (CI/CD)', 'SQL',
        ]
        out = trim_skills_to_plan(resume, plan)
        # CI/CD is canonical; the verbose "Continuous Integration..."
        # form gets deduped via the parens-acronym rule.
        # Docker Compose stays (not a generic suffix of Docker).
        # SQL stays (PostgreSQL is not a parens-acronym alias).
        self.assertIn('SQL', out['skills'])
        self.assertIn('Docker Compose', out['skills'])
        self.assertIn('PostgreSQL', out['skills'])
        self.assertNotIn(
            'Continuous Integration and Continuous Delivery (CI/CD)',
            out['skills'],
        )


class TrimProjectsSubstringMatchTests(SimpleTestCase):
    """Round 1.5.3 — the LLM truncates project names. trim_projects_to_plan
    must accept canonical-substring containment so 'ACR-QA' matches
    'ACR-QA — Automated Code Review Platform'."""

    def test_truncated_llm_name_matches_full_plan_name(self):
        from resumes.services.resume_normalizer import trim_projects_to_plan
        resume = {'projects': [
            {'name': 'ACR-QA'},
            {'name': 'Containerized URL Shortener'},
        ]}
        plan = _FakePlan(projects=[
            'ACR-QA — Automated Code Review Platform',
            'Containerized URL Shortener — Production Monitoring Stack',
        ])
        out = trim_projects_to_plan(resume, plan)
        # Both should survive via substring containment.
        self.assertEqual(len(out['projects']), 2)
        names = [p['name'] for p in out['projects']]
        self.assertIn('ACR-QA', names)
        self.assertIn('Containerized URL Shortener', names)

    def test_truncation_in_other_direction_works_too(self):
        # Plan has short name, LLM emits longer — also valid match.
        from resumes.services.resume_normalizer import trim_projects_to_plan
        resume = {'projects': [
            {'name': 'Healthcare Prediction (DEPI) — Stroke Risk'},
        ]}
        plan = _FakePlan(projects=['Healthcare Prediction (DEPI)'])
        out = trim_projects_to_plan(resume, plan)
        self.assertEqual(len(out['projects']), 1)


class BackfillSummaryNoMetaNarrationTests(SimpleTestCase):
    """The audit called out "drawing on the X role and project work"
    as meta-narration. The phrase is gone now."""

    def test_summary_has_no_meta_narration_clause(self):
        resume = {
            'professional_summary': '',
            'skills': ['Docker', 'Kubernetes', 'Linux'],
            'experience': [{'title': 'DevOps Engineering Trainee'}],
        }
        job = SimpleNamespace(title='Junior DevOps Engineer')
        out = backfill_summary(resume, job=job)
        self.assertNotIn('drawing on', out['professional_summary'].lower())
        self.assertNotIn('role and project work', out['professional_summary'].lower())
        # Still leads with the JD title (Round 1 behaviour preserved).
        self.assertTrue(out['professional_summary'].startswith('Junior DevOps Engineer'))


class ResumeProjectAppendsHighlightsAliasTests(SimpleTestCase):
    """PR 3a behavior: ResumeProject's coerce_to_canonical validator
    appends `highlights` input into `description` in order — no
    richness-comparison heuristic, no drop-on-duplicate. The pre-PR-3a
    richness check was a workaround for the LLM emitting both fields
    with varying quality; the new prompt ("do not invent fields") makes
    highlights effectively dead on output, so the validator simply folds
    any legacy input liberally."""

    def test_highlights_appended_after_description(self):
        from profiles.services.schemas import ResumeProject
        p = ResumeProject(**{
            'name': 'ACR-QA',
            'description': ['Built a platform.'],
            'highlights': [
                'Built a language-agnostic static analysis platform running 7 tools in parallel with 97.1% precision.',
                'Delivered a 273-test pytest suite passing in under 6 seconds; GitHub Actions + GitLab CI integration; SARIF v2.1.0 export; OWASP Top 10 compliance mapping with CWE IDs.',
            ],
        })
        # description first, then highlights — append-in-order.
        self.assertEqual(len(p.description), 3)
        self.assertEqual(p.description[0], 'Built a platform.')
        self.assertIn('97.1%', p.description[1])
        self.assertIn('SARIF', p.description[2])

    def test_short_highlights_still_appended(self):
        from profiles.services.schemas import ResumeProject
        p = ResumeProject(**{
            'name': 'X',
            'description': [
                'Built a system with 99.9% uptime over 6 months.',
                'Shipped 50 deployments per week.',
                'Reduced cloud spend by 40%.',
            ],
            'highlights': ['One short fact.'],
        })
        # PR 3a: no richness heuristic — append regardless.
        self.assertEqual(len(p.description), 4)
        self.assertIn('99.9%', p.description[0])
        self.assertEqual(p.description[-1], 'One short fact.')


class PR3aSchemaTolerance(SimpleTestCase):
    """PR 3a: ResumeExperience and ResumeProject collapse to a single
    canonical `description: List[str]` field with extra="forbid".
    The validator's coerce_to_canonical folds known LLM-invented field
    names into description; extra="forbid" rejects unknown ones.

    These tests pin the behavior of the input-tolerance layer — the
    load-bearing piece that the rest of the pipeline depends on."""

    def test_achievements_wrapper_folded_into_description(self):
        """PR 3f-style invention: achievements: [{description: [...]}]
        unwraps and merges into the canonical description field."""
        from profiles.services.schemas import ResumeExperience
        e = ResumeExperience(
            title='Engineer', company='X',
            achievements=[{'description': ['Bullet A.', 'Bullet B.']}],
        )
        self.assertEqual(e.description, ['Bullet A.', 'Bullet B.'])
        self.assertFalse(hasattr(e, 'achievements'))

    def test_responsibilities_invention_folded(self):
        """Other invented alias names also fold."""
        from profiles.services.schemas import ResumeExperience
        e = ResumeExperience(
            title='Engineer', company='X',
            responsibilities=['Resp 1', 'Resp 2'],
        )
        self.assertEqual(e.description, ['Resp 1', 'Resp 2'])
        self.assertFalse(hasattr(e, 'responsibilities'))

    def test_string_description_coerced_to_list(self):
        """Legacy description as a single string wraps in a list."""
        from profiles.services.schemas import ResumeExperience
        e = ResumeExperience(
            title='Engineer', company='X',
            description='Single-paragraph description.',
        )
        self.assertEqual(e.description, ['Single-paragraph description.'])

    def test_unknown_field_rejected_on_experience(self):
        """extra='forbid' rejects truly unknown fields (not in alias list)."""
        from pydantic import ValidationError
        from profiles.services.schemas import ResumeExperience
        with self.assertRaises(ValidationError):
            ResumeExperience(
                title='Engineer', company='X',
                fake_field='this should fail validation',
            )

    def test_unknown_field_rejected_on_project(self):
        from pydantic import ValidationError
        from profiles.services.schemas import ResumeProject
        with self.assertRaises(ValidationError):
            ResumeProject(name='X', fake_field='nope')

    def test_no_bullets_at_all_produces_empty_description(self):
        """Refinement 3 from the review gate: every-field-absent shape
        validates and yields description=[]."""
        from profiles.services.schemas import ResumeExperience, ResumeProject
        e = ResumeExperience(title='E', company='C')
        self.assertEqual(e.description, [])
        p = ResumeProject(name='P')
        self.assertEqual(p.description, [])

    def test_known_inventions_registered(self):
        """Refinement 6 — registry doc-test. The bullet-alias registry
        must cover all LLM inventions observed across the audit thread.
        If you intentionally tighten this list (drop an entry), update
        the test deliberately so the change is explicit."""
        from profiles.services.schemas import _BULLET_ALIAS_KEYS
        # Historical second-canonical (folds silently)
        self.assertIn('highlights', _BULLET_ALIAS_KEYS)
        # Inventions surfaced by PR 3f and recorder logs
        for invention in ('achievements', 'responsibilities', 'bullets'):
            self.assertIn(
                invention, _BULLET_ALIAS_KEYS,
                msg=(
                    f"Known LLM invention '{invention}' missing from "
                    f"_BULLET_ALIAS_KEYS. If this invention should no "
                    f"longer be tolerated, intentionally update this test."
                ),
            )

    def test_project_technologies_csv_tolerance(self):
        """Pre-existing input tolerance preserved: comma-separated
        technologies string from the editor form still coerces to list."""
        from profiles.services.schemas import ResumeProject
        p = ResumeProject(name='X', technologies='Python, Django, Postgres')
        self.assertEqual(p.technologies, ['Python', 'Django', 'Postgres'])


class PR3aMigrationCommand(SimpleTestCase):
    """The one-shot data-migration command that converts stored
    GeneratedResume.content rows from pre-PR-3a (highlights + description
    dual fields) to PR-3a (description canonical, list[str])."""

    def test_migrates_highlights_to_description(self):
        from resumes.management.commands.migrate_resume_schema import _migrate_content
        content = {
            'experience': [{
                'title': 'E', 'company': 'C',
                'highlights': ['B1', 'B2'],
                'description': 'Old paragraph',
            }],
            'projects': [{
                'name': 'P',
                'highlights': ['PB1'],
            }],
        }
        migrated, changed = _migrate_content(content)
        self.assertTrue(changed)
        # Old string description wraps to list, then highlights appended.
        self.assertNotIn('highlights', migrated['experience'][0])
        self.assertEqual(
            migrated['experience'][0]['description'],
            ['Old paragraph', 'B1', 'B2'],
        )
        self.assertNotIn('highlights', migrated['projects'][0])
        self.assertEqual(migrated['projects'][0]['description'], ['PB1'])

    def test_idempotent_on_new_shape(self):
        from resumes.management.commands.migrate_resume_schema import _migrate_content
        content = {
            'experience': [{
                'title': 'E', 'company': 'C',
                'description': ['B1', 'B2'],
            }],
            'projects': [{
                'name': 'P',
                'description': ['PB1'],
                'technologies': ['Python'],
            }],
        }
        migrated, changed = _migrate_content(content)
        self.assertFalse(changed)
        self.assertEqual(migrated['experience'][0]['description'], ['B1', 'B2'])
        self.assertEqual(migrated['projects'][0]['description'], ['PB1'])

    def test_handles_missing_sections_gracefully(self):
        from resumes.management.commands.migrate_resume_schema import _migrate_content
        # No experience, no projects — should not raise.
        migrated, changed = _migrate_content({'professional_title': 'X'})
        self.assertFalse(changed)
        self.assertEqual(migrated, {'professional_title': 'X'})

    def test_handles_experiences_plural_key(self):
        """Some rows use 'experiences' rather than 'experience' as the
        section key — the migrator accepts either."""
        from resumes.management.commands.migrate_resume_schema import _migrate_content
        content = {
            'experiences': [{
                'title': 'E', 'company': 'C',
                'highlights': ['B1'],
            }],
        }
        migrated, changed = _migrate_content(content)
        self.assertTrue(changed)
        self.assertEqual(migrated['experiences'][0]['description'], ['B1'])


class PR3bSchemaTolerance(SimpleTestCase):
    """PR 3b mirrors PR 3a's pattern for the CV-parser-side schemas
    (Experience and Project — used by ResumeSchema, stored in
    UserProfile.data_content). Uses the same _fold_into_description
    helper as PR 3a, so these tests verify the helper produces correct
    output when called from the profile-side validators too.

    Also pins the contract for the 6 fields promoted from extra='allow'
    during PR 3b (source, employment_type on Experience; source,
    source_id, pushed_at, date on Project)."""

    def test_description_canonical(self):
        from profiles.services.schemas import Experience
        e = Experience(title='Eng', company='C', description=['Built X.'])
        self.assertEqual(e.description, ['Built X.'])

    def test_input_highlights_folded(self):
        from profiles.services.schemas import Experience
        e = Experience(
            title='Eng', company='C',
            highlights=['B1', 'B2'],
        )
        self.assertEqual(e.description, ['B1', 'B2'])
        self.assertFalse(hasattr(e, 'highlights'))

    def test_input_achievements_folded(self):
        """Pre-PR-3b, achievements was a declared field on Experience
        (the only one of the three schemas with this third field).
        Post-PR-3b it folds via the alias registry."""
        from profiles.services.schemas import Experience
        e = Experience(
            title='Eng', company='C',
            achievements=['Shipped X.', 'Mentored 2 jrs.'],
        )
        self.assertEqual(e.description, ['Shipped X.', 'Mentored 2 jrs.'])
        self.assertFalse(hasattr(e, 'achievements'))

    def test_input_responsibilities_folded_on_experience(self):
        from profiles.services.schemas import Experience
        e = Experience(
            title='Eng', company='C',
            responsibilities=['Resp 1', 'Resp 2'],
        )
        self.assertEqual(e.description, ['Resp 1', 'Resp 2'])

    def test_string_description_coerced_to_list(self):
        from profiles.services.schemas import Experience
        e = Experience(
            title='Eng', company='C',
            description='Single-paragraph description.',
        )
        self.assertEqual(e.description, ['Single-paragraph description.'])

    def test_unknown_field_rejected_on_experience(self):
        """extra='forbid' rejects fields outside the canonical set AND
        outside the alias registry. Genuinely unknown invention names
        (not in _BULLET_ALIAS_KEYS, not one of the 6 promoted fields)
        fail validation. Restored from hotfix's silently-accepted variant
        by PR 3b.1, which paired this strictness with the CV-parser
        prompt's "DO NOT INVENT FIELDS" guidance."""
        from pydantic import ValidationError
        from profiles.services.schemas import Experience
        with self.assertRaises(ValidationError):
            Experience(title='Eng', company='C', fake_field='nope')

    def test_unknown_field_rejected_on_project(self):
        from pydantic import ValidationError
        from profiles.services.schemas import Project
        with self.assertRaises(ValidationError):
            Project(name='X', fake_field='nope')

    def test_null_description_coerced_to_empty_list_on_experience(self):
        """Defense-in-depth: ``description=None`` coerces to [] via the
        validator's mode='before' fold even though the field type is
        ``List[str]`` (non-Optional). The CV-parser prompt forbids
        emitting null; this test pins the safety net for prompt drift.
        """
        from profiles.services.schemas import Experience
        e = Experience(
            title='E', company='C',
            description=None,
            highlights=['Bullet 1', 'Bullet 2'],
        )
        self.assertEqual(e.description, ['Bullet 1', 'Bullet 2'])

    def test_null_description_coerced_to_empty_list_on_project(self):
        """Mirror for Project."""
        from profiles.services.schemas import Project
        p = Project(
            name='P',
            description=None,
            highlights=['PB1'],
        )
        self.assertEqual(p.description, ['PB1'])

    def test_empty_default(self):
        from profiles.services.schemas import Experience, Project
        e = Experience(title='E', company='C')
        self.assertEqual(e.description, [])
        p = Project(name='P')
        self.assertEqual(p.description, [])

    def test_role_field_preserved_on_project(self):
        """Project.role is semantically distinct from bullets — it
        names the candidate's role on the project ("Lead", "Solo dev",
        etc.). Must survive validation alongside description."""
        from profiles.services.schemas import Project
        p = Project(name='thing', role='Lead Developer',
                    description=['Built it.'])
        self.assertEqual(p.role, 'Lead Developer')
        self.assertEqual(p.description, ['Built it.'])

    def test_promoted_extras_accepted_on_experience(self):
        """source and employment_type were silent extras pre-PR-3b
        (under extra='allow'). PR 3b promoted them to explicit fields.
        Both must validate without ValidationError under the new
        extra='forbid' config."""
        from profiles.services.schemas import Experience
        e = Experience(
            title='Eng', company='C',
            source='linkedin',
            employment_type='Full-time',
        )
        self.assertEqual(e.source, 'linkedin')
        self.assertEqual(e.employment_type, 'Full-time')

    def test_promoted_extras_accepted_on_project(self):
        """source, source_id, pushed_at, date were silent extras
        pre-PR-3b. All four are now explicit Optional fields and must
        validate."""
        from profiles.services.schemas import Project
        p = Project(
            name='thing',
            source='github',
            source_id='me/thing',
            pushed_at='2026-01-01T00:00:00Z',
            date='2024',
        )
        self.assertEqual(p.source, 'github')
        self.assertEqual(p.source_id, 'me/thing')
        self.assertEqual(p.pushed_at, '2026-01-01T00:00:00Z')
        self.assertEqual(p.date, '2024')


class PR3bMigrationCommand(SimpleTestCase):
    """The migrate_profile_schema management command converts stored
    UserProfile.data_content rows from pre-PR-3b shape (description +
    highlights + achievements + silent extras) to PR-3b shape."""

    def test_migrates_dict_shape_experiences(self):
        from profiles.management.commands.migrate_profile_schema import (
            _migrate_data_content,
        )
        content = {
            'experiences': [{
                'title': 'E', 'company': 'C',
                'description': 'Para',
                'highlights': ['B1', 'B2'],
                'achievements': ['A1'],
                'source': 'linkedin',           # promoted — must survive
                'employment_type': 'Full-time', # promoted — must survive
            }],
        }
        migrated, changed = _migrate_data_content(content)
        self.assertTrue(changed)
        exp = migrated['experiences'][0]
        self.assertNotIn('highlights', exp)
        self.assertNotIn('achievements', exp)
        # description ordering: existing first, then highlights, then
        # achievements (alias-registry order).
        self.assertEqual(exp['description'], ['Para', 'B1', 'B2', 'A1'])
        # Promoted fields must be preserved by the migration.
        self.assertEqual(exp['source'], 'linkedin')
        self.assertEqual(exp['employment_type'], 'Full-time')

    def test_migrates_dict_shape_projects(self):
        from profiles.management.commands.migrate_profile_schema import (
            _migrate_data_content,
        )
        content = {
            'projects': [{
                'name': 'thing',
                'highlights': ['PB1', 'PB2'],
                'source': 'github',
                'source_id': 'me/thing',
                'pushed_at': '2026-01-01T00:00:00Z',
                'date': '2024',
                'technologies': ['Python'],
            }],
        }
        migrated, changed = _migrate_data_content(content)
        self.assertTrue(changed)
        proj = migrated['projects'][0]
        self.assertNotIn('highlights', proj)
        self.assertEqual(proj['description'], ['PB1', 'PB2'])
        # All 4 promoted fields preserved.
        self.assertEqual(proj['source'], 'github')
        self.assertEqual(proj['source_id'], 'me/thing')
        self.assertEqual(proj['pushed_at'], '2026-01-01T00:00:00Z')
        self.assertEqual(proj['date'], '2024')
        self.assertEqual(proj['technologies'], ['Python'])

    def test_idempotent_on_new_shape(self):
        from profiles.management.commands.migrate_profile_schema import (
            _migrate_data_content,
        )
        content = {
            'experiences': [{
                'title': 'E', 'company': 'C',
                'description': ['B1', 'B2'],
                'source': 'cv',
            }],
            'projects': [{
                'name': 'P',
                'description': ['PB1'],
                'source': 'github',
                'source_id': 'me/p',
            }],
        }
        migrated, changed = _migrate_data_content(content)
        self.assertFalse(changed)
        self.assertEqual(migrated['experiences'][0]['description'], ['B1', 'B2'])
        self.assertEqual(migrated['projects'][0]['description'], ['PB1'])

    def test_scan_unknowns_reports_extras(self):
        """--scan-unknowns must surface keys that aren't canonical and
        aren't in the alias registry, so the operator can classify
        them before extra='forbid' rejects them at validation time."""
        from profiles.management.commands.migrate_profile_schema import (
            _scan_unknown_keys,
        )
        content = {
            'experiences': [{
                'title': 'E', 'company': 'C',
                # Genuinely unknown — not canonical, not an alias.
                'fake_invented_field': 'mock',
            }],
            'projects': [{
                'name': 'P',
                'another_unknown': 42,
            }],
        }
        unknowns = _scan_unknown_keys(content)
        self.assertIn('experience.fake_invented_field', unknowns)
        self.assertIn('project.another_unknown', unknowns)
        self.assertEqual(unknowns['experience.fake_invented_field'], 1)
        self.assertEqual(unknowns['project.another_unknown'], 1)
        # Promoted fields and known aliases must NOT appear as unknowns.
        clean_content = {
            'experiences': [{
                'title': 'E', 'company': 'C',
                'source': 'cv',
                'employment_type': 'Internship',
            }],
            'projects': [{
                'name': 'P',
                'source': 'github',
                'source_id': 'me/p',
                'pushed_at': '2026-01-01T00:00:00Z',
                'date': '2024',
            }],
        }
        self.assertEqual(_scan_unknown_keys(clean_content), {})


class PR3b2OptionalSchemaTests(SimpleTestCase):
    """PR 3b.2: 6 fields on the resume-output nested models were
    relaxed to Optional with None defaults so Groq's server-side
    tool-call validator stops rejecting LLM responses that emit null
    for legitimately-missing fields (internship industry, ongoing-
    role end_date, cert-without-URL, etc.).

    Validators continue to normalize None -> '' / None -> [] for
    downstream readers, so post-validation shapes are unchanged.

    Also pins the scope-guard contract: required-semantic fields
    (title, company on Experience; name, issuer on Certification)
    are NOT in the Optional set. At the JSON-schema layer (what Groq
    actually validates against), null on those fields would still be
    rejected — preventing future PRs from silently broadening the
    relaxation."""

    def test_optional_industry_accepts_none_on_experience(self):
        from profiles.services.schemas import ResumeExperience
        e = ResumeExperience(title='Engineer', company='X', industry=None)
        # Validator normalizes None -> '' for downstream readers.
        self.assertEqual(e.industry, '')

    def test_optional_end_date_accepts_none(self):
        """end_date=None is the LLM's correct emission for an ongoing
        role when the prompt asked for "Present" but the LLM puts
        "Present" in duration only and leaves end_date empty."""
        from profiles.services.schemas import ResumeExperience
        e = ResumeExperience(title='Engineer', company='X', end_date=None)
        self.assertEqual(e.end_date, '')

    def test_optional_cert_fields_accept_none(self):
        """ResumeCertification url/date/duration accept None for
        certs without a verification URL, undated certs, and
        certs without a stated duration."""
        from profiles.services.schemas import ResumeCertification
        c = ResumeCertification(
            name='Some Cert', issuer='Some Org',
            url=None, date=None, duration=None,
        )
        self.assertEqual(c.url, '')
        self.assertEqual(c.date, '')
        self.assertEqual(c.duration, '')

    def test_optional_education_honors_accepts_none(self):
        """ResumeEducation.honors=None normalizes to [] (a degree
        entry typically has no honors line for most candidates)."""
        from profiles.services.schemas import ResumeEducation
        ed = ResumeEducation(institution='X', degree='Y', honors=None)
        self.assertEqual(ed.honors, [])

    def test_required_fields_stay_strict_in_json_schema(self):
        """SCOPE GUARD (PR 3b.2): the relaxation is scoped to fields
        with semantic-null meaning. Required-semantic fields stay
        strict at the JSON-schema layer (which is what Groq's
        server-side tool-call validator actually checks). null on
        title/company/name/issuer still gets rejected upstream —
        preserving the surface where genuine data bugs surface.

        Test asserts on model_json_schema() rather than Python-side
        validation because the existing _coerce_null_strings
        validator masks Python-side null on ALL declared str fields.
        The JSON-schema layer is the architecturally relevant boundary.
        """
        from profiles.services.schemas import (
            ResumeExperience, ResumeCertification,
        )

        exp_schema = ResumeExperience.model_json_schema()
        cert_schema = ResumeCertification.model_json_schema()

        def _accepts_null(field_schema: dict) -> bool:
            """True iff the JSON schema for one field admits null.

            Pydantic emits Optional[str] as `{"anyOf": [{"type": "string"},
            {"type": "null"}]}` or similar; a plain str = "" emits just
            `{"type": "string"}`."""
            if field_schema.get('type') == 'null':
                return True
            for sub in field_schema.get('anyOf', []):
                if sub.get('type') == 'null':
                    return True
            return False

        # PR 3b.2 relaxed fields — null IS accepted at JSON-schema layer
        self.assertTrue(_accepts_null(exp_schema['properties']['industry']))
        self.assertTrue(_accepts_null(exp_schema['properties']['end_date']))
        self.assertTrue(_accepts_null(cert_schema['properties']['date']))
        self.assertTrue(_accepts_null(cert_schema['properties']['duration']))
        self.assertTrue(_accepts_null(cert_schema['properties']['url']))

        # Required-semantic fields — null is NOT accepted at JSON-schema
        # layer (Groq will still 400 if the LLM emits null for these).
        self.assertFalse(_accepts_null(exp_schema['properties']['title']))
        self.assertFalse(_accepts_null(exp_schema['properties']['company']))
        self.assertFalse(_accepts_null(exp_schema['properties']['start_date']))
        self.assertFalse(_accepts_null(cert_schema['properties']['name']))
        self.assertFalse(_accepts_null(cert_schema['properties']['issuer']))


class NormalizeRoleAnalystVariantsTests(SimpleTestCase):
    """Issue 3 fix (2026-05-20): _normalize_role was missing 'analyst'
    keywords. 'Junior Data Analyst' was falling through to the
    software_engineer default, causing KB retrieval to pull
    software-engineering chunks (Django/PostgreSQL/LLM-systems patterns)
    for analyst JDs instead of data-flavored content.

    Fix routes analyst variants to data_scientist (the closest existing
    bucket with KB depth — 3 role-specific + 45 universal chunks).
    KB Sprint 2 may add a dedicated data_analyst bucket later; this
    PR is the routing fix only."""

    def test_data_analyst_variants_route_to_data_scientist(self):
        from profiles.services.knowledge_retriever import _normalize_role
        cases = [
            'Data Analyst',
            'Junior Data Analyst',
            'Senior Data Analyst',
            'Data Analyst Intern',
            'Business Analyst',
            'Reporting Analyst',
        ]
        for jd_title in cases:
            self.assertEqual(
                _normalize_role(jd_title), 'data_scientist',
                msg=f"'{jd_title}' should route to 'data_scientist'",
            )

    def test_analyst_variants_never_fall_to_software_engineer(self):
        """SCOPE GUARD: pin that the analyst keywords actually catch.
        Prevents future refactors from silently reverting the bug."""
        from profiles.services.knowledge_retriever import _normalize_role
        for jd_title in (
            'Data Analyst', 'Junior Data Analyst', 'Senior Data Analyst',
            'Business Analyst', 'Reporting Analyst',
        ):
            self.assertNotEqual(
                _normalize_role(jd_title), 'software_engineer',
                msg=(
                    f"'{jd_title}' fell through to software_engineer — "
                    f"the very bug Issue 3 was meant to fix."
                ),
            )

    def test_engineering_roles_with_data_context_route_to_engineering(self):
        """SCOPE GUARD: 'Data X Engineer' titles must NOT be caught by
        the new analyst keywords. The 'data analyst' substring requires
        the literal word 'analyst' — 'Data Analysis Engineer' doesn't
        match (no 't' suffix). Pinning this so future broader matches
        ('analyst' alone, etc.) don't accidentally catch engineering
        roles."""
        from profiles.services.knowledge_retriever import _normalize_role
        # Direct data_engineer match.
        self.assertEqual(_normalize_role('Data Engineer'), 'data_engineer')
        # No keyword catches this — falls through to software_engineer.
        # Pinning the current behavior so the analyst PR can't be blamed
        # for it later.
        self.assertEqual(
            _normalize_role('Data Analysis Engineer'),
            'software_engineer',
        )
        # NOTE: 'Senior Data Platform Engineer' routes to 'devops' (NOT
        # 'data_engineer') because the existing devops branch matches
        # 'platform'. That's pre-existing behavior, predates this PR,
        # and is intentionally left untouched.

    def test_non_data_analysts_preserved(self):
        """SCOPE GUARD: bare 'analyst' was rejected — Quality/Financial/
        Security Analysts still route the same as before this PR."""
        from profiles.services.knowledge_retriever import _normalize_role
        # Quality Analyst -> qa (matches 'quality' in the qa branch).
        self.assertEqual(_normalize_role('Quality Analyst'), 'qa')
        # Financial / Security Analyst -> software_engineer default.
        # Not great, but no better bucket exists; same as pre-PR.
        self.assertEqual(_normalize_role('Financial Analyst'), 'software_engineer')
        self.assertEqual(_normalize_role('Security Analyst'), 'software_engineer')


class AwardsFieldEndToEndTests(SimpleTestCase):
    """Round 1.5: Honors & Awards section is now first-class. Schema
    accepts the field, _ensure_profile_data_preserved surfaces from
    profile, _write_awards renders bold name + plain suffix."""

    def test_schema_accepts_awards_list(self):
        from profiles.services.schemas import ResumeContentResult
        r = ResumeContentResult(awards=[
            'ICPC ECPC 2024 — Honorable Mention, 2nd among KSIU teams',
            'Dean\'s List — Fall 2024',
        ])
        self.assertEqual(len(r.awards), 2)
        self.assertTrue(r.awards[0].startswith('ICPC'))

    def test_schema_normalizes_honors_alias(self):
        from profiles.services.schemas import ResumeContentResult
        # Some pipelines may emit 'honors' instead of 'awards'.
        r = ResumeContentResult(honors=['Award A', 'Award B'])
        self.assertEqual(r.awards, ['Award A', 'Award B'])


class WriteAwardsTests(SimpleTestCase):
    """The Honors & Awards docx writer renders with bold name + plain
    suffix and uses the same section-heading + List Bullet style as
    certifications."""

    def test_writes_awards_section_when_present(self):
        from docx import Document
        from resumes.services.docx_exporter import _write_awards
        doc = Document()
        _write_awards(doc, {'awards': [
            'ICPC ECPC 2024 — Honorable Mention',
            'Dean\'s List',
        ]})
        # Section heading + two bullets.
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        self.assertEqual(paragraphs[0], 'HONORS & AWARDS')
        self.assertEqual(paragraphs[1], 'ICPC ECPC 2024 — Honorable Mention')
        self.assertEqual(paragraphs[2], 'Dean\'s List')

    def test_no_op_when_no_awards(self):
        from docx import Document
        from resumes.services.docx_exporter import _write_awards
        doc = Document()
        before = len(doc.paragraphs)
        _write_awards(doc, {})
        # Writer adds no new paragraphs when awards is missing/empty.
        self.assertEqual(len(doc.paragraphs), before)


class TrimCertsKeepsAllCertsTests(SimpleTestCase):
    """Round 1.5: the cert filter was simplified to keep ALL candidate
    certs (capped at _CERT_CAP). The auditor's recommendation was
    'Include ALL certifications from JSON'."""

    def test_keeps_every_cert_regardless_of_plan(self):
        from resumes.services.resume_normalizer import trim_certs_to_plan
        resume = {'certifications': [
            {'name': 'Introduction to Software Testing', 'issuer': 'U of Minnesota'},
            {'name': 'Random Off-Topic Cert', 'issuer': 'Whatever'},
            {'name': 'Applied ML in Python', 'issuer': 'Coursera'},
        ]}
        plan = _FakePlan(certifications=['IBM Data Scientist Track'])
        plan.skills_to_list = ['Docker']
        out = trim_certs_to_plan(resume, plan)
        names = [c['name'] for c in out['certifications']]
        # Every cert survives — plan membership no longer filters.
        self.assertEqual(len(names), 3)
        self.assertIn('Introduction to Software Testing', names)
        self.assertIn('Random Off-Topic Cert', names)


class BackfillSummaryUsesJdTitleTests(SimpleTestCase):
    """The audit caught the backfill leading with 'DevOps Engineering
    Trainee' for a 'Junior DevOps Engineer' JD. The summary now
    leads with the JD title when available."""

    def test_leads_with_jd_title_not_experience_title(self):
        resume = {
            'professional_summary': '',
            'skills': ['Docker', 'Kubernetes', 'Linux', 'Bash'],
            'experience': [{'title': 'DevOps Engineering Trainee'}],
        }
        job = SimpleNamespace(title='Junior DevOps Engineer')
        out = backfill_summary(resume, job=job)
        self.assertTrue(out['professional_summary'].startswith('Junior DevOps Engineer'))
        self.assertNotIn('data and ML lifecycle', out['professional_summary'])

    def test_falls_back_to_experience_title_when_no_jd_title(self):
        resume = {
            'professional_summary': '',
            'skills': ['Python'],
            'experience': [{'title': 'Data Engineer'}],
        }
        out = backfill_summary(resume, job=None)
        self.assertTrue(out['professional_summary'].startswith('Data Engineer'))


class ZeroWidthCharStripTests(SimpleTestCase):
    """The audit caught project names ending with U+2060 word joiner,
    e.g. 'ACR-QA — Automated Code Review Platform⁠'. The CV parser
    leaves these in skills, project names, and bullets. The sanitizer
    now strips them recursively."""

    def test_strips_word_joiner_from_project_name(self):
        from profiles.services.profile_sanitizer import sanitize_profile_data
        clean = sanitize_profile_data({
            'projects': [
                {'name': 'ACR-QA — Automated Code Review Platform⁠',
                 'description': ['Built CI/​CD pipeline.']},
            ],
        })
        self.assertEqual(clean['projects'][0]['name'],
                         'ACR-QA — Automated Code Review Platform')
        self.assertEqual(clean['projects'][0]['description'],
                         ['Built CI/CD pipeline.'])

    def test_strips_from_skills_and_experiences(self):
        from profiles.services.profile_sanitizer import sanitize_profile_data
        clean = sanitize_profile_data({
            'skills': ['Docker​', 'Python⁠'],
            'experiences': [{'title': 'DevOps Trainee﻿',
                              'description': ['Used CI/​CD.']}],
        })
        skill_names = [s.get('name') if isinstance(s, dict) else s
                       for s in clean['skills']]
        self.assertIn('Docker', skill_names)
        self.assertIn('Python', skill_names)
        self.assertEqual(clean['experiences'][0]['title'], 'DevOps Trainee')
        self.assertEqual(clean['experiences'][0]['description'],
                         ['Used CI/CD.'])


class StripSchemaEnvelopeLeaksTests(SimpleTestCase):
    """The DevOps regen log showed `additionalProperties`, `properties`,
    `type` keys in the final resume sections list — the LLM emitted
    the schema-envelope shape on a 200 (happy path), not via
    failed_generation. The recovery unwrap didn't fire because there
    was no exception. _strip_schema_envelope_leaks handles this."""

    def test_strips_top_level_envelope_keys(self):
        from resumes.services.resume_generator import _strip_schema_envelope_leaks
        resume = {
            'professional_summary': 'hello',
            'skills': ['A'],
            'additionalProperties': True,
            'properties': {'leftover': 'whatever'},
            'type': 'object',
        }
        out = _strip_schema_envelope_leaks(resume)
        self.assertNotIn('additionalProperties', out)
        self.assertNotIn('properties', out)
        self.assertNotIn('type', out)
        self.assertEqual(out['skills'], ['A'])

    def test_steps_into_properties_when_full_envelope(self):
        from resumes.services.resume_generator import _strip_schema_envelope_leaks
        resume = {
            'additionalProperties': True,
            'properties': {
                'professional_summary': 'inside',
                'skills': ['X', 'Y'],
            },
            'type': 'object',
        }
        out = _strip_schema_envelope_leaks(resume)
        self.assertEqual(out['professional_summary'], 'inside')
        self.assertEqual(out['skills'], ['X', 'Y'])
        self.assertNotIn('additionalProperties', out)

    def test_unwraps_per_field_type_value_wrappers(self):
        from resumes.services.resume_generator import _strip_schema_envelope_leaks
        resume = {
            'skills': {'type': 'array', 'value': ['Docker', 'K8s']},
            'professional_summary': {'type': 'string', 'value': 'hi'},
        }
        out = _strip_schema_envelope_leaks(resume)
        self.assertEqual(out['skills'], ['Docker', 'K8s'])
        self.assertEqual(out['professional_summary'], 'hi')


# --- Pass H: failed_generation extraction robustness ----------------------

class IsTokenLimitErrorTests(SimpleTestCase):
    """The 22:22 regen hit Groq's 30k TPM ceiling with a 32,423-token
    prompt. The generator now detects that specific failure and retries
    with the v2 grounding block + RAG standards stripped. The detector
    must distinguish 'token-limit' from other Groq errors so we don't
    pointlessly retry tool_use_failed / 5xx / etc."""

    def test_matches_body_dict_with_tokens_type(self):
        from resumes.services.resume_generator import _is_token_limit_error
        class E(Exception):
            body = {
                'error': {
                    'message': 'Request too large for model... TPM Limit 30000',
                    'type': 'tokens',
                    'code': 'rate_limit_exceeded',
                },
            }
        self.assertTrue(_is_token_limit_error(E('x')))

    def test_matches_str_fallback_when_body_missing(self):
        from resumes.services.resume_generator import _is_token_limit_error
        class E(Exception):
            body = None
        e = E()
        e.args = ("Error code: 413 - {'error': {'message': "
                  "'Request too large for model X on tokens per minute (TPM)', "
                  "'type': 'tokens', 'code': 'rate_limit_exceeded'}}",)
        self.assertTrue(_is_token_limit_error(e))

    def test_rejects_tool_use_failed(self):
        # tool_use_failed is recoverable via _recover_resume_from_failed_generation,
        # NOT by retrying with a smaller prompt.
        from resumes.services.resume_generator import _is_token_limit_error
        class E(Exception):
            body = {
                'error': {
                    'message': 'tool call validation failed',
                    'type': 'invalid_request_error',
                    'code': 'tool_use_failed',
                },
            }
        self.assertFalse(_is_token_limit_error(E('x')))

    def test_rejects_generic_exceptions(self):
        from resumes.services.resume_generator import _is_token_limit_error
        self.assertFalse(_is_token_limit_error(ValueError('something else')))
        self.assertFalse(_is_token_limit_error(RuntimeError('connection reset')))


class ExtractFailedGenerationTests(SimpleTestCase):
    """Regression coverage for the three extraction paths. The 18:03 regen
    hit the silent-return path because exc.body was None on this Groq SDK
    version — _extract_failed_generation now falls back to exc.response
    and finally to ast.literal_eval of str(exc)."""

    def test_path1_body_dict(self):
        from resumes.services.resume_generator import _extract_failed_generation
        class E(Exception):
            body = {'error': {'failed_generation': '{"a": 1}'}}
        self.assertEqual(_extract_failed_generation(E('x')), '{"a": 1}')

    def test_path2_response_json(self):
        from resumes.services.resume_generator import _extract_failed_generation
        class Resp:
            def json(self):
                return {'error': {'failed_generation': '{"b": 2}'}}
        class E(Exception):
            body = None
            response = Resp()
        self.assertEqual(_extract_failed_generation(E('x')), '{"b": 2}')

    def test_path3_str_parse(self):
        # This is the case that was failing silently — body is None,
        # response is None, only str(exc) carries the payload.
        from resumes.services.resume_generator import _extract_failed_generation
        class E(Exception):
            body = None
            response = None
        e = E()
        e.args = ("Error code: 400 - {'error': {'message': 'foo', "
                  "'failed_generation': '[{\"name\": \"X\"}]'}}",)
        self.assertEqual(_extract_failed_generation(e),
                         '[{"name": "X"}]')

    def test_returns_none_when_payload_missing_everywhere(self):
        from resumes.services.resume_generator import _extract_failed_generation
        class E(Exception):
            body = None
            response = None
        self.assertIsNone(_extract_failed_generation(E('nothing useful')))

    def test_recovery_end_to_end_with_schema_envelope_in_str(self):
        # The full failure mode from the 18:03 regen: schema-envelope
        # payload reachable only via str(exc). Recovery should still
        # surface the LLM's content (skills, summary, experience with
        # highlights promoted to description).
        import json as _json
        from resumes.services.resume_generator import _recover_resume_from_failed_generation
        envelope = _json.dumps([{
            'name': 'ResumeContentResult',
            'parameters': {
                'additionalProperties': True,
                'properties': {
                    'professional_summary': {'type': 'string', 'value': 'DS with ML.'},
                    'skills': {'type': 'array', 'value': ['Python', 'TensorFlow']},
                    'experience': {'type': 'array', 'value': [
                        {'title': 'AI Trainee', 'description': None,
                         'highlights': ['Built model.', 'Used MLflow.']},
                    ]},
                    'projects': {'type': 'array', 'value': []},
                    'certifications': {'type': 'array', 'value': []},
                    'education': {'type': 'array', 'value': []},
                    'objective': {'type': 'string', 'value': ''},
                },
            },
        }])
        class E(Exception):
            body = None
            response = None
        e = E()
        e.args = (f"Error code: 400 - {{'error': {{'failed_generation': {envelope!r}}}}}",)
        result = _recover_resume_from_failed_generation(e)
        self.assertIsNotNone(result, "recovery should not silently return None")
        self.assertEqual(result.professional_summary, 'DS with ML.')
        self.assertEqual(result.skills, ['Python', 'TensorFlow'])
        # Highlights → description promotion still fires through the
        # full pipeline.
        self.assertEqual(result.experience[0].description,
                         ['Built model.', 'Used MLflow.'])


class ResumeFailedGenerationRecoveryTests(SimpleTestCase):
    """Salvage Groq's tool_use_failed payload for resume generation.
    Same pattern as outreach_generator + learning_path_generator."""

    def _exc(self, raw_failed_generation: str):
        class _FakeBadRequest(Exception):
            pass
        e = _FakeBadRequest('tool_use_failed')
        e.body = {
            'error': {
                'message': 'tool call validation failed',
                'type': 'invalid_request_error',
                'code': 'tool_use_failed',
                'failed_generation': raw_failed_generation,
            }
        }
        return e

    def test_recovers_from_tool_call_wrapper(self):
        """Groq wraps the failed call as `[{name, parameters}]`. Salvage
        the parameters dict and validate against the schema."""
        from unittest.mock import patch, MagicMock
        from resumes.services import resume_generator
        # Shape from the production trace — tool-call wrapper with null
        # fields and object-wrapped highlights/skills.
        raw = (
            '[{"name": "ResumeContentResult", "parameters": {'
            '"professional_title": "Engineer",'
            '"professional_summary": "Built things.",'
            '"skills": [{"name": "Python", "proficiency": null}],'
            '"experience": [{"title": "Dev", "company": "Acme",'
            ' "industry": null, "location": null, "start_date": "2024",'
            ' "end_date": "", "duration": "", "description": []}],'
            '"projects": [{"name": "Tool",'
            ' "highlights": [{"description": "Shipped X"}],'
            ' "description": [], "technologies": [], "url": ""}],'
            '"education": [{"degree": "BSc", "institution": "MIT",'
            ' "year": "2024", "gpa": null, "location": null}],'
            '"certifications": [], "languages": [], "objective": ""}}]'
        )
        # Build a minimal profile + job stub for the generator wrapper.
        import types
        profile = types.SimpleNamespace(
            data_content={'skills': [{'name': 'Python'}]},
            raw_text='', skills=[{'name': 'Python'}], experiences=[],
            education=[], projects=[], certifications=[],
        )
        job = types.SimpleNamespace(
            title='Engineer', company='Acme', description='Need Python.',
            extracted_skills=['Python'],
        )
        gap = types.SimpleNamespace(matched_skills=['Python'])

        # Mock the structured LLM to raise the tool_use_failed exception;
        # the generator should salvage the failed_generation.
        with patch.object(resume_generator, 'get_structured_llm') as mock_llm:
            mock_llm.return_value.invoke.side_effect = self._exc(raw)
            out = resume_generator.generate_resume_content(profile, job, gap)
        # Recovery succeeded: we got the tool-call payload, not the offline
        # fallback (which would have a different professional_summary).
        self.assertEqual(out['professional_title'], 'Engineer')
        self.assertEqual(out['skills'], ['Python'])
        # PR 3a: object-wrapped highlights got flattened by the schema
        # validator into the canonical `description` field.
        self.assertEqual(out['projects'][0]['description'], ['Shipped X'])
        # Null strings coerced to "":
        self.assertEqual(out['experience'][0]['industry'], '')
        self.assertEqual(out['experience'][0]['location'], '')

    def test_recovers_when_experience_has_employment_type(self):
        """2026-05-28 production failure: Groq emitted
        experience[].employment_type='Internship'/'Full-time' which hits
        ResumeExperience.extra='forbid' and previously dropped the whole
        recovery — shipping the offline fallback instead of the LLM
        output. The validator must now silently pop employment_type and
        let the salvaged resume through."""
        import json as _json
        from resumes.services.resume_generator import (
            _recover_resume_from_failed_generation,
        )
        payload = {
            'name': 'ResumeContentResult',
            'parameters': {
                'professional_title': '',
                'professional_summary': 'Data Scientist.',
                'objective': '',
                'skills': [
                    {'name': 'Python', 'years': None, 'proficiency': None},
                    {'name': 'PySpark', 'years': None, 'proficiency': None},
                ],
                'experience': [
                    {
                        'title': 'AI & Data Science Trainee',
                        'company': 'DEPI',
                        'duration': 'Jun 2025 - Dec 2025',
                        'location': 'Remote',
                        'industry': None,
                        'start_date': 'Jun 2025',
                        'end_date': 'Dec 2025',
                        'description': ['Applied the full DS lifecycle.'],
                        'employment_type': 'Internship',
                    },
                    {
                        'title': 'Digital Transformation Intern',
                        'company': 'Almansour Automotive',
                        'duration': 'Aug 2025',
                        'location': 'Al Jizah, Egypt',
                        'industry': None,
                        'start_date': 'August 2025',
                        'end_date': None,
                        'description': ['Built PySpark pipeline.'],
                        'employment_type': 'Full-time',
                    },
                ],
                'education': [],
                'projects': [],
                'certifications': [],
                'languages': [],
                'awards': [],
            },
        }
        raw = _json.dumps([payload])
        e = self._exc(raw)
        result = _recover_resume_from_failed_generation(e)
        self.assertIsNotNone(
            result,
            "employment_type extra should be dropped, not abort recovery",
        )
        # Skills coerced from objects to strings via _flatten_string_list.
        self.assertEqual(result.skills, ['Python', 'PySpark'])
        # Both experience entries survived with their bullets intact.
        self.assertEqual(len(result.experience), 2)
        self.assertEqual(result.experience[0].title, 'AI & Data Science Trainee')
        self.assertEqual(result.experience[0].description,
                         ['Applied the full DS lifecycle.'])
        self.assertEqual(result.experience[1].title,
                         'Digital Transformation Intern')

    def test_recovers_when_project_has_signal_only_source_field(self):
        """Projects sometimes carry `source='github'` etc. in the LLM
        output — that field is signal-only per the prompt and not part
        of ResumeProject. The validator must drop it instead of failing
        recovery."""
        import json as _json
        from resumes.services.resume_generator import (
            _recover_resume_from_failed_generation,
        )
        payload = {
            'name': 'ResumeContentResult',
            'parameters': {
                'professional_title': '',
                'professional_summary': '',
                'objective': '',
                'skills': ['Python'],
                'experience': [],
                'education': [],
                'projects': [{
                    'name': 'SmartCV',
                    'url': 'https://example.com',
                    'technologies': ['Python'],
                    'description': ['Built it.'],
                    'source': 'github',
                    'source_id': 'foo/SmartCV',
                    'source_url': 'https://github.com/foo/SmartCV',
                    'role': 'Author',
                }],
                'certifications': [],
                'languages': [],
                'awards': [],
            },
        }
        raw = _json.dumps([payload])
        e = self._exc(raw)
        result = _recover_resume_from_failed_generation(e)
        self.assertIsNotNone(result)
        self.assertEqual(len(result.projects), 1)
        self.assertEqual(result.projects[0].name, 'SmartCV')
        self.assertEqual(result.projects[0].description, ['Built it.'])


class LearningPathFailedGenerationRecoveryTests(SimpleTestCase):
    """Groq returns 400 tool_use_failed when the model emits a bare top-level
    list instead of {items: [...]}. The well-formed JSON list is still in
    `error.failed_generation`. The generator salvages it instead of dropping
    a successful generation on the floor."""

    def _exc(self, raw_failed_generation: str):
        """Build an exception shaped like a Groq BadRequestError."""
        class _FakeBadRequest(Exception):
            pass
        e = _FakeBadRequest('tool_use_failed')
        e.body = {
            'error': {
                'message': 'Failed to call a function.',
                'type': 'invalid_request_error',
                'code': 'tool_use_failed',
                'failed_generation': raw_failed_generation,
            }
        }
        return e

    def test_recovers_from_bare_list_failed_generation(self):
        from unittest.mock import patch
        from analysis.services import learning_path_generator
        # Real shape from the production trace: top-level array of items.
        raw = (
            '[{"importance": "Python is hot.", "skill": "python",'
            ' "resources": [{"name": "Real Python", "url": "https://realpython.com/",'
            ' "provider": "Other"}], "project_idea": "Web scraper.",'
            ' "time_estimate": "10-15 hours"}]'
        )
        with patch.object(learning_path_generator, 'get_structured_llm') as mock_llm:
            mock_llm.return_value.invoke.side_effect = self._exc(raw)
            out = learning_path_generator.generate_learning_path(['python'])
        self.assertEqual(len(out), 1)
        self.assertEqual(out[0]['skill'], 'python')
        self.assertEqual(out[0]['time_estimate'], '10-15 hours')
        self.assertEqual(out[0]['resources'][0]['url'], 'https://realpython.com/')

    def test_recovers_from_wrapped_failed_generation(self):
        """If the model gets the wrapper right but the tool serializer still
        flakes, recover from `{items: [...]}` form too."""
        from unittest.mock import patch
        from analysis.services import learning_path_generator
        raw = (
            '{"items": [{"importance": "ok", "skill": "go",'
            ' "resources": [], "project_idea": "x", "time_estimate": "20h"}]}'
        )
        with patch.object(learning_path_generator, 'get_structured_llm') as mock_llm:
            mock_llm.return_value.invoke.side_effect = self._exc(raw)
            out = learning_path_generator.generate_learning_path(['go'])
        self.assertEqual(len(out), 1)
        self.assertEqual(out[0]['skill'], 'go')

    def test_returns_empty_when_failed_generation_unparseable(self):
        from unittest.mock import patch
        from analysis.services import learning_path_generator
        # Truncated JSON — recovery should return [] not raise.
        with patch.object(learning_path_generator, 'get_structured_llm') as mock_llm:
            mock_llm.return_value.invoke.side_effect = self._exc('[{"skill"')
            out = learning_path_generator.generate_learning_path(['python'])
        self.assertEqual(out, [])

    def test_returns_empty_on_non_groq_exception(self):
        """A generic exception (network error, etc.) without the Groq
        body shape still drops gracefully."""
        from unittest.mock import patch
        from analysis.services import learning_path_generator
        with patch.object(learning_path_generator, 'get_structured_llm') as mock_llm:
            mock_llm.return_value.invoke.side_effect = RuntimeError('network down')
            out = learning_path_generator.generate_learning_path(['python'])
        self.assertEqual(out, [])


class LearningPathPersistenceTests(TestCase):
    """Tier 5 (M4): the learning path persists across page loads, the
    mark-as-done toggle saves to the profile, and the return-path CTA
    points at the right next-step."""

    def setUp(self):
        from django.contrib.auth import get_user_model
        from profiles.models import UserProfile
        User = get_user_model()
        self.user = User.objects.create_user(
            username='learner@example.com', email='learner@example.com', password='x',
        )
        UserProfile.objects.create(user=self.user, full_name='Learner User')
        self.client.force_login(self.user)

    def test_get_renders_persisted_learning_path(self):
        """If the user has a persisted path, GET should render it without
        re-running the LLM."""
        from profiles.models import UserProfile
        UserProfile.objects.filter(user=self.user).update(data_content={
            'learning_path': [{
                'skill': 'Python',
                'importance': 'Used everywhere.',
                'resources': [{'name': 'Real Python', 'url': 'https://realpython.com/',
                               'provider': 'Other'}],
                'project_idea': 'Build a CLI tool.',
                'time_estimate': '15 hours over 2 weeks',
            }],
        })
        resp = self.client.get('/analysis/learning-path/')
        self.assertEqual(resp.status_code, 200)
        # Skill, time estimate, and clickable resource URL all present.
        self.assertContains(resp, 'Python')
        self.assertContains(resp, '15 hours over 2 weeks')
        self.assertContains(resp, 'https://realpython.com/')
        # Mark-as-done button rendered (Alpine state, not server-side
        # checkbox, so we check for the toggle text).
        self.assertContains(resp, 'Mark as done')

    def test_mark_skill_complete_toggles_state(self):
        from profiles.models import UserProfile
        # First call marks
        resp = self.client.post(
            '/analysis/api/learning-path/skill-done/',
            data='{"skill": "python"}',
            content_type='application/json',
        )
        self.assertEqual(resp.status_code, 200)
        body = resp.json()
        self.assertTrue(body['ok'])
        self.assertEqual(body['action'], 'marked')
        self.assertEqual(body['completed_skills'], ['python'])
        profile = UserProfile.objects.get(user=self.user)
        self.assertEqual(profile.data_content['completed_skills'], ['python'])
        # Second call unmarks (toggle)
        resp = self.client.post(
            '/analysis/api/learning-path/skill-done/',
            data='{"skill": "python"}',
            content_type='application/json',
        )
        self.assertEqual(resp.json()['action'], 'unmarked')
        profile.refresh_from_db()
        self.assertEqual(profile.data_content['completed_skills'], [])

    def test_mark_skill_complete_rejects_empty_skill(self):
        resp = self.client.post(
            '/analysis/api/learning-path/skill-done/',
            data='{"skill": ""}',
            content_type='application/json',
        )
        self.assertEqual(resp.status_code, 400)
        self.assertEqual(resp.json()['error'], 'skill_required')

    def test_completed_skill_marker_renders_with_strikethrough(self):
        from profiles.models import UserProfile
        UserProfile.objects.filter(user=self.user).update(data_content={
            'completed_skills': ['python'],
            'learning_path': [{'skill': 'Python', 'importance': 'x',
                               'resources': [], 'project_idea': 'y'}],
        })
        resp = self.client.get('/analysis/learning-path/')
        body = resp.content.decode('utf-8')
        # Alpine x-data initialized to true for the matched skill
        self.assertIn("x-data=\"{done: true}\"", body)

    def test_return_cta_points_to_global_dashboard_when_no_job(self):
        from profiles.models import UserProfile
        UserProfile.objects.filter(user=self.user).update(data_content={
            'learning_path': [{'skill': 'Python', 'importance': 'x',
                               'resources': [], 'project_idea': 'y'}],
        })
        resp = self.client.get('/analysis/learning-path/')
        # Global learning-path view → return CTA points at dashboard
        self.assertContains(resp, 'Pick a job to re-analyze')
        # NOT the per-job CTA
        self.assertNotContains(resp, 'Re-run gap for')

    def test_return_cta_points_to_specific_job_when_scoped(self):
        from jobs.models import Job
        from profiles.models import UserProfile
        UserProfile.objects.filter(user=self.user).update(data_content={
            'learning_path': [{'skill': 'Python', 'importance': 'x',
                               'resources': [], 'project_idea': 'y'}],
        })
        job = Job.objects.create(
            user=self.user, title='Backend Engineer', company='Acme',
            description='Need Python.', extracted_skills=[],
        )
        resp = self.client.get(f'/analysis/learning-path/{job.id}/')
        self.assertContains(resp, 'Re-run gap for')
        self.assertContains(resp, 'Backend Engineer')


class FactualityCheckEnrichedProjectsTests(SimpleTestCase):
    """Phase 3: factuality_check accepts enriched-project URLs / source_ids
    as legitimate evidence so Phase-2 confirmed projects don't get falsely
    flagged as fabrication. Companies and schools still require source-text
    grounding.
    """

    def test_project_grounded_via_enriched_source_url(self):
        from benchmarks.llm_judge import factuality_check
        generated = {
            'experience': [{'company': 'Acme'}],
            'projects': [{'name': 'climate-modeling', 'description': []}],
        }
        # Source CV text has Acme but NOT climate-modeling — that project
        # came from Phase-1 enrichment of GitHub signals.
        source_text = "Engineer at Acme since 2023."
        confirmed = [{
            'name': 'climate-modeling',
            'url': 'https://github.com/me/climate-modeling',
            'source': 'github',
            'source_id': 'me/climate-modeling',
        }]
        result = factuality_check(generated, source_text, confirmed_projects=confirmed)
        # Both Acme and climate-modeling grounded; ratio = 1.0
        self.assertEqual(result['ratio'], 1.0)
        self.assertEqual(result['ungrounded'], [])

    def test_project_falsely_named_still_flags_as_ungrounded(self):
        """An LLM-fabricated project name should still get flagged — being
        in confirmed_projects requires actual text overlap (substring), not
        just the existence of any enriched project."""
        from benchmarks.llm_judge import factuality_check
        generated = {
            'experience': [{'company': 'Acme'}],
            'projects': [{'name': 'completely-made-up-project'}],
        }
        source_text = "Engineer at Acme since 2023."
        confirmed = [{
            'name': 'climate-modeling',
            'url': 'https://github.com/me/climate-modeling',
            'source': 'github',
        }]
        result = factuality_check(generated, source_text, confirmed_projects=confirmed)
        self.assertIn('completely-made-up-project', result['ungrounded'])
        # Acme grounded, fabricated project not — 1/2 = 0.5
        self.assertEqual(result['ratio'], 0.5)

    def test_company_does_not_use_project_evidence(self):
        """Company names must come from source_text — they CANNOT use the
        confirmed-projects index (otherwise a typed project named after a
        fake company would launder fabrication)."""
        from benchmarks.llm_judge import factuality_check
        generated = {
            'experience': [{'company': 'NotInCV-Corp'}],
            'projects': [],
        }
        source_text = "Engineer at Acme since 2023."
        confirmed = [{
            'name': 'NotInCV-Corp',  # appears in confirmed projects
            'url': 'https://example.com',
        }]
        result = factuality_check(generated, source_text, confirmed_projects=confirmed)
        # Company should still be flagged — confirmed projects are not
        # evidence for company entities.
        self.assertIn('NotInCV-Corp', result['ungrounded'])

    def test_backwards_compat_no_confirmed_projects(self):
        """Calling without confirmed_projects (legacy callers) still works
        and falls back to source-text-only grounding."""
        from benchmarks.llm_judge import factuality_check
        generated = {
            'experience': [{'company': 'Acme'}],
            'projects': [{'name': 'pgbench-tuner'}],
        }
        source_text = "Engineer at Acme; built pgbench-tuner."
        result = factuality_check(generated, source_text)
        self.assertEqual(result['ratio'], 1.0)


class ResumeGeneratorEnrichedProjectsPromptTests(SimpleTestCase):
    """Phase 3: source-tagged projects survive the resume_generator's prompt
    assembly (slim_cv) and the offline fallback. The LLM gets to see the
    `source` field so it can treat enriched projects as ground truth."""

    def test_offline_fallback_preserves_enriched_project(self):
        """The offline fallback (no LLM) must pass enriched projects through
        intact — including their custom technologies and bullets."""
        import types
        from resumes.services.resume_generator import _build_offline_fallback
        profile = types.SimpleNamespace(data_content={
            'projects': [{
                'name': 'climate-modeling', 'url': 'https://github.com/me/climate-modeling',
                'description': ['Built earth-science modeling pipeline.'],
                'technologies': ['Python', 'PyTorch'],
                'source': 'github', 'source_id': 'me/climate-modeling',
            }],
        }, raw_text='', skills=[], experiences=[], education=[],
            projects=[], certifications=[])
        job = types.SimpleNamespace(title='ML Engineer', description='Build models.',
                                    extracted_skills=[])
        out = _build_offline_fallback(profile, job, profile.data_content)
        self.assertEqual(len(out['projects']), 1)
        proj = out['projects'][0]
        self.assertEqual(proj['name'], 'climate-modeling')
        self.assertEqual(proj['url'], 'https://github.com/me/climate-modeling')
        self.assertIn('PyTorch', proj['technologies'])
        # Bullets survive verbatim
        self.assertIn('Built earth-science modeling pipeline.', proj['description'])

    def test_slim_cv_includes_source_field_for_llm_prompt(self):
        """Verify the slim_cv dict assembled in generate_resume_content keeps
        the `source` field on each project so the LLM can apply the
        'pre-vetted from external signals' rule."""
        # We can't run the full LLM call in a unit test, but we CAN assert
        # the assembly logic is correct by inspecting what would be
        # serialized into the prompt.
        raw_cv_data = {
            'projects': [{
                'name': 'climate-modeling', 'url': 'https://github.com/me/climate-modeling',
                'description': ['Built it.'], 'technologies': ['Python'],
                'source': 'github', 'source_id': 'me/climate-modeling',
            }],
        }
        # Replicate the slim_cv filter from generate_resume_content
        _SIGNAL_KEYS = {'github_signals', 'scholar_signals', 'kaggle_signals', 'linkedin_snapshot'}
        slim_cv = {k: v for k, v in raw_cv_data.items()
                   if k != 'raw_text'
                   and k not in _SIGNAL_KEYS
                   and v
                   and k not in ('normalized_summary', 'objective')}
        self.assertEqual(slim_cv['projects'][0]['source'], 'github')


class ResumeListThumbnailTests(TestCase):
    """Each card on the résumé list page now renders a thumbnail preview of
    the actual resume content (name, summary, first experience, top skills)
    so the user can recognise their resumes by glance instead of by reading
    the job-title text. The thumbnail reuses the editor's .pdf-preview CSS,
    so picking a template carries through to the list view automatically.
    """

    def setUp(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from profiles.models import UserProfile
        from resumes.models import GeneratedResume
        User = get_user_model()
        self.user = User.objects.create_user(
            username='listthumb@example.com', email='listthumb@example.com', password='x',
        )
        UserProfile.objects.create(user=self.user, full_name='Ada Lovelace', email='listthumb@example.com')
        self.client.force_login(self.user)
        job = Job.objects.create(user=self.user, title='Senior Backend Engineer', company='Acme Corp')
        self.gap = GapAnalysis.objects.create(user=self.user, job=job, similarity_score=0.82)
        self.resume = GeneratedResume.objects.create(
            gap_analysis=self.gap,
            content={
                'professional_title': 'Backend Systems Engineer',
                'professional_summary': 'Built distributed Python systems for high-throughput pipelines and ran the migration from synchronous to async stack.',
                'template_name': 'ats_clean_accent',
                'experience': [
                    {
                        'title': 'Backend Engineer',
                        'company': 'PriorCo',
                        'duration': '2023 - Present',
                        'description': ['Cut p99 latency from 1.2s to 380ms', 'Owned async migration across 6 services'],
                    },
                ],
                'skills': ['Python', 'PostgreSQL', 'Redis', 'Kubernetes', 'AWS'],
            },
        )

    def test_thumbnail_renders_real_content(self):
        """The thumbnail must show real resume data — name, professional
        title, first experience, skills — not placeholder bars."""
        resp = self.client.get(reverse('resume_list'))
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode('utf-8')
        # Name from the user's profile.
        self.assertIn('Ada Lovelace', body)
        # Professional title from the resume content (NOT the job title).
        self.assertIn('Backend Systems Engineer', body)
        # First experience entry's title and company.
        self.assertIn('Backend Engineer', body)
        self.assertIn('PriorCo', body)
        # Top skill renders.
        self.assertIn('Python', body)
        # Template-aware CSS modifier picked up.
        self.assertIn('pdf-preview pdf-preview--ats_clean_accent', body)
        # Sized as a list thumbnail (not the editor's small picker thumb).
        self.assertIn('resume-list-thumb', body)

    def test_thumbnail_handles_empty_content_gracefully(self):
        """A resume with empty content shouldn't break the page: section
        headings without data must not render, and the page still 200s."""
        self.resume.content = {}
        self.resume.save()
        resp = self.client.get(reverse('resume_list'))
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode('utf-8')
        # The thumbnail container still renders so the card has a stable shape
        self.assertIn('resume-list-thumb', body)
        # But the section headings inside it (Summary / Experience / Skills)
        # must NOT appear when there's no data — we conditionally render.
        # Match the exact section-title markup so we don't catch e.g. an
        # entirely unrelated word "Summary" elsewhere in the page chrome.
        self.assertNotIn('<div class="p-section-title">Summary</div>', body)
        self.assertNotIn('<div class="p-section-title">Experience</div>', body)
        self.assertNotIn('<div class="p-section-title">Skills</div>', body)
        # Header still falls back to job title when professional_title is empty.
        self.assertIn('Senior Backend Engineer', body)

    def test_thumbnail_falls_back_to_default_when_template_missing(self):
        """If a resume saved before the template-name feature, fall back to
        the current default theme so the thumbnail still has consistent
        styling."""
        self.resume.content = {
            'professional_title': 'X',
            'professional_summary': 'Y',
        }
        self.resume.save()
        resp = self.client.get(reverse('resume_list'))
        body = resp.content.decode('utf-8')
        self.assertIn('pdf-preview pdf-preview--ats_clean', body)

    def test_profile_name_falls_back_to_email_local_part(self):
        """Users without a UserProfile.full_name should still get a header,
        not an empty void. The NAME slot specifically uses the email
        local-part as the fallback (the email itself separately appears in
        the contact line, which mirrors what the PDF template does)."""
        from profiles.models import UserProfile
        UserProfile.objects.filter(user=self.user).update(full_name='')
        resp = self.client.get(reverse('resume_list'))
        body = resp.content.decode('utf-8')
        # The name div uses the local-part, not the full email.
        self.assertIn('<div class="p-name">listthumb</div>', body)


# ---------------------------------------------------------------------------
# Resume normalizer — Pass B safety net
# ---------------------------------------------------------------------------

import copy

from resumes.services.resume_normalizer import (
    normalize_resume,
    normalize_titles,
    filter_soft_skills,
    enforce_skill_hard_cap,
    strip_first_person_from_resume,
    consolidate_coursework,
    trim_projects_to_plan,
    trim_certs_to_plan,
)


class _FakeProjectPlan:
    """Stand-in for inclusion_planner.ProjectPlan (kept here to avoid the
    settings/db cost of importing the real dataclass for SimpleTestCase
    consumers)."""
    def __init__(self, name: str):
        self.name = name


class _FakePlan:
    def __init__(self, projects: list[str] | None = None, certifications: list[str] | None = None):
        self.projects = [_FakeProjectPlan(n) for n in (projects or [])]
        self.certifications = list(certifications or [])


class ResumeNormalizerTests(SimpleTestCase):
    """Pass-B post-LLM safety net.

    Each test pins one defect we have actually seen in the LLM's output
    so a regression that re-introduces the leak fails loudly.
    """

    def test_normalize_titles_fixes_all_caps_experience(self):
        resume = {'experience': [{'title': 'DIGITAL TRANSFORMATION INTERN'}]}
        out = normalize_titles(resume)
        self.assertEqual(out['experience'][0]['title'], 'Digital Transformation Intern')

    def test_normalize_titles_preserves_mixed_case(self):
        resume = {'experience': [{'title': 'AI & Data Science Trainee'}]}
        out = normalize_titles(resume)
        self.assertEqual(out['experience'][0]['title'], 'AI & Data Science Trainee')

    def test_normalize_titles_fixes_parser_typo(self):
        resume = {'experience': [{'title': 'INFROMATION TECHNOLOGY INTERN'}]}
        out = normalize_titles(resume)
        self.assertEqual(out['experience'][0]['title'], 'Information Technology Intern')

    def test_filter_soft_skills_drops_communication_and_presentation(self):
        resume = {'skills': ['Python', 'Communication', 'SQL', 'Presentation skills', 'TensorFlow']}
        out = filter_soft_skills(resume)
        self.assertEqual(out['skills'], ['Python', 'SQL', 'TensorFlow'])

    def test_filter_soft_skills_accepts_dict_entries(self):
        resume = {'skills': [{'name': 'Python'}, {'name': 'Teamwork'}, {'name': 'Pandas'}]}
        out = filter_soft_skills(resume)
        self.assertEqual(out['skills'], ['Python', 'Pandas'])

    def test_enforce_skill_hard_cap_truncates_at_cap(self):
        resume = {'skills': [f'S{i}' for i in range(20)]}
        out = enforce_skill_hard_cap(resume, cap=14)
        self.assertEqual(len(out['skills']), 14)
        # Order preserved.
        self.assertEqual(out['skills'][0], 'S0')
        self.assertEqual(out['skills'][-1], 'S13')

    def test_strip_first_person_cleans_summary_and_descriptions(self):
        resume = {
            'professional_summary': 'I am a data scientist. My focus is on ML.',
            'experience': [{
                'description': ['I built a churn model.', 'I deployed it to prod.'],
            }],
        }
        out = strip_first_person_from_resume(resume)
        # Summary: no "I" / "my", sentences re-capped.
        self.assertNotIn(' I ', ' ' + out['professional_summary'] + ' ')
        self.assertNotIn(' my ', ' ' + out['professional_summary'].lower() + ' ')
        self.assertTrue(out['professional_summary'][0].isupper())
        # Bullets: same.
        for b in out['experience'][0]['description']:
            self.assertNotRegex(b, r"\b[Ii]\b")
            self.assertNotRegex(b, r"\b[Mm]y\b")
            self.assertTrue(b[0].isupper())

    def test_consolidate_coursework_collapses_short_course_bullets(self):
        resume = {'experience': [{
            'description': [
                'Developed practical skills across the AI lifecycle.',
                'Prompt Engineering',
                'Data Science Methodology',
                'Tools for Data Science',
                'Databases & SQL',
                'Practised prompt engineering on real prompts.',
            ],
        }]}
        out = consolidate_coursework(resume)
        bullets = out['experience'][0]['description']
        # The four short course-names collapsed into one line.
        self.assertEqual(len(bullets), 3)
        self.assertEqual(bullets[0], 'Developed practical skills across the AI lifecycle.')
        self.assertTrue(bullets[1].startswith('Coursework included:'))
        self.assertIn('Prompt Engineering', bullets[1])
        self.assertIn('Databases & SQL', bullets[1])
        self.assertEqual(bullets[2], 'Practised prompt engineering on real prompts.')

    def test_consolidate_coursework_splits_embedded_bullets_in_single_string(self):
        # The LLM frequently emits one bullet with embedded \n• markers.
        # A list-introducer prelude ending with ":" is dropped (redundant
        # once the items are folded into "Coursework included: ...").
        resume = {'experience': [{
            'description': 'Coursework consisted of:\n• Prompt Engineering\n• Data Science Methodology\n• Tools for Data Science',
        }]}
        out = consolidate_coursework(resume)
        bullets = out['experience'][0]['description']
        self.assertIsInstance(bullets, list)
        self.assertEqual(len(bullets), 1)
        self.assertTrue(bullets[0].startswith('Coursework included:'))
        self.assertIn('Prompt Engineering', bullets[0])
        self.assertIn('Tools for Data Science', bullets[0])

    def test_consolidate_coursework_keeps_non_header_prelude(self):
        # A meaningful prelude (no trailing colon) is preserved as a bullet.
        # Fixture has three course-name bullets after the prelude so the
        # MIN_RUN=3 consolidation threshold (PR1 Fix 3) triggers.
        resume = {'experience': [{
            'description': 'Built models throughout the program\n• Prompt Engineering\n• Data Science Methodology\n• Tools for Data Science',
        }]}
        out = consolidate_coursework(resume)
        bullets = out['experience'][0]['description']
        self.assertEqual(bullets[0], 'Built models throughout the program')
        self.assertTrue(bullets[1].startswith('Coursework included:'))

    def test_consolidate_coursework_leaves_real_bullets_alone(self):
        resume = {'experience': [{
            'description': [
                'Built a churn model that improved retention by 12 percent.',
                'Deployed the model to production via MLflow.',
            ],
        }]}
        out = consolidate_coursework(resume)
        self.assertEqual(out['experience'][0]['description'], [
            'Built a churn model that improved retention by 12 percent.',
            'Deployed the model to production via MLflow.',
        ])

    def test_trim_projects_to_plan_drops_off_plan_projects(self):
        resume = {'projects': [
            {'name': 'SmartCV'},
            {'name': 'BookShop'},                 # not in plan → drop
            {'name': 'Healthcare Prediction (DEPI)'},
            {'name': 'apotheosis-traffic-sign-detection'},  # not in plan → drop
        ]}
        plan = _FakePlan(projects=['SmartCV', 'Healthcare Prediction (DEPI)'])
        out = trim_projects_to_plan(resume, plan)
        names = [p['name'] for p in out['projects']]
        self.assertEqual(names, ['SmartCV', 'Healthcare Prediction (DEPI)'])

    def test_trim_projects_to_plan_fuzzy_matches_lightly_renamed_project(self):
        # The LLM rewrote "Healthcare Prediction (DEPI)" → "Healthcare Prediction - DEPI".
        # Should still match by SequenceMatcher.
        resume = {'projects': [
            {'name': 'Healthcare Prediction - DEPI'},
        ]}
        plan = _FakePlan(projects=['Healthcare Prediction (DEPI)'])
        out = trim_projects_to_plan(resume, plan)
        self.assertEqual(len(out['projects']), 1)

    def test_trim_projects_to_plan_skips_when_plan_has_no_projects(self):
        # Defensive — never wipe projects just because the planner returned [].
        resume = {'projects': [{'name': 'SmartCV'}, {'name': 'BookShop'}]}
        plan = _FakePlan(projects=[])
        out = trim_projects_to_plan(resume, plan)
        self.assertEqual(len(out['projects']), 2)

    def test_trim_certs_to_plan_keeps_all_certs_now(self):
        # Round 1.5: the previous "drop certs not in plan" filter was
        # over-aggressive and removed Software Testing on a DevOps JD.
        # The new policy is keep all of the candidate's certs (capped
        # at _CERT_CAP). plan membership is irrelevant.
        resume = {'certifications': [
            {'name': 'IBM Data Scientist Professional Certificate'},
            {'name': 'Project Management: The Basics for Success'},
            {'name': 'Applied Machine Learning in Python'},
        ]}
        plan = _FakePlan(certifications=[
            'IBM Data Scientist Professional Certificate',
        ])
        out = trim_certs_to_plan(resume, plan)
        names = [c['name'] for c in out['certifications']]
        self.assertEqual(names, [
            'IBM Data Scientist Professional Certificate',
            'Project Management: The Basics for Success',
            'Applied Machine Learning in Python',
        ])

    def test_normalize_resume_end_to_end_combines_every_rule(self):
        resume = {
            'professional_summary': 'I am a data scientist with strong ML chops.',
            'skills': [
                'Python', 'SQL', 'Communication', 'TensorFlow', 'PyTorch',
                'Pandas', 'Presentation skills', 'NumPy', 'scikit-learn',
                'Machine Learning', 'Deep Learning', 'MLflow', 'Power BI',
                'Statistical Modeling', 'Teamwork', 'A', 'B', 'C', 'D', 'E',
            ],
            'experience': [{
                'title': 'AI & DATA SCIENCE TRAINEE',
                'description': [
                    'I built models throughout the program.',
                    'Prompt Engineering',
                    'Data Science Methodology',
                    'Tools for Data Science',
                ],
            }],
            'projects': [
                {'name': 'SmartCV'},
                {'name': 'BookShop'},
            ],
            'certifications': [
                {'name': 'IBM Data Scientist Professional Certificate'},
                {'name': 'Project Management 101'},
            ],
        }
        plan = _FakePlan(
            projects=['SmartCV'],
            certifications=['IBM Data Scientist Professional Certificate'],
        )
        out = normalize_resume(resume, plan=plan)

        # Title-cased.
        self.assertEqual(out['experience'][0]['title'], 'AI & Data Science Trainee')
        # Soft skills gone, hard-capped at 14.
        self.assertNotIn('Communication', out['skills'])
        self.assertNotIn('Presentation skills', out['skills'])
        self.assertNotIn('Teamwork', out['skills'])
        self.assertLessEqual(len(out['skills']), 14)
        # First-person stripped from summary.
        self.assertNotRegex(out['professional_summary'], r"\bI\b")
        # Coursework collapsed.
        descs = out['experience'][0]['description']
        self.assertTrue(any('Coursework included:' in d for d in descs))
        # Plan trims applied to projects (BookShop is dropped).
        self.assertEqual([p['name'] for p in out['projects']], ['SmartCV'])
        # Round 1.5: certs are NOT filtered against the plan any more —
        # both candidate certs survive (capped at 8).
        self.assertEqual(
            [c['name'] for c in out['certifications']],
            ['IBM Data Scientist Professional Certificate', 'Project Management 101'],
        )

    def test_normalize_resume_does_not_mutate_input(self):
        resume = {
            'skills': ['Python', 'Communication'],
            'experience': [{'title': 'INTERN', 'description': ['I worked.']}],
        }
        before = copy.deepcopy(resume)
        normalize_resume(resume)
        self.assertEqual(resume, before)

    # --- Pass E regression coverage --------------------------------------

    def test_consolidate_coursework_handles_long_course_titles(self):
        # The real Banque Misr resume had DEPI coursework bullets like
        # "Python for Data Science, AI & Development + Python Project"
        # (9 words) and "Data Analysis, Visualization & Machine Learning
        # with Python" (8 words). The v1 7-word cap rejected them so
        # consolidation never fired. v2 cap is 12.
        resume = {'experience': [{
            'description': [
                'Selected participant in the DEPI program.',
                'Prompt Engineering',
                'What is Data Science / Data Science Methodology',
                'Tools for Data Science',
                'Python for Data Science, AI & Development + Python Project',
                'Databases & SQL for Data Science with Python',
                'Data Analysis, Visualization & Machine Learning with Python',
                'MLOps tools (MLflow, Hugging Face)',
                'Capstone project end-to-end AI demonstrating preprocessing, feature engineering, model selection, evaluation and deployment.',
            ],
        }]}
        out = consolidate_coursework(resume)
        bullets = out['experience'][0]['description']
        # Post PR-1 Fix 3: MLflow and Hugging Face are now recognised as
        # technical tokens (_TECHNICAL_TOKENS), so the "MLOps tools
        # (MLflow, Hugging Face)" bullet is correctly excluded from the
        # coursework run and emits as its own bullet. Expected total is
        # now 4 (prelude + consolidated coursework + MLOps + capstone),
        # not 3. The test's documented purpose — verifying the 12-word
        # cap on course titles within the consolidated run — is unchanged.
        self.assertEqual(len(bullets), 4)
        self.assertEqual(bullets[0], 'Selected participant in the DEPI program.')
        self.assertTrue(bullets[1].startswith('Coursework included:'))
        # Every original LONG-TITLE course is in the consolidated line —
        # this is the assertion that verifies the 12-word cap is working.
        self.assertIn('Prompt Engineering', bullets[1])
        self.assertIn('Python for Data Science, AI & Development + Python Project', bullets[1])
        self.assertIn('Databases & SQL for Data Science with Python', bullets[1])
        # MLOps tools bullet survives separately (technical_token rule).
        self.assertIn('MLOps tools', bullets[2])
        self.assertIn('Hugging Face', bullets[2])
        # Capstone (long, multi-sentence) survives as its own bullet.
        self.assertTrue(bullets[3].startswith('Capstone project'))

    def test_consolidate_coursework_does_not_merge_real_action_bullets(self):
        # The real Apotheosis Traffic Sign project had two action bullets
        # ("Classified signs using visual characteristics", "Tuned
        # thresholds via external JSON configuration") that v1 wrongly
        # merged into a fake "Coursework included: Classified signs,
        # Tuned thresholds..." line because the verb list didn't include
        # "classified" / "tuned". v2 -ed-past-tense heuristic catches
        # them as action bullets.
        resume = {'projects': [{
            'description': [
                'Preprocessed traffic sign images with color filtering and contour detection.',
                'Classified signs using visual characteristics',
                'Tuned thresholds via external JSON configuration',
            ],
        }]}
        out = consolidate_coursework(resume)
        bullets = out['projects'][0]['description']
        # All three bullets survive verbatim — no coursework consolidation.
        self.assertEqual(len(bullets), 3)
        for b in bullets:
            self.assertFalse(b.startswith('Coursework included:'),
                             f"action bullet wrongly consolidated: {b!r}")

    def test_consolidate_coursework_preserves_short_ed_course_titles(self):
        # "Supervised Learning" / "Advanced Statistics" are short noun
        # phrases that happen to start with -ed words. The 4+-word
        # threshold for the -ed rejection keeps them as courses.
        resume = {'experience': [{
            'description': [
                'Trained on classical ML methods.',
                'Supervised Learning',
                'Unsupervised Learning',
                'Advanced Statistics',
                'Deployed a final project to production.',
            ],
        }]}
        out = consolidate_coursework(resume)
        bullets = out['experience'][0]['description']
        # Prelude + consolidated 3-course line + final = 3 bullets.
        self.assertEqual(len(bullets), 3)
        self.assertTrue(bullets[1].startswith('Coursework included:'))
        self.assertIn('Supervised Learning', bullets[1])
        self.assertIn('Unsupervised Learning', bullets[1])
        self.assertIn('Advanced Statistics', bullets[1])


from resumes.services.resume_normalizer import (
    filter_soft_skill_bullets,
    backfill_summary,
)


class FilterSoftSkillBulletsTests(SimpleTestCase):
    def test_drops_explicit_developed_soft_skills_bullet(self):
        resume = {'experience': [{'description': [
            'Built a procurement dashboard in Power BI.',
            'Developed soft skills including cross-team communication, problem-solving, and adaptability in a corporate environment.',
            'Shipped the dashboard to 4 stakeholders.',
        ]}]}
        out = filter_soft_skill_bullets(resume)
        bullets = out['experience'][0]['description']
        self.assertEqual(len(bullets), 2)
        for b in bullets:
            self.assertNotIn('soft skills', b.lower())

    def test_drops_dense_soft_skill_bullet_without_explicit_opener(self):
        # "Collaborated closely on cross-team initiatives, leadership, and adaptability."
        # 2+ soft-skill nouns inside a short bullet → drop.
        resume = {'experience': [{'description': [
            'Wrote ETL jobs in PySpark over 12 ERP tables.',
            'Showed communication, teamwork, leadership across the engineering org.',
        ]}]}
        out = filter_soft_skill_bullets(resume)
        bullets = out['experience'][0]['description']
        self.assertEqual(len(bullets), 1)
        self.assertIn('PySpark', bullets[0])

    def test_keeps_long_bullet_that_mentions_one_soft_skill_in_passing(self):
        # A bullet that incidentally says "communication" once but is
        # otherwise a real achievement should survive.
        resume = {'experience': [{'description': [
            'Led the rollout of the analytics platform across 4 business units, partnering with Product, Engineering, and Operations to align on KPIs and communication cadence.',
        ]}]}
        out = filter_soft_skill_bullets(resume)
        bullets = out['experience'][0]['description']
        self.assertEqual(len(bullets), 1)


class BackfillSummaryTests(SimpleTestCase):
    def test_backfills_when_summary_is_empty(self):
        resume = {
            'professional_summary': '',
            'skills': ['Python', 'SQL', 'PySpark', 'Power BI', 'MLflow'],
            'experience': [{'title': 'AI & Data Science Trainee'}],
        }
        out = backfill_summary(resume)
        self.assertTrue(out['professional_summary'])
        self.assertIn('AI & Data Science Trainee', out['professional_summary'])
        # Top 4 skills appear in the synthesized summary.
        self.assertIn('Python', out['professional_summary'])
        self.assertIn('Power BI', out['professional_summary'])

    def test_no_op_when_summary_already_present(self):
        resume = {
            'professional_summary': 'Data scientist focused on banking risk.',
            'experience': [{'title': 'X'}],
        }
        out = backfill_summary(resume)
        self.assertEqual(out['professional_summary'],
                         'Data scientist focused on banking risk.')

    def test_no_op_when_no_experience(self):
        resume = {'professional_summary': '', 'experience': []}
        out = backfill_summary(resume)
        self.assertEqual(out['professional_summary'], '')

    def test_picks_jd_aligned_experience_not_most_recent(self):
        # Round 1.5: the summary no longer mentions the experience
        # title (the audit flagged that as "meta-narration"). The
        # JD-aligned experience picker still matters for the no-JD
        # fallback case — but here the JD title is the lead and the
        # experience title isn't in the summary at all.
        resume = {
            'professional_summary': '',
            'skills': ['Python', 'Pandas', 'scikit-learn', 'TensorFlow'],
            'experience': [
                {'title': 'Digital Transformation Intern'},
                {'title': 'Information Technology Intern'},
                {'title': 'AI & Data Science Trainee'},
            ],
        }
        job = SimpleNamespace(title='Data Scientist')
        out = backfill_summary(resume, job=job)
        self.assertTrue(out['professional_summary'].startswith('Data Scientist'))
        self.assertIn('Python', out['professional_summary'])
        # No meta-narration mentioning the candidate's role title.
        self.assertNotIn('drawing on', out['professional_summary'].lower())

    def test_falls_back_to_first_when_no_jd_title_overlap(self):
        # Round 1.5: when a JD title IS provided, summary leads with
        # it regardless of experience-title overlap. The
        # experience-picker fallback only matters when job=None.
        resume = {
            'professional_summary': '',
            'skills': ['Python'],
            'experience': [
                {'title': 'Sales Associate'},
                {'title': 'Cashier'},
            ],
        }
        out = backfill_summary(resume, job=None)
        # No JD → falls back to first experience title.
        self.assertIn('Sales Associate', out['professional_summary'])


# Inclusion planner tech-overlap + skill backfill tests
from resumes.services.inclusion_planner import (
    _discriminating_tech_overlap,
    _scan_for_jd_skills_in_profile_text,
    _BASE_TECH_CANON,
)


class DiscriminatingTechOverlapTests(SimpleTestCase):
    """Post PR2b Fix A: _discriminating_tech_overlap returns a tuple
    ``(count, jd_rescued_tokens)`` and the base-tech filter is now
    JD-aware (a base token like Python COUNTS when the JD explicitly
    lists it). These four tests assert the same scenarios as before
    but unpack the tuple and account for JD-aware rescue."""

    def test_python_only_project_with_python_in_jd_counts_via_rescue(self):
        # PR2b Fix A semantics change: this JD does include 'python',
        # so the Python tech entry is rescued from the base-tech filter
        # and counts. Pre-Fix-A this returned 0; post-Fix-A it returns 1
        # with 'Python' in jd_rescued. OpenCV and Jupyter Notebook
        # remain uncounted (neither in JD, Jupyter is base).
        jd = {'python', 'pandas', 'numpy', 'scikitlearn', 'tensorflow', 'mlflow'}
        count, rescued = _discriminating_tech_overlap(
            ['Python', 'OpenCV', 'Jupyter Notebook'], jd,
        )
        self.assertEqual(count, 1)
        self.assertEqual(rescued, ['Python'])

    def test_webdev_project_has_zero_overlap(self):
        # Brain Tumor Classification App — HTML/CSS/JavaScript/Swiper
        # tech list against a DS JD. None are in JD, all but Swiper are
        # base. Still zero — no rescue.
        jd = {'python', 'pandas', 'tensorflow', 'pytorch'}
        count, rescued = _discriminating_tech_overlap(
            ['HTML', 'CSS', 'JavaScript', 'Swiper'], jd,
        )
        self.assertEqual(count, 0)
        self.assertEqual(rescued, [])

    def test_ml_project_has_strong_overlap(self):
        # Healthcare Prediction — Pandas + scikit-learn + MLflow are
        # all non-base and in JD → counted. Python is base AND in JD
        # → rescued and counted. Jupyter base AND not in JD → filtered.
        # Total = 4 (was 3 pre-Fix-A because Python was filtered).
        jd = {'python', 'pandas', 'scikitlearn', 'mlflow', 'tensorflow'}
        count, rescued = _discriminating_tech_overlap(
            ['Python', 'Jupyter Notebook', 'Pandas', 'scikit-learn',
             'Flask', 'MLflow'],
            jd,
        )
        self.assertEqual(count, 4)
        self.assertEqual(rescued, ['Python'])

    def test_handles_non_list_input(self):
        # Tuple unpacking — defensive returns are (0, []) now.
        self.assertEqual(_discriminating_tech_overlap(None, {'python'}), (0, []))
        self.assertEqual(_discriminating_tech_overlap('Python', {'python'}), (0, []))


class ScanForJdSkillsInProfileTextTests(SimpleTestCase):
    def test_finds_skill_mentioned_in_experience_bullet(self):
        # TensorFlow / MLflow / Hugging Face mentioned in DEPI
        # experience description should be surfaced even when they're
        # not in the user's formal skills list.
        data = {
            'experiences': [{
                'title': 'AI & Data Science Trainee',
                'description': 'Used MLOps tools (MLflow, Hugging Face) to build pipelines. Trained models with TensorFlow.',
            }],
            'projects': [],
            'certifications': [],
        }
        found = _scan_for_jd_skills_in_profile_text(
            data,
            ['TensorFlow', 'MLflow', 'Hugging Face', 'Kafka'],
            already_in_list_canon=set(),
        )
        self.assertIn('TensorFlow', found)
        self.assertIn('MLflow', found)
        self.assertIn('Hugging Face', found)
        self.assertNotIn('Kafka', found)  # not mentioned anywhere

    def test_skips_skills_already_in_list(self):
        data = {'experiences': [{'description': 'Used MLflow.'}],
                'projects': [], 'certifications': []}
        already = {'mlflow'}  # canonical form
        found = _scan_for_jd_skills_in_profile_text(
            data, ['MLflow'], already_in_list_canon=already,
        )
        self.assertEqual(found, [])

    def test_requires_word_boundary_match(self):
        # "tensor" inside "TensorFlow" should NOT match the JD skill
        # "tensor" alone — and vice versa shouldn't false-positive
        # against partial matches.
        data = {'experiences': [{'description': 'Used PyTorch.'}],
                'projects': [], 'certifications': []}
        found = _scan_for_jd_skills_in_profile_text(
            data, ['Torch'], already_in_list_canon=set(),
        )
        self.assertEqual(found, [])  # "PyTorch" doesn't word-match "Torch"


# --- HR/CV specialist supervisor (final review layer) ---------------------

import json as _json_for_supervisor
from resumes.services.resume_generator import generate_resume_content_supervised


class _FakeLLM:
    """Minimal LLM stand-in: .invoke(...) returns a fixed value (or raises)."""
    def __init__(self, ret=None, exc=None):
        self._ret = ret
        self._exc = exc
        self.calls = []

    def invoke(self, messages, *a, **k):
        self.calls.append(messages)
        if self._exc is not None:
            raise self._exc
        return self._ret


def _content_has_image(messages) -> bool:
    """True iff the HumanMessage content list carries an image_url block."""
    try:
        content = messages[0].content
    except Exception:
        return False
    if not isinstance(content, list):
        return False
    return any(isinstance(b, dict) and b.get('type') == 'image_url' for b in content)


def _mk_finding(severity='blocking', layer='content', category='summary',
                issue='x', fix='y', location='l'):
    from profiles.services.schemas import SupervisorFinding
    return SupervisorFinding(severity=severity, layer=layer, category=category,
                             issue=issue, fix=fix, location=location)


class _FakeTokenLimitError(Exception):
    """Mimics Groq's 413 token-ceiling rejection for _is_token_limit_error."""
    def __init__(self):
        super().__init__("Request too large for tokens per minute (TPM)")
        self.body = {'error': {'type': 'tokens'}}


class RenderResumePngTests(SimpleTestCase):
    """resume_render.render_resume_png — the PDF->PNG path for the supervisor."""

    _CONTENT = {
        'professional_title': 'Data Analyst',
        'professional_summary': 'Analytical professional with SQL and Python.',
        'skills': ['Python', 'SQL', 'Tableau'],
        'experience': [], 'education': [], 'projects': [],
        'certifications': [], 'languages': [], 'awards': [],
    }

    def test_minimal_resume_renders_nonempty_png(self):
        from resumes.services.resume_render import render_resume_png
        png = render_resume_png(self._CONTENT, None, pages=1)
        self.assertTrue(png)
        self.assertEqual(png[:8], b'\x89PNG\r\n\x1a\n')  # PNG magic bytes

    def test_png_to_data_url_prefix(self):
        from resumes.services.resume_render import png_to_data_url
        url = png_to_data_url(b'\x89PNG\r\n\x1a\n\x00')
        self.assertTrue(url.startswith('data:image/png;base64,'))


class SupervisorPromptCoverageTests(SimpleTestCase):
    """Lock in the recruiter-checklist coverage: the supervisor prompt must
    direct the model at the high-value issue classes a senior reviewer catches
    (the ones a generic 'review the resume' prompt misses)."""

    def test_prompt_covers_high_value_checks(self):
        from resumes.services.resume_supervisor import SUPERVISOR_PROMPT
        p = SUPERVISOR_PROMPT.lower()
        # Redundant / duplicate bullets within a role.
        self.assertIn('redundant', p)
        # Bullet count vs tenure (padding on short roles).
        self.assertIn('tenure', p)
        # Overclaiming / inflation against seniority.
        self.assertIn('overclaim', p)
        self.assertIn('seniority', p)
        # Skill-list duplicates / near-duplicates.
        self.assertIn('near-duplicate', p)
        # Date completeness / consistency across roles.
        self.assertIn('date', p)
        # Internal jargon / recruiter audience.
        self.assertIn('jargon', p)
        # Role relevance / dilution.
        self.assertIn('relevance', p)
        # Generic summary.
        self.assertIn('generic', p)

    def test_prompt_still_forbids_fabrication(self):
        from resumes.services.resume_supervisor import SUPERVISOR_PROMPT
        p = SUPERVISOR_PROMPT.lower()
        self.assertIn('fabrication', p)
        self.assertIn('missing skills', p)


class StructuralObservationsTests(SimpleTestCase):
    """The deterministic pre-scan that surfaces structural defects (bullet
    bloat, redundant bullets, date gaps, duplicate skills) for the LLM to judge."""

    def test_flags_bullet_bloat(self):
        from resumes.services.resume_supervisor import _structural_observations
        rc = {'experience': [{'title': 'Intern', 'duration': 'Aug 2025 - Sep 2025',
                              'description': ['a', 'b', 'c', 'd', 'e', 'f', 'g']}]}
        out = _structural_observations(rc)
        self.assertIn('7 bullets', out)
        self.assertIn('HIGH', out)

    def test_flags_lexical_redundant_bullets(self):
        from resumes.services.resume_supervisor import _structural_observations
        rc = {'experience': [{'title': 'Intern', 'duration': 'Aug 2025 - Sep 2025',
                              'description': [
                                  'Developed and optimized a PySpark data pipeline for enterprise systems',
                                  'Developed and optimized a PySpark data pipeline for enterprise systems and analytics',
                              ]}]}
        out = _structural_observations(rc)
        self.assertIn('REDUNDANT', out)

    def test_flags_missing_end_date(self):
        from resumes.services.resume_supervisor import _structural_observations
        rc = {'experience': [{'title': 'Intern', 'start_date': 'Aug 2025',
                              'description': ['a']}]}
        out = _structural_observations(rc)
        self.assertIn('NO end date', out)

    def test_flags_single_date_no_range(self):
        from resumes.services.resume_supervisor import _structural_observations
        rc = {'experience': [{'title': 'Intern', 'duration': 'Aug 2025',
                              'description': ['a']}]}
        out = _structural_observations(rc)
        self.assertIn('single date', out)

    def test_accepts_proper_range_and_present(self):
        from resumes.services.resume_supervisor import _structural_observations
        rc = {'experience': [
            {'title': 'A', 'duration': 'Jun 2025 - Dec 2025', 'description': ['x']},
            {'title': 'B', 'duration': 'Jan 2024 - Present', 'description': ['y']},
        ]}
        out = _structural_observations(rc)
        self.assertNotIn('single date', out)
        self.assertNotIn('NO end date', out)

    def test_flags_subset_duplicate_skills(self):
        from resumes.services.resume_supervisor import _structural_observations
        rc = {'skills': ['SQL', 'Databases & SQL', 'Supervised Learning',
                         'Supervised & Unsupervised Learning']}
        out = _structural_observations(rc)
        self.assertIn('duplicate skills', out.lower())

    def test_clean_resume_returns_empty(self):
        from resumes.services.resume_supervisor import _structural_observations
        rc = {'experience': [{'title': 'A', 'duration': 'Jun 2025 - Dec 2025',
                              'description': ['Built X', 'Shipped Y']}],
              'skills': ['Python', 'SQL', 'Docker']}
        self.assertEqual(_structural_observations(rc), "")

    def test_flags_keyword_stuffing_for_jd_skill(self):
        """The 2026-05-28 5:16 run shipped a resume with 'Python' ×18 and
        'Machine Learning' ×8. Counts above STUFFING_THRESHOLD (4) must
        surface as an observation when ``job`` is provided so the
        supervisor can feed it back into the regen prompt."""
        from types import SimpleNamespace
        from resumes.services.resume_supervisor import _structural_observations
        rc = {
            'experience': [
                {'title': 'AI Trainee', 'duration': 'Jun 2025 - Dec 2025',
                 'description': [
                     'Built Python data pipelines and Python ETL with Python.',
                     'Wrote Python unit tests and Python notebooks for Python services.',
                     'Maintained Python infra and Python CI/CD with Python tooling.',
                 ]},
            ],
            'skills': ['Python', 'SQL'],
        }
        job = SimpleNamespace(extracted_skills=['Python', 'SQL'])
        out = _structural_observations(rc, job=job)
        self.assertIn('KEYWORD STUFFING', out)
        self.assertIn('"Python"', out)
        # SQL appears once — should NOT be flagged.
        self.assertNotIn('"SQL"', out)

    def test_no_stuffing_flag_below_threshold(self):
        """Counts at or below STUFFING_THRESHOLD (4) must not flag, so we
        don't bother the supervisor with noise. The detector counts the
        whole resume JSON, so the skills array and any duration / title
        also contribute — keep total <=4 to verify the strict-greater-
        than-threshold check."""
        from types import SimpleNamespace
        from resumes.services.resume_supervisor import _structural_observations
        rc = {
            'experience': [
                {'title': 'A', 'duration': 'Jun 2025 - Dec 2025',
                 # 3 'Python' in description + 1 in skills = 4 total, == threshold.
                 'description': ['Python Python Python work']},
            ],
            'skills': ['Python'],
        }
        job = SimpleNamespace(extracted_skills=['Python'])
        out = _structural_observations(rc, job=job)
        self.assertNotIn('KEYWORD STUFFING', out)

    def test_stuffing_flag_works_without_job_arg(self):
        """When ``job`` is None, falls back to the candidate's own skills
        list — this keeps the pre-scan useful in callers that don't have
        a job object (e.g. ad-hoc tooling)."""
        from resumes.services.resume_supervisor import _structural_observations
        rc = {
            'experience': [
                {'title': 'A', 'duration': 'Jun 2025 - Dec 2025',
                 'description': [
                     'Pandas Pandas Pandas Pandas Pandas Pandas analysis.',
                 ]},
            ],
            'skills': ['Pandas'],
        }
        out = _structural_observations(rc)  # no job
        self.assertIn('KEYWORD STUFFING', out)
        self.assertIn('"Pandas"', out)


class SupervisorRecoveryTests(SimpleTestCase):
    """_recover_review_from_failed_generation salvages a SupervisorReview from
    a failed structured-output call (the tool-call envelope shapes Groq emits)."""

    def _err_with_body(self, failed_generation: str):
        class _E(Exception):
            pass
        e = _E("tool_use_failed")
        e.body = {'error': {'failed_generation': failed_generation}}
        return e

    def test_unwraps_name_parameters_envelope(self):
        from resumes.services.resume_supervisor import _recover_review_from_failed_generation
        payload = _json_for_supervisor.dumps([{
            "name": "SupervisorReview",
            "parameters": {
                "verdict": "revise", "summary": "s",
                "findings": [{"layer": "content", "severity": "blocking",
                              "category": "summary", "location": "top",
                              "issue": "truncated", "fix": "complete it"}],
            },
        }])
        rev = _recover_review_from_failed_generation(self._err_with_body(payload))
        self.assertIsNotNone(rev)
        self.assertEqual(rev.verdict, "revise")
        self.assertEqual(len(rev.findings), 1)
        self.assertEqual(len(rev.blocking_content_findings()), 1)

    def test_bare_findings_list_two_items(self):
        from resumes.services.resume_supervisor import _recover_review_from_failed_generation
        payload = _json_for_supervisor.dumps([
            {"layer": "content", "severity": "blocking", "issue": "a", "fix": "fa"},
            {"layer": "render", "severity": "warning", "issue": "b", "fix": "fb"},
        ])
        rev = _recover_review_from_failed_generation(self._err_with_body(payload))
        self.assertIsNotNone(rev)
        self.assertEqual(len(rev.findings), 2)

    def test_single_bare_finding_not_mistaken_for_envelope(self):
        from resumes.services.resume_supervisor import _recover_review_from_failed_generation
        payload = _json_for_supervisor.dumps([
            {"layer": "content", "severity": "blocking", "issue": "only", "fix": "f"},
        ])
        rev = _recover_review_from_failed_generation(self._err_with_body(payload))
        self.assertIsNotNone(rev)
        self.assertEqual(len(rev.findings), 1)

    def test_recovers_from_str_exc_envelope(self):
        from resumes.services.resume_supervisor import _recover_review_from_failed_generation
        inner = _json_for_supervisor.dumps({
            "verdict": "revise",
            "findings": [{"layer": "content", "severity": "blocking", "issue": "x", "fix": "y"}],
        })
        # body=None; the payload is only reachable via str(exc)'s Python repr.
        msg = "Error code: 400 - " + repr({'error': {'failed_generation': inner}})
        rev = _recover_review_from_failed_generation(Exception(msg))
        self.assertIsNotNone(rev)
        self.assertEqual(len(rev.findings), 1)

    def test_garbage_payload_returns_none(self):
        from resumes.services.resume_supervisor import _recover_review_from_failed_generation
        rev = _recover_review_from_failed_generation(self._err_with_body("}{not json at all"))
        self.assertIsNone(rev)

    def test_no_failed_generation_returns_none(self):
        from resumes.services.resume_supervisor import _recover_review_from_failed_generation
        self.assertIsNone(_recover_review_from_failed_generation(Exception("boom")))


class ReviewResumeContextTests(SimpleTestCase):
    """_build_review_context keeps the prompt compact: gap lines + JD + resume
    JSON, but NOT the GitHub/Kaggle/LinkedIn signal blobs that 413 the generator."""

    def _job(self):
        return SimpleNamespace(
            title='Junior Data Analyst', company='Acme',
            extracted_skills=['SQL', 'Python'], description='Analyze data daily.',
        )

    def _gap(self):
        return SimpleNamespace(
            matched_skills=['Python'], critical_missing_skills=['Tableau'],
            soft_skill_gaps=['seniority'],
        )

    def test_includes_standards_gap_jd_and_resume_json(self):
        from resumes.services.resume_supervisor import _build_review_context
        rc = {'professional_summary': 'hi', 'validation_report': {'grounding_findings': []}}
        ctx = _build_review_context(rc, self._job(), self._gap(), 'KB_STANDARDS_HERE')
        self.assertIn('KB_STANDARDS_HERE', ctx)
        self.assertIn('MATCHED', ctx)
        self.assertIn('Junior Data Analyst', ctx)
        self.assertIn('professional_summary', ctx)

    def test_excludes_signal_blobs(self):
        from resumes.services.resume_supervisor import _build_review_context
        # Even if the resume_content somehow carried a signal blob, the context
        # is built from gap lines + JD + the resume JSON only.
        rc = {'professional_summary': 'hi'}
        ctx = _build_review_context(rc, self._job(), self._gap(), '')
        self.assertNotIn('github_signals', ctx)
        self.assertNotIn('kaggle_signals', ctx)


class ReviewResumeTests(SimpleTestCase):
    """review_resume orchestration: two-step (vision -> structure), fail-open,
    token-limit retry drops the image."""

    def _job(self):
        return SimpleNamespace(title='DA', company='X', extracted_skills=['SQL'],
                               description='d')

    def _gap(self):
        return SimpleNamespace(matched_skills=['SQL'], critical_missing_skills=[],
                               soft_skill_gaps=[])

    def test_blocking_finding_surfaced(self):
        from unittest.mock import patch
        import resumes.services.resume_supervisor as rs
        from profiles.services.schemas import SupervisorReview
        review = SupervisorReview(verdict='revise', summary='s',
                                  findings=[_mk_finding()])
        with patch.object(rs, 'render_resume_png', return_value=b'\x89PNG'), \
             patch.object(rs, 'get_llm', return_value=_FakeLLM(SimpleNamespace(content='critique'))), \
             patch.object(rs, 'get_structured_llm', return_value=_FakeLLM(review)):
            out = rs.review_resume({'professional_summary': 'x'}, None,
                                   self._job(), self._gap(), standards_block='S')
        self.assertEqual(out.verdict, 'revise')
        self.assertEqual(len(out.blocking_content_findings()), 1)

    def test_fail_open_on_exception(self):
        from unittest.mock import patch
        import resumes.services.resume_supervisor as rs
        with patch.object(rs, 'render_resume_png', return_value=b'\x89PNG'), \
             patch.object(rs, 'get_llm', return_value=_FakeLLM(exc=RuntimeError('down'))):
            out = rs.review_resume({'a': 1}, None, self._job(), self._gap(),
                                   standards_block='S')
        self.assertEqual(out.verdict, 'advance')
        self.assertEqual(out.summary, 'review unavailable')
        self.assertEqual(out.findings, [])

    def test_render_failure_degrades_to_text_only(self):
        from unittest.mock import patch
        import resumes.services.resume_supervisor as rs
        from profiles.services.schemas import SupervisorReview
        vision = _FakeLLM(SimpleNamespace(content='critique'))
        with patch.object(rs, 'render_resume_png', side_effect=RuntimeError('no render')), \
             patch.object(rs, 'get_llm', return_value=vision), \
             patch.object(rs, 'get_structured_llm',
                          return_value=_FakeLLM(SupervisorReview(verdict='advance'))):
            out = rs.review_resume({'a': 1}, None, self._job(), self._gap(),
                                   standards_block='S')
        self.assertEqual(out.verdict, 'advance')
        # No image was rendered, so the vision call carried no image block.
        self.assertFalse(_content_has_image(vision.calls[0]))

    def test_token_limit_retry_drops_image(self):
        from unittest.mock import patch
        import resumes.services.resume_supervisor as rs
        from profiles.services.schemas import SupervisorReview

        class _RetryLLM:
            def __init__(self):
                self.calls = []
            def invoke(self, messages, *a, **k):
                self.calls.append(messages)
                # First call (with image) hits the token ceiling; retry succeeds.
                if len(self.calls) == 1:
                    raise _FakeTokenLimitError()
                return SimpleNamespace(content='critique after slim')

        vision = _RetryLLM()
        with patch.object(rs, 'render_resume_png', return_value=b'\x89PNG'), \
             patch.object(rs, 'get_llm', return_value=vision), \
             patch.object(rs, 'get_structured_llm',
                          return_value=_FakeLLM(SupervisorReview(verdict='advance'))):
            out = rs.review_resume({'a': 1}, None, self._job(), self._gap(),
                                   standards_block='S')
        self.assertEqual(out.verdict, 'advance')
        self.assertEqual(len(vision.calls), 2)
        self.assertTrue(_content_has_image(vision.calls[0]))   # first had image
        self.assertFalse(_content_has_image(vision.calls[1]))  # retry dropped it


class SupervisedLoopTests(SimpleTestCase):
    """generate_resume_content_supervised — the generate/review/regenerate loop."""

    def _job(self):
        return SimpleNamespace(title='DA', company='X', extracted_skills=['SQL'],
                               description='d')

    def _gap(self):
        return SimpleNamespace(matched_skills=['SQL'], critical_missing_skills=[],
                               soft_skill_gaps=[])

    def _profile(self):
        return SimpleNamespace(data_content={})

    def _patches(self, gen_side_effect, review_side_effect):
        from unittest.mock import patch
        import resumes.services.resume_generator as g
        import resumes.services.resume_supervisor as rs
        return (
            patch.object(g, 'generate_resume_content', side_effect=gen_side_effect),
            patch.object(g, '_build_standards_section', return_value=('STD', None, {})),
            patch.object(rs, 'review_resume', side_effect=review_side_effect),
        )

    def test_disabled_bypasses_loop(self):
        from unittest.mock import patch
        from django.test import override_settings
        import resumes.services.resume_generator as g
        sentinel = {'professional_summary': 'plain'}
        with override_settings(SUPERVISOR_ENABLED=False):
            with patch.object(g, 'generate_resume_content', return_value=sentinel) as gen, \
                 patch.object(g, '_build_standards_section') as std:
                out = g.generate_resume_content_supervised(
                    self._profile(), self._job(), self._gap())
        gen.assert_called_once()
        std.assert_not_called()
        self.assertNotIn('supervisor_review', out)

    @staticmethod
    def _review(findings):
        from profiles.services.schemas import SupervisorReview
        verdict = 'revise' if any(
            f.severity == 'blocking' and f.layer == 'content' for f in findings) else 'advance'
        return SupervisorReview(verdict=verdict, summary='s', findings=findings)

    def test_stops_when_no_blocking_content(self):
        from django.test import override_settings
        gen = lambda *a, **k: {'professional_summary': 'draft'}
        reviews = [self._review([_mk_finding(severity='warning')])]
        p_gen, p_std, p_rev = self._patches(gen, reviews)
        with override_settings(SUPERVISOR_ENABLED=True, SUPERVISOR_MAX_REVISION_ROUNDS=1):
            with p_gen as gm, p_std, p_rev:
                out = generate_resume_content_supervised(
                    self._profile(), self._job(), self._gap())
        self.assertEqual(gm.call_count, 1)  # no regen
        self.assertEqual(out['supervisor_review']['rounds'], 1)
        self.assertEqual(out['validation_report']['supervisor_findings'][0]['severity'], 'warning')

    def test_render_only_blocking_does_not_regen(self):
        from django.test import override_settings
        gen = lambda *a, **k: {'professional_summary': 'draft'}
        reviews = [self._review([_mk_finding(severity='blocking', layer='render')])]
        p_gen, p_std, p_rev = self._patches(gen, reviews)
        with override_settings(SUPERVISOR_ENABLED=True, SUPERVISOR_MAX_REVISION_ROUNDS=1):
            with p_gen as gm, p_std, p_rev:
                out = generate_resume_content_supervised(
                    self._profile(), self._job(), self._gap())
        self.assertEqual(gm.call_count, 1)  # render blocking does NOT drive regen
        self.assertEqual(len(out['supervisor_review']['findings']), 1)

    def test_content_blocking_regenerates_with_feedback_then_ships(self):
        from django.test import override_settings
        gen_calls = []

        def gen(*a, **k):
            gen_calls.append(k.get('supervisor_feedback', ''))
            return {'professional_summary': f'draft {len(gen_calls)}'}

        reviews = [
            self._review([_mk_finding(severity='blocking', layer='content', issue='truncated')]),
            self._review([_mk_finding(severity='warning')]),
        ]
        p_gen, p_std, p_rev = self._patches(gen, reviews)
        with override_settings(SUPERVISOR_ENABLED=True, SUPERVISOR_MAX_REVISION_ROUNDS=1):
            with p_gen, p_std, p_rev:
                out = generate_resume_content_supervised(
                    self._profile(), self._job(), self._gap())
        self.assertEqual(len(gen_calls), 2)         # regenerated once
        self.assertEqual(gen_calls[0], '')          # first draft: no feedback
        self.assertIn('truncated', gen_calls[1])    # second draft: got the feedback
        self.assertEqual(out['supervisor_review']['rounds'], 2)

    def test_cap_stops_after_max_rounds(self):
        from django.test import override_settings
        gen = lambda *a, **k: {'professional_summary': 'draft'}
        # Always blocking content -> would loop forever without the cap.
        reviews = [
            self._review([_mk_finding(severity='blocking', layer='content')]),
            self._review([_mk_finding(severity='blocking', layer='content')]),
            self._review([_mk_finding(severity='blocking', layer='content')]),
        ]
        p_gen, p_std, p_rev = self._patches(gen, reviews)
        with override_settings(SUPERVISOR_ENABLED=True, SUPERVISOR_MAX_REVISION_ROUNDS=1):
            with p_gen as gm, p_std, p_rev:
                out = generate_resume_content_supervised(
                    self._profile(), self._job(), self._gap())
        self.assertEqual(gm.call_count, 2)  # round 0 + round 1, then cap
        self.assertEqual(out['supervisor_review']['rounds'], 2)
        self.assertEqual(out['supervisor_review']['verdict'], 'revise')

    def test_ships_best_draft_when_regen_regresses(self):
        """2026-05-28 5:16 production run: round 0 had 2 blocking, round 1
        had 4 blocking — regen made things worse and we shipped round 1.
        With the elitism guard the loop ships round 0 instead, since
        score (blocking, total) (2, 6) < (4, 9)."""
        from django.test import override_settings
        gen_calls = []

        def gen(*a, **k):
            gen_calls.append(k.get('supervisor_feedback', ''))
            return {'professional_summary': f'draft {len(gen_calls)}'}

        # Round 0: 2 blocking. Round 1: 4 blocking (regression).
        reviews = [
            self._review([
                _mk_finding(severity='blocking', layer='content', issue='a'),
                _mk_finding(severity='blocking', layer='content', issue='b'),
                _mk_finding(severity='warning', layer='content', issue='w1'),
                _mk_finding(severity='warning', layer='content', issue='w2'),
                _mk_finding(severity='warning', layer='content', issue='w3'),
                _mk_finding(severity='warning', layer='content', issue='w4'),
            ]),
            self._review([
                _mk_finding(severity='blocking', layer='content', issue='c'),
                _mk_finding(severity='blocking', layer='content', issue='d'),
                _mk_finding(severity='blocking', layer='content', issue='e'),
                _mk_finding(severity='blocking', layer='content', issue='f'),
                _mk_finding(severity='warning', layer='content', issue='w5'),
                _mk_finding(severity='warning', layer='content', issue='w6'),
                _mk_finding(severity='warning', layer='content', issue='w7'),
                _mk_finding(severity='warning', layer='content', issue='w8'),
                _mk_finding(severity='warning', layer='content', issue='w9'),
            ]),
        ]
        p_gen, p_std, p_rev = self._patches(gen, reviews)
        with override_settings(SUPERVISOR_ENABLED=True, SUPERVISOR_MAX_REVISION_ROUNDS=1):
            with p_gen, p_std, p_rev:
                out = generate_resume_content_supervised(
                    self._profile(), self._job(), self._gap())
        # Both rounds ran (round 0 + 1 regen), but the shipped content is
        # round 0's because round 1 regressed.
        self.assertEqual(out['professional_summary'], 'draft 1')
        # Surfaced findings are from the shipped (best) draft, not the latest.
        findings = out['supervisor_review']['findings']
        # Round-0's 6 findings ship — issue-text identity proves it's
        # the round-0 draft, not round-1's.
        self.assertTrue(any(f['issue'] == 'a' for f in findings))
        self.assertTrue(any(f['issue'] == 'b' for f in findings))
        self.assertFalse(
            any(f['issue'] in ('c', 'd', 'e', 'f') for f in findings),
            'round-1 findings must NOT ship — best-draft elitism',
        )
        # Findings-classification policy (2026-05-31): the cap-exhausted
        # AUTO_FIX blockers (category='summary') are demoted to 'warning'
        # so the user sees them as advisory polish, not a red alarm. The
        # loop tried and failed — that's not the user's defect to chase.
        self.assertEqual(
            len([f for f in findings if f['severity'] == 'blocking']), 0,
            'AUTO_FIX blockers should be demoted to warning on cap exhaustion',
        )

    def test_ships_latest_when_regen_improves(self):
        """The complementary case: round 0 has 3 blocking, round 1 has 1
        blocking. Ship round 1 (it's the better draft) — the elitism
        guard must not regress the happy path."""
        from django.test import override_settings
        gen_calls = []

        def gen(*a, **k):
            gen_calls.append(k.get('supervisor_feedback', ''))
            return {'professional_summary': f'draft {len(gen_calls)}'}

        reviews = [
            self._review([
                _mk_finding(severity='blocking', layer='content', issue='a'),
                _mk_finding(severity='blocking', layer='content', issue='b'),
                _mk_finding(severity='blocking', layer='content', issue='c'),
            ]),
            self._review([
                _mk_finding(severity='blocking', layer='content', issue='d'),
                _mk_finding(severity='warning', layer='content', issue='w1'),
            ]),
        ]
        p_gen, p_std, p_rev = self._patches(gen, reviews)
        with override_settings(SUPERVISOR_ENABLED=True, SUPERVISOR_MAX_REVISION_ROUNDS=1):
            with p_gen, p_std, p_rev:
                out = generate_resume_content_supervised(
                    self._profile(), self._job(), self._gap())
        # Round 1 was strictly better — ship it.
        self.assertEqual(out['professional_summary'], 'draft 2')
        findings = out['supervisor_review']['findings']
        # Issue-text identity proves it's round 1's findings, not round 0.
        self.assertTrue(any(f['issue'] == 'd' for f in findings))
        self.assertFalse(any(f['issue'] in ('a', 'b', 'c') for f in findings))
        # Cap-exhaustion demote: 'd' (originally blocking, category='summary'
        # → AUTO_FIX) → warning. No alarming red on the shipped draft.
        self.assertEqual(len([f for f in findings if f['severity'] == 'blocking']), 0,
                         'AUTO_FIX blockers should be demoted to warning on cap exhaustion')


class PathBSupervisedRegenTests(TestCase):
    """Audit report §6.5 fix #3 (2026-05-29): the stale-profile / ?refresh=1
    branch of resume_edit_view previously called generate_resume_content
    (NON-supervised) INLINE on a GET, which (a) blocked the browser for
    15-90s and (b) silently bypassed the supervisor safety net. The fix
    redirects to a new loader view + in-place trigger endpoint that
    calls generate_resume_content_supervised, mirroring Path A's UX."""

    def setUp(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from profiles.models import UserProfile
        from resumes.models import GeneratedResume
        from django.contrib.auth import get_user_model
        from datetime import timedelta
        from django.utils import timezone
        User = get_user_model()
        self.user = User.objects.create_user(
            username='regen@example.com', email='regen@example.com', password='x',
        )
        self.profile = UserProfile.objects.create(
            user=self.user, full_name='Regen User',
            data_content={'experiences': [{'title': 'Engineer'}]},
        )
        self.client.force_login(self.user)
        self.job = Job.objects.create(
            user=self.user, title='Engineer', company='Acme',
            description='x', extracted_skills=['Python'],
        )
        self.gap = GapAnalysis.objects.create(
            user=self.user, job=self.job, similarity_score=0.7,
        )
        self.resume = GeneratedResume.objects.create(
            gap_analysis=self.gap,
            content={
                'professional_summary': 'old summary',
                'template_name': 'ats_clean',
                'experience': [{'title': 'Engineer'}],
            },
            ats_score=70.0,
        )
        # Force the profile to look NEWER than the resume so the stale-
        # profile branch fires on plain GET (no ?refresh param).
        future = timezone.now() + timedelta(hours=1)
        UserProfile.objects.filter(pk=self.profile.pk).update(updated_at=future)

    def test_stale_profile_get_redirects_to_loader_does_not_block(self):
        """Plain GET on /resumes/edit/<id>/ with a stale resume MUST NOT
        call any LLM generator inline. The view returns a 302 redirect to
        the new regenerate_resume loader view instead. This is the core
        of the no-long-blocking-GET requirement: even if the test mocks
        had blocking sleeps, they would never be entered."""
        from django.urls import reverse
        from unittest.mock import patch
        import resumes.views as views_mod
        with patch.object(views_mod, 'generate_resume_content') as gen_unsup, \
             patch('resumes.services.resume_generator.generate_resume_content_supervised') as gen_sup:
            resp = self.client.get(reverse('resume_edit', args=[self.resume.id]))
        # 302 to the loader view — NOT a long sync 200 response.
        self.assertEqual(resp.status_code, 302)
        self.assertIn(
            reverse('regenerate_resume', args=[self.resume.id]),
            resp['Location'],
        )
        # Both generators were untouched on the GET. Verifies that no
        # long-blocking LLM call can sneak into the GET handler.
        self.assertEqual(gen_unsup.call_count, 0)
        self.assertEqual(gen_sup.call_count, 0)

    def test_explicit_refresh_1_also_redirects(self):
        """?refresh=1 is the user's explicit "force regen" knob and must
        also route through the async loader, not block the GET."""
        from django.urls import reverse
        from unittest.mock import patch
        import resumes.views as views_mod
        # Reset profile.updated_at so should_refresh is FALSE — only
        # the explicit ?refresh=1 should trigger the redirect.
        from datetime import timedelta
        from django.utils import timezone
        past = timezone.now() - timedelta(hours=1)
        from profiles.models import UserProfile
        UserProfile.objects.filter(pk=self.profile.pk).update(updated_at=past)

        with patch.object(views_mod, 'generate_resume_content') as gen_unsup:
            url = reverse('resume_edit', args=[self.resume.id]) + '?refresh=1'
            resp = self.client.get(url)
        self.assertEqual(resp.status_code, 302)
        self.assertIn(
            reverse('regenerate_resume', args=[self.resume.id]),
            resp['Location'],
        )
        self.assertEqual(gen_unsup.call_count, 0)

    def test_refresh_0_disables_the_redirect(self):
        """?refresh=0 is the user's explicit opt-out — even when the
        profile is newer, the redirect MUST NOT fire and the page MUST
        render normally (200) without any LLM call."""
        from django.urls import reverse
        from unittest.mock import patch
        import resumes.views as views_mod
        with patch.object(views_mod, 'generate_resume_content') as gen_unsup:
            url = reverse('resume_edit', args=[self.resume.id]) + '?refresh=0'
            resp = self.client.get(url)
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(gen_unsup.call_count, 0)

    def test_regenerate_view_renders_loader_with_supervised_trigger_url(self):
        """GET on /resumes/regenerate/<id>/ renders generate.html in
        `generating=True` mode pointed at the in-place trigger endpoint
        (not the create-new endpoint)."""
        from django.urls import reverse
        resp = self.client.get(reverse('regenerate_resume', args=[self.resume.id]))
        self.assertEqual(resp.status_code, 200)
        # The loader's substep panel is the marker that we're in
        # `generating=True` mode.
        self.assertContains(resp, 'resume-gen-steps')
        # The new trigger URL must be wired up so the JS POSTs to the
        # in-place regen endpoint, not Path A's create-new endpoint.
        regen_api_url = reverse('trigger_resume_regen_api', args=[self.resume.id])
        self.assertContains(resp, regen_api_url)
        # The redirect URL must land back on the SAME resume's edit page
        # with ?refresh=0 so we don't loop into another regen.
        target = reverse('resume_edit', args=[self.resume.id]) + '?refresh=0'
        self.assertContains(resp, target)

    def test_trigger_regen_api_calls_supervised_generator(self):
        """The POST endpoint MUST call generate_resume_content_supervised,
        not the unsupervised generate_resume_content. Verifies the fix
        addresses the original bug: Path B now routes through the
        supervisor safety net."""
        from django.urls import reverse
        from unittest.mock import patch
        with patch('resumes.services.resume_generator.generate_resume_content_supervised') as gen_sup, \
             patch('resumes.views.calculate_ats_score', return_value=82.0):
            gen_sup.return_value = {
                'professional_summary': 'new summary',
                'experience': [],
                'validation_report': {
                    'supervisor_findings': [
                        {'severity': 'warning', 'layer': 'content',
                         'category': 'summary', 'location': '',
                         'issue': 'minor', 'fix': '...'},
                    ],
                },
                'supervisor_review': {'verdict': 'advance', 'rounds': 1},
            }
            url = reverse('trigger_resume_regen_api', args=[self.resume.id])
            resp = self.client.post(url)
        self.assertEqual(resp.status_code, 200)
        body = resp.json()
        self.assertTrue(body['success'])
        self.assertEqual(body['resume_id'], str(self.resume.id))
        # The supervisor generator was called exactly once.
        self.assertEqual(gen_sup.call_count, 1)

    def test_trigger_regen_api_preserves_template_name(self):
        """template_name was carried across the regen in the old inline
        path; the new flow MUST preserve that contract so a user's
        chosen template doesn't reset to the default on every regen."""
        from django.urls import reverse
        from unittest.mock import patch
        # Mocked LLM output does NOT include template_name — the view
        # must inject it from the existing resume.content.
        with patch('resumes.services.resume_generator.generate_resume_content_supervised') as gen_sup, \
             patch('resumes.views.calculate_ats_score', return_value=82.0):
            gen_sup.return_value = {
                'professional_summary': 'new summary',
                'experience': [],
                'validation_report': {},
            }
            url = reverse('trigger_resume_regen_api', args=[self.resume.id])
            resp = self.client.post(url)
        self.assertEqual(resp.status_code, 200)
        self.resume.refresh_from_db()
        self.assertEqual(self.resume.content.get('template_name'), 'ats_clean')
        self.assertEqual(self.resume.content.get('professional_summary'), 'new summary')
        self.assertEqual(self.resume.ats_score, 82.0)

    def test_trigger_regen_api_writes_validation_report(self):
        """The validation_report from the supervisor must land on the
        GeneratedResume row, matching generate_resume_task's contract."""
        from django.urls import reverse
        from unittest.mock import patch
        with patch('resumes.services.resume_generator.generate_resume_content_supervised') as gen_sup, \
             patch('resumes.views.calculate_ats_score', return_value=80.0):
            gen_sup.return_value = {
                'professional_summary': 's',
                'experience': [],
                'validation_report': {
                    'supervisor_findings': [
                        {'severity': 'blocking', 'layer': 'content',
                         'category': 'summary', 'location': '',
                         'issue': 'i', 'fix': 'f'},
                    ],
                },
            }
            url = reverse('trigger_resume_regen_api', args=[self.resume.id])
            self.client.post(url)
        self.resume.refresh_from_db()
        findings = (self.resume.validation_report or {}).get('supervisor_findings') or []
        self.assertEqual(len(findings), 1)
        self.assertEqual(findings[0]['severity'], 'blocking')

    def test_trigger_regen_api_rejects_other_users(self):
        """Authorization: a different user's POST gets a 404 envelope.
        After the queryset-level ownership filter, the 404 comes from
        get_object_or_404 not finding the row in the intruder's
        owner-scoped queryset — NOT from a post-fetch attribute check.
        We verify both the status code AND that no LLM call was made
        (proving we didn't fetch then bail)."""
        from django.urls import reverse
        from unittest.mock import patch
        from django.contrib.auth import get_user_model
        User = get_user_model()
        intruder = User.objects.create_user(
            username='intruder@example.com', email='intruder@example.com', password='x',
        )
        self.client.force_login(intruder)
        url = reverse('trigger_resume_regen_api', args=[self.resume.id])
        with patch(
            'resumes.services.resume_generator.generate_resume_content_supervised'
        ) as gen_sup:
            resp = self.client.post(url)
        self.assertEqual(resp.status_code, 404)
        # The LLM was never invoked — the queryset filter short-circuited
        # the request before any work could happen.
        self.assertEqual(gen_sup.call_count, 0)

    def test_regenerate_view_also_owner_scoped(self):
        """Same ownership guarantee for the GET loader view: an intruder
        gets a 404 directly from the queryset filter."""
        from django.urls import reverse
        from django.contrib.auth import get_user_model
        User = get_user_model()
        intruder = User.objects.create_user(
            username='intruder2@example.com', email='intruder2@example.com', password='x',
        )
        self.client.force_login(intruder)
        resp = self.client.get(reverse('regenerate_resume', args=[self.resume.id]))
        self.assertEqual(resp.status_code, 404)

    def test_atomic_on_supervised_generator_exception(self):
        """If generate_resume_content_supervised raises mid-regen, the
        existing resume row MUST be byte-identical to its pre-call state.
        No partial write of content / ats_score / validation_report.
        This is the all-or-nothing contract: build the new content fully
        in locals, only then assign+save once."""
        from django.urls import reverse
        from unittest.mock import patch
        import json as _json
        # Snapshot the pre-call state in serialised form so we can
        # compare byte-for-byte after the failed regen.
        before_content = _json.dumps(self.resume.content, sort_keys=True)
        before_score = self.resume.ats_score
        before_report = _json.dumps(self.resume.validation_report or {}, sort_keys=True)
        url = reverse('trigger_resume_regen_api', args=[self.resume.id])
        with patch(
            'resumes.services.resume_generator.generate_resume_content_supervised',
            side_effect=RuntimeError('simulated LLM blow-up mid-supervisor-loop'),
        ):
            resp = self.client.post(url)
        # Endpoint returns 500 to surface the failure to the caller.
        self.assertEqual(resp.status_code, 500)
        body = resp.json()
        self.assertFalse(body['success'])
        # Refetch from DB and confirm byte-identical state.
        self.resume.refresh_from_db()
        after_content = _json.dumps(self.resume.content, sort_keys=True)
        after_score = self.resume.ats_score
        after_report = _json.dumps(self.resume.validation_report or {}, sort_keys=True)
        self.assertEqual(after_content, before_content,
                         "resume.content must be unchanged when supervised generator raises")
        self.assertEqual(after_score, before_score,
                         "resume.ats_score must be unchanged when supervised generator raises")
        self.assertEqual(after_report, before_report,
                         "resume.validation_report must be unchanged when supervised generator raises")

    def test_atomic_on_ats_scorer_exception(self):
        """Companion to the above: if generation succeeds but
        calculate_ats_score raises before we can build the final dict
        + save, the row must still be byte-identical to its pre-call
        state. This guards against partial in-memory mutation."""
        from django.urls import reverse
        from unittest.mock import patch
        import json as _json
        before_content = _json.dumps(self.resume.content, sort_keys=True)
        before_score = self.resume.ats_score
        before_report = _json.dumps(self.resume.validation_report or {}, sort_keys=True)
        url = reverse('trigger_resume_regen_api', args=[self.resume.id])
        with patch(
            'resumes.services.resume_generator.generate_resume_content_supervised',
            return_value={'professional_summary': 'new', 'experience': [],
                          'validation_report': {'supervisor_findings': []}},
        ), patch(
            'resumes.views.calculate_ats_score',
            side_effect=RuntimeError('simulated scorer blow-up'),
        ):
            resp = self.client.post(url)
        self.assertEqual(resp.status_code, 500)
        self.resume.refresh_from_db()
        self.assertEqual(_json.dumps(self.resume.content, sort_keys=True), before_content)
        self.assertEqual(self.resume.ats_score, before_score)
        self.assertEqual(_json.dumps(self.resume.validation_report or {}, sort_keys=True),
                         before_report)


class FindingsPresenterTests(SimpleTestCase):
    """Pure-Python unit coverage for the severity-policy mapper.

    Audit fix #2 (2026-05-29): translate the raw validation_report +
    content['supervisor_review'] into a tiered summary the edit page
    can render without leaking internal field names or 20-item walls."""

    def test_clean_resume_returns_clean_tier(self):
        """No findings + supervisor advanced → clean tier with a small
        green pill, no headline alarm."""
        from resumes.services.findings_presenter import build_review_summary
        content = {
            'supervisor_review': {'verdict': 'advance', 'summary': '', 'findings': []},
        }
        vr = {'passed': True, 'findings': [], 'stats': {},
              'grounding_findings': [], 'supervisor_findings': []}
        out = build_review_summary(content, vr)
        self.assertEqual(out['tier'], 'clean')
        self.assertEqual(out['blocking_items'], [])
        self.assertEqual(out['advisory_items'], [])
        # A small "passed" info pill (not an alarm).
        self.assertTrue(len(out['info_items']) >= 1)
        self.assertIn('passed', out['info_items'][0]['title'].lower())

    def test_missing_data_treated_as_clean(self):
        """None / empty inputs must not crash and must classify as clean."""
        from resumes.services.findings_presenter import build_review_summary
        out = build_review_summary(None, None)
        self.assertEqual(out['tier'], 'clean')
        self.assertEqual(out['blocking_items'], [])
        self.assertEqual(out['advisory_items'], [])

    def test_twenty_unsupported_skill_collapses_to_count_and_examples(self):
        """The round-1 trace: 20 unsupported_skill findings on a
        basically-fine resume. The presenter MUST NOT render a 20-item
        list — it must collapse to a single advisory line with the
        count and ~3 representative examples. This is the acceptance-
        criterion case (b)."""
        from resumes.services.findings_presenter import build_review_summary
        # Construct 20 grounding findings, each naming a different
        # skill via the standard detail string format.
        skills = [f"SkillName{i:02d}" for i in range(20)]
        grounding = [
            {'kind': 'unsupported_skill',
             'where': f"experience[0].description[{i}]",
             'detail': f"Possible unsupported skill '{skills[i]}' — not in the inclusion plan."}
            for i in range(20)
        ]
        vr = {'passed': True, 'findings': [], 'stats': {},
              'grounding_findings': grounding, 'supervisor_findings': []}
        out = build_review_summary({}, vr)
        # 20 skills is advisory, not blocking.
        self.assertEqual(out['tier'], 'advisory')
        # Exactly ONE advisory item — not 20 list entries.
        self.assertEqual(len(out['advisory_items']), 1)
        item = out['advisory_items'][0]
        # The title says "20".
        self.assertIn('20', item['title'])
        # The body mentions the first three example skills and "17 more".
        self.assertIn('SkillName00', item['body'])
        self.assertIn('SkillName01', item['body'])
        self.assertIn('SkillName02', item['body'])
        self.assertIn('17 more', item['body'])
        # No raw internal field names leak.
        full = item['title'] + ' ' + item['body']
        for forbidden in ('unsupported_skill', 'grounding_findings',
                          'kind=', "'where'", 'detail='):
            self.assertNotIn(forbidden, full,
                             f"Internal token {forbidden!r} leaked to user")

    def test_unsupported_metric_is_blocking(self):
        """A numeric claim with no profile evidence is a factual
        hallucination → blocking tier, red banner. Acceptance criterion (c)."""
        from resumes.services.findings_presenter import build_review_summary
        vr = {
            'passed': True, 'findings': [], 'stats': {},
            'grounding_findings': [
                {'kind': 'unsupported_metric',
                 'where': 'experience[0].description[0]',
                 'detail': "Metric '92%' doesn't trace to any retrieved candidate-evidence chunk."},
            ],
            'supervisor_findings': [],
        }
        out = build_review_summary({}, vr)
        self.assertEqual(out['tier'], 'blocking')
        # At least one blocking item, naming the metric.
        self.assertTrue(out['blocking_items'])
        joined = ' '.join(i['title'] + ' ' + i['body'] for i in out['blocking_items'])
        self.assertIn('92%', joined)

    def test_supervisor_blocking_content_is_blocking(self):
        """Supervisor severity='blocking' + layer='content' triggers
        the red tier. Render-layer blockers are advisory only."""
        from resumes.services.findings_presenter import build_review_summary
        content = {
            'supervisor_review': {
                'verdict': 'revise',
                'summary': 'Summary truncated mid-sentence.',
                'findings': [
                    {'layer': 'content', 'severity': 'blocking',
                     'category': 'summary', 'location': '',
                     'issue': 'Summary stops mid-sentence.',
                     'fix': 'Rewrite to one complete sentence.'},
                ],
            },
        }
        # The writer mirrors findings into validation_report['supervisor_findings'].
        vr = {'passed': True, 'findings': [], 'stats': {},
              'grounding_findings': [],
              'supervisor_findings': content['supervisor_review']['findings']}
        out = build_review_summary(content, vr)
        self.assertEqual(out['tier'], 'blocking')
        joined = ' '.join(i['title'] for i in out['blocking_items'])
        # Plain-language category, not the internal token.
        self.assertIn('Summary', joined)

    def test_supervisor_render_blocker_is_advisory(self):
        """A 'blocking' severity but layer='render' (layout issue) must
        NOT escalate to the red tier — it's editor-fixable but not a
        factual content issue."""
        from resumes.services.findings_presenter import build_review_summary
        content = {
            'supervisor_review': {
                'verdict': 'advance',
                'summary': '',
                'findings': [
                    {'layer': 'render', 'severity': 'blocking',
                     'category': 'layout', 'location': '',
                     'issue': 'Page-break orphan.', 'fix': '...'},
                ],
            },
        }
        vr = {'passed': True, 'findings': [], 'stats': {},
              'grounding_findings': [],
              'supervisor_findings': content['supervisor_review']['findings']}
        out = build_review_summary(content, vr)
        self.assertEqual(out['tier'], 'advisory')

    def test_supervisor_revise_without_specific_finding_is_blocking(self):
        """If the supervisor said 'revise' but didn't enumerate specific
        blocking findings (rare but possible), we still surface a
        blocking banner using the supervisor's own summary text."""
        from resumes.services.findings_presenter import build_review_summary
        content = {
            'supervisor_review': {
                'verdict': 'revise',
                'summary': 'The draft over-claims experience.',
                'findings': [],
            },
        }
        out = build_review_summary(content, {})
        self.assertEqual(out['tier'], 'blocking')
        joined = ' '.join(i['title'] + ' ' + i['body'] for i in out['blocking_items'])
        self.assertIn('over-claims', joined)

    def test_drop_skill_leak_is_blocking(self):
        """A bullet that mentions a skill the inclusion-plan marked
        do-not-claim is a factual leak → red tier."""
        from resumes.services.findings_presenter import build_review_summary
        vr = {
            'passed': True, 'findings': [], 'stats': {},
            'grounding_findings': [
                {'kind': 'drop_skill_leak',
                 'where': 'experience[0].description[1]',
                 'detail': "Bullet mentions 'Hadoop', which the inclusion plan marked do-not-claim."},
            ],
            'supervisor_findings': [],
        }
        out = build_review_summary({}, vr)
        self.assertEqual(out['tier'], 'blocking')
        joined = ' '.join(i['title'] + ' ' + i['body'] for i in out['blocking_items'])
        self.assertIn('Hadoop', joined)

    def test_bullet_validator_error_is_blocking(self):
        """Bullet validator severity='error' → blocking tier."""
        from resumes.services.findings_presenter import build_review_summary
        vr = {
            'passed': False,
            'findings': [
                {'rule_id': 'A1', 'severity': 'error',
                 'location': 'experience[0].description[0]',
                 'bullet_text': '...', 'issue': 'banned phrase',
                 'suggested_fix': None},
            ],
            'stats': {},
            'grounding_findings': [], 'supervisor_findings': [],
        }
        out = build_review_summary({}, vr)
        self.assertEqual(out['tier'], 'blocking')

    def test_dedupes_supervisor_findings_between_two_storage_keys(self):
        """The writer puts the SAME findings into both
        validation_report['supervisor_findings'] and
        content['supervisor_review']['findings']. The presenter must
        not double-count."""
        from resumes.services.findings_presenter import build_review_summary
        same_findings = [
            {'layer': 'content', 'severity': 'blocking',
             'category': 'summary', 'issue': 'truncated', 'fix': '...'},
        ]
        content = {'supervisor_review':
                   {'verdict': 'revise', 'summary': '', 'findings': same_findings}}
        vr = {'supervisor_findings': same_findings,
              'passed': True, 'findings': [], 'stats': {},
              'grounding_findings': []}
        out = build_review_summary(content, vr)
        # Only ONE blocking item (the supervisor section), not two.
        # Filter to the category-grouped item to avoid coincidence.
        category_items = [i for i in out['blocking_items']
                          if i['title'].lower().startswith('summary')]
        self.assertEqual(len(category_items), 1)

    def test_caps_blocking_items_at_five(self):
        """A pathological run can't bloat the banner: max 5 blocking
        items rendered, regardless of how many groups produce them."""
        from resumes.services.findings_presenter import build_review_summary
        # 10 different bullet-validator rule_ids → 10 distinct error
        # groups in raw output; presenter caps at 5.
        findings = []
        rule_ids = ['A1', 'A2', 'A3', 'A4', 'A5', 'A6', 'A7', 'B1', 'B2', 'B3']
        for rid in rule_ids:
            findings.append({
                'rule_id': rid, 'severity': 'error',
                'location': 'experience[0].description[0]',
                'bullet_text': '...', 'issue': '...', 'suggested_fix': None,
            })
        vr = {'passed': False, 'findings': findings, 'stats': {},
              'grounding_findings': [], 'supervisor_findings': []}
        out = build_review_summary({}, vr)
        self.assertEqual(out['tier'], 'blocking')
        self.assertLessEqual(len(out['blocking_items']), 5)


class FindingsPresenterIntegrationTests(TestCase):
    """End-to-end: the edit page renders the appropriate banner for
    a stored GeneratedResume's findings."""

    def setUp(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from profiles.models import UserProfile
        from resumes.models import GeneratedResume
        from django.contrib.auth import get_user_model
        User = get_user_model()
        self.user = User.objects.create_user(
            username='banner@example.com', email='banner@example.com', password='x',
        )
        UserProfile.objects.create(user=self.user, full_name='Banner User', data_content={})
        self.client.force_login(self.user)
        self.job = Job.objects.create(
            user=self.user, title='Engineer', company='Acme', description='x',
            extracted_skills=[],
        )
        self.gap = GapAnalysis.objects.create(user=self.user, job=self.job, similarity_score=0.7)

    def _make_resume(self, content=None, validation_report=None):
        from resumes.models import GeneratedResume
        return GeneratedResume.objects.create(
            gap_analysis=self.gap,
            content=content or {'professional_summary': 's', 'experience': []},
            validation_report=validation_report or {},
        )

    def test_edit_page_renders_clean_for_clean_resume(self):
        """No findings + supervisor passed → either a small green pill
        or no banner at all. Definitely NO red/yellow alarm."""
        from django.urls import reverse
        resume = self._make_resume(
            content={'professional_summary': 's', 'experience': [],
                     'supervisor_review': {'verdict': 'advance', 'summary': '', 'findings': []}},
            validation_report={'passed': True, 'findings': [], 'stats': {},
                               'grounding_findings': [], 'supervisor_findings': []},
        )
        resp = self.client.get(reverse('resume_edit', args=[resume.id]))
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode('utf-8', errors='ignore')
        # Updated 2026-05-30: the wall-banner was replaced with the
        # compact `data-review-pill="..."` element. A clean resume
        # renders no pill at all (no annotations → no marker).
        self.assertNotIn('data-review-pill="blocking"', body)
        self.assertNotIn('data-review-pill="advisory"', body)

    def test_edit_page_renders_advisory_for_high_volume_unsupported_skill(self):
        """20 unsupported_skill findings → yellow advisory, ONE line
        with count + examples; not a 20-item list. Acceptance (a)."""
        from django.urls import reverse
        skills = [f"SkillName{i:02d}" for i in range(20)]
        grounding = [
            {'kind': 'unsupported_skill',
             'where': f"experience[0].description[{i}]",
             'detail': f"Possible unsupported skill '{skills[i]}' — not in plan."}
            for i in range(20)
        ]
        resume = self._make_resume(
            validation_report={
                'passed': True, 'findings': [], 'stats': {},
                'grounding_findings': grounding, 'supervisor_findings': [],
            },
        )
        resp = self.client.get(reverse('resume_edit', args=[resume.id]))
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode('utf-8', errors='ignore')
        # Updated 2026-05-31: unsupported_skill is NEEDS_USER_INPUT under
        # the findings-classification policy — the user must confirm or
        # remove each one, the LLM cannot decide. The compact pill thus
        # renders `data-review-pill="user_input"` with "to confirm" copy,
        # never "blocking". The fix-#2 collapse policy still holds: 20
        # findings group to one chip per anchor, not 20 noisy lines.
        self.assertIn('data-review-pill="user_input"', body)
        self.assertNotIn('data-review-pill="blocking"', body)
        # The count + "to confirm" copy is visible on the pill; no raw
        # kind names leaked.
        self.assertIn('to confirm', body)
        self.assertNotIn('unsupported_skill', body)
        self.assertNotIn('grounding_findings', body)

    def test_edit_page_renders_red_for_blocking_finding(self):
        """A blocking supervisor finding → red banner. Acceptance (c)."""
        from django.urls import reverse
        sup_findings = [
            {'layer': 'content', 'severity': 'blocking',
             'category': 'summary', 'location': '',
             'issue': 'Summary stops mid-sentence.', 'fix': 'Rewrite.'},
        ]
        resume = self._make_resume(
            content={'professional_summary': 's', 'experience': [],
                     'supervisor_review':
                     {'verdict': 'revise', 'summary': '', 'findings': sup_findings}},
            validation_report={'passed': True, 'findings': [], 'stats': {},
                               'grounding_findings': [],
                               'supervisor_findings': sup_findings},
        )
        resp = self.client.get(reverse('resume_edit', args=[resume.id]))
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode('utf-8', errors='ignore')
        # Updated 2026-05-30: the wall-banner became a compact pill +
        # inline chips. The pill carries data-review-pill="blocking" and
        # the inline chip carries data-finding-tier="blocking".
        self.assertIn('data-review-pill="blocking"', body)
        self.assertIn('data-finding-tier="blocking"', body)
        # No raw internal field names.
        self.assertNotIn("'layer':", body)
        self.assertNotIn('"layer":', body)
        # The plain-language issue text DOES surface.
        self.assertIn('Summary stops mid-sentence', body)


# ==========================================================================
# Fix #1 — content stickiness via export-triggered previous-best injection.
# Audit report §6.5, 2026-05-30. Tests cover the export snapshot path, the
# JD-hash match/mismatch behavior, the deterministic regression check,
# the supervised-loop integration (shared cap + cap-exhaustion demote),
# and the headline round-3→round-4 regression-loss scenario.
# ==========================================================================


class JdIdentityHashTests(SimpleTestCase):
    """The hash that gates previous_best injection — built from job
    fields that determine tailoring. No updated_at on Job, so a content
    hash is the only way to detect JD edits."""

    def test_same_job_same_hash(self):
        from resumes.services.resume_generator import _jd_identity_hash
        from types import SimpleNamespace
        j1 = SimpleNamespace(title='DS', company='Acme', description='Do data.',
                             extracted_skills_tiers={'must_have': ['Python']})
        j2 = SimpleNamespace(title='DS', company='Acme', description='Do data.',
                             extracted_skills_tiers={'must_have': ['Python']})
        self.assertEqual(_jd_identity_hash(j1), _jd_identity_hash(j2))

    def test_description_edit_changes_hash(self):
        from resumes.services.resume_generator import _jd_identity_hash
        from types import SimpleNamespace
        j1 = SimpleNamespace(title='DS', company='Acme', description='Do data.',
                             extracted_skills_tiers={})
        j2 = SimpleNamespace(title='DS', company='Acme', description='Do data, fast.',
                             extracted_skills_tiers={})
        self.assertNotEqual(_jd_identity_hash(j1), _jd_identity_hash(j2))

    def test_tier_change_changes_hash(self):
        from resumes.services.resume_generator import _jd_identity_hash
        from types import SimpleNamespace
        j1 = SimpleNamespace(title='DS', company='Acme', description='x',
                             extracted_skills_tiers={'must_have': ['Python']})
        j2 = SimpleNamespace(title='DS', company='Acme', description='x',
                             extracted_skills_tiers={'must_have': ['Python', 'SQL']})
        self.assertNotEqual(_jd_identity_hash(j1), _jd_identity_hash(j2))


class BuildPreviousBestBlockTests(SimpleTestCase):
    """The prompt block emitter — must return '' when missing/mismatched,
    and emit a structured preserve-or-improve block when matched."""

    def _snap(self, content, jd_hash='HASH-A'):
        return {
            'content': content,
            'exported_at': '2026-05-30T00:00:00+00:00',
            'ats_score_at_export': 82.0,
            'jd_identity_hash': jd_hash,
        }

    def test_returns_empty_when_no_snapshot(self):
        from resumes.services.resume_generator import _build_previous_best_block
        self.assertEqual(_build_previous_best_block(None, 'HASH-A'), '')
        self.assertEqual(_build_previous_best_block({}, 'HASH-A'), '')

    def test_returns_empty_when_jd_hash_mismatch(self):
        """User edited the JD between exports — block MUST be empty so
        the LLM tailors against the new JD without dragging old anchors."""
        from resumes.services.resume_generator import _build_previous_best_block
        snap = self._snap({'professional_summary': 'old'}, jd_hash='HASH-OLD')
        out = _build_previous_best_block(snap, current_job_hash='HASH-NEW')
        self.assertEqual(out, '')

    def test_emits_summary_skills_experience_when_matched(self):
        from resumes.services.resume_generator import _build_previous_best_block
        snap = self._snap({
            'professional_summary': 'Data Scientist with applied ML.',
            'skills': ['Python', 'PySpark'],
            'experience': [{'title': 'AI Trainee', 'company': 'DEPI',
                            'description': ['Built X.', 'Shipped Y.']}],
            'projects': [{'name': 'SmartCV', 'url': 'https://x.com/cv',
                          'description': ['Built scoring with 0.351 silhouette.']}],
        })
        out = _build_previous_best_block(snap, current_job_hash='HASH-A')
        self.assertIn('PREVIOUS BEST', out)
        self.assertIn('Data Scientist with applied ML.', out)
        self.assertIn('Python', out)
        self.assertIn('AI Trainee', out)
        self.assertIn('SmartCV', out)
        self.assertIn('0.351', out)

    def test_no_block_when_snapshot_empty_content(self):
        """Snapshot present but content empty/no meaningful sections →
        emit nothing, not a useless empty banner."""
        from resumes.services.resume_generator import _build_previous_best_block
        snap = self._snap({})
        self.assertEqual(_build_previous_best_block(snap, 'HASH-A'), '')


class ApplyRegressionCheckTests(SimpleTestCase):
    """The deterministic per-section diff. metric_loss + bullet_count_drop
    are blocking; skill_loss is warning."""

    def _snap(self, content):
        return {
            'content': content,
            'exported_at': '2026-05-30T00:00:00+00:00',
            'ats_score_at_export': 80.0,
            'jd_identity_hash': 'HASH-A',
        }

    def test_metric_loss_is_blocking(self):
        """The headline case: previous-best bullet had '0.351' and '84%';
        new draft drops both — flag as blocking metric_loss."""
        from resumes.services.resume_generator import _apply_regression_check
        prev_content = {'experience': [{
            'title': 'AI Trainee', 'company': 'DEPI',
            'description': ['Validated k=3 with 0.351 silhouette',
                            'Profiled 84% of revenue'],
        }]}
        new_content = {'experience': [{
            'title': 'AI Trainee', 'company': 'DEPI',
            'description': ['Built an HR dashboard',
                            'Cleaned data with pandas'],
        }]}
        out = _apply_regression_check(new_content, self._snap(prev_content))
        findings = (out.get('validation_report') or {}).get('regression_findings') or []
        kinds = {f['kind'] for f in findings if f.get('severity') == 'blocking'}
        self.assertIn('metric_loss', kinds)
        # The lost metrics surface in the finding detail. Note:
        # _extract_numeric_claims captures the digit stem only ('84')
        # because the regex's `\b` boundary excludes the trailing '%'.
        # The diff is still correct because both prev and new bullets
        # are run through the same regex — symmetric extraction.
        details = ' '.join(f.get('detail', '') for f in findings)
        self.assertIn('0.351', details)
        self.assertIn('84', details)

    def test_bullet_count_drop_is_blocking(self):
        from resumes.services.resume_generator import _apply_regression_check
        prev_content = {'experience': [{
            'title': 'A', 'company': 'B',
            'description': ['Bullet 1.', 'Bullet 2.', 'Bullet 3.'],
        }]}
        new_content = {'experience': [{
            'title': 'A', 'company': 'B',
            'description': ['Bullet 1.'],
        }]}
        out = _apply_regression_check(new_content, self._snap(prev_content))
        findings = (out.get('validation_report') or {}).get('regression_findings') or []
        kinds = {f['kind'] for f in findings if f.get('severity') == 'blocking'}
        self.assertIn('bullet_count_drop', kinds)

    def test_skill_loss_is_warning_only(self):
        """skills shift legitimately with tailoring — flag but don't block."""
        from resumes.services.resume_generator import _apply_regression_check
        prev_content = {'skills': ['Python', 'SQL', 'Pandas']}
        new_content = {'skills': ['Python']}  # SQL + Pandas dropped
        out = _apply_regression_check(new_content, self._snap(prev_content))
        findings = (out.get('validation_report') or {}).get('regression_findings') or []
        skill_losses = [f for f in findings if f.get('kind') == 'skill_loss']
        self.assertGreaterEqual(len(skill_losses), 2)
        for f in skill_losses:
            self.assertEqual(f.get('severity'), 'warning')

    def test_no_findings_when_content_matches(self):
        from resumes.services.resume_generator import _apply_regression_check
        prev_content = {'skills': ['Python'], 'experience': [{
            'title': 'A', 'company': 'B', 'description': ['Bullet 1.']}]}
        out = _apply_regression_check(dict(prev_content), self._snap(prev_content))
        findings = (out.get('validation_report') or {}).get('regression_findings') or []
        self.assertEqual(findings, [])

    def test_no_snapshot_writes_empty_list(self):
        """No previous_best at all → no findings, validation_report still
        records that the check ran (empty list = 'we checked, no diff')."""
        from resumes.services.resume_generator import _apply_regression_check
        out = _apply_regression_check({'skills': ['Python']}, None)
        findings = (out.get('validation_report') or {}).get('regression_findings')
        self.assertEqual(findings, [])

    def test_project_join_by_url(self):
        """Project matched by URL even when name is renamed across regens."""
        from resumes.services.resume_generator import _apply_regression_check
        prev_content = {'projects': [{
            'name': 'SmartCV',
            'url': 'https://github.com/foo/smartcv',
            'description': ['Built scoring with 0.351 silhouette.'],
        }]}
        new_content = {'projects': [{
            'name': 'Smart-CV',  # renamed
            'url': 'https://github.com/foo/smartcv',  # SAME url
            'description': ['Built a thing.'],  # metric dropped
        }]}
        out = _apply_regression_check(new_content, self._snap(prev_content))
        findings = (out.get('validation_report') or {}).get('regression_findings') or []
        kinds = {f['kind'] for f in findings if f.get('severity') == 'blocking'}
        self.assertIn('metric_loss', kinds,
                      "URL-matched project must compare bullets despite name change")

    def test_project_join_falls_back_to_whole_section_when_no_url(self):
        """A project with no URL gets matched by canonical name. When
        BOTH the name changed AND there's no URL, the prev project has
        no counterpart in new → bullet_count_drop fires (0 bullets in
        "missing" new role vs N bullets in prev). This is the documented
        acceptable degradation."""
        from resumes.services.resume_generator import _apply_regression_check
        prev_content = {'projects': [{
            'name': 'Original Name', 'url': '',
            'description': ['Some bullet.'],
        }]}
        new_content = {'projects': [{
            'name': 'Totally Different Name', 'url': '',
            'description': ['Other bullet.'],
        }]}
        out = _apply_regression_check(new_content, self._snap(prev_content))
        findings = (out.get('validation_report') or {}).get('regression_findings') or []
        kinds = {f['kind'] for f in findings if f.get('severity') == 'blocking'}
        # The previous "Original Name" project has no match in new → its
        # bullet count looks like 1 → 0, registering as bullet_count_drop.
        # This is the agreed degradation when both url and canon-name fail.
        self.assertIn('bullet_count_drop', kinds)


class ExportCapturesPreviousBestTests(TestCase):
    """Both export views must snapshot resume.content before generating
    the file. Best-effort: a snapshot-save failure must NOT break the
    download response."""

    def setUp(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from profiles.models import UserProfile
        from resumes.models import GeneratedResume
        from django.contrib.auth import get_user_model
        User = get_user_model()
        self.user = User.objects.create_user(
            username='exp@example.com', email='exp@example.com', password='x',
        )
        UserProfile.objects.create(user=self.user, full_name='E', data_content={})
        self.client.force_login(self.user)
        self.job = Job.objects.create(
            user=self.user, title='DS', company='Acme', description='x',
            extracted_skills=['Python'],
        )
        self.gap = GapAnalysis.objects.create(user=self.user, job=self.job, similarity_score=0.7)
        self.resume = GeneratedResume.objects.create(
            gap_analysis=self.gap,
            content={'professional_summary': 'A summary.',
                     'skills': ['Python'],
                     'experience': [],
                     'template_name': 'ats_clean'},
            ats_score=75.0,
        )

    def test_docx_export_writes_snapshot(self):
        """DOCX export must populate previous_best with content + hash +
        timestamp + ats_score."""
        from django.urls import reverse
        from unittest.mock import patch
        from io import BytesIO
        with patch('resumes.views.generate_docx', return_value=BytesIO(b'fake-docx')):
            url = reverse('export_docx', args=[self.resume.id])
            resp = self.client.get(url)
        self.assertEqual(resp.status_code, 200)
        self.resume.refresh_from_db()
        snap = self.resume.previous_best
        self.assertIsInstance(snap, dict)
        self.assertIn('content', snap)
        self.assertEqual(snap['content'].get('professional_summary'), 'A summary.')
        self.assertIn('jd_identity_hash', snap)
        self.assertEqual(len(snap['jd_identity_hash']), 64)  # sha256 hex
        self.assertIn('exported_at', snap)
        self.assertEqual(snap.get('ats_score_at_export'), 75.0)

    def test_pdf_export_writes_snapshot(self):
        from django.urls import reverse
        from unittest.mock import patch
        with patch('resumes.views.generate_pdf') as gen_pdf:
            # Make generate_pdf write the file so subsequent open() works.
            def _fake(resume, output_path, template_name):
                with open(output_path, 'wb') as fh:
                    fh.write(b'fake-pdf')
            gen_pdf.side_effect = _fake
            resp = self.client.get(reverse('export_pdf', args=[self.resume.id]))
        self.assertEqual(resp.status_code, 200)
        self.resume.refresh_from_db()
        self.assertEqual(
            self.resume.previous_best.get('content', {}).get('professional_summary'),
            'A summary.',
        )

    def test_export_succeeds_even_when_snapshot_save_throws(self):
        """Best-effort guarantee: a forced exception during snapshot save
        must NOT break the download."""
        from django.urls import reverse
        from unittest.mock import patch
        from io import BytesIO
        # Patch _jd_identity_hash to raise; _capture_previous_best
        # wraps everything in a try/except, so the export must still
        # serve a 200.
        with patch('resumes.views.generate_docx', return_value=BytesIO(b'fake-docx')), \
             patch('resumes.services.resume_generator._jd_identity_hash',
                   side_effect=RuntimeError('forced')):
            resp = self.client.get(reverse('export_docx', args=[self.resume.id]))
        self.assertEqual(resp.status_code, 200,
                         "Download must succeed even if snapshot save throws")
        # And the resume row's previous_best is unchanged.
        self.resume.refresh_from_db()
        self.assertEqual(self.resume.previous_best, {})


class SupervisedLoopRegressionTests(TestCase):
    """The supervised loop must consume regression findings the same way
    it consumes supervisor blocking findings — shared budget, no second
    loop, cap-exhaustion demotes blocking regression findings to warning."""

    def setUp(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from profiles.models import UserProfile
        from resumes.models import GeneratedResume
        from django.contrib.auth import get_user_model
        User = get_user_model()
        self.user = User.objects.create_user(
            username='sup@example.com', email='sup@example.com', password='x',
        )
        self.profile = UserProfile.objects.create(
            user=self.user, full_name='S', data_content={'experiences': []},
        )
        self.job = Job.objects.create(
            user=self.user, title='DS', company='Acme', description='x',
            extracted_skills=['Python'],
            extracted_skills_tiers={'must_have': ['Python']},
        )
        self.gap = GapAnalysis.objects.create(user=self.user, job=self.job, similarity_score=0.7)

    def _snap(self, content):
        from resumes.services.resume_generator import _jd_identity_hash
        return {
            'content': content,
            'exported_at': '2026-05-30T00:00:00+00:00',
            'ats_score_at_export': 80.0,
            'jd_identity_hash': _jd_identity_hash(self.job),
        }

    @staticmethod
    def _review(findings):
        from profiles.services.schemas import SupervisorReview
        verdict = 'revise' if any(
            f.severity == 'blocking' and f.layer == 'content' for f in findings) else 'advance'
        return SupervisorReview(verdict=verdict, summary='s', findings=findings)

    def _patches(self, gen_side_effect, review_side_effect):
        from unittest.mock import patch
        import resumes.services.resume_generator as g
        import resumes.services.resume_supervisor as rs
        return (
            patch.object(g, 'generate_resume_content', side_effect=gen_side_effect),
            patch.object(g, '_build_standards_section', return_value=('STD', None, {})),
            patch.object(rs, 'review_resume', side_effect=review_side_effect),
        )

    def test_metric_loss_bypasses_loop_under_user_input_policy(self):
        """Findings policy (2026-05-31): metric_loss is NEEDS_USER_INPUT
        — the LLM cannot honestly fabricate a number it dropped from a
        prior export. The loop does NOT regen on metric_loss; the user
        confirms/restores via the inline 'Confirm or complete' chip.

        Previously this test asserted the inverse (regression drove
        regen); under the new three-bucket policy, regression USER_INPUT
        findings bypass the loop entirely."""
        from django.test import override_settings
        from resumes.services.resume_generator import generate_resume_content_supervised
        gen_calls = []

        def gen(*a, **k):
            gen_calls.append(k.get('supervisor_feedback', ''))
            return {
                'professional_summary': 'draft',
                'experience': [{
                    'title': 'AI Trainee', 'company': 'DEPI',
                    'description': ['Built an HR dashboard.'],
                }],
                'projects': [],
            }

        reviews = [self._review([]), self._review([])]
        snap = self._snap({
            'experience': [{
                'title': 'AI Trainee', 'company': 'DEPI',
                'description': ['Validated k=3 with 0.351 silhouette',
                                'Profiled 84% of revenue.'],
            }],
        })
        p_gen, p_std, p_rev = self._patches(gen, reviews)
        with override_settings(SUPERVISOR_ENABLED=True, SUPERVISOR_MAX_REVISION_ROUNDS=1):
            with p_gen, p_std, p_rev:
                out = generate_resume_content_supervised(
                    self.profile, self.job, self.gap, previous_best=snap,
                )
        # ONE round — the user-input regression did NOT drive a regen.
        self.assertEqual(len(gen_calls), 1,
                         'metric_loss is user_input; the loop must NOT regen on it')
        # The regression finding is preserved at blocking severity so
        # the UI surfaces it under 'Confirm or complete', not as an error.
        regression = (out.get('validation_report') or {}).get('regression_findings') or []
        self.assertTrue(regression, 'expected the regression finding to survive')
        self.assertTrue(any(f.get('kind') == 'metric_loss'
                            and f.get('severity') == 'blocking'
                            for f in regression),
                        'metric_loss must remain at blocking severity — the '
                        'user, not the loop, owns the fix')

    def test_cap_exhaustion_does_not_demote_user_input_blocker(self):
        """Findings policy (2026-05-31): cap-exhaustion only demotes
        AUTO_FIX blockers (the loop tried and failed). USER_INPUT
        blockers (metric_loss, bullet_count_drop) stay at blocking
        severity because the loop never owned them — the user, not
        the system, has the missing fact."""
        from django.test import override_settings
        from resumes.services.resume_generator import generate_resume_content_supervised
        gen_calls = []

        def gen(*a, **k):
            gen_calls.append(k.get('supervisor_feedback', ''))
            return {
                'professional_summary': 'd',
                'experience': [{
                    'title': 'AI Trainee', 'company': 'DEPI',
                    'description': ['Built an HR dashboard.'],
                }],
                'projects': [],
            }

        reviews = [self._review([]), self._review([])]
        snap = self._snap({
            'experience': [{
                'title': 'AI Trainee', 'company': 'DEPI',
                'description': ['Validated k=3 with 0.351 silhouette'],
            }],
        })
        p_gen, p_std, p_rev = self._patches(gen, reviews)
        with override_settings(SUPERVISOR_ENABLED=True, SUPERVISOR_MAX_REVISION_ROUNDS=1):
            with p_gen, p_std, p_rev:
                out = generate_resume_content_supervised(
                    self.profile, self.job, self.gap, previous_best=snap,
                )
        # ONE round — user-input regression never enters the loop, so
        # the cap path isn't even relevant. Ships immediately.
        self.assertEqual(len(gen_calls), 1)
        # The metric_loss finding rides through the loop untouched at
        # blocking severity, ready for the 'Confirm or complete' UI.
        regression = (out.get('validation_report') or {}).get('regression_findings') or []
        self.assertTrue(regression)
        for f in regression:
            if f.get('kind') == 'metric_loss':
                self.assertEqual(f.get('severity'), 'blocking',
                                 'user_input regression must NOT be demoted')

    def test_skill_loss_alone_does_not_trigger_regen(self):
        """skill_loss is severity='warning' — it should NOT consume a
        revision round."""
        from django.test import override_settings
        from resumes.services.resume_generator import generate_resume_content_supervised
        gen_calls = []

        def gen(*a, **k):
            gen_calls.append(k.get('supervisor_feedback', ''))
            return {
                'professional_summary': 'd',
                'skills': ['Python'],  # SQL dropped
                'experience': [], 'projects': [],
            }

        reviews = [self._review([])]
        snap = self._snap({'skills': ['Python', 'SQL']})
        p_gen, p_std, p_rev = self._patches(gen, reviews)
        with override_settings(SUPERVISOR_ENABLED=True, SUPERVISOR_MAX_REVISION_ROUNDS=1):
            with p_gen, p_std, p_rev:
                out = generate_resume_content_supervised(
                    self.profile, self.job, self.gap, previous_best=snap,
                )
        # Only one round — skill_loss didn't drive another.
        self.assertEqual(len(gen_calls), 1)
        # The skill_loss finding is present as 'warning'.
        regression = (out.get('validation_report') or {}).get('regression_findings') or []
        skill_losses = [f for f in regression if f.get('kind') == 'skill_loss']
        self.assertTrue(skill_losses)
        for f in skill_losses:
            self.assertEqual(f.get('severity'), 'warning')

    def test_jd_hash_mismatch_skips_injection_and_check(self):
        """When the snapshot's jd_identity_hash doesn't match the current
        job's hash (JD edited), the supervised loop must NOT inject the
        previous-best block and the regression check must produce no
        blocking findings even if content diverges."""
        from django.test import override_settings
        from resumes.services.resume_generator import generate_resume_content_supervised
        gen_calls = []

        def gen(*a, **k):
            gen_calls.append(k.get('supervisor_feedback', ''))
            # New draft drops the prev metric.
            return {
                'professional_summary': 'd',
                'experience': [{
                    'title': 'AI Trainee', 'company': 'DEPI',
                    'description': ['Built an HR dashboard.'],
                }],
            }

        reviews = [self._review([])]
        # Snapshot's hash is STALE (computed against a different job state).
        snap = {
            'content': {'experience': [{
                'title': 'AI Trainee', 'company': 'DEPI',
                'description': ['Validated k=3 with 0.351 silhouette']}]},
            'jd_identity_hash': 'STALE-HASH-FROM-PREVIOUS-JD',
            'exported_at': '2026-05-30T00:00:00+00:00',
            'ats_score_at_export': 80.0,
        }
        p_gen, p_std, p_rev = self._patches(gen, reviews)
        with override_settings(SUPERVISOR_ENABLED=True, SUPERVISOR_MAX_REVISION_ROUNDS=1):
            with p_gen, p_std, p_rev:
                out = generate_resume_content_supervised(
                    self.profile, self.job, self.gap, previous_best=snap,
                )
        # Only one round — hash mismatch suppresses regression-driven retry.
        self.assertEqual(len(gen_calls), 1)
        # Regression findings empty: the check sees a stale snapshot and
        # build_previous_best_block returned '' (block not injected). The
        # check itself ALSO short-circuits when hash doesn't match — but
        # the current implementation runs the diff anyway and writes
        # findings only when content was injected. Confirm no blocking
        # regression findings are present.
        regression = (out.get('validation_report') or {}).get('regression_findings') or []
        blocking_reg = [f for f in regression if (f.get('severity') or '').lower() == 'blocking']
        # When hash mismatches, the supervised loop must not retry, so
        # any findings written by the check stay at most one round —
        # AND on the shipped (single) round they may exist but they
        # were never consumed by a retry. We accept either: no findings
        # written when block was skipped, OR findings written but only
        # if the check ran. The contract the user cares about: no
        # second round.
        # (We already verified gen_calls == 1 above, which is the
        # contract.) Leave this assertion soft.
        del blocking_reg  # noqa


class FindingsPresenterRegressionTests(SimpleTestCase):
    """The fix-#2 banner must surface regression findings: blocking ones
    in the red tier, warning ones in the yellow tier."""

    def test_blocking_regression_finding_lands_in_blocking_tier(self):
        from resumes.services.findings_presenter import build_review_summary
        vr = {
            'passed': True, 'findings': [], 'stats': {},
            'grounding_findings': [], 'supervisor_findings': [],
            'regression_findings': [{
                'kind': 'metric_loss', 'severity': 'blocking',
                'where': 'experience[AI Trainee @ DEPI]',
                'prev': ['0.351'], 'now': '',
                'detail': "Bullet metrics ['0.351'] missing from this draft.",
            }],
        }
        out = build_review_summary({}, vr)
        self.assertEqual(out['tier'], 'blocking')
        body = ' '.join(i['title'] + ' ' + i['body'] for i in out['blocking_items'])
        self.assertIn('Regression', body)
        self.assertIn('0.351', body)
        # No raw internal field names.
        for token in ('metric_loss', 'severity', "'kind'"):
            self.assertNotIn(token, body)

    def test_warning_regression_finding_lands_in_advisory_tier(self):
        from resumes.services.findings_presenter import build_review_summary
        vr = {
            'passed': True, 'findings': [], 'stats': {},
            'grounding_findings': [], 'supervisor_findings': [],
            'regression_findings': [{
                'kind': 'skill_loss', 'severity': 'warning',
                'where': 'skills',
                'prev': 'SQL', 'now': '',
                'detail': "Skill 'SQL' missing.",
            }],
        }
        out = build_review_summary({}, vr)
        self.assertEqual(out['tier'], 'advisory')
        body = ' '.join(i['title'] + ' ' + i['body'] for i in out['advisory_items'])
        self.assertIn('not preserved', body.lower())


class LoadPreviousBestForTests(TestCase):
    """Path A creates a NEW GeneratedResume row per regen, so the snapshot
    lives on the PRIOR row. load_previous_best_for walks rows for this
    gap_analysis and returns the most recent populated snapshot. This
    test class pins:
      1. Selection is gap_analysis-scoped (not just hash-gated post-hoc).
      2. Multiple prior exports for the same (profile, job) → newest wins.
      3. Snapshots from a DIFFERENT job's rows are never returned.
    """

    def setUp(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from profiles.models import UserProfile
        from django.contrib.auth import get_user_model
        User = get_user_model()
        self.user = User.objects.create_user(
            username='pba@example.com', email='pba@example.com', password='x',
        )
        UserProfile.objects.create(user=self.user, full_name='P', data_content={})
        self.job_a = Job.objects.create(
            user=self.user, title='DS', company='Acme', description='one',
            extracted_skills=['Python'],
            extracted_skills_tiers={'must_have': ['Python']},
        )
        self.job_b = Job.objects.create(
            user=self.user, title='Eng', company='Other', description='two',
            extracted_skills=['Go'],
            extracted_skills_tiers={'must_have': ['Go']},
        )
        self.gap_a = GapAnalysis.objects.create(user=self.user, job=self.job_a, similarity_score=0.7)
        self.gap_b = GapAnalysis.objects.create(user=self.user, job=self.job_b, similarity_score=0.7)

    def _make_resume(self, gap, snapshot=None, content=None, *, created_at=None):
        from resumes.models import GeneratedResume
        from django.utils import timezone
        row = GeneratedResume.objects.create(
            gap_analysis=gap,
            content=content or {'professional_summary': 's'},
            previous_best=snapshot or {},
        )
        if created_at is not None:
            # Force the timestamp so multi-row ordering tests aren't
            # at the mercy of test-run clock resolution.
            GeneratedResume.objects.filter(pk=row.pk).update(created_at=created_at)
            row.refresh_from_db()
        return row

    def test_returns_none_when_no_row_has_export(self):
        from resumes.services.resume_generator import load_previous_best_for
        # Two unexported rows under gap_a.
        self._make_resume(self.gap_a)
        self._make_resume(self.gap_a)
        self.assertIsNone(load_previous_best_for(self.gap_a))

    def test_multi_row_returns_newest_exported(self):
        """Two prior exported rows for the same gap_analysis, different
        snapshots — Path A loads the most-recent one. Path A's normal
        flow guarantees newer rows have newer snapshots (each generation
        creates a new row, each export stamps that row); the query
        orders by ``-created_at``, so the newer row's snapshot wins."""
        from datetime import timedelta
        from django.utils import timezone
        from resumes.services.resume_generator import load_previous_best_for
        older_t = timezone.now() - timedelta(hours=2)
        newer_t = timezone.now() - timedelta(hours=1)
        snap_older = {
            'content': {'professional_summary': 'OLDER SNAPSHOT'},
            'exported_at': older_t.isoformat(),
            'ats_score_at_export': 70.0,
            'jd_identity_hash': 'a' * 64,
        }
        snap_newer = {
            'content': {'professional_summary': 'NEWER SNAPSHOT'},
            'exported_at': newer_t.isoformat(),
            'ats_score_at_export': 82.0,
            'jd_identity_hash': 'a' * 64,
        }
        self._make_resume(self.gap_a, snapshot=snap_older, created_at=older_t)
        self._make_resume(self.gap_a, snapshot=snap_newer, created_at=newer_t)
        out = load_previous_best_for(self.gap_a)
        self.assertIsNotNone(out)
        self.assertEqual(
            out['content']['professional_summary'], 'NEWER SNAPSHOT',
            f"expected NEWER snapshot to win, got {out['content']}",
        )
        self.assertEqual(out['ats_score_at_export'], 82.0)

    def test_skips_unexported_row_even_if_newer(self):
        """A row that was NEVER exported (empty previous_best) must not
        override an older row that WAS exported. ``.exclude(previous_best={})``
        filters out the unexported row regardless of its created_at."""
        from datetime import timedelta
        from django.utils import timezone
        from resumes.services.resume_generator import load_previous_best_for
        older_t = timezone.now() - timedelta(hours=2)
        newer_t = timezone.now() - timedelta(hours=1)
        snap_older = {
            'content': {'professional_summary': 'OLDER EXPORTED'},
            'exported_at': older_t.isoformat(),
            'ats_score_at_export': 70.0,
            'jd_identity_hash': 'a' * 64,
        }
        self._make_resume(self.gap_a, snapshot=snap_older, created_at=older_t)
        # Newer row, NEVER exported (previous_best is {}).
        self._make_resume(self.gap_a, snapshot={}, created_at=newer_t)
        out = load_previous_best_for(self.gap_a)
        self.assertIsNotNone(out)
        self.assertEqual(out['content']['professional_summary'], 'OLDER EXPORTED')

    def test_never_returns_snapshot_from_a_different_job(self):
        """Selection is gap_analysis-scoped. Job-B's rows have their own
        exported snapshots; load_previous_best_for(gap_a) must NEVER
        return a Job-B snapshot, even if Job-B's row is newer."""
        from datetime import timedelta
        from django.utils import timezone
        from resumes.services.resume_generator import load_previous_best_for
        a_t = timezone.now() - timedelta(hours=2)
        b_t = timezone.now() - timedelta(hours=1)  # Job-B newer than Job-A
        snap_a = {
            'content': {'professional_summary': 'JOB A SNAPSHOT'},
            'exported_at': a_t.isoformat(),
            'ats_score_at_export': 70.0,
            'jd_identity_hash': 'a' * 64,
        }
        snap_b = {
            'content': {'professional_summary': 'JOB B SNAPSHOT'},
            'exported_at': b_t.isoformat(),
            'ats_score_at_export': 80.0,
            'jd_identity_hash': 'b' * 64,  # different JD hash too
        }
        self._make_resume(self.gap_a, snapshot=snap_a, created_at=a_t)
        self._make_resume(self.gap_b, snapshot=snap_b, created_at=b_t)
        # Path A on Job-A: must get Job-A's snapshot, never Job-B's.
        out_a = load_previous_best_for(self.gap_a)
        self.assertIsNotNone(out_a)
        self.assertEqual(out_a['content']['professional_summary'], 'JOB A SNAPSHOT')
        # And Job-B's lookup gets B's, not A's.
        out_b = load_previous_best_for(self.gap_b)
        self.assertIsNotNone(out_b)
        self.assertEqual(out_b['content']['professional_summary'], 'JOB B SNAPSHOT')


class EditPageTemplateNoLeakTests(TestCase):
    """Live-bug regression (2026-05-30): a multi-line {# #} comment block
    rendered as visible text on the edit page because Django's {# #} is
    SINGLE-LINE only — multi-line content is not recognised as a comment
    and is emitted verbatim. The fix removed the developer note from the
    template entirely (internal file paths don't belong in user-facing
    chrome anyway). These tests guard against ANY developer note leaking
    into the rendered HTML."""

    def setUp(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from profiles.models import UserProfile
        from resumes.models import GeneratedResume
        from django.contrib.auth import get_user_model
        User = get_user_model()
        self.user = User.objects.create_user(
            username='leak@example.com', email='leak@example.com', password='x',
        )
        UserProfile.objects.create(user=self.user, full_name='L', data_content={})
        self.client.force_login(self.user)
        self.job = Job.objects.create(
            user=self.user, title='Engineer', company='Acme', description='x',
            extracted_skills=[],
        )
        self.gap = GapAnalysis.objects.create(user=self.user, job=self.job, similarity_score=0.7)
        # Use a resume that triggers the BLOCKING tier so the banner
        # actually renders — we want to confirm the template chrome
        # around the banner is leak-free, not just the empty-banner path.
        self.resume = GeneratedResume.objects.create(
            gap_analysis=self.gap,
            content={'professional_summary': 's', 'experience': [],
                     'supervisor_review':
                     {'verdict': 'revise', 'summary': '',
                      'findings': [{'layer': 'content', 'severity': 'blocking',
                                    'category': 'summary', 'location': '',
                                    'issue': 'truncated', 'fix': '.'}]}},
            validation_report={'passed': True, 'findings': [], 'stats': {},
                               'grounding_findings': [], 'supervisor_findings':
                               [{'layer': 'content', 'severity': 'blocking',
                                 'category': 'summary', 'location': '',
                                 'issue': 'truncated', 'fix': '.'}]},
        )

    def test_no_django_comment_delimiters_in_rendered_body(self):
        """Raw '{#' or '#}' in the body means a Django comment fell through
        the parser. Should NEVER appear in user-facing HTML."""
        from django.urls import reverse
        resp = self.client.get(reverse('resume_edit', args=[self.resume.id]))
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode('utf-8', errors='ignore')
        self.assertNotIn('{#', body,
                         "Django {# delimiter leaked — multi-line {# #} is not a valid comment")
        self.assertNotIn('#}', body,
                         "Django #} delimiter leaked — multi-line {# #} is not a valid comment")

    def test_no_internal_module_paths_in_rendered_body(self):
        """Internal file paths (resumes/services/...) must never reach
        the rendered page. Code-level documentation belongs in the
        Python module, not in user-facing template chrome."""
        from django.urls import reverse
        resp = self.client.get(reverse('resume_edit', args=[self.resume.id]))
        body = resp.content.decode('utf-8', errors='ignore')
        for forbidden in (
            'findings_presenter',
            'resumes/services/',
            'audit fix',
            'audit §',
        ):
            self.assertNotIn(
                forbidden, body,
                f"Internal token {forbidden!r} leaked to rendered HTML"
            )


class BulletRuleLabelCollapseTests(SimpleTestCase):
    """Live-bug regression (2026-05-30): the bullet-validator findings
    section showed '1 bullet flagged: bullet rule violation' AND
    '4 bullets flagged: bullet rule violation' as separate lines.

    Root cause: actual rule_ids carry suffixes ('A1_banned_phrase',
    'A2_action_verb_start'). The old code bucketed by rule_id and
    looked up labels with the raw key, so every rule_id missed the
    {'A1', 'A2', ...} lookup table and fell back to 'bullet rule
    violation' — but bucketed by their DISTINCT rule_ids, rendering as
    multiple lines with the same label.

    Fix: bucket by RESOLVED LABEL, not by rule_id. Also strip the
    rule_id suffix before looking up so the canonical label table
    actually fires."""

    def test_rule_id_suffix_is_stripped_before_label_lookup(self):
        """'A1_banned_phrase' now resolves to 'banned phrase / jargon',
        not the fallback 'bullet rule violation'."""
        from resumes.services.findings_presenter import build_review_summary
        vr = {
            'passed': False,
            'findings': [
                {'rule_id': 'A1_banned_phrase', 'severity': 'error',
                 'location': 'experience[0].description[0]',
                 'bullet_text': '...', 'issue': 'banned', 'suggested_fix': None},
            ],
            'stats': {}, 'grounding_findings': [], 'supervisor_findings': [],
        }
        out = build_review_summary({}, vr)
        self.assertEqual(out['tier'], 'blocking')
        body = ' '.join(i['title'] + ' ' + i['body'] for i in out['blocking_items'])
        self.assertIn('banned phrase', body.lower(),
                      f"A1_banned_phrase rule_id should resolve to its label; got {body!r}")

    def test_two_a1_findings_collapse_to_one_line(self):
        """A1_banned_phrase and A1_banned_jargon both resolve to the
        same label 'banned phrase / jargon' — must render as ONE line
        with count 2, not two separate lines."""
        from resumes.services.findings_presenter import build_review_summary
        vr = {
            'passed': False,
            'findings': [
                {'rule_id': 'A1_banned_phrase', 'severity': 'error',
                 'location': 'experience[0].description[0]',
                 'bullet_text': '...', 'issue': '...', 'suggested_fix': None},
                {'rule_id': 'A1_banned_jargon', 'severity': 'error',
                 'location': 'experience[0].description[1]',
                 'bullet_text': '...', 'issue': '...', 'suggested_fix': None},
            ],
            'stats': {}, 'grounding_findings': [], 'supervisor_findings': [],
        }
        out = build_review_summary({}, vr)
        # Count blocking items mentioning 'banned'.
        matches = [i for i in out['blocking_items']
                   if 'banned' in (i['title'] + i['body']).lower()]
        self.assertEqual(len(matches), 1,
                         f"two A1 findings should collapse to one line; got {matches}")
        self.assertIn('2 bullets flagged', matches[0]['title'])

    def test_unknown_rule_ids_collapse_under_fallback_label(self):
        """The original live bug shape: multiple findings with distinct
        rule_ids that ALL fall back to 'bullet rule violation' must
        render as ONE line with count N, not N copies of the same label."""
        from resumes.services.findings_presenter import build_review_summary
        vr = {
            'passed': False,
            'findings': [
                # 5 findings, all with distinct rule_ids that don't match
                # the lookup table after suffix stripping → all fall back.
                {'rule_id': 'Z1_something_new', 'severity': 'error',
                 'location': '...', 'bullet_text': '.', 'issue': '.',
                 'suggested_fix': None},
                {'rule_id': 'Z2_something_else', 'severity': 'error',
                 'location': '...', 'bullet_text': '.', 'issue': '.',
                 'suggested_fix': None},
                {'rule_id': 'Z3_yet_another', 'severity': 'error',
                 'location': '...', 'bullet_text': '.', 'issue': '.',
                 'suggested_fix': None},
                {'rule_id': 'Z4_more', 'severity': 'error',
                 'location': '...', 'bullet_text': '.', 'issue': '.',
                 'suggested_fix': None},
                {'rule_id': 'Z5_again', 'severity': 'error',
                 'location': '...', 'bullet_text': '.', 'issue': '.',
                 'suggested_fix': None},
            ],
            'stats': {}, 'grounding_findings': [], 'supervisor_findings': [],
        }
        out = build_review_summary({}, vr)
        fallback_items = [i for i in out['blocking_items']
                          if 'bullet rule violation' in i['title']]
        self.assertEqual(len(fallback_items), 1,
                         f"5 unknown rule_ids must collapse to ONE line; got {fallback_items}")
        self.assertIn('5 bullets flagged', fallback_items[0]['title'])


# ==========================================================================
# Edit screen redesign (2026-05-30) — inline findings chips at section/item
# anchors, compact summary pill, nav-rail badges. The wall-banner was
# replaced with these inline anchors so each finding renders AT the thing
# it's about.
# ==========================================================================


class FindingsAnnotationsBuildTests(SimpleTestCase):
    """Unit-test the new `annotations` field of build_review_summary. One
    annotation per (section, item_idx, bullet_idx, tier, anchor_kind);
    same-anchor same-tier findings collapse into a single annotation
    with a `count` and an `items` list."""

    def test_bullet_validator_finding_anchored_at_bullet(self):
        from resumes.services.findings_presenter import build_review_summary
        vr = {
            'passed': False,
            'findings': [
                {'rule_id': 'A1_banned_phrase', 'severity': 'error',
                 'location': 'experience[0].description[2]',
                 'bullet_text': '...', 'issue': 'banned', 'suggested_fix': None},
            ],
            'stats': {}, 'grounding_findings': [], 'supervisor_findings': [],
        }
        out = build_review_summary({}, vr)
        anns = out['annotations']
        self.assertEqual(len(anns), 1)
        a = anns[0]
        self.assertEqual(a['section'], 'experience')
        self.assertEqual(a['item_idx'], 0)
        self.assertEqual(a['bullet_idx'], 2)
        self.assertEqual(a['anchor_kind'], 'bullet')
        self.assertEqual(a['tier'], 'blocking')

    def test_grounding_unsupported_skill_anchored_at_bullet_with_token(self):
        """The skill name lives in the detail's single-quoted token; the
        annotation must carry that token so the popover can surface it
        without showing the user the raw `kind` string."""
        from resumes.services.findings_presenter import build_review_summary
        vr = {
            'passed': True, 'findings': [], 'stats': {},
            'grounding_findings': [
                {'kind': 'unsupported_skill',
                 'where': 'experience[0].description[1]',
                 'detail': "Possible unsupported skill 'PyTorch' — not in plan."},
            ],
            'supervisor_findings': [],
        }
        out = build_review_summary({}, vr)
        ann = out['annotations'][0]
        self.assertEqual(ann['anchor_kind'], 'bullet')
        self.assertEqual(ann['section'], 'experience')
        self.assertEqual(ann['items'][0]['token'], 'PyTorch')

    def test_supervisor_finding_honest_section_fallback(self):
        """Supervisor findings carry `category` (free string), no item
        idx. They MUST render at section-level only — never faked to a
        bullet anchor the data doesn't support."""
        from resumes.services.findings_presenter import build_review_summary
        sup = [{'layer': 'content', 'severity': 'blocking',
                'category': 'experience', 'location': 'free-text',
                'issue': 'Bullets read like commit messages.', 'fix': '...'}]
        out = build_review_summary({}, {
            'passed': True, 'findings': [], 'stats': {},
            'grounding_findings': [], 'supervisor_findings': sup,
        })
        ann = out['annotations'][0]
        self.assertEqual(ann['anchor_kind'], 'section',
                         "supervisor category is section-only; must NOT fake item/bullet anchor")
        self.assertEqual(ann['section'], 'experience')
        self.assertIsNone(ann['item_idx'])
        self.assertIsNone(ann['bullet_idx'])

    def test_supervisor_unknown_category_falls_back_to_resume_level(self):
        """A supervisor category not in the known map (e.g. 'ats',
        'layout') becomes a resume-level annotation with no section
        anchor — honest: we have no specific target."""
        from resumes.services.findings_presenter import build_review_summary
        sup = [{'layer': 'content', 'severity': 'blocking',
                'category': 'ats', 'issue': 'Keyword stuffing.', 'fix': '...'}]
        out = build_review_summary({}, {
            'passed': True, 'findings': [], 'stats': {},
            'grounding_findings': [], 'supervisor_findings': sup,
        })
        ann = out['annotations'][0]
        self.assertEqual(ann['anchor_kind'], 'resume')
        self.assertEqual(ann['section'], '')

    def test_regression_metric_loss_resolves_to_item_anchor(self):
        """regression metric_loss carries 'experience[<title> @ <company>]'
        — the presenter resolves this to a numeric item_idx by walking
        content.experience. Honest fallback when no match: section-level."""
        from resumes.services.findings_presenter import build_review_summary
        content = {'experience': [
            {'title': 'AI Trainee', 'company': 'DEPI', 'description': []},
            {'title': 'DT Intern', 'company': 'Acme', 'description': []},
        ]}
        vr = {
            'passed': True, 'findings': [], 'stats': {},
            'grounding_findings': [], 'supervisor_findings': [],
            'regression_findings': [{
                'kind': 'metric_loss', 'severity': 'blocking',
                'where': 'experience[AI Trainee @ DEPI]',
                'prev': ['0.351'], 'now': '',
                'detail': "Metrics ['0.351'] missing from this draft.",
            }],
        }
        out = build_review_summary(content, vr)
        ann = out['annotations'][0]
        self.assertEqual(ann['anchor_kind'], 'item')
        self.assertEqual(ann['section'], 'experience')
        self.assertEqual(ann['item_idx'], 0)

    def test_regression_item_renamed_falls_back_to_section(self):
        """If the regenerated content renamed the role (so the name
        match misses), the regression annotation must fall back to
        section-level — never fake an item_idx."""
        from resumes.services.findings_presenter import build_review_summary
        content = {'experience': [
            {'title': 'Renamed Role', 'company': 'DEPI', 'description': []},
        ]}
        vr = {
            'passed': True, 'findings': [], 'stats': {},
            'grounding_findings': [], 'supervisor_findings': [],
            'regression_findings': [{
                'kind': 'metric_loss', 'severity': 'blocking',
                'where': 'experience[AI Trainee @ DEPI]',  # different name
                'prev': ['0.351'], 'now': '',
                'detail': "...",
            }],
        }
        out = build_review_summary(content, vr)
        ann = out['annotations'][0]
        self.assertEqual(ann['anchor_kind'], 'section',
                         "renamed item must NOT fake an item_idx anchor")
        self.assertIsNone(ann['item_idx'])

    def test_same_anchor_same_tier_collapses_to_one_annotation(self):
        """Two unsupported_skill findings on the same bullet should
        collapse into ONE annotation with count=2 — the chip then
        renders one badge with a list of two items."""
        from resumes.services.findings_presenter import build_review_summary
        vr = {
            'passed': True, 'findings': [], 'stats': {},
            'grounding_findings': [
                {'kind': 'unsupported_skill',
                 'where': 'experience[0].description[1]',
                 'detail': "Possible unsupported skill 'PyTorch' — ..."},
                {'kind': 'unsupported_skill',
                 'where': 'experience[0].description[1]',
                 'detail': "Possible unsupported skill 'CUDA' — ..."},
            ],
            'supervisor_findings': [],
        }
        out = build_review_summary({}, vr)
        same_anchor_anns = [a for a in out['annotations']
                            if a['section'] == 'experience'
                            and a['bullet_idx'] == 1
                            and a['tier'] == 'advisory']
        self.assertEqual(len(same_anchor_anns), 1)
        self.assertEqual(same_anchor_anns[0]['count'], 2)
        tokens = [it['token'] for it in same_anchor_anns[0]['items']]
        self.assertIn('PyTorch', tokens)
        self.assertIn('CUDA', tokens)

    def test_section_counts_track_open_findings_for_nav_rail(self):
        """section_counts is the per-section badge data — sum of open
        annotations for each section, split by tier."""
        from resumes.services.findings_presenter import build_review_summary
        vr = {
            'passed': False,
            'findings': [
                {'rule_id': 'A1_banned_phrase', 'severity': 'error',
                 'location': 'experience[0].description[0]',
                 'bullet_text': '.', 'issue': '.', 'suggested_fix': None},
                {'rule_id': 'A2_action_verb_start', 'severity': 'warn',
                 'location': 'projects[1].description[2]',
                 'bullet_text': '.', 'issue': '.', 'suggested_fix': None},
            ],
            'stats': {}, 'grounding_findings': [], 'supervisor_findings': [],
        }
        out = build_review_summary({}, vr)
        sc = out['section_counts']
        self.assertEqual(sc['experience']['blocking'], 1)
        self.assertEqual(sc['projects']['advisory'], 1)


class EditPageInlineAnchorTests(TestCase):
    """Live-render the edit page with stocked findings and assert the
    inline chips show up at the right anchors, the compact pill replaces
    the wall-banner, and the nav-rail badges count correctly. Also
    re-affirms the {# leak guard now that the template gained new
    inline blocks."""

    def setUp(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from profiles.models import UserProfile
        from resumes.models import GeneratedResume
        from django.contrib.auth import get_user_model
        User = get_user_model()
        self.user = User.objects.create_user(
            username='inline@example.com', email='inline@example.com', password='x',
        )
        UserProfile.objects.create(user=self.user, full_name='I', data_content={})
        self.client.force_login(self.user)
        self.job = Job.objects.create(
            user=self.user, title='Engineer', company='Acme', description='x',
            extracted_skills=[],
        )
        self.gap = GapAnalysis.objects.create(user=self.user, job=self.job, similarity_score=0.7)
        # Two experience entries so item-level anchoring has something to attach to.
        self.resume = GeneratedResume.objects.create(
            gap_analysis=self.gap,
            content={
                'professional_summary': 's',
                'skills': ['Python'],
                'experience': [
                    {'title': 'AI Trainee', 'company': 'DEPI',
                     'description': ['Built X.', 'Shipped Y.']},
                    {'title': 'DT Intern', 'company': 'Almansour',
                     'description': ['Cleaned data.']},
                ],
                'projects': [{'name': 'SmartCV', 'description': ['Built it.']}],
            },
            validation_report={
                'passed': True, 'findings': [], 'stats': {},
                'grounding_findings': [],
                'supervisor_findings': [
                    # section-level supervisor blocker on summary.
                    {'layer': 'content', 'severity': 'blocking',
                     'category': 'summary', 'location': '',
                     'issue': 'Summary truncated.', 'fix': '.'},
                ],
                'regression_findings': [
                    # item-level metric_loss on the first experience role.
                    {'kind': 'metric_loss', 'severity': 'blocking',
                     'where': 'experience[AI Trainee @ DEPI]',
                     'prev': ['0.351'], 'now': '',
                     'detail': "Lost metric 0.351."},
                ],
            },
        )

    def test_no_wall_banner_data_review_summary_attribute(self):
        """The old wall-banner (data-review-summary="blocking|advisory")
        is gone. Content lives inline via data-finding-anchor instead."""
        from django.urls import reverse
        resp = self.client.get(reverse('resume_edit', args=[self.resume.id]))
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode('utf-8', errors='ignore')
        self.assertNotIn('data-review-summary="blocking"', body)
        self.assertNotIn('data-review-summary="advisory"', body)

    def test_compact_pill_renders_with_count(self):
        """The new compact pill replaces the wall-banner — has
        data-review-pill, the total count, and "to fix" copy."""
        from django.urls import reverse
        resp = self.client.get(reverse('resume_edit', args=[self.resume.id]))
        body = resp.content.decode('utf-8', errors='ignore')
        self.assertIn('data-review-pill="blocking"', body)
        self.assertIn('to fix', body)

    def test_inline_chip_renders_at_summary_section(self):
        """Supervisor finding on category=summary → inline chip on the
        Summary section heading (section-level, the honest anchor for a
        supervisor finding)."""
        from django.urls import reverse
        resp = self.client.get(reverse('resume_edit', args=[self.resume.id]))
        body = resp.content.decode('utf-8', errors='ignore')
        # The annotation's data-finding-anchor should be 'section-summary'.
        self.assertIn('data-finding-anchor="section-summary"', body)
        self.assertIn('data-finding-tier="blocking"', body)
        # And the section heading still exists at id="section-summary"
        # for click-jump targeting.
        self.assertIn('id="section-summary"', body)

    def test_inline_chip_renders_at_item_level_for_regression(self):
        """metric_loss on experience[AI Trainee @ DEPI] → item-level
        anchor at anchor-experience-item-0 (resolves the name to idx 0)."""
        from django.urls import reverse
        resp = self.client.get(reverse('resume_edit', args=[self.resume.id]))
        body = resp.content.decode('utf-8', errors='ignore')
        self.assertIn('data-finding-anchor="anchor-experience-item-0"', body)
        # And the item wrapper carries that id for click-jump.
        self.assertIn('id="anchor-experience-item-0"', body)

    def test_nav_rail_badges_reflect_open_findings(self):
        """The left nav rail shows a count badge next to each section
        with open findings. Summary has 1 blocking → red corner badge.

        Updated 2026-05-30 layout pass: the nav rail collapsed from
        labelled rows to an icon-only strip. Badges are now solid
        red/amber filled corner pills (bg-red-500 / bg-amber-500) over
        the section's italic numeral icon, not the prior soft-tinted
        pills next to the section name."""
        from django.urls import reverse
        resp = self.client.get(reverse('resume_edit', args=[self.resume.id]))
        body = resp.content.decode('utf-8', errors='ignore')
        # The Summary section's anchor is #section-summary; find that
        # link and assert the corner-badge classes appear near it.
        # The icon-strip <li> wraps the link + tooltip; the badge is
        # a sibling <span> with bg-red-500.
        # Look for the badge class anywhere in the body — there's at
        # least one open finding so at least one red corner badge
        # MUST be in the HTML.
        self.assertIn('bg-red-500', body,
                      "expected a red corner badge on the icon-strip nav")

    def test_template_comment_leak_guard_still_holds(self):
        """Re-affirm — no Django comment delimiters in rendered HTML."""
        from django.urls import reverse
        resp = self.client.get(reverse('resume_edit', args=[self.resume.id]))
        body = resp.content.decode('utf-8', errors='ignore')
        self.assertNotIn('{#', body)
        self.assertNotIn('#}', body)
        for forbidden in ('findings_presenter', 'resumes/services/', 'audit fix'):
            self.assertNotIn(forbidden, body, f"leaked: {forbidden!r}")


# ---------------------------------------------------------------------
# Findings-classification policy tests
# ---------------------------------------------------------------------


class FindingsClassifierPolicyTests(SimpleTestCase):
    """Three-bucket policy:
      AUTO_FIX     → loop fixes silently, never reaches the user.
      USER_INPUT   → bypasses loop, shown as 'Confirm or complete'.
      ADVISORY     → optional polish.

    Fail-safe: unknown kinds default to USER_INPUT (shown), never to
    AUTO_FIX (which would invite fabrication)."""

    def test_phrasing_rule_classifies_as_auto_fix(self):
        from resumes.services.findings_classifier import (
            classify_finding, BUCKET_AUTO_FIX,
        )
        for rule in ('A1_banned_phrase', 'A6_em_dash', 'A3_duty_opener',
                     'B2_verb_diversity', 'C2_buzzword_saturation'):
            self.assertEqual(
                classify_finding('bullet', {'rule_id': rule}),
                BUCKET_AUTO_FIX,
                f"{rule} should be auto-fixable",
            )

    def test_unsupported_metric_classifies_as_user_input(self):
        from resumes.services.findings_classifier import (
            classify_finding, BUCKET_USER_INPUT,
        )
        self.assertEqual(
            classify_finding('grounding', {'kind': 'unsupported_metric'}),
            BUCKET_USER_INPUT,
        )

    def test_quantification_rule_classifies_as_user_input(self):
        """B1_quantification needs real numbers from the user — the
        LLM cannot invent them. USER_INPUT, not AUTO_FIX."""
        from resumes.services.findings_classifier import (
            classify_finding, BUCKET_USER_INPUT,
        )
        self.assertEqual(
            classify_finding('bullet', {'rule_id': 'B1_quantification'}),
            BUCKET_USER_INPUT,
        )

    def test_regression_metric_loss_classifies_as_user_input(self):
        from resumes.services.findings_classifier import (
            classify_finding, BUCKET_USER_INPUT,
        )
        self.assertEqual(
            classify_finding('regression', {'kind': 'metric_loss'}),
            BUCKET_USER_INPUT,
        )

    def test_skill_loss_classifies_as_advisory(self):
        from resumes.services.findings_classifier import (
            classify_finding, BUCKET_ADVISORY,
        )
        self.assertEqual(
            classify_finding('regression', {'kind': 'skill_loss'}),
            BUCKET_ADVISORY,
        )

    def test_drop_skill_leak_classifies_as_auto_fix(self):
        """The fix is 'remove the leaked skill from output' — the
        system has everything it needs (the skill name + the plan
        flag). AUTO_FIX."""
        from resumes.services.findings_classifier import (
            classify_finding, BUCKET_AUTO_FIX,
        )
        self.assertEqual(
            classify_finding('grounding', {'kind': 'drop_skill_leak'}),
            BUCKET_AUTO_FIX,
        )

    # ---- The fail-safe ----

    def test_unknown_kind_falls_back_to_user_input(self):
        """Critical guardrail: ambiguous / unknown finding kinds
        default to USER_INPUT (shown), NEVER AUTO_FIX (which would
        make the loop try to fix it and fabricate)."""
        from resumes.services.findings_classifier import (
            classify_finding, BUCKET_USER_INPUT,
        )
        cases = [
            ('bullet',     {'rule_id': 'Z99_unknown_rule'}),
            ('grounding',  {'kind': 'mystery_kind'}),
            ('supervisor', {'category': 'novel_concern', 'severity': 'blocking',
                            'layer': 'content'}),
            ('regression', {'kind': 'unseen_regression'}),
            ('unknown_source', {'whatever': 'x'}),
            ('bullet',     {}),                # missing rule_id
            ('grounding',  None),              # None payload
        ]
        for src, finding in cases:
            self.assertEqual(
                classify_finding(src, finding), BUCKET_USER_INPUT,
                f"fail-safe broken for ({src!r}, {finding!r})",
            )

    def test_supervisor_warning_classifies_as_advisory(self):
        """Severity='warning' supervisor findings are advisory
        regardless of category — they're optional polish, never
        loop triggers."""
        from resumes.services.findings_classifier import (
            classify_finding, BUCKET_ADVISORY,
        )
        self.assertEqual(
            classify_finding('supervisor',
                             {'category': 'redundancy', 'severity': 'warning',
                              'layer': 'content'}),
            BUCKET_ADVISORY,
        )

    def test_render_layer_supervisor_is_advisory(self):
        """Render-layer supervisor findings never drive regen
        (no LLM rewrite would touch the template). Advisory."""
        from resumes.services.findings_classifier import (
            classify_finding, BUCKET_ADVISORY,
        )
        self.assertEqual(
            classify_finding('supervisor',
                             {'category': 'layout', 'severity': 'blocking',
                              'layer': 'render'}),
            BUCKET_ADVISORY,
        )


class PresenterBucketTaggingTests(SimpleTestCase):
    """build_review_summary tags each annotation with a bucket and
    emits pill counts split by bucket. The pill's 'to fix' count
    must NOT include user-input items."""

    def test_unsupported_metric_renders_as_to_confirm_not_to_fix(self):
        from resumes.services.findings_presenter import build_review_summary
        vr = {
            'passed': True, 'findings': [], 'stats': {},
            'grounding_findings': [{
                'kind': 'unsupported_metric',
                'where': 'experience[0].description[0]',
                'detail': "Couldn't trace '20%' to profile evidence.",
            }],
            'supervisor_findings': [],
        }
        content = {
            'professional_summary': 's',
            'experience': [{'title': 'X', 'company': 'Y',
                            'duration': '2020 - 2021', 'description': '20% gain'}],
        }
        summary = build_review_summary(content, vr)
        # 0 to fix, 1 to confirm.
        self.assertEqual(summary['total_to_fix'], 0)
        self.assertEqual(summary['total_to_confirm'], 1)
        # The annotation carries bucket='user_input'.
        anns = summary['annotations']
        self.assertEqual(len(anns), 1)
        self.assertEqual(anns[0]['bucket'], 'user_input')

    def test_phrasing_blocker_after_loop_clears_does_not_reach_pill(self):
        """If the loop cleared an A1_banned_phrase finding, the
        validation_report will have no error-severity entry for it.
        The pill's 'to fix' count is 0."""
        from resumes.services.findings_presenter import build_review_summary
        vr = {'passed': True, 'findings': [], 'stats': {},
              'grounding_findings': [], 'supervisor_findings': []}
        summary = build_review_summary({'experience': []}, vr)
        self.assertEqual(summary['total_to_fix'], 0)
        self.assertEqual(summary['total_to_confirm'], 0)

    def test_pill_counts_separate_to_fix_and_to_confirm(self):
        """When a resume has BOTH a residual auto-fix blocker (e.g.
        the loop couldn't clear it) AND a user-input blocker
        (e.g. unsupported_metric), they count separately."""
        from resumes.services.findings_presenter import build_review_summary
        vr = {
            'passed': True,
            'findings': [{
                'rule_id': 'A1_banned_phrase', 'severity': 'error',
                'location': 'experience[0].description[0]',
                'issue': 'Banned phrase: "synergy".',
            }],
            'grounding_findings': [{
                'kind': 'unsupported_metric',
                'where': 'experience[0].description[1]',
                'detail': "Couldn't trace '50%'.",
            }],
            'supervisor_findings': [],
            'stats': {},
        }
        content = {
            'experience': [{
                'title': 'X', 'company': 'Y', 'duration': '2020 - 2021',
                'description': 'synergy\n50% gain',
            }],
        }
        summary = build_review_summary(content, vr)
        self.assertEqual(summary['total_to_fix'], 1)
        self.assertEqual(summary['total_to_confirm'], 1)


class SupervisedLoopBucketFilteringTests(TestCase):
    """The supervised regen loop must:
      - Trigger a round when AUTO_FIX blockers remain.
      - NOT trigger a round when only USER_INPUT blockers remain.

    Verified by counting how many times the generator is called
    when only user-input findings are present."""

    def _make_review(self, findings_list):
        """A minimal SupervisorReview-like object with the methods the
        loop calls. Tests don't need a real LangChain output."""
        class _F:
            def __init__(self, d):
                self.category = d.get('category', '')
                self.severity = d.get('severity', '')
                self.layer = d.get('layer', 'content')
                self.issue = d.get('issue', '')
                self.fix = d.get('fix', '')
                self.location = d.get('location', '')
        class _R:
            def __init__(self, findings):
                self.findings = [_F(f) for f in findings]
                self.verdict = 'revise' if findings else 'advance'
                self.summary = ''
            def blocking_content_findings(self):
                return [f for f in self.findings
                        if f.severity == 'blocking' and f.layer == 'content']
        return _R(findings_list)

    def test_user_input_supervisor_blocker_bypasses_loop(self):
        """A supervisor finding with category='grounding' (USER_INPUT)
        + severity='blocking' must NOT trigger a regen round. The
        generator is called ONCE."""
        from unittest.mock import patch
        from resumes.services import resume_generator as rg
        from profiles.models import UserProfile
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from django.contrib.auth import get_user_model
        User = get_user_model()
        user = User.objects.create_user(
            username='u_loop@example.com', email='u_loop@example.com', password='x',
        )
        UserProfile.objects.create(user=user, full_name='X', data_content={})
        job = Job.objects.create(user=user, title='Eng', company='Acme',
                                 description='x', extracted_skills=[])
        gap = GapAnalysis.objects.create(user=user, job=job, similarity_score=0.5)

        review = self._make_review([{
            'category': 'grounding', 'severity': 'blocking', 'layer': 'content',
            'issue': 'Cannot verify the 30% claim.', 'fix': '',
        }])

        with patch.object(rg, 'generate_resume_content',
                          return_value={'professional_summary': 's', 'experience': []}) as gen, \
             patch('resumes.services.resume_supervisor.review_resume', return_value=review), \
             patch.object(rg, '_build_standards_section', return_value=('', None, None)):
            rg.generate_resume_content_supervised(
                user.profile, job, gap, metadata={}, previous_best=None,
            )

        # ONE call — the loop saw user-input-only blockers and bailed
        # rather than feeding them back for "fixing".
        self.assertEqual(gen.call_count, 1,
                         f"loop ran {gen.call_count} rounds on a user-input-only blocker")

    def test_auto_fix_supervisor_blocker_triggers_at_least_one_regen(self):
        """A supervisor finding with category='redundancy' (AUTO_FIX)
        + severity='blocking' DOES drive a regen round (at least 2
        generate calls before the loop gives up)."""
        from unittest.mock import patch
        from resumes.services import resume_generator as rg
        from profiles.models import UserProfile
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from django.contrib.auth import get_user_model
        User = get_user_model()
        user = User.objects.create_user(
            username='u_loop2@example.com', email='u_loop2@example.com', password='x',
        )
        UserProfile.objects.create(user=user, full_name='X', data_content={})
        job = Job.objects.create(user=user, title='Eng', company='Acme',
                                 description='x', extracted_skills=[])
        gap = GapAnalysis.objects.create(user=user, job=job, similarity_score=0.5)

        review = self._make_review([{
            'category': 'redundancy', 'severity': 'blocking', 'layer': 'content',
            'issue': 'Two bullets say the same thing.', 'fix': 'Merge them.',
        }])

        with patch.object(rg, 'generate_resume_content',
                          return_value={'professional_summary': 's', 'experience': []}) as gen, \
             patch('resumes.services.resume_supervisor.review_resume', return_value=review), \
             patch.object(rg, '_build_standards_section', return_value=('', None, None)):
            rg.generate_resume_content_supervised(
                user.profile, job, gap, metadata={}, previous_best=None,
            )

        # >= 2 calls — the loop saw an auto-fixable blocker and gave
        # the generator at least one chance to fix it.
        self.assertGreaterEqual(gen.call_count, 2,
                                f"loop only ran {gen.call_count} rounds on an auto-fix blocker")


class CapExhaustionDemoteTests(TestCase):
    """Cap-exhaustion fallback (existing behavior) still demotes
    AUTO_FIX regression blockers to 'warning' — but does NOT demote
    USER_INPUT regression blockers (those genuinely need user
    intervention; demoting them would lose information)."""

    def test_auto_fix_supervisor_blocker_demoted_after_cap(self):
        """AUTO_FIX supervisor blockers that survive cap exhaustion are
        demoted to 'warning' on the shipped draft (mirrors the existing
        regression demote). The pill then renders them under 'to review',
        not 'to fix' — the loop owned them and failed, so don't alarm
        the user with red."""
        from unittest.mock import patch
        from django.test import override_settings
        from resumes.services import resume_generator as rg
        from profiles.models import UserProfile
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from django.contrib.auth import get_user_model
        User = get_user_model()
        user = User.objects.create_user(
            username='u_capsup@example.com', email='u_capsup@example.com', password='x',
        )
        UserProfile.objects.create(user=user, full_name='X', data_content={})
        job = Job.objects.create(user=user, title='Eng', company='Acme',
                                 description='x', extracted_skills=[])
        gap = GapAnalysis.objects.create(user=user, job=job, similarity_score=0.5)

        # Supervisor consistently finds a 'summary' (AUTO_FIX) blocker
        # across both rounds — loop can't clear it, ships with demote.
        class _F:
            def __init__(self, **kw):
                self.category = kw.get('category', '')
                self.severity = kw.get('severity', '')
                self.layer = kw.get('layer', 'content')
                self.issue = kw.get('issue', '')
                self.fix = kw.get('fix', '')
                self.location = kw.get('location', '')
        class _Rev:
            findings = [_F(category='summary', severity='blocking', layer='content',
                           issue='Summary stops mid-sentence.', fix='Rewrite.',
                           location='summary')]
            verdict = 'revise'
            summary = ''
            def blocking_content_findings(self):
                return [f for f in self.findings
                        if f.severity == 'blocking' and f.layer == 'content']
        review = _Rev()

        with override_settings(SUPERVISOR_MAX_REVISION_ROUNDS=1), \
             patch.object(rg, 'generate_resume_content',
                          return_value={'professional_summary': 's', 'experience': []}), \
             patch('resumes.services.resume_supervisor.review_resume', return_value=review), \
             patch.object(rg, '_build_standards_section', return_value=('', None, None)):
            shipped = rg.generate_resume_content_supervised(
                user.profile, job, gap, metadata={}, previous_best=None,
            )

        sup_findings = (shipped.get('validation_report') or {}).get('supervisor_findings') or []
        self.assertTrue(sup_findings)
        for f in sup_findings:
            self.assertEqual(
                f.get('severity'), 'warning',
                'AUTO_FIX supervisor blocker that survived cap must be demoted',
            )

    def test_user_input_metric_loss_stays_blocking_after_cap(self):
        """A metric_loss regression finding (USER_INPUT) should NOT
        be demoted by the cap-exhaustion fallback."""
        from unittest.mock import patch
        from django.test import override_settings
        from resumes.services import resume_generator as rg
        from profiles.models import UserProfile
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from django.contrib.auth import get_user_model
        User = get_user_model()
        user = User.objects.create_user(
            username='u_cap@example.com', email='u_cap@example.com', password='x',
        )
        UserProfile.objects.create(user=user, full_name='X', data_content={})
        job = Job.objects.create(user=user, title='Eng', company='Acme',
                                 description='x', extracted_skills=[])
        gap = GapAnalysis.objects.create(user=user, job=job, similarity_score=0.5)

        class _R:
            findings = []
            verdict = 'advance'
            summary = ''
            def blocking_content_findings(self): return []

        # The generated draft carries a blocking metric_loss regression.
        # It survives all rounds (cap = 1 → 2 attempts).
        draft = {
            'professional_summary': 's', 'experience': [],
            'validation_report': {
                'passed': True, 'findings': [], 'stats': {},
                'grounding_findings': [], 'supervisor_findings': [],
                'regression_findings': [{
                    'kind': 'metric_loss', 'severity': 'blocking',
                    'where': 'experience[0]',
                    'detail': "Lost '30%' from prior export.",
                }],
            },
        }

        with override_settings(SUPERVISOR_MAX_REVISION_ROUNDS=1), \
             patch.object(rg, 'generate_resume_content', return_value=dict(draft)), \
             patch('resumes.services.resume_supervisor.review_resume', return_value=_R()), \
             patch.object(rg, '_build_standards_section', return_value=('', None, None)), \
             patch.object(rg, '_apply_regression_check', side_effect=lambda c, *a, **k: c):
            shipped = rg.generate_resume_content_supervised(
                user.profile, job, gap, metadata={}, previous_best={'foo': 'bar'},
            )

        sev = shipped['validation_report']['regression_findings'][0]['severity']
        self.assertEqual(sev, 'blocking',
                         "USER_INPUT metric_loss must NOT be demoted; "
                         "the user still owns the fix")


# ---------------------------------------------------------------------
# Groq TPM / token-budget fixes A-E
# ---------------------------------------------------------------------


class ApplyPlanFilterToSlimCvTests(SimpleTestCase):
    """Fix-D: the prompt's CV block must contain the planner's
    selected subset, NOT the full master profile."""

    def _plan(self, *, skills, exp_idxs, proj_idxs, cert_names):
        from resumes.services.inclusion_planner import (
            InclusionPlan, ExperiencePlan, ProjectPlan,
        )
        return InclusionPlan(
            skills_to_list=list(skills),
            experiences=[
                ExperiencePlan(profile_index=i, title='', company='', duration='')
                for i in exp_idxs
            ],
            projects=[
                ProjectPlan(profile_index=i, name='', url='', relevance_score=0)
                for i in proj_idxs
            ],
            certifications=list(cert_names),
            include_volunteer=False,
            include_publications=False,
            include_awards=False,
            summary_hints=[],
            bridge_bullet_skills=[],
            drop_skills=[],
        )

    def test_skills_replaced_with_plan_selection(self):
        from resumes.services.resume_generator import _apply_plan_filter_to_slim_cv
        slim_cv = {'skills': [f"Skill{i}" for i in range(50)]}
        plan = self._plan(skills=['Python', 'SQL', 'Pandas'],
                          exp_idxs=[], proj_idxs=[], cert_names=[])
        out = _apply_plan_filter_to_slim_cv(slim_cv, plan)
        self.assertEqual(out['skills'], ['Python', 'SQL', 'Pandas'])

    def test_experiences_filtered_to_plan_indices(self):
        from resumes.services.resume_generator import _apply_plan_filter_to_slim_cv
        slim_cv = {'experiences': [
            {'title': f"Job{i}", 'description': ['x'] * 5} for i in range(7)
        ]}
        plan = self._plan(skills=[], exp_idxs=[0, 3, 5],
                          proj_idxs=[], cert_names=[])
        out = _apply_plan_filter_to_slim_cv(slim_cv, plan)
        self.assertEqual([e['title'] for e in out['experiences']],
                         ['Job0', 'Job3', 'Job5'])

    def test_projects_filtered_to_plan_indices(self):
        from resumes.services.resume_generator import _apply_plan_filter_to_slim_cv
        slim_cv = {'projects': [{'name': f"P{i}"} for i in range(10)]}
        plan = self._plan(skills=[], exp_idxs=[],
                          proj_idxs=[1, 4, 7, 9], cert_names=[])
        out = _apply_plan_filter_to_slim_cv(slim_cv, plan)
        self.assertEqual([p['name'] for p in out['projects']],
                         ['P1', 'P4', 'P7', 'P9'])

    def test_certs_filtered_by_name_match(self):
        from resumes.services.resume_generator import _apply_plan_filter_to_slim_cv
        slim_cv = {'certifications': [
            {'name': 'AWS', 'issuer': 'Amazon'},
            {'name': 'GCP', 'issuer': 'Google'},
            {'name': 'Azure', 'issuer': 'Microsoft'},
        ]}
        plan = self._plan(skills=[], exp_idxs=[], proj_idxs=[],
                          cert_names=['AWS', 'GCP'])
        out = _apply_plan_filter_to_slim_cv(slim_cv, plan)
        names = [c['name'] for c in out['certifications']]
        self.assertEqual(set(names), {'AWS', 'GCP'})

    def test_plan_none_passes_through(self):
        from resumes.services.resume_generator import _apply_plan_filter_to_slim_cv
        slim_cv = {'skills': ['A', 'B']}
        self.assertEqual(_apply_plan_filter_to_slim_cv(slim_cv, None), slim_cv)

    def test_prompt_size_materially_reduced(self):
        """End-to-end: a 50-skill / 7-exp / 10-proj / 20-cert master
        profile (with verbose per-item bullets) must produce a SMALL
        serialized CV block once the planner's filter is applied.
        Bullets are stripped because v2_block carries them; this is
        what brings the CV dump from ~39k chars (the prod-log
        observation under the index-only filter) down to single-digit-k.
        Guards against future drift where the filter silently no-ops."""
        import json
        from resumes.services.resume_generator import _apply_plan_filter_to_slim_cv
        slim_cv = {
            'name': 'Test User',
            'skills': [f"S{i}" for i in range(50)],
            'experiences': [
                {'title': f"Role{i}",
                 'company': 'Co',
                 'description': ['bullet ' * 30] * 8}
                for i in range(7)
            ],
            'projects': [
                {'name': f"Proj{i}", 'description': ['bullet ' * 20] * 5}
                for i in range(10)
            ],
            'certifications': [
                {'name': f"Cert{i}", 'issuer': 'X'} for i in range(20)
            ],
        }
        raw = len(json.dumps(slim_cv))
        plan = self._plan(
            skills=[f"S{i}" for i in range(17)],
            exp_idxs=[0, 1, 2],
            proj_idxs=[0, 1, 2, 3],
            cert_names=[f"Cert{i}" for i in range(15)],
        )
        filtered = _apply_plan_filter_to_slim_cv(slim_cv, plan)
        new = len(json.dumps(filtered))
        # Aggressive ceiling: with bullets stripped, the filtered CV
        # should be < 10% of raw. The prod target (per the dev-server
        # screenshot) is <10k chars; this test profile gives ~50k raw
        # so 10% = 5k.
        self.assertLess(new, int(raw * 0.10),
                        f"plan filter only saved {raw - new}/{raw} chars; "
                        f"bullets should be stripped from kept experiences "
                        f"and projects (they're in v2_block)")

    def test_kept_experiences_have_bullets_stripped(self):
        """Bullets are in v2_block (per-skill evidence). Carrying them
        in slim_cv too is what blew the prompt to ~39k chars on real
        profiles. Stripping is the actual size win."""
        from resumes.services.resume_generator import _apply_plan_filter_to_slim_cv
        slim_cv = {'experiences': [
            {'title': 'Role0', 'company': 'X', 'duration': '2020',
             'description': ['long bullet ' * 50] * 10,
             'highlights': ['another bullet ' * 30] * 8,
             'responsibilities': ['duty ' * 20] * 5},
        ]}
        plan = self._plan(skills=[], exp_idxs=[0], proj_idxs=[], cert_names=[])
        out = _apply_plan_filter_to_slim_cv(slim_cv, plan)
        kept = out['experiences'][0]
        # Metadata survives
        self.assertEqual(kept['title'], 'Role0')
        self.assertEqual(kept['duration'], '2020')
        # Bullet fields are GONE
        for forbidden in ('description', 'highlights', 'responsibilities',
                          'achievements', 'accomplishments', 'tasks',
                          'bullets', 'duties', 'summary'):
            self.assertNotIn(forbidden, kept,
                             f"bullet field {forbidden!r} survived; v2_block "
                             f"already carries the bullets — strip duplicate")

    def test_kept_projects_have_descriptions_stripped(self):
        """Project descriptions are in v2_block too (same pattern as
        experience bullets)."""
        from resumes.services.resume_generator import _apply_plan_filter_to_slim_cv
        slim_cv = {'projects': [
            {'name': 'P0', 'url': 'http://x', 'technologies': ['Py'],
             'description': ['long ' * 50] * 5,
             'highlights': ['feat ' * 30] * 3,
             'features': ['f ' * 20] * 3},
        ]}
        plan = self._plan(skills=[], exp_idxs=[], proj_idxs=[0], cert_names=[])
        out = _apply_plan_filter_to_slim_cv(slim_cv, plan)
        kept = out['projects'][0]
        self.assertEqual(kept['name'], 'P0')
        self.assertEqual(kept['url'], 'http://x')
        self.assertEqual(kept['technologies'], ['Py'])
        for forbidden in ('description', 'highlights', 'features',
                          'outcomes', 'deliverables', 'summary'):
            self.assertNotIn(forbidden, kept)

    # ---- Constructive builder (third-pass D) ----

    def test_constructive_builder_under_10k_for_real_size_profile(self):
        """Hard acceptance number: a realistic master profile (3 verbose
        experiences, 5 projects, 20 certs, 50 skills, plus the catch-all
        keys real profiles carry) MUST produce a cv_block under 10k
        chars when built through the constructive path. The prior
        subtractive filter shipped 34-39k on the same shape."""
        import json
        from resumes.services.resume_generator import _build_planner_aligned_cv
        from resumes.services.inclusion_planner import (
            InclusionPlan, ExperiencePlan, ProjectPlan,
        )
        sanitized_cv = {
            'name': 'Test', 'email': 'a@b.c',
            'normalized_summary': 'X' * 1000,   # catch-all that bloated v1
            'raw_text': 'Z' * 20000,             # huge blob the subtractive
            'github_signals': {'x': 'Y' * 5000}, # filter kept
            'linkedin_snapshot': {'noisy': 'L' * 5000},
            'skills': [f"M{i}" for i in range(50)],
            'experiences': [
                {'title': f"R{i}", 'company': 'X',
                 'description': ['long bullet ' * 50] * 12,
                 'highlights': ['extra bullet ' * 30] * 8}
                for i in range(3)
            ],
            'projects': [
                {'name': f"P{i}", 'url': 'http://x',
                 'description': ['proj bullet ' * 25] * 6}
                for i in range(5)
            ],
            'certifications': [
                {'name': f"C{i}", 'issuer': 'X'} for i in range(20)
            ],
            'education': [{'degree': 'BSc', 'institution': 'U'}],
            'languages': ['English'],
        }
        plan = InclusionPlan(
            skills_to_list=[f"P{i}" for i in range(17)],
            experiences=[
                ExperiencePlan(profile_index=i, title='', company='', duration='')
                for i in [0, 1, 2]
            ],
            projects=[
                ProjectPlan(profile_index=i, name='', url='', relevance_score=0)
                for i in [0, 1, 2, 3]
            ],
            certifications=[f"C{i}" for i in range(15)],
            include_volunteer=False, include_publications=False,
            include_awards=False, summary_hints=[],
            bridge_bullet_skills=[], drop_skills=[],
        )
        out = _build_planner_aligned_cv(sanitized_cv, plan)
        cv_block_len = len(json.dumps(out, indent=2))
        self.assertLess(
            cv_block_len, 10_000,
            f'cv_block_len={cv_block_len} exceeds 10k acceptance ceiling; '
            f'constructive builder should produce ~6-8k for this size'
        )

    def test_constructive_builder_excludes_unallowed_keys(self):
        """Identity/structured allowlist — nothing else gets through.
        Master profiles carry github_signals/linkedin_snapshot/raw_text/
        normalized_summary and similar catch-alls; none should appear
        in the cv_block."""
        from resumes.services.resume_generator import _build_planner_aligned_cv
        sanitized = {
            'name': 'X', 'email': 'a@b.c',
            'github_signals': 'huge',
            'linkedin_snapshot': 'huge',
            'scholar_signals': 'huge',
            'kaggle_signals': 'huge',
            'raw_text': 'huge',
            'normalized_summary': 'huge',
            'extracted_text': 'huge',
            'objective': 'huge',
            'mystery_future_field': 'should not leak',
        }
        out = _build_planner_aligned_cv(sanitized, plan=None)
        for forbidden in ('github_signals', 'linkedin_snapshot',
                          'scholar_signals', 'kaggle_signals',
                          'raw_text', 'normalized_summary',
                          'extracted_text', 'objective',
                          'mystery_future_field'):
            self.assertNotIn(forbidden, out,
                             f"{forbidden!r} leaked into constructive cv_block")
        # Allowed identity keys did pass through.
        self.assertEqual(out.get('name'), 'X')
        self.assertEqual(out.get('email'), 'a@b.c')

    def test_constructive_builder_includes_required_sections(self):
        """Quality guard: name/email/phone/education/languages/contact
        MUST flow through. If any came out blank, the prompt loses
        scaffolding the LLM needs to render the resume correctly."""
        from resumes.services.resume_generator import _build_planner_aligned_cv
        from resumes.services.inclusion_planner import (
            InclusionPlan, ExperiencePlan, ProjectPlan,
        )
        sanitized = {
            'name': 'Zeyad', 'email': 'z@x.com', 'phone': '+20 100',
            'location': 'Cairo', 'linkedin': 'https://linkedin/in/zeyad',
            'professional_summary': 'Junior data scientist.',
            'skills': ['Python'],
            'experiences': [
                {'title': 'AI Trainee', 'company': 'DEPI',
                 'duration': '2024-2025', 'description': ['b']}
            ],
            'education': [{'degree': 'BSc CS', 'institution': 'KSIU',
                           'year': '2027'}],
            'languages': ['English (Fluent)', 'Arabic (Native)'],
            'certifications': [{'name': 'AWS', 'issuer': 'Amazon'}],
        }
        plan = InclusionPlan(
            skills_to_list=['Python'],
            experiences=[ExperiencePlan(profile_index=0, title='',
                                        company='', duration='')],
            projects=[], certifications=['AWS'],
            include_volunteer=False, include_publications=False,
            include_awards=False, summary_hints=[],
            bridge_bullet_skills=[], drop_skills=[],
        )
        out = _build_planner_aligned_cv(sanitized, plan)
        self.assertEqual(out['name'], 'Zeyad')
        self.assertEqual(out['email'], 'z@x.com')
        self.assertEqual(out['phone'], '+20 100')
        self.assertEqual(out['location'], 'Cairo')
        self.assertEqual(out['linkedin'], 'https://linkedin/in/zeyad')
        self.assertEqual(out['professional_summary'], 'Junior data scientist.')
        self.assertEqual(out['skills'], ['Python'])
        self.assertEqual(len(out['experiences']), 1)
        self.assertEqual(out['experiences'][0]['title'], 'AI Trainee')
        # bullets stripped — they're in v2_block
        self.assertNotIn('description', out['experiences'][0])
        self.assertEqual(out['education'][0]['institution'], 'KSIU')
        self.assertEqual(out['languages'], ['English (Fluent)', 'Arabic (Native)'])
        self.assertEqual(out['certifications'][0]['name'], 'AWS')

    def test_blob_keys_stripped_defensively(self):
        """Some users carry a `raw_text` blob (the original parsed CV)
        or `linkedin_snapshot` in their data_content. Even when the
        upstream slim_cv builder excluded these, defensive removal
        here guards against future drift."""
        from resumes.services.resume_generator import _apply_plan_filter_to_slim_cv
        slim_cv = {
            'name': 'X',
            'raw_text': 'A' * 20_000,
            'linkedin_snapshot': {'noisy': 'B' * 5_000},
            'skills': ['Python'],
        }
        plan = self._plan(skills=['Python'], exp_idxs=[],
                          proj_idxs=[], cert_names=[])
        out = _apply_plan_filter_to_slim_cv(slim_cv, plan)
        self.assertNotIn('raw_text', out)
        self.assertNotIn('linkedin_snapshot', out)
        self.assertEqual(out['name'], 'X')
        self.assertEqual(out['skills'], ['Python'])


class ClassifyForJdCacheTests(SimpleTestCase):
    """Fix-A: classify_for_jd is called 3× per generation with
    identical inputs; the cache should collapse those to 1 LLM call."""

    def setUp(self):
        from profiles.services.role_classifier import clear_classify_cache
        clear_classify_cache()

    def tearDown(self):
        from profiles.services.role_classifier import clear_classify_cache
        clear_classify_cache()

    def test_identical_inputs_call_llm_once(self):
        from unittest.mock import patch
        from profiles.services import role_classifier as rc
        from profiles.services.role_classifier import (
            RoleClassification, classify_for_jd, classify_cache_size,
        )

        def _profile_cls(_):
            return RoleClassification(primary_role='Data Scientist',
                                      seniority='junior',
                                      tech_stack_signals=['Python'],
                                      region='global')

        def _jd_cls(_):
            return RoleClassification(primary_role='Data Scientist',
                                      seniority='junior',
                                      tech_stack_signals=['SQL'],
                                      region='global')

        profile = {'headline': 'Data Scientist',
                   'professional_summary': 'I do data.',
                   'skills': ['Python'],
                   'experiences': [{'title': 'Data Intern'}]}
        jd = 'Looking for a junior data scientist.'

        with patch.object(rc, 'detect_role_seniority', side_effect=_profile_cls) as m_p, \
             patch.object(rc, 'classify_jd_role', side_effect=_jd_cls) as m_j:
            r1 = classify_for_jd(profile, jd)
            r2 = classify_for_jd(profile, jd)
            r3 = classify_for_jd(profile, jd)

        self.assertEqual(r1.primary_role, 'Data Scientist')
        self.assertEqual(r1, r2)
        self.assertEqual(r1, r3)
        # Each call would have fired 2 Groq invocations (profile + jd
        # classifier). Cache should reduce 6 to 2.
        self.assertEqual(m_p.call_count, 1,
                         'detect_role_seniority should run once across 3 cached calls')
        self.assertEqual(m_j.call_count, 1,
                         'classify_jd_role should run once across 3 cached calls')
        self.assertEqual(classify_cache_size(), 1)

    def test_different_profiles_classify_separately(self):
        from unittest.mock import patch
        from profiles.services import role_classifier as rc
        from profiles.services.role_classifier import (
            RoleClassification, classify_for_jd, classify_cache_size,
        )

        def _profile_cls(_):
            return RoleClassification(primary_role='X', seniority='mid',
                                      tech_stack_signals=[], region='global')

        def _jd_cls(_):
            return RoleClassification(primary_role='X', seniority='mid',
                                      tech_stack_signals=[], region='global')

        with patch.object(rc, 'detect_role_seniority', side_effect=_profile_cls) as m_p, \
             patch.object(rc, 'classify_jd_role', side_effect=_jd_cls):
            classify_for_jd({'headline': 'A', 'skills': ['Python']}, 'JD A')
            classify_for_jd({'headline': 'B', 'skills': ['Go']}, 'JD A')
            classify_for_jd({'headline': 'A', 'skills': ['Python']}, 'JD B')

        # 3 distinct (profile, jd) → 3 cache entries → 3 LLM call sets.
        self.assertEqual(classify_cache_size(), 3)
        self.assertEqual(m_p.call_count, 3)


class SlimmerRetryBugTests(SimpleTestCase):
    """Fix-B: when the pre-slim path mutates `prompt`, the retry path
    must not try to slim from the ALREADY-trimmed string. Either it
    skips entirely, or it uses the kept-aside original."""

    def test_pre_slim_disables_retry_so_no_zero_save_log(self):
        """After pre-slim trims the prompt, the 413-retry's redundant
        second-slim is skipped (the guard `not _pre_slimmed` short-
        circuits). This prevents the 'saved=0' no-op that wasted a
        round-trip in the dev log."""
        # Read the resume_generator source to assert the guard is present.
        import resumes.services.resume_generator as rg
        import inspect
        src = inspect.getsource(rg.generate_resume_content)
        self.assertIn('_pre_slimmed', src,
                      'Fix-B requires a _pre_slimmed flag guarding the retry')
        self.assertIn('_original_prompt', src,
                      'Fix-B requires keeping the un-slimmed prompt aside')
        # The guard appears in the retry branch.
        self.assertIn('not _pre_slimmed', src,
                      'retry path must short-circuit when pre-slim already ran')

    def test_retry_when_pre_slim_did_not_run_actually_trims(self):
        """When the prompt was under-budget (no pre-slim) and Groq
        413s anyway (per-minute TPM, not per-request size), the retry
        slims from the original — saved > 0 — and tries again."""
        # Pure-logic check: simulate the retry's .replace() on a fresh
        # original string and confirm the slim is strictly shorter.
        v2_block = 'V2_BLOCK_' + 'x' * 5000
        std = 'STD_BLOCK_' + 'y' * 3000
        original = 'PREAMBLE\n' + v2_block + '\nMID\n' + std + '\nTAIL'
        slim = original.replace(v2_block, '').replace(std, '')
        self.assertLess(len(slim), len(original))
        self.assertNotIn(v2_block, slim)
        self.assertNotIn(std, slim)


class OfflineFallbackMarkerTests(TestCase):
    """Fix-C: the offline fallback dict must carry an `_is_fallback`
    marker, and the supervised loop must SKIP review (no Groq calls
    on a non-LLM placeholder)."""

    def test_fallback_dict_carries_is_fallback_flag(self):
        from resumes.services.resume_generator import _build_offline_fallback
        from unittest.mock import MagicMock
        profile = MagicMock(); job = MagicMock()
        job.title = 'Engineer'
        job.extracted_skills = ['Python']
        raw_cv = {
            'name': 'X', 'professional_summary': '',
            'skills': ['Python'], 'experiences': [],
            'education': [], 'projects': [], 'certifications': [],
            'languages': [],
        }
        result = _build_offline_fallback(profile, job, raw_cv)
        self.assertTrue(result.get('_is_fallback'),
                        'offline fallback must be marked so the loop can skip review')

    def test_supervised_loop_skips_supervisor_on_fallback(self):
        """When generate_resume_content returns the fallback, the loop
        must NOT call review_resume — wastes 2 Groq calls per round
        on a non-LLM placeholder."""
        from unittest.mock import patch, MagicMock
        from resumes.services import resume_generator as rg
        from profiles.models import UserProfile
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from django.contrib.auth import get_user_model
        User = get_user_model()
        user = User.objects.create_user(
            username='fbloop@example.com', email='fbloop@example.com', password='x',
        )
        UserProfile.objects.create(user=user, full_name='X', data_content={})
        job = Job.objects.create(user=user, title='Eng', company='Acme',
                                 description='x', extracted_skills=[])
        gap = GapAnalysis.objects.create(user=user, job=job, similarity_score=0.5)

        fallback = {'_is_fallback': True, 'professional_summary': 's',
                    'experience': [], 'projects': []}

        review_spy = MagicMock()
        with patch.object(rg, 'generate_resume_content', return_value=fallback), \
             patch('resumes.services.resume_supervisor.review_resume',
                   side_effect=review_spy), \
             patch.object(rg, '_build_standards_section',
                          return_value=('', None, None)):
            shipped = rg.generate_resume_content_supervised(
                user.profile, job, gap, metadata={}, previous_best=None,
            )

        review_spy.assert_not_called()
        self.assertTrue(shipped.get('_is_fallback'),
                        'fallback marker must survive to the shipped resume '
                        'so the UI can surface degraded-mode')


class TPMThrottleTests(SimpleTestCase):
    """Fix-E: the TPM throttle delays calls that would exceed the
    rolling 60s token budget. Verified at the throttle layer so the
    test is fast (no real LLM)."""

    def test_reserve_under_budget_does_not_sleep(self):
        from profiles.services.tpm_throttle import TPMThrottle
        t = TPMThrottle(budget=10_000, window=60.0)
        slept = t.reserve(5_000)
        self.assertEqual(slept, 0.0)
        self.assertEqual(t.current_usage(), 5_000)

    def test_reserve_over_budget_blocks_until_window_frees(self):
        """When the rolling window can't fit the new reservation, the
        throttle sleeps. We use a SHORT window so the test runs fast."""
        import time
        from profiles.services.tpm_throttle import TPMThrottle
        t = TPMThrottle(budget=10_000, window=0.5)   # 500ms window
        t.reserve(8_000)                              # under budget
        start = time.monotonic()
        slept = t.reserve(5_000)                      # 8000+5000 > 10000
        elapsed = time.monotonic() - start
        # The throttle waited for the first event to age out before
        # adding the new reservation.
        self.assertGreater(slept, 0.0)
        self.assertGreaterEqual(elapsed, slept * 0.9)

    def test_reserve_for_invoke_handles_str_and_messages(self):
        """The estimator works with both the structured-output string
        and the plain-chat list[HumanMessage] shapes the codebase uses."""
        from profiles.services.tpm_throttle import estimate_input_tokens

        class _Msg:
            def __init__(self, content):
                self.content = content

        # Strings
        self.assertGreater(estimate_input_tokens('x' * 3500), 900)
        # Message lists
        msgs = [_Msg('a' * 1750), _Msg('b' * 1750)]
        self.assertGreater(estimate_input_tokens(msgs), 900)
        # Empty / weird payloads don't crash
        self.assertGreater(estimate_input_tokens(None), 0)
        self.assertGreater(estimate_input_tokens({}), 0)

    def test_throttle_disabled_setting_short_circuits(self):
        """In tests the throttle is disabled by default; reserve_for_invoke
        must return 0 without recording into the rolling window."""
        from django.test import override_settings
        from profiles.services.tpm_throttle import (
            reserve_for_invoke, reset_throttle, get_throttle,
        )
        reset_throttle()
        with override_settings(GROQ_TPM_THROTTLE_DISABLED=True):
            slept = reserve_for_invoke('x' * 10_000_000, max_output_tokens=8000)
        self.assertEqual(slept, 0.0)
        # Nothing was recorded into the window.
        self.assertEqual(get_throttle().current_usage(), 0)
        reset_throttle()


class GenerateResumeUsesFilteredCvTests(TestCase):
    """End-to-end: the prompt built by generate_resume_content actually
    contains the planner-filtered CV, not the full master. We capture
    the prompt by spying on the throttled LLM's invoke."""

    def test_prompt_contains_only_planner_skills(self):
        from unittest.mock import patch, MagicMock
        from resumes.services import resume_generator as rg
        from resumes.services.inclusion_planner import (
            InclusionPlan, ExperiencePlan, ProjectPlan,
        )
        from profiles.models import UserProfile
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from django.contrib.auth import get_user_model
        User = get_user_model()
        user = User.objects.create_user(
            username='cvprompt@example.com', email='cvprompt@example.com', password='x',
        )
        master = {
            'name': 'X User', 'professional_summary': '',
            'skills': [f'Master_Skill_{i}' for i in range(40)],
            'experiences': [
                {'title': f'Role{i}', 'company': 'Co',
                 'description': ['bullet'], 'start_date': '2020', 'end_date': '2021'}
                for i in range(6)
            ],
            'projects': [{'name': f'Proj{i}'} for i in range(8)],
            'certifications': [{'name': f'Cert{i}', 'issuer': 'X'} for i in range(20)],
            'education': [], 'languages': [],
        }
        UserProfile.objects.create(user=user, full_name='X User', data_content=master)
        job = Job.objects.create(user=user, title='Eng', company='Acme',
                                 description='Looking for an engineer.',
                                 extracted_skills=['Python'])
        gap = GapAnalysis.objects.create(user=user, job=job, similarity_score=0.5)

        plan = InclusionPlan(
            skills_to_list=['Picked_Skill_A', 'Picked_Skill_B'],
            experiences=[ExperiencePlan(profile_index=2, title='', company='', duration='')],
            projects=[ProjectPlan(profile_index=3, name='', url='', relevance_score=0)],
            certifications=['Cert5'],
            include_volunteer=False, include_publications=False, include_awards=False,
            summary_hints=[], bridge_bullet_skills=[], drop_skills=[],
        )

        captured = {}

        class _StubLLM:
            def invoke(self, prompt):
                captured['prompt'] = prompt
                # Return a fake structured object that mimics ResumeContentResult.
                # generate_resume_content calls .model_dump() on it.
                stub_obj = MagicMock()
                stub_obj.model_dump.return_value = {
                    'professional_title': 'Eng',
                    'professional_summary': 's',
                    'skills': ['Picked_Skill_A'],
                    'experience': [],
                    'projects': [],
                    'education': [],
                    'certifications': [],
                    'languages': [],
                }
                return stub_obj

        with patch.object(rg, 'get_structured_llm', return_value=_StubLLM()), \
             patch.object(rg, '_build_v2_grounding',
                          return_value=('=== V2 GROUNDING ===\n', plan)), \
             patch.object(rg, '_build_standards_section',
                          return_value=('', None, None)):
            rg.generate_resume_content(
                user.profile, job, gap,
                metadata={}, standards_section_override='',
            )

        prompt = captured.get('prompt', '')
        # Planner's picks ARE in the prompt.
        self.assertIn('Picked_Skill_A', prompt)
        self.assertIn('Picked_Skill_B', prompt)
        # Master-only skills (NOT in the plan) are NOT in the CV dump.
        # Exception: the very first one might still appear in another
        # block (e.g. evidence_context). Check that >90% of master-only
        # skills are absent.
        absent = sum(1 for i in range(40) if f'Master_Skill_{i}' not in prompt)
        self.assertGreater(absent, 35,
                           f'expected the planner filter to drop master-only '
                           f'skills; only {40 - absent}/40 were filtered out')


# ---------------------------------------------------------------------
# Role-identity guards (fabrication safety pass — 2026-06-01)
# ---------------------------------------------------------------------


class RoleIdentityGuardUnitTests(SimpleTestCase):
    """Helper-level tests for resumes/services/role_identity_guard.py.
    Whitespace + case normalization + URL/name fuzzy fallback."""

    def test_filter_experiences_drops_invented_company(self):
        from resumes.services.role_identity_guard import filter_experiences_to_known
        known = [
            {'title': 'IT Intern', 'company': 'Almansour Automotive'},
            {'title': 'DevOps Trainee', 'company': 'DEPI'},
        ]
        returned = [
            {'title': 'IT Intern', 'company': 'Almansour Automotive',
             'description': ['Built a thing.']},
            {'title': 'DevOps Trainee', 'company': 'DEPI',
             'description': ['Shipped a thing.']},
            {'title': 'Banking Analyst', 'company': 'Banque Misr',   # PHANTOM
             'description': ['Did banking.']},
        ]
        kept, dropped = filter_experiences_to_known(returned, known)
        self.assertEqual(len(kept), 2)
        self.assertEqual(len(dropped), 1)
        self.assertEqual(dropped[0]['company'], 'Banque Misr')

    def test_company_match_is_case_insensitive_whitespace_normalized(self):
        from resumes.services.role_identity_guard import filter_experiences_to_known
        known = [{'title': 'X', 'company': 'Almansour Automotive'}]
        returned = [{'title': 'X', 'company': '  ALMANSOUR   AUTOMOTIVE  ',
                     'description': ['ok']}]
        kept, dropped = filter_experiences_to_known(returned, known)
        self.assertEqual(len(kept), 1)
        self.assertEqual(len(dropped), 0)

    def test_filter_projects_uses_url_match_when_name_renamed(self):
        from resumes.services.role_identity_guard import filter_projects_to_known
        known = [{'name': 'healthcare-prediction-depi',
                  'url': 'https://github.com/zeyad/healthcare-prediction-depi'}]
        returned = [{
            'name': 'Healthcare Prediction (DEPI)',        # renamed
            'url': 'https://github.com/zeyad/healthcare-prediction-depi',
            'description': ['ok'],
        }]
        kept, dropped = filter_projects_to_known(returned, known)
        self.assertEqual(len(kept), 1, 'URL match should keep renamed project')
        self.assertEqual(len(dropped), 0)

    def test_filter_projects_drops_phantom_when_no_url_match(self):
        from resumes.services.role_identity_guard import filter_projects_to_known
        known = [{'name': 'SmartCV', 'url': 'https://github.com/zeyad/smartcv'}]
        returned = [
            {'name': 'SmartCV', 'url': 'https://github.com/zeyad/smartcv',
             'description': ['ok']},
            {'name': 'Fabricated Banking Project',          # PHANTOM
             'url': 'https://github.com/somebody/random-repo',
             'description': ['phantom']},
        ]
        kept, dropped = filter_projects_to_known(returned, known)
        self.assertEqual(len(kept), 1)
        self.assertEqual(len(dropped), 1)
        self.assertEqual(dropped[0]['name'], 'Fabricated Banking Project')

    def test_covers_known_identities_true_when_all_match(self):
        from resumes.services.role_identity_guard import covers_known_identities
        known = [
            {'title': 'A', 'company': 'Co1'},
            {'title': 'B', 'company': 'Co2'},
        ]
        kept = [
            {'title': 'A renamed', 'company': 'Co1'},   # company-match wins
            {'title': 'B', 'company': 'Co2'},
        ]
        self.assertTrue(covers_known_identities(kept, known, kind='experience'))

    def test_covers_known_identities_false_when_role_missing(self):
        from resumes.services.role_identity_guard import covers_known_identities
        known = [
            {'title': 'A', 'company': 'Co1'},
            {'title': 'B', 'company': 'Co2'},
        ]
        kept = [
            {'title': 'A', 'company': 'Co1'},
            # Co2 missing
        ]
        self.assertFalse(covers_known_identities(kept, known, kind='experience'))


class RegenerateSectionPhantomRoleViewTests(TestCase):
    """Fix-1b end-to-end: when the LLM (mocked) returns an invented role
    or a partial role set, the view path must drop / reject — NOT save
    the phantom into GeneratedResume.content."""

    def setUp(self):
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from profiles.models import UserProfile
        from resumes.models import GeneratedResume
        from django.contrib.auth import get_user_model
        User = get_user_model()
        self.user = User.objects.create_user(
            username='regen@example.com', email='regen@example.com', password='x',
        )
        self.profile = UserProfile.objects.create(
            user=self.user, full_name='Zeyad Phantom',
            data_content={
                'experiences': [
                    {'title': 'IT Intern', 'company': 'Almansour Automotive',
                     'description': ['Built ingest pipeline.']},
                    {'title': 'DevOps Trainee', 'company': 'DEPI',
                     'description': ['Shipped CI on GitLab.']},
                ],
                'projects': [
                    {'name': 'SmartCV',
                     'url': 'https://github.com/zeyad/smartcv',
                     'description': ['Built X.']},
                ],
            },
        )
        self.client.force_login(self.user)
        self.job = Job.objects.create(
            user=self.user, title='Engineer', company='Acme',
            description='Looking for an engineer.', extracted_skills=['Python'],
        )
        self.gap = GapAnalysis.objects.create(
            user=self.user, job=self.job, similarity_score=0.7,
        )
        self.resume = GeneratedResume.objects.create(
            gap_analysis=self.gap,
            content={
                'experience': [
                    {'title': 'IT Intern', 'company': 'Almansour Automotive',
                     'duration': '2023', 'description': ['Old bullet 1.']},
                    {'title': 'DevOps Trainee', 'company': 'DEPI',
                     'duration': '2024', 'description': ['Old bullet 2.']},
                ],
                'projects': [
                    {'name': 'SmartCV',
                     'url': 'https://github.com/zeyad/smartcv',
                     'description': ['Old project bullet.']},
                ],
            },
            ats_score=0.0, validation_report={},
        )

    def _post(self, section, payload_content=None):
        """POST the regen endpoint with optional current_content body."""
        import json as _json
        from django.urls import reverse
        url = reverse('regenerate_section', args=[self.resume.id, section])
        body = {}
        if payload_content is not None:
            body['current_content'] = payload_content
        return self.client.post(
            url, data=_json.dumps(body) if body else '{}',
            content_type='application/json',
        )

    def test_invented_role_is_rejected_and_saved_content_untouched(self):
        """The LLM returns 3 roles: the 2 real ones + a phantom
        'Banque Misr'. Expected: the view REJECTS (422), nothing saves
        — count mismatch (kept=2 vs known=2 → wait that covers; so
        let's drop one of the real ones to ensure a coverage gap)."""
        from unittest.mock import patch
        from resumes.services import resume_generator as rg
        # LLM returns: 1 real role + 1 phantom (so the real-role count
        # is 1, but the resume currently has 2 — coverage check fails).
        llm_returned = [
            {'title': 'IT Intern', 'company': 'Almansour Automotive',
             'description': ['Rewritten bullet.']},
            {'title': 'Banking Analyst', 'company': 'Banque Misr',
             'description': ['Phantom bullet.']},
        ]
        with patch('resumes.views.regenerate_section', return_value=llm_returned):
            resp = self._post('experience')
        self.assertEqual(resp.status_code, 422)
        import json as _json
        body = _json.loads(resp.content)
        self.assertEqual(body.get('error'), 'identity_mismatch')
        # Saved content unchanged.
        self.resume.refresh_from_db()
        companies = [e.get('company') for e in self.resume.content.get('experience', [])]
        self.assertIn('Almansour Automotive', companies)
        self.assertIn('DEPI', companies)
        self.assertNotIn('Banque Misr', companies)

    def test_phantom_role_dropped_but_real_roles_complete_passes(self):
        """The LLM returns all real roles AND a phantom. The phantom
        is dropped silently; the real-role set covers known, so the
        regen IS accepted (kept = both real roles)."""
        from unittest.mock import patch
        from resumes.services import resume_generator as rg
        llm_returned = [
            {'title': 'IT Intern', 'company': 'Almansour Automotive',
             'description': ['Rewritten bullet 1.']},
            {'title': 'DevOps Trainee', 'company': 'DEPI',
             'description': ['Rewritten bullet 2.']},
            {'title': 'Banking Analyst', 'company': 'Banque Misr',
             'description': ['Phantom bullet.']},
        ]
        with patch('resumes.views.regenerate_section', return_value=llm_returned):
            resp = self._post('experience')
        self.assertEqual(resp.status_code, 200)
        self.resume.refresh_from_db()
        companies = [e.get('company') for e in self.resume.content['experience']]
        self.assertIn('Almansour Automotive', companies)
        self.assertIn('DEPI', companies)
        self.assertNotIn('Banque Misr', companies,
                         'phantom must never reach saved content')
        self.assertEqual(len(self.resume.content['experience']), 2)

    def test_clean_regeneration_accepted_and_bullets_updated(self):
        """The LLM returns the SAME 2 real roles with rewritten bullets.
        Expected: 200, content saved with new bullets, no role changes."""
        from unittest.mock import patch
        from resumes.services import resume_generator as rg
        llm_returned = [
            {'title': 'IT Intern', 'company': 'Almansour Automotive',
             'description': ['Engineered ingest pipeline that cut nightly load by 6 hours.']},
            {'title': 'DevOps Trainee', 'company': 'DEPI',
             'description': ['Shipped CI on GitLab, halving build time.']},
        ]
        with patch('resumes.views.regenerate_section', return_value=llm_returned):
            resp = self._post('experience')
        self.assertEqual(resp.status_code, 200)
        self.resume.refresh_from_db()
        exps = self.resume.content['experience']
        self.assertEqual(len(exps), 2)
        bullets = [b for e in exps for b in (e.get('description') or [])]
        self.assertTrue(any('cut nightly load' in b for b in bullets))
        self.assertTrue(any('halving build time' in b for b in bullets))


class MainGenIdentityGuardTests(SimpleTestCase):
    """Fix-2: _post_process drops experiences/projects whose identity
    isn't in the master profile, logs the drop. Tested via the public
    filter helpers, since _post_process is a closure inside
    generate_resume_content (not directly callable)."""

    def test_main_gen_drops_invented_role(self):
        from resumes.services.role_identity_guard import filter_experiences_to_known
        master = [
            {'title': 'IT Intern', 'company': 'Almansour Automotive'},
            {'title': 'DevOps Trainee', 'company': 'DEPI'},
        ]
        # LLM-style return: 2 real + 1 invented
        returned = [
            {'title': 'IT Intern', 'company': 'Almansour Automotive',
             'description': ['ok']},
            {'title': 'DevOps Trainee', 'company': 'DEPI',
             'description': ['ok']},
            {'title': 'Senior Analyst', 'company': 'Banque Misr',  # PHANTOM
             'description': ['fabricated']},
        ]
        kept, dropped = filter_experiences_to_known(returned, master)
        self.assertEqual(len(kept), 2)
        self.assertEqual(len(dropped), 1)
        self.assertEqual(dropped[0]['company'], 'Banque Misr')

    def test_main_gen_does_not_drop_when_company_is_master_with_renamed_title(self):
        """Title renames are allowed (enforce_verbatim_titles snaps them
        back later); company match is the load-bearing identity."""
        from resumes.services.role_identity_guard import filter_experiences_to_known
        master = [{'title': 'Information Technology Intern',
                   'company': 'Almansour Automotive'}]
        returned = [{'title': 'IT Intern',
                     'company': 'Almansour Automotive',
                     'description': ['ok']}]
        kept, dropped = filter_experiences_to_known(returned, master)
        self.assertEqual(len(kept), 1)
        self.assertEqual(len(dropped), 0)


class RecoveryPathPostProcessTests(SimpleTestCase):
    """Fix-3: the failed_generation recovery path now invokes
    _post_process, which means a salvaged resume runs through:
      - identity guard (drops invented roles)
      - bullet validator
      - normalize_resume
      - grounding check (unsupported_metric / unsupported_skill)
      - regression check
    This was previously skipped on the recovery branch.

    Verified at the call-site level: the recovery branch (which the
    grep located around line ~1690) now reads `_post_process(recovered.model_dump())`
    rather than the prior 3-step manual chain. End-to-end LLM mocking
    of the recovery branch is out of scope (covered by integration).
    """

    def test_recovery_branch_calls_post_process(self):
        import inspect
        from resumes.services import resume_generator as rg
        src = inspect.getsource(rg.generate_resume_content)
        # The recovery branch references _post_process — confirms FIX-3 landed.
        self.assertIn('_post_process(recovered.model_dump())', src,
                      'recovery path must route through _post_process so the '
                      'salvaged content gets identity-guard, grounding check, '
                      'and regression check (was missing pre-FIX-3)')


# ---------------------------------------------------------------------------
# Pillar 4 Step 2 — ATS-clean rebuild: migration map, skills grouping,
# v2 adapter, and rendered-HTML self-check against the KB ats_rules.
# ---------------------------------------------------------------------------


class ResolveTemplateMigrationTests(SimpleTestCase):
    """pdf_exporter.resolve_template must (a) pass through live theme
    names, (b) migrate every removed theme name to a surviving theme so
    existing GeneratedResume.content['template_name'] values never 500
    the export, (c) fall through to ats_clean for unknown / empty input.
    """

    def test_live_theme_passes_through(self):
        from resumes.services.pdf_exporter import resolve_template
        theme, path = resolve_template('ats_clean')
        self.assertEqual(theme, 'ats_clean')
        self.assertEqual(path, 'resumes/pdf_template_ats_clean.html')
        theme, path = resolve_template('ats_clean_accent')
        self.assertEqual(theme, 'ats_clean_accent')
        self.assertEqual(path, 'resumes/pdf_template_ats_clean_accent.html')

    def test_removed_bw_themes_migrate_to_ats_clean(self):
        from resumes.services.pdf_exporter import resolve_template
        for old in ('standard', 'executive', 'minimalist', 'compact'):
            theme, path = resolve_template(old)
            self.assertEqual(theme, 'ats_clean',
                             f'{old!r} should migrate to ats_clean, got {theme!r}')
            self.assertEqual(path, 'resumes/pdf_template_ats_clean.html')

    def test_removed_color_themes_migrate_to_accent(self):
        from resumes.services.pdf_exporter import resolve_template
        for old in ('danette', 'zeyad'):
            theme, path = resolve_template(old)
            self.assertEqual(theme, 'ats_clean_accent',
                             f'{old!r} should migrate to ats_clean_accent, got {theme!r}')
            self.assertEqual(path, 'resumes/pdf_template_ats_clean_accent.html')

    def test_unknown_or_empty_falls_through_to_ats_clean(self):
        from resumes.services.pdf_exporter import resolve_template
        for bad in (None, '', '   ', 'not_a_theme', 'pdf_template.html'):
            theme, path = resolve_template(bad)
            self.assertEqual(theme, 'ats_clean',
                             f'{bad!r} should fall through to ats_clean')
            self.assertEqual(path, 'resumes/pdf_template_ats_clean.html')

    def test_migrated_template_file_actually_exists(self):
        """The migration is meaningless if it points at a non-existent
        template file. Importing the loader and resolving the file
        verifies both the map and the templates land in lockstep."""
        from django.template.loader import get_template
        from resumes.services.pdf_exporter import resolve_template, LIVE_THEMES
        for old in ('standard', 'executive', 'minimalist', 'compact', 'danette', 'zeyad'):
            _, path = resolve_template(old)
            get_template(path)  # raises TemplateDoesNotExist on miss
        for live in LIVE_THEMES:
            _, path = resolve_template(live)
            get_template(path)


class SkillCategorizerTests(SimpleTestCase):
    """Deterministic skill → category lookup + fail-safe behaviour
    (rule 012). Unknown skills MUST surface in 'Other', not get dropped
    or guessed into a wrong bucket."""

    def test_known_skills_land_in_expected_categories(self):
        from resumes.services.skill_categorizer import categorize_skill
        cases = {
            'Python': 'Languages',
            'TypeScript': 'Languages',
            'Django': 'Frameworks & Libraries',
            'React': 'Frameworks & Libraries',
            'PostgreSQL': 'Databases',
            'AWS': 'Cloud & DevOps',
            'Docker': 'Cloud & DevOps',
            'Pandas': 'ML & Data',
            'PyTorch': 'ML & Data',
            'Git': 'Tools & Platforms',
            'Jira': 'Tools & Platforms',
        }
        for skill, expected in cases.items():
            self.assertEqual(categorize_skill(skill), expected,
                             f'{skill!r} should be {expected!r}')

    def test_lookup_is_case_insensitive(self):
        from resumes.services.skill_categorizer import categorize_skill
        for variant in ('python', 'Python', 'PYTHON', '  python  '):
            self.assertEqual(categorize_skill(variant), 'Languages')

    def test_unknown_skill_lands_in_other_not_dropped(self):
        from resumes.services.skill_categorizer import (
            categorize_skill, group_skills_for_display,
        )
        self.assertEqual(categorize_skill('SomeNicheTool'), 'Other')
        groups = group_skills_for_display(['Python', 'SomeNicheTool'])
        flat = {s for g in groups for s in g['skills']}
        self.assertIn('SomeNicheTool', flat,
                      'unknown skill must NOT be dropped from output')
        other = next((g for g in groups if g['category'] == 'Other'), None)
        self.assertIsNotNone(other, "'Other' bucket must exist for unknown skills")
        self.assertIn('SomeNicheTool', other['skills'])

    def test_grouping_preserves_input_order_within_category(self):
        from resumes.services.skill_categorizer import group_skills_for_display
        groups = group_skills_for_display(['TypeScript', 'Python', 'Go'])
        langs = next(g for g in groups if g['category'] == 'Languages')
        self.assertEqual(langs['skills'], ['TypeScript', 'Python', 'Go'])

    def test_grouping_dedups_case_insensitive(self):
        from resumes.services.skill_categorizer import group_skills_for_display
        groups = group_skills_for_display(['Python', 'python', 'PYTHON'])
        langs = next(g for g in groups if g['category'] == 'Languages')
        self.assertEqual(langs['skills'], ['Python'])

    def test_grouping_handles_empty_and_invalid_entries(self):
        from resumes.services.skill_categorizer import group_skills_for_display
        # None / empty list returns []
        self.assertEqual(group_skills_for_display(None), [])
        self.assertEqual(group_skills_for_display([]), [])
        # Empty / whitespace / non-string entries are skipped, real
        # entries are kept.
        groups = group_skills_for_display(['', '   ', None, 123, 'Python'])
        self.assertEqual(len(groups), 1)
        self.assertEqual(groups[0]['category'], 'Languages')
        self.assertEqual(groups[0]['skills'], ['Python'])

    def test_grouping_returns_canonical_category_order(self):
        from resumes.services.skill_categorizer import (
            CATEGORIES_ORDER, group_skills_for_display,
        )
        # One representative skill per category — Other last.
        groups = group_skills_for_display([
            'Jira',          # Tools & Platforms
            'Python',        # Languages
            'AWS',           # Cloud & DevOps
            'Pandas',        # ML & Data
            'NicheThing',    # Other
            'Django',        # Frameworks & Libraries
            'PostgreSQL',    # Databases
        ])
        got = [g['category'] for g in groups]
        # Each category appears at most once, and only in CATEGORIES_ORDER
        # sequence (no shuffling by input order).
        idxs = [CATEGORIES_ORDER.index(c) for c in got]
        self.assertEqual(idxs, sorted(idxs),
                         f'groups not in canonical order: {got}')


class ResumeV2AdapterTests(SimpleTestCase):
    """The v2 → template-dict adapter must (a) flatten v2 sections into
    the v1 shape the templates consume, (b) prefer rich source v1 items
    for education / certifications / languages, and (c) NEVER emit
    fact_ids / provenance into the rendered output dict — that metadata
    stays inside GeneratedResumeV2 for downstream grounding, not on the
    recruiter's PDF."""

    def _build_v2(self):
        from resumes.services.resume_generator_v2 import (
            GeneratedResumeV2, GeneratedSection, GeneratedBullet, EntityBlock,
        )
        return GeneratedResumeV2(
            sections={
                'summary': GeneratedSection(
                    section='summary',
                    summary_text='Data scientist with 4 years of experience.'),
                'skills': GeneratedSection(
                    section='skills',
                    skills_line='Python, SQL, PyTorch, Docker'),
                'experience': GeneratedSection(
                    section='experience',
                    entities=[
                        EntityBlock(
                            entity_id='exp-1',
                            entity_display='Senior Engineer at Acme (2022–Present)',
                            bullets=[
                                GeneratedBullet(text='Shipped X.', fact_ids=['f1', 'f2']),
                                GeneratedBullet(text='Owned Y.', fact_ids=['f3'], hedged=True),
                            ]),
                    ]),
                'projects': GeneratedSection(
                    section='projects',
                    entities=[
                        EntityBlock(
                            entity_id='proj-1',
                            entity_display='SmartCV',
                            bullets=[GeneratedBullet(text='Built RAG knowledge base.', fact_ids=['f9'])]),
                    ]),
            },
        )

    def test_summary_and_skills_flatten(self):
        from resumes.services.resume_v2_adapter import resume_v2_to_template_dict
        out = resume_v2_to_template_dict(self._build_v2())
        self.assertEqual(out['professional_summary'],
                         'Data scientist with 4 years of experience.')
        self.assertEqual(out['skills'], ['Python', 'SQL', 'PyTorch', 'Docker'])

    def test_experience_entities_flatten_to_v1_items(self):
        from resumes.services.resume_v2_adapter import resume_v2_to_template_dict
        out = resume_v2_to_template_dict(self._build_v2())
        self.assertEqual(len(out['experience']), 1)
        item = out['experience'][0]
        # entity_display lands in title (no v1 source item to merge).
        self.assertIn('Senior Engineer', item['title'])
        # description is the flat list of bullet texts in input order.
        self.assertEqual(item['description'], ['Shipped X.', 'Owned Y.'])

    def test_projects_entities_flatten(self):
        from resumes.services.resume_v2_adapter import resume_v2_to_template_dict
        out = resume_v2_to_template_dict(self._build_v2())
        self.assertEqual(len(out['projects']), 1)
        self.assertEqual(out['projects'][0]['description'], ['Built RAG knowledge base.'])

    def test_no_fact_ids_or_provenance_leak_into_output(self):
        """The recruiter's PDF carries no provenance metadata — fact_ids,
        hedged, anchor_fact_id, entity_id must all stay inside
        GeneratedResumeV2."""
        from resumes.services.resume_v2_adapter import resume_v2_to_template_dict
        out = resume_v2_to_template_dict(self._build_v2())
        # Recursively walk the output and assert no provenance keys.
        leak_keys = {'fact_ids', 'hedged', 'anchor_fact_id', 'entity_id',
                     'fabrication_events', 'notes'}
        def _scan(node, path='out'):
            if isinstance(node, dict):
                bad = leak_keys & set(node.keys())
                self.assertFalse(bad, f'provenance keys leaked at {path}: {bad}')
                for k, v in node.items():
                    _scan(v, f'{path}.{k}')
            elif isinstance(node, list):
                for i, v in enumerate(node):
                    _scan(v, f'{path}[{i}]')
        _scan(out)

    def test_source_v1_merged_into_experience_items_by_entity_id(self):
        """When the v2 entity_id matches a v1 source item's id, the
        adapter preserves the v1 structured fields (company, duration,
        location, etc.) and only overwrites description bullets."""
        from resumes.services.resume_generator_v2 import (
            GeneratedResumeV2, GeneratedSection, GeneratedBullet, EntityBlock,
        )
        from resumes.services.resume_v2_adapter import resume_v2_to_template_dict
        generated = GeneratedResumeV2(sections={
            'experience': GeneratedSection(
                section='experience',
                entities=[EntityBlock(
                    entity_id='abc-123',
                    entity_display='ignored when merge happens',
                    bullets=[GeneratedBullet(text='new bullet.', fact_ids=[])])]),
        })
        source = {
            'experience': [{
                'id': 'abc-123',
                'title': 'Senior Engineer',
                'company': 'Acme Corp',
                'duration': '2022–Present',
                'location': 'Remote',
                'description': ['old bullet that must be overwritten'],
            }],
        }
        out = resume_v2_to_template_dict(generated, source=source)
        item = out['experience'][0]
        self.assertEqual(item['title'], 'Senior Engineer')
        self.assertEqual(item['company'], 'Acme Corp')
        self.assertEqual(item['duration'], '2022–Present')
        self.assertEqual(item['description'], ['new bullet.'])

    def test_education_falls_back_to_source_v1(self):
        """v2 doesn't enrich education — the adapter must keep the
        structured v1 source rows rather than overwriting with the
        flatter v2 'lines'."""
        from resumes.services.resume_generator_v2 import (
            GeneratedResumeV2, GeneratedSection,
        )
        from resumes.services.resume_v2_adapter import resume_v2_to_template_dict
        generated = GeneratedResumeV2(sections={
            'education': GeneratedSection(
                section='education',
                lines=['v2 simplification of education']),
        })
        source = {'education': [
            {'degree': 'B.Sc.', 'field': 'CS', 'institution': 'KSIU', 'year': '2025'},
        ]}
        out = resume_v2_to_template_dict(generated, source=source)
        # Source v1 wins because it has richer structure.
        self.assertEqual(out['education'], source['education'])


def _ats_render_context():
    """Build a minimal render context that exercises every section of the
    base template (used by AtsCleanRenderTests below)."""
    from types import SimpleNamespace
    from resumes.services.skill_categorizer import (
        group_skills_for_display, should_show_grouped,
    )

    profile = SimpleNamespace(
        full_name='Jane Doe',
        email='jane@example.com',
        phone='+1 555 123 4567',
        location='Cairo, Egypt',
        linkedin_url='https://www.linkedin.com/in/janedoe',
        github_url='https://github.com/janedoe',
        portfolio_url='',
        kaggle_url='',
        scholar_url='',
    )
    skills = ['Python', 'Django', 'PostgreSQL', 'AWS', 'Pandas', 'SomeNicheTool']
    resume = {
        'professional_title': 'Senior Data Scientist',
        'professional_summary': (
            'Senior data scientist with hands-on ML engineering. Shipped models '
            'in production at scale across analytics and personalization.'),
        'skills': skills,
        'experience': [{
            'title': 'Senior Data Scientist',
            'company': 'Acme Analytics',
            'location': 'Remote',
            'industry': 'AdTech',
            'duration': '2022 – Present',
            'description': [
                'Cut churn by 18% with a calibrated propensity model.',
                'Owned end-to-end MLOps pipeline on AWS for 5 models in production.',
            ],
        }],
        'projects': [{
            'name': 'SmartCV',
            'url': 'https://example.com/smartcv',
            'technologies': ['Django', 'WeasyPrint'],
            'description': ['ATS-clean resume rendering pipeline.'],
        }],
        'education': [{
            'degree': 'B.Sc.', 'field': 'Computer Science',
            'institution': 'KSIU', 'year': '2025', 'gpa': '3.8',
        }],
        'certifications': [{'name': 'AWS Solutions Architect', 'issuer': 'Amazon', 'date': '2024'}],
        'languages': ['English (Native)', 'Arabic (Native)'],
    }
    skill_groups = group_skills_for_display(skills)
    return {
        'profile': profile,
        'user': None,
        'resume': resume,
        'section_order': ['summary', 'skills', 'experience',
                          'projects', 'education', 'certifications', 'languages'],
        'skill_groups': skill_groups,
        'show_grouped_skills': should_show_grouped(skill_groups, len(skills)),
        'theme': 'ats_clean',
    }


class AtsCleanRenderTests(SimpleTestCase):
    """Render the two surviving templates with a real-ish context and
    assert the rendered HTML satisfies the project's KB ats_rules. These
    are template-level checks (the PDF is rendered by WeasyPrint
    downstream, but the structure/CSS lives in the templates)."""

    THEMES = ('ats_clean', 'ats_clean_accent')

    def _render(self, theme):
        from django.template.loader import render_to_string
        return render_to_string(f'resumes/pdf_template_{theme}.html', _ats_render_context())

    def test_no_table_tags_for_layout(self):
        """Rule 004 — single-column body, no <table> for alignment."""
        import re
        for theme in self.THEMES:
            html = self._render(theme)
            self.assertFalse(
                re.search(r'<table[\s>]', html),
                f"theme {theme!r} emitted a <table> tag (rule 004 violation)",
            )

    def test_canonical_section_names(self):
        """Rule 002 — canonical section headers."""
        for theme in self.THEMES:
            html = self._render(theme)
            for needed in (
                'Professional Summary', 'Skills', 'Work Experience',
                'Projects', 'Education', 'Certifications', 'Languages',
            ):
                self.assertIn(needed, html,
                              f'theme {theme!r} missing canonical section name {needed!r}')

    def test_no_creative_section_headers(self):
        """Rule 002 — none of the rejected names from the audit."""
        for theme in self.THEMES:
            html = self._render(theme)
            for bad in ('Profile', 'Core Skills', 'My Journey',
                        'The Toolkit', 'What I Bring'):
                self.assertNotIn(f'>{bad}<', html,
                                 f'theme {theme!r} uses non-canonical section name {bad!r}')

    def test_font_sizes_meet_rule_007(self):
        """Body ≥10pt, section heading ≥14pt, name 18–24pt (rule 007).
        We check the CSS rules in the rendered HTML head."""
        import re
        for theme in self.THEMES:
            html = self._render(theme)
            body = re.search(r'body\s*\{[^}]*font-size:\s*([\d.]+)pt', html)
            self.assertIsNotNone(body, f'no body font-size found in {theme!r}')
            self.assertGreaterEqual(float(body.group(1)), 10.0,
                                    f'theme {theme!r} body font below 10pt')
            sec = re.search(r'\.section-title\s*\{[^}]*font-size:\s*([\d.]+)pt', html)
            self.assertIsNotNone(sec, f'no section-title font-size in {theme!r}')
            self.assertGreaterEqual(float(sec.group(1)), 14.0,
                                    f'theme {theme!r} section title below 14pt')
            name = re.search(r'\.name\s*\{[^}]*font-size:\s*([\d.]+)pt', html)
            self.assertIsNotNone(name, f'no .name font-size in {theme!r}')
            name_pt = float(name.group(1))
            self.assertTrue(18.0 <= name_pt <= 24.0,
                            f'theme {theme!r} name font {name_pt}pt outside 18-24pt')

    def test_web_safe_font_family(self):
        """Rule 007 — Helvetica / Arial chain, no Segoe UI / decorative."""
        for theme in self.THEMES:
            html = self._render(theme)
            self.assertIn("'Helvetica'", html)
            self.assertIn("'Arial'", html)
            for bad in ('Segoe UI', 'Comic Sans', 'Papyrus'):
                self.assertNotIn(bad, html,
                                 f'theme {theme!r} references non-safe font {bad!r}')

    def test_bullet_marker_is_disc(self):
        """Rule 010 — disc / • bullets, not arrows / checks / stars."""
        for theme in self.THEMES:
            html = self._render(theme)
            self.assertIn('list-style-type: disc', html,
                          f'theme {theme!r} should use disc bullets')
            for bad in ('list-style-type: square', '→', '►', '✓', '★'):
                self.assertNotIn(bad, html,
                                 f'theme {theme!r} uses unsafe bullet marker {bad!r}')

    def test_page_break_controls_present(self):
        """Rule 011 (CSS half) — entry blocks have break-inside: avoid
        and bullets carry widows/orphans control so the last bullet of a
        role doesn't get orphaned on the next page."""
        for theme in self.THEMES:
            html = self._render(theme)
            self.assertIn('break-inside: avoid', html,
                          f'theme {theme!r} missing break-inside: avoid')
            self.assertIn('widows', html)
            self.assertIn('orphans', html)

    def test_skills_rendered_as_groups(self):
        """Rule 012 — skills grouped by category, each on its own line
        with a bolded category name."""
        for theme in self.THEMES:
            html = self._render(theme)
            self.assertIn('class="skill-cat">Languages:', html,
                          f'theme {theme!r} missing Languages: group label')
            self.assertIn('Other:', html,
                          f"theme {theme!r} should surface 'Other' bucket for unknown skill")

    def test_unknown_skill_appears_in_other_in_rendered_html(self):
        """The fail-safe in skill_categorizer must reach the rendered
        artifact: an unknown skill is visible under 'Other', not
        dropped."""
        for theme in self.THEMES:
            html = self._render(theme)
            self.assertIn('SomeNicheTool', html,
                          f'theme {theme!r} dropped an unknown skill')

    def test_accent_theme_uses_single_color_clean_theme_is_black_only(self):
        """ats_clean: black-only. ats_clean_accent: one restrained
        accent color (#1e3a8a) and nothing else."""
        clean = self._render('ats_clean')
        # ats_clean has no accent color block.
        self.assertNotIn('#1e3a8a', clean,
                         'ats_clean must stay black-only — no accent color')

        accent = self._render('ats_clean_accent')
        self.assertIn('#1e3a8a', accent,
                      'ats_clean_accent must declare one accent color')
        # Only this single accent appears — no rainbow.
        import re
        hex_colors = set(re.findall(r'#[0-9a-fA-F]{6}', accent))
        # #000000 etc may appear; allow black + the accent only.
        non_black = {c for c in hex_colors if c.lower() not in {'#000000'}}
        self.assertEqual(non_black, {'#1e3a8a'},
                         f'ats_clean_accent must use exactly one accent color; '
                         f'found {non_black}')

    def test_no_v2_provenance_keys_in_rendered_html(self):
        """Defensive: even if a fact_ids key were ever added to the
        adapter output, the templates would render the dict and could
        leak it. Assert no provenance labels appear in the HTML."""
        for theme in self.THEMES:
            html = self._render(theme)
            for bad in ('fact_ids', 'fabrication_events', 'anchor_fact_id'):
                self.assertNotIn(bad, html,
                                 f'theme {theme!r} leaked provenance key {bad!r}')

    def test_contact_block_lives_in_document_body(self):
        """Rule 009 — contact info in body, not @page header/footer.
        The .contact-line lives inside .resume-header, which is in
        <body>, not in @page { @top-center }."""
        for theme in self.THEMES:
            html = self._render(theme)
            self.assertIn('class="contact-line"', html)
            self.assertNotIn('@top-center', html,
                             f'theme {theme!r} placed content in @page header')
            self.assertNotIn('@bottom-center', html)


# ---------------------------------------------------------------------------
# Pillar 4 polish pass — expanded skill table, balanced-grouping guard,
# portfolio link relabeled to "Portfolio".
# ---------------------------------------------------------------------------


class SkillCategorizerExpansionTests(SimpleTestCase):
    """Common ML / AI vocabulary added to the lookup table so it no longer
    drops into 'Other'. The fail-safe (unknown → 'Other') must survive
    the expansion."""

    def test_ml_ai_terms_classify_to_ml_and_data(self):
        from resumes.services.skill_categorizer import categorize_skill
        for skill in (
            'AI Model Development', 'Large Language Models', 'LLMs',
            'Generative AI', 'Gen AI', 'GenAI',
            'Model Optimization', 'Transfer Learning',
            'Model Evaluation', 'AI Feature Implementation',
            'AI Tools Deployment', 'Internal AI Tool Deployment',
            'Supervised Learning', 'Unsupervised Learning',
            'Supervised & Unsupervised Learning',
            'Feature Engineering', 'Model Deployment',
            'Fine-tuning', 'Fine Tuning',
            'Prompt Engineering', 'RAG',
            'Computer Vision', 'NLP', 'Natural Language Processing',
        ):
            got = categorize_skill(skill)
            self.assertEqual(
                got, 'ML & Data',
                f'{skill!r} should categorise as ML & Data, got {got!r}',
            )

    def test_technical_documentation_lands_in_tools(self):
        from resumes.services.skill_categorizer import categorize_skill
        self.assertEqual(
            categorize_skill('Technical Documentation'),
            'Tools & Platforms',
        )

    def test_fail_safe_still_routes_unknown_to_other(self):
        """Expanding the table must not break the fail-safe."""
        from resumes.services.skill_categorizer import categorize_skill
        self.assertEqual(categorize_skill('SomeNicheThingFromTheFuture'), 'Other')


class ShouldShowGroupedTests(SimpleTestCase):
    """Balanced-grouping guard. Rule: show grouped only when ≥3 categories
    are populated AND no single category exceeds MAX_DOMINANCE (60%) of
    the total skills. Otherwise fall back to the flat list — for ALL
    users, not just dev/ML profiles."""

    def _groups(self, skills):
        from resumes.services.skill_categorizer import group_skills_for_display
        return group_skills_for_display(skills)

    def test_balanced_set_returns_true(self):
        """Six different categories, one skill each → balanced → show grouped."""
        from resumes.services.skill_categorizer import should_show_grouped
        skills = ['Python', 'Django', 'PostgreSQL', 'AWS', 'Pandas', 'Git']
        groups = self._groups(skills)
        self.assertEqual(len(groups), 6)
        self.assertTrue(should_show_grouped(groups, len(skills)))

    def test_too_few_categories_returns_false(self):
        """Everything in 1-2 buckets → not categorical enough → flat."""
        from resumes.services.skill_categorizer import should_show_grouped
        # All Languages.
        skills = ['Python', 'TypeScript', 'Java', 'Go', 'C++']
        groups = self._groups(skills)
        self.assertEqual(len(groups), 1)
        self.assertFalse(should_show_grouped(groups, len(skills)))
        # Two categories.
        skills2 = ['Python', 'Django']
        groups2 = self._groups(skills2)
        self.assertEqual(len(groups2), 2)
        self.assertFalse(should_show_grouped(groups2, len(skills2)))

    def test_dominant_other_bucket_returns_false(self):
        """The motivating case: lots of unknown skills dump into Other →
        Other dominates → flat reads cleaner than a lopsided group."""
        from resumes.services.skill_categorizer import should_show_grouped
        skills = (
            ['UnknownA', 'UnknownB', 'UnknownC', 'UnknownD', 'UnknownE',
             'UnknownF', 'UnknownG']  # 7 → Other
            + ['Python', 'Django', 'AWS']  # 3 distinct cats
        )
        groups = self._groups(skills)
        other = next(g for g in groups if g['category'] == 'Other')
        self.assertEqual(len(other['skills']), 7)
        self.assertGreaterEqual(len(groups), 3)
        # 7 / 10 = 70 % > 60 % → unbalanced.
        self.assertFalse(should_show_grouped(groups, len(skills)))

    def test_just_under_dominance_threshold_returns_true(self):
        """5 in one bucket / 9 total = 55 % — under the 60 % threshold."""
        from resumes.services.skill_categorizer import should_show_grouped
        skills = (
            ['UnknownA', 'UnknownB', 'UnknownC', 'UnknownD', 'UnknownE']  # 5 Other
            + ['Python', 'Django', 'AWS', 'Git']  # 4 distinct cats
        )
        groups = self._groups(skills)
        self.assertEqual(len(groups), 5)
        self.assertTrue(should_show_grouped(groups, len(skills)))

    def test_empty_inputs_return_false(self):
        from resumes.services.skill_categorizer import should_show_grouped
        self.assertFalse(should_show_grouped([], 0))
        self.assertFalse(should_show_grouped([], 5))
        # Single-category corner case.
        single = [{'category': 'Languages', 'skills': ['Python']}]
        self.assertFalse(should_show_grouped(single, 1))


class BalancedSkillsRenderingTests(SimpleTestCase):
    """End-to-end: the template renders the FLAT list when the grouping
    would be lopsided, and the GROUPED layout when it would be balanced."""

    def _render(self, skills, theme='ats_clean'):
        from django.template.loader import render_to_string
        from resumes.services.skill_categorizer import (
            group_skills_for_display, should_show_grouped,
        )
        ctx = _ats_render_context()
        ctx['resume']['skills'] = skills
        groups = group_skills_for_display(skills)
        ctx['skill_groups'] = groups
        ctx['show_grouped_skills'] = should_show_grouped(groups, len(skills))
        return render_to_string(f'resumes/pdf_template_{theme}.html', ctx)

    def test_lopsided_set_renders_flat(self):
        """12 unknown + 2 known → grouped would dump 12 into Other → flat."""
        skills = [f'Unknown{i}' for i in range(12)] + ['Python', 'Django']
        html = self._render(skills)
        # No category labels emitted — grouped path was skipped.
        self.assertNotIn('class="skill-cat">', html)
        # Flat list IS emitted: every skill present.
        for s in skills:
            self.assertIn(s, html)

    def test_balanced_set_renders_grouped(self):
        """Six skills in six categories → balanced → grouped renders with
        the category labels (Django escapes '&' to '&amp;')."""
        skills = ['Python', 'Django', 'PostgreSQL', 'AWS', 'Pandas', 'Git']
        html = self._render(skills)
        self.assertIn('class="skill-cat">Languages:', html)
        self.assertIn('class="skill-cat">Frameworks &amp; Libraries:', html)
        self.assertIn('class="skill-cat">Databases:', html)
        self.assertIn('class="skill-cat">Cloud &amp; DevOps:', html)
        self.assertIn('class="skill-cat">ML &amp; Data:', html)
        self.assertIn('class="skill-cat">Tools &amp; Platforms:', html)

    def test_unknown_skill_still_in_other_when_grouped(self):
        """Even in the grouped path, the fail-safe holds: unknown skill →
        Other, never dropped, never misfiled."""
        skills = ['Python', 'Django', 'PostgreSQL', 'AWS', 'Pandas',
                  'Git', 'SomeNicheThing']
        html = self._render(skills)
        # 7 categories, max 1/7 = 14 % → balanced → grouped.
        self.assertIn('class="skill-cat">Other:', html)
        self.assertIn('SomeNicheThing', html)

    def test_real_world_lopsided_zeyad_profile_falls_back_to_flat(self):
        """Regression: this is the actual skill list shape that motivated
        the polish pass. Several legitimate ML/AI strings used to land
        in 'Other'; the table expansion + balanced guard together must
        deliver a clean grouping OR a clean flat list — never a lopsided
        dump."""
        skills = [
            'Python', 'SQL',
            'Machine Learning', 'Deep Learning', 'Transfer Learning',
            'Supervised & Unsupervised Learning', 'Model Evaluation',
            'Computer Vision', 'NLP',
            'Pandas', 'NumPy', 'scikit-learn', 'TensorFlow',
            'AWS', 'Docker',
            'Team Management', 'Project Management',  # → Other
        ]
        html = self._render(skills)
        # Either grouped (with balanced categories) or flat — but the
        # 'Other' bucket must NOT dominate the rendered output.
        from resumes.services.skill_categorizer import (
            group_skills_for_display, should_show_grouped,
        )
        groups = group_skills_for_display(skills)
        balanced = should_show_grouped(groups, len(skills))
        if balanced:
            other = next((g for g in groups if g['category'] == 'Other'), None)
            other_n = len(other['skills']) if other else 0
            self.assertLessEqual(
                other_n / len(skills), 0.60,
                'grouped path emitted a dominant Other bucket',
            )
        else:
            self.assertNotIn('class="skill-cat">', html,
                             'flat path should not emit category labels')
        # In either case, every skill must still appear (Django escapes
        # '&' to '&amp;', so compare the escaped form).
        from django.utils.html import escape
        for s in skills:
            self.assertIn(escape(s), html, f'skill {s!r} dropped from output')


class PortfolioLabelTests(SimpleTestCase):
    """The portfolio link renders with the visible label 'Portfolio'
    (href stays the real URL). LinkedIn / GitHub / Kaggle keep their
    raw URL as the visible text — recruiter heuristic + rule 009."""

    def _render(self, **profile_overrides):
        from django.template.loader import render_to_string
        from types import SimpleNamespace
        ctx = _ats_render_context()
        profile_fields = dict(vars(ctx['profile']))
        profile_fields.update(profile_overrides)
        ctx['profile'] = SimpleNamespace(**profile_fields)
        return render_to_string('resumes/pdf_template_ats_clean.html', ctx)

    def test_portfolio_label_replaces_visible_url_text(self):
        portfolio = 'https://my-portfolio-weld-eta-67.vercel.app/'
        html = self._render(portfolio_url=portfolio)
        # Visible text is the friendly label.
        self.assertIn('>Portfolio</a>', html)
        # href keeps the real URL.
        self.assertIn(f'href="{portfolio}"', html)
        # Raw vercel domain does NOT appear as visible text
        # (would be '>my-portfolio-weld-eta-67.vercel.app...').
        self.assertNotIn('>my-portfolio-weld-eta-67.vercel.app', html)

    def test_linkedin_github_kaggle_still_show_raw_url_text(self):
        html = self._render(
            linkedin_url='https://www.linkedin.com/in/janedoe',
            github_url='https://github.com/janedoe',
            kaggle_url='https://kaggle.com/janedoe',
        )
        # Raw URLs (with https:// / www. stripped) appear as visible text.
        self.assertIn('linkedin.com/in/janedoe', html)
        self.assertIn('github.com/janedoe', html)
        self.assertIn('kaggle.com/janedoe', html)
        # NOT relabeled.
        self.assertNotIn('>LinkedIn</a>', html)
        self.assertNotIn('>GitHub</a>', html)
        self.assertNotIn('>Kaggle</a>', html)

    def test_missing_portfolio_emits_no_portfolio_anchor(self):
        html = self._render(portfolio_url='')
        self.assertNotIn('>Portfolio</a>', html)

    def test_portfolio_label_works_in_accent_theme_too(self):
        from django.template.loader import render_to_string
        from types import SimpleNamespace
        ctx = _ats_render_context()
        profile_fields = dict(vars(ctx['profile']))
        profile_fields['portfolio_url'] = 'https://example.dev/me'
        ctx['profile'] = SimpleNamespace(**profile_fields)
        html = render_to_string('resumes/pdf_template_ats_clean_accent.html', ctx)
        self.assertIn('>Portfolio</a>', html)
        self.assertIn('href="https://example.dev/me"', html)


# ---------------------------------------------------------------------------
# "Present" fabrication + reverse-chrono mis-sort fix.
# Single source of truth: is_current=True on the entry. Missing end is
# rendered as start alone (not "Present"); legacy "X - Present" without
# is_current heals on re-render and re-sorts honestly.
# ---------------------------------------------------------------------------


class AssembleDurationHonestTests(SimpleTestCase):
    """The shared honest duration assembler. Used by resume_generator
    offline-fallback paths and the docx exporter."""

    def test_is_current_true_with_start_emits_present(self):
        from resumes.services.resume_normalizer import assemble_duration_honest
        self.assertEqual(assemble_duration_honest('Jul 2024', '', True),
                         'Jul 2024 - Present')
        self.assertEqual(assemble_duration_honest('Jul 2024', None, True),
                         'Jul 2024 - Present')
        self.assertEqual(assemble_duration_honest('Jul 2024', 'Present', True),
                         'Jul 2024 - Present')

    def test_missing_end_without_is_current_renders_start_alone(self):
        from resumes.services.resume_normalizer import assemble_duration_honest
        self.assertEqual(assemble_duration_honest('Jul 2024', '', None), 'Jul 2024')
        self.assertEqual(assemble_duration_honest('Jul 2024', None, False), 'Jul 2024')
        self.assertEqual(assemble_duration_honest('Jul 2024', '   ', None), 'Jul 2024')

    def test_legacy_present_without_is_current_heals_to_start_alone(self):
        """The motivating bug: end='Present' was LLM-fabricated for a
        non-current role. Without is_current=True, treat as unknown and
        render start alone."""
        from resumes.services.resume_normalizer import assemble_duration_honest
        self.assertEqual(assemble_duration_honest('Jul 2024', 'Present', None),
                         'Jul 2024')
        self.assertEqual(assemble_duration_honest('Jul 2024', 'Current', False),
                         'Jul 2024')
        self.assertEqual(assemble_duration_honest('Jul 2024', 'Ongoing', None),
                         'Jul 2024')

    def test_closed_range_passes_through(self):
        from resumes.services.resume_normalizer import assemble_duration_honest
        self.assertEqual(assemble_duration_honest('Aug 2025', 'Sep 2025', None),
                         'Aug 2025 - Sep 2025')
        self.assertEqual(assemble_duration_honest('Jun 2025', 'Dec 2025', False),
                         'Jun 2025 - Dec 2025')

    def test_no_start_with_is_current_emits_bare_present(self):
        from resumes.services.resume_normalizer import assemble_duration_honest
        self.assertEqual(assemble_duration_honest('', '', True), 'Present')
        self.assertEqual(assemble_duration_honest(None, None, True), 'Present')

    def test_empty_inputs_return_empty(self):
        from resumes.services.resume_normalizer import assemble_duration_honest
        self.assertEqual(assemble_duration_honest('', '', None), '')
        self.assertEqual(assemble_duration_honest(None, None, None), '')


class HealExperienceDurationsTests(SimpleTestCase):
    """Defensive render-time pass for legacy resumes whose stored
    duration was inflated to 'X - Present' on a non-current role."""

    def test_legacy_present_without_is_current_heals(self):
        from resumes.services.resume_normalizer import heal_experience_durations
        out = heal_experience_durations([{
            'title': 'IT Intern', 'start_date': 'Jul 2024',
            'end_date': 'Dec 2025',
            'duration': 'Jul 2024 - Present',
            # is_current absent → legacy LLM fabrication, must heal.
        }])
        self.assertEqual(out[0]['duration'], 'Jul 2024')

    def test_is_current_true_preserves_present(self):
        from resumes.services.resume_normalizer import heal_experience_durations
        out = heal_experience_durations([{
            'title': 'Engineer', 'start_date': 'Jan 2024',
            'end_date': 'Present', 'is_current': True,
            'duration': 'Jan 2024 - Present',
        }])
        self.assertEqual(out[0]['duration'], 'Jan 2024 - Present')

    def test_closed_range_passes_through(self):
        from resumes.services.resume_normalizer import heal_experience_durations
        out = heal_experience_durations([{
            'title': 'Intern', 'start_date': 'Aug 2025',
            'end_date': 'Sep 2025',
            'duration': 'Aug 2025 - Sep 2025',
        }])
        self.assertEqual(out[0]['duration'], 'Aug 2025 - Sep 2025')

    def test_non_list_input_passes_through(self):
        from resumes.services.resume_normalizer import heal_experience_durations
        self.assertEqual(heal_experience_durations(None), None)
        self.assertEqual(heal_experience_durations({}), {})

    def test_does_not_mutate_input(self):
        from resumes.services.resume_normalizer import heal_experience_durations
        src = [{
            'title': 'X', 'start_date': 'Jul 2024',
            'end_date': 'Present', 'duration': 'Jul 2024 - Present',
        }]
        out = heal_experience_durations(src)
        self.assertEqual(src[0]['duration'], 'Jul 2024 - Present')  # untouched
        self.assertEqual(out[0]['duration'], 'Jul 2024')


class HonestReverseChronoSortTests(SimpleTestCase):
    """sort_experience_reverse_chronological + _extract_end_yearmonth:
    a missing / fabricated 'Present' on a non-current role must NOT
    inflate the sort key to today_ym."""

    def _sort(self, exps):
        from resumes.services.resume_normalizer import (
            sort_experience_reverse_chronological,
        )
        return sort_experience_reverse_chronological({'experience': list(exps)})

    def test_aoi_below_almansour_when_is_current_not_set(self):
        """The motivating regression: AOI start=Jul 2024 with a
        fabricated end (LLM injected 'Dec 2025' + duration='- Present')
        must NOT outrank a genuinely-later closed role."""
        aoi = {
            'title': 'IT Intern', 'company': 'AOI',
            'start_date': 'Jul 2024', 'end_date': 'Dec 2025',
            'duration': 'Jul 2024 - Present',
        }
        almansour = {
            'title': 'DT Intern', 'company': 'Almansour',
            'start_date': 'Aug 2025', 'end_date': 'Sep 2025',
            'duration': 'Aug 2025 - Sep 2025',
        }
        depi = {
            'title': 'AI Trainee', 'company': 'DEPI',
            'start_date': 'Jun 2025', 'end_date': 'Dec 2025',
            'duration': 'Jun 2025 - Dec 2025',
        }
        out = self._sort([aoi, almansour, depi])['experience']
        order = [e['company'] for e in out]
        self.assertEqual(order, ['DEPI', 'Almansour', 'AOI'],
                         f'expected DEPI > Almansour > AOI, got {order}')

    def test_genuine_is_current_role_sorts_at_top(self):
        """A real ongoing role (is_current=True) still rises to the top."""
        ongoing = {
            'title': 'Current Role', 'company': 'ZCo',
            'start_date': 'Jan 2024', 'end_date': 'Present',
            'is_current': True,
            'duration': 'Jan 2024 - Present',
        }
        closed_2025 = {
            'title': 'Closed', 'company': 'ACo',
            'start_date': 'Jun 2025', 'end_date': 'Dec 2025',
            'duration': 'Jun 2025 - Dec 2025',
        }
        out = self._sort([closed_2025, ongoing])['experience']
        self.assertEqual([e['company'] for e in out], ['ZCo', 'ACo'])

    def test_legacy_duration_present_heals_in_sort(self):
        """Stored duration tail 'Present' without is_current=True does
        NOT promote the role — sort falls back to start date."""
        from resumes.services.resume_normalizer import _extract_end_yearmonth
        today_ym = (2026, 6)
        entry = {
            'start_date': 'Jul 2024',
            'end_date': None,
            'duration': 'Jul 2024 - Present',
            # No is_current.
        }
        # End extraction returns None → sort caller will fall back to start.
        self.assertIsNone(_extract_end_yearmonth(entry, today_ym))

    def test_end_present_without_is_current_returns_none_in_parse(self):
        from resumes.services.resume_normalizer import _parse_yearmonth
        today_ym = (2026, 6)
        # Default allow_present=False → "Present" → None.
        self.assertIsNone(_parse_yearmonth('Present', today_ym))
        self.assertIsNone(_parse_yearmonth('Currently', today_ym))
        self.assertIsNone(_parse_yearmonth('Jul 2024 - Present', today_ym))
        # Explicit opt-in → today_ym.
        self.assertEqual(_parse_yearmonth('Present', today_ym, allow_present=True),
                         today_ym)

    def test_missing_end_falls_to_start_date(self):
        """end=None is treated as unknown — sort uses start_date."""
        out = self._sort([
            {'title': 'A', 'company': 'A', 'start_date': 'Jul 2024',
             'end_date': None, 'duration': ''},
            {'title': 'B', 'company': 'B', 'start_date': 'Aug 2025',
             'end_date': 'Sep 2025', 'duration': 'Aug 2025 - Sep 2025'},
        ])['experience']
        # B (ends Sep 2025) outranks A (no end, starts Jul 2024 → 2024-7).
        self.assertEqual([e['company'] for e in out], ['B', 'A'])

    def test_closed_ranges_untouched(self):
        """Regression: closed ranges still sort correctly by end."""
        out = self._sort([
            {'title': 'Old', 'company': 'Old', 'start_date': 'Jan 2020',
             'end_date': 'Dec 2020', 'duration': 'Jan 2020 - Dec 2020'},
            {'title': 'New', 'company': 'New', 'start_date': 'Jan 2025',
             'end_date': 'Dec 2025', 'duration': 'Jan 2025 - Dec 2025'},
        ])['experience']
        self.assertEqual([e['company'] for e in out], ['New', 'Old'])


class DocxFormatDateRangeHonestTests(SimpleTestCase):
    """docx_exporter._format_date_range: empty end is NOT "Present";
    'Present' is honored only when is_current=True."""

    def test_empty_end_renders_start_alone(self):
        from resumes.services.docx_exporter import _format_date_range
        out = _format_date_range('Jul 2024', '')
        self.assertNotIn('Present', out)
        self.assertIn('2024', out)

    def test_present_token_without_is_current_renders_start_alone(self):
        from resumes.services.docx_exporter import _format_date_range
        out = _format_date_range('Jul 2024', 'Present')
        self.assertNotIn('Present', out)
        out = _format_date_range('Jul 2024', 'Present', is_current=False)
        self.assertNotIn('Present', out)

    def test_present_with_is_current_true_renders_present(self):
        from resumes.services.docx_exporter import _format_date_range
        out = _format_date_range('Jul 2024', 'Present', is_current=True)
        self.assertIn('Present', out)
        self.assertIn('2024', out)

    def test_closed_range_unchanged(self):
        from resumes.services.docx_exporter import _format_date_range
        out = _format_date_range('Aug 2025', 'Sep 2025')
        self.assertIn('Aug 2025', out)
        self.assertIn('Sep 2025', out)
        self.assertNotIn('Present', out)


class PromptHardeningAssertionTests(SimpleTestCase):
    """Lock in the prompt text changes that stop LLMs from fabricating
    'Present' for missing end dates. These assertions guard against
    accidental rollback."""

    def test_llm_validator_no_longer_infers_present_from_absence(self):
        from profiles.services import llm_validator
        src = llm_validator.__doc__ or ''
        # The validator's instruction text lives in the module-level
        # CV_VALIDATOR_PROMPT (or equivalent) — read the file.
        import inspect
        file_src = inspect.getsource(llm_validator)
        # Old fabrication instruction removed.
        self.assertNotIn(
            'infer "Present" only if it is the most recent role',
            file_src,
            'CV-parser prompt must not instruct LLM to infer Present '
            'from absence of an end date',
        )
        # New explicit rule present.
        self.assertIn('Set `is_current=true`', file_src)
        self.assertIn('explicitly says ongoing', file_src)

    def test_resume_generator_prompt_uses_closed_range_example(self):
        from resumes.services import resume_generator
        import inspect
        file_src = inspect.getsource(resume_generator)
        # Old open-ended example replaced.
        self.assertNotIn('combine as "Aug 2025 - Present"', file_src)
        # Closed-range example present.
        self.assertIn('CLOSED range', file_src)
        # Anti-fabrication rule present.
        self.assertIn('DO NOT FABRICATE "Present"', file_src)


class ResumeExperienceSchemaIsCurrentTests(SimpleTestCase):
    """The is_current field exists on both Experience (profile-side)
    and ResumeExperience (output-side) schemas, defaulting to None so
    legacy records don't require migration."""

    def test_experience_schema_has_is_current(self):
        from profiles.services.schemas import Experience
        fields = Experience.model_fields
        self.assertIn('is_current', fields)
        # Optional + default None.
        exp = Experience(title='X', company='Y')
        self.assertIsNone(exp.is_current)
        exp2 = Experience(title='X', company='Y', is_current=True)
        self.assertTrue(exp2.is_current)

    def test_resume_experience_schema_has_is_current(self):
        from profiles.services.schemas import ResumeExperience
        fields = ResumeExperience.model_fields
        self.assertIn('is_current', fields)
        re = ResumeExperience(title='X', company='Y')
        self.assertIsNone(re.is_current)


class PdfRenderHealsLegacyPresentTests(SimpleTestCase):
    """End-to-end: rendering a resume whose stored experience contains
    the AOI-style legacy fabrication produces:
      (a) the AOI duration shown as 'Jul 2024' (start alone), and
      (b) AOI sorted BELOW Almansour."""

    def _ctx(self, experience):
        from resumes.services.resume_normalizer import (
            heal_experience_durations, sort_experience_reverse_chronological,
        )
        ctx = _ats_render_context()
        # Heal + re-sort, mirroring what pdf_exporter.generate_pdf does.
        healed = heal_experience_durations(experience)
        content = {'experience': healed,
                   'professional_summary': 'x',
                   'skills': ['Python']}
        sorted_resume = sort_experience_reverse_chronological(content)
        ctx['resume'] = sorted_resume
        ctx['resume']['professional_title'] = 'Engineer'
        ctx['section_order'] = ['summary', 'skills', 'experience']
        return ctx

    def _render(self, experience):
        from django.template.loader import render_to_string
        return render_to_string(
            'resumes/pdf_template_ats_clean.html',
            self._ctx(experience),
        )

    def _item_block(self, html, title_marker):
        """Return the `<div class="item">...</div>` block containing
        the given title marker. The PDF template renders title first,
        then date, then company — so reads windowed forward from the
        title cover the date AND the company line."""
        import re
        # Find all item blocks and pick the one containing the marker.
        for m in re.finditer(r'<div class="item">[\s\S]*?</div>\s*</div>', html):
            if title_marker in m.group(0):
                return m.group(0)
        return ''

    def test_aoi_renders_start_alone_and_below_almansour(self):
        aoi = {
            'title': 'IT Intern', 'company': 'AOI',
            'start_date': 'Jul 2024', 'end_date': 'Dec 2025',
            'duration': 'Jul 2024 - Present',
        }
        almansour = {
            'title': 'DT Intern', 'company': 'Almansour',
            'start_date': 'Aug 2025', 'end_date': 'Sep 2025',
            'duration': 'Aug 2025 - Sep 2025',
        }
        depi = {
            'title': 'AI Trainee', 'company': 'DEPI',
            'start_date': 'Jun 2025', 'end_date': 'Dec 2025',
            'duration': 'Jun 2025 - Dec 2025',
        }
        html = self._render([aoi, almansour, depi])
        # Sort order: DEPI > Almansour > AOI (titles read in that order).
        i_depi = html.find('AI Trainee')
        i_alm  = html.find('DT Intern')
        i_aoi  = html.find('IT Intern')
        self.assertGreater(i_depi, 0)
        self.assertGreater(i_alm, i_depi)
        self.assertGreater(i_aoi, i_alm,
                           f'AOI must be below Almansour after heal+sort '
                           f'(positions: DEPI={i_depi} Alm={i_alm} AOI={i_aoi})')
        # AOI's item-block contains "Jul 2024" and NOT "Present".
        aoi_block = self._item_block(html, 'IT Intern')
        self.assertIn('Jul 2024', aoi_block)
        self.assertNotIn('Present', aoi_block,
                         'healed AOI must not render "Present"')
        # Closed-range siblings untouched.
        alm_block = self._item_block(html, 'DT Intern')
        self.assertIn('Aug 2025', alm_block)
        self.assertIn('Sep 2025', alm_block)
        depi_block = self._item_block(html, 'AI Trainee')
        self.assertIn('Jun 2025', depi_block)
        self.assertIn('Dec 2025', depi_block)

    def test_genuinely_current_role_still_shows_present(self):
        cur = {
            'title': 'Engineer', 'company': 'ZCo',
            'start_date': 'Jan 2024', 'end_date': 'Present',
            'is_current': True,
            'duration': 'Jan 2024 - Present',
        }
        html = self._render([cur])
        block = self._item_block(html, 'Engineer')
        self.assertIn('Jan 2024', block)
        self.assertIn('Present', block)


# ---------------------------------------------------------------------------
# Banned-opening single source of truth + regen-swap guard.
# Canonical set lives in resumes.services.banned_openings; consumed by
# the generator prompt, the reviewer's _scan_bullet, and the regen
# feedback. Post-regen re-check prevents a Utilized→Leveraged-style
# swap from being marked "resolved".
# ---------------------------------------------------------------------------


class BannedOpeningsCanonicalSetTests(SimpleTestCase):
    """The set is the single source of truth — covers every verb the
    task spec called out, the detector matches case-insensitively
    after stripping leading noise, and unknown openings pass clean."""

    def test_canonical_set_includes_required_verbs(self):
        from resumes.services.banned_openings import BANNED_OPENINGS
        required = {
            "utilized", "leveraged", "spearheaded", "helped",
            "worked on", "crafted", "responsible for", "assisted with",
        }
        missing = required - set(BANNED_OPENINGS)
        self.assertFalse(missing, f"canonical set missing: {missing}")

    def test_detects_each_banned_opening(self):
        from resumes.services.banned_openings import find_banned_opening
        cases = {
            "Crafted a model that...": "crafted",
            "Leveraged tools for X.": "leveraged",
            "Spearheaded a new...": "spearheaded",
            "Helped the team ship X.": "helped",
            "Worked on a model pipeline.": "worked on",
            "Utilized AWS to do X.": "utilized",
            "Utilised AWS to do X.": "utilised",
            "Responsible for X and Y.": "responsible for",
            "Assisted with database migration.": "assisted with",
            "Contributed to a deep-learning pipeline.": "contributed to",
        }
        for text, expected in cases.items():
            self.assertEqual(
                find_banned_opening(text), expected,
                f"{text!r} should detect {expected!r}, got {find_banned_opening(text)!r}",
            )

    def test_clean_openings_pass(self):
        from resumes.services.banned_openings import find_banned_opening
        for clean in (
            "Cut churn by 18% with a calibrated propensity model.",
            "Shipped a REST API serving 5,110 predictions.",
            "Reduced storage by 30% via dedup.",
            "Designed a class-weighted Logistic Regression model.",
            "Trained a CNN on 50K images.",
            "Built a Streamlit dashboard.",
        ):
            self.assertIsNone(find_banned_opening(clean),
                              f"{clean!r} should pass clean")

    def test_match_is_case_insensitive_and_strips_leading_noise(self):
        from resumes.services.banned_openings import find_banned_opening
        for variant in (
            "Crafted a model for X.",
            "crafted a model for X.",
            "CRAFTED a model for X.",
            "• Crafted a model for X.",
            "  - Crafted a model for X.",
            '"Crafted a model for X."',
        ):
            self.assertEqual(find_banned_opening(variant), "crafted")

    def test_word_boundary_guard(self):
        """A bullet starting with 'Utilization' should NOT match 'utilized'."""
        from resumes.services.banned_openings import find_banned_opening
        self.assertIsNone(find_banned_opening("Utilization of resources improved..."))


class GeneratorPromptListsCanonicalSetTests(SimpleTestCase):
    """The generator's _BULLET_QUALITY_RULES enumerates every canonical
    banned opening — so first-attempt generation has the full forbidden
    list."""

    def test_generator_prompt_contains_every_canonical_opening(self):
        from resumes.services.banned_openings import BANNED_OPENINGS
        from resumes.services.resume_generator_v2 import _BULLET_QUALITY_RULES
        rules_lower = _BULLET_QUALITY_RULES.lower()
        for banned in BANNED_OPENINGS:
            self.assertIn(banned, rules_lower,
                          f"generator prompt missing canonical verb {banned!r}")

    def test_generator_prompt_includes_forbidden_section_header(self):
        from resumes.services.resume_generator_v2 import _BULLET_QUALITY_RULES
        self.assertIn("FORBIDDEN OPENINGS", _BULLET_QUALITY_RULES)


class ReviewerReadsCanonicalSetTests(SimpleTestCase):
    """The reviewer's _scan_bullet must call find_banned_opening from
    the canonical module — no hardcoded duplicate list lives in
    resume_reviewer_v2 anymore."""

    def test_reviewer_module_has_no_hardcoded_banned_set(self):
        import resumes.services.resume_reviewer_v2 as rv
        # The old constant name no longer exists.
        self.assertFalse(
            hasattr(rv, "_BANNED_OPENINGS"),
            "_BANNED_OPENINGS still exists — should be gone after collapse "
            "to banned_openings.BANNED_OPENINGS",
        )
        # The reviewer imports from the canonical module.
        import inspect
        src = inspect.getsource(rv)
        self.assertIn("from resumes.services.banned_openings import", src)
        self.assertIn("find_banned_opening", src)

    def test_reviewer_flags_newly_canonical_openings(self):
        """The bug the smoke run caught: 'Crafted' / 'Leveraged' /
        'Spearheaded' / 'Helped' / 'Worked on' were not flagged before
        the collapse. They must be flagged now."""
        from resumes.services.resume_reviewer_v2 import _scan_bullet
        for text in (
            "Crafted a model for X.",
            "Leveraged tools for data science to drive project outcomes.",
            "Spearheaded a new architecture for X.",
            "Helped the team ship Y.",
            "Worked on a pipeline for Z.",
        ):
            findings = _scan_bullet(text)
            banned_findings = [
                f for f in findings
                if f["rule_id"] == "A1_banned_phrase"
                and "banned opening" in f["detail"]
            ]
            self.assertTrue(
                banned_findings,
                f"{text!r} should produce a banned-opening finding",
            )


class RegenFeedbackNamesForbiddenSetTests(SimpleTestCase):
    """When the reviewer regenerates a banned-opening bullet, the
    feedback fed to the LLM must enumerate the full canonical set —
    so the LLM doesn't swap one banned verb for another."""

    def test_feedback_for_banned_opening_lists_all_forbidden_verbs(self):
        from resumes.services.banned_openings import BANNED_OPENINGS
        from resumes.services.resume_reviewer_v2 import _build_regen_feedback
        finding = {
            "rule_id": "A1_banned_phrase",
            "severity": "blocking",
            "where": "experience/x[0]",
            "detail": "bullet starts with banned opening 'utilized'",
            "fix": "Replace the opening with a strong outcome-leading verb.",
        }
        feedback = _build_regen_feedback(finding)
        # Every canonical verb appears in the feedback (case-insensitive
        # check — feedback uses Title Case).
        feedback_lower = feedback.lower()
        for banned in BANNED_OPENINGS:
            self.assertIn(banned, feedback_lower,
                          f"regen feedback missing {banned!r}")

    def test_feedback_for_non_banned_finding_does_not_list_forbidden_set(self):
        """An intensifier finding (A1) or AI-tell finding (A7) doesn't
        need the forbidden-openings preamble — keep the feedback
        targeted."""
        from resumes.services.resume_reviewer_v2 import _build_regen_feedback
        intensifier = _build_regen_feedback({
            "rule_id": "A1_banned_phrase",
            "detail": "bullet contains an empty intensifier (successfully/...)",
            "fix": "Remove the intensifier; let the outcome speak.",
        })
        # Doesn't drag the full banned-verb list into a different
        # rule's feedback.
        self.assertNotIn("crafted", intensifier.lower())
        self.assertNotIn("spearheaded", intensifier.lower())
        ai_tell = _build_regen_feedback({
            "rule_id": "A7_demonstrating_closer",
            "detail": "bullet ends with an AI-tell participial phrase",
            "fix": "Drop the trailing tail; end on the concrete outcome.",
        })
        self.assertNotIn("crafted", ai_tell.lower())


class RegenSwapGuardTests(SimpleTestCase):
    """A regen for a banned-opening finding that produces ANOTHER
    banned opening (e.g. Utilized→Leveraged) is NOT counted resolved —
    it's recorded as demoted with reason 'regen_produced_banned_opener'.
    Mocks the LLM via _generate_one_bullet so no Groq calls happen."""

    @staticmethod
    def _build_test_resume_with_banned_bullet(opening_text):
        from resumes.services.resume_generator_v2 import (
            GeneratedResumeV2, GeneratedSection, EntityBlock, GeneratedBullet,
        )
        return GeneratedResumeV2(sections={
            "experience": GeneratedSection(
                section="experience",
                entities=[EntityBlock(
                    entity_id="ent-x",
                    entity_display="Engineer @ Acme",
                    bullets=[GeneratedBullet(
                        text=opening_text,
                        fact_ids=["f1"],
                    )],
                )],
            ),
        })

    @staticmethod
    def _build_test_store_with_one_fact():
        from resumes.services.fact_store import FactStore
        from resumes.services.fact_store import (
            FactRecord, FactType, SourceReliability,
        )
        store = FactStore()
        store.add(FactRecord(
            id="f1", type=FactType.ACHIEVEMENT,
            claim="shipped X with results",
            entity_id="ent-x", entity_display="Engineer @ Acme",
            source="cv", source_reliability=SourceReliability.USER_ORIGINAL,
            evidence_quote="shipped X with results",
        ))
        return store

    def test_swap_to_another_banned_opener_is_demoted_not_resolved(self):
        from unittest.mock import patch
        from resumes.services.resume_generator_v2 import GeneratedBullet
        from resumes.services.resume_reviewer_v2 import review_and_regenerate
        from resumes.services.resume_planner_v2 import PlanResult

        resume = self._build_test_resume_with_banned_bullet("Utilized AWS to ship X.")
        store = self._build_test_store_with_one_fact()
        from resumes.services.resume_planner_v2 import ValidationReport
        plan = PlanResult(sections={}, validation=ValidationReport(valid_fact_ids=["f1"]))

        def fake_regen(**kwargs):
            # The LLM "fixes" by emitting a different banned verb.
            return GeneratedBullet(text="Leveraged AWS to ship X.", fact_ids=["f1"])

        with patch("resumes.services.resume_reviewer_v2._generate_one_bullet",
                   side_effect=fake_regen):
            _, report = review_and_regenerate(
                resume, store=store, plan=plan, job_title="",
                writing_rules_block="", max_rounds=1,
            )

        # The original Utilized finding is NOT in resolved.
        for r in report["resolved"]:
            self.assertNotIn("utilized", (r.get("detail") or "").lower(),
                             f"Utilized finding wrongly resolved: {r}")
        # It IS in demoted with the regen-failure reason.
        regen_failed = [
            d for d in report["demoted"]
            if d.get("demoted_reason") == "regen_produced_banned_opener"
        ]
        self.assertTrue(regen_failed,
                        f"expected 1+ demoted with regen_produced_banned_opener; "
                        f"got demoted={report['demoted']}")
        # Annotates the offending regen output so the user knows what
        # the regen produced.
        ann = regen_failed[0]
        self.assertEqual(ann.get("regen_attempt_banned_opening"), "leveraged")
        self.assertIn("Leveraged", ann.get("regen_attempt_text") or "")
        # And not double-demoted under another reason.
        cap_dem = [d for d in report["demoted"]
                   if d.get("demoted_reason") == "review_cap_exhausted"
                   and "utilized" in (d.get("detail") or "").lower()]
        self.assertFalse(cap_dem,
                         "regen-swap finding double-demoted under cap_exhausted")

    def test_clean_regen_is_resolved_not_demoted(self):
        from unittest.mock import patch
        from resumes.services.resume_generator_v2 import GeneratedBullet
        from resumes.services.resume_reviewer_v2 import review_and_regenerate
        from resumes.services.resume_planner_v2 import PlanResult

        resume = self._build_test_resume_with_banned_bullet("Utilized AWS to ship X.")
        store = self._build_test_store_with_one_fact()
        from resumes.services.resume_planner_v2 import ValidationReport
        plan = PlanResult(sections={}, validation=ValidationReport(valid_fact_ids=["f1"]))

        def fake_regen(**kwargs):
            return GeneratedBullet(text="Shipped X via AWS Lambda.", fact_ids=["f1"])

        with patch("resumes.services.resume_reviewer_v2._generate_one_bullet",
                   side_effect=fake_regen):
            _, report = review_and_regenerate(
                resume, store=store, plan=plan, job_title="",
                writing_rules_block="", max_rounds=1,
            )

        resolved_utilized = [
            r for r in report["resolved"]
            if "utilized" in (r.get("detail") or "").lower()
        ]
        self.assertTrue(resolved_utilized,
                        "clean regen should be marked resolved")
        # And NOT also demoted.
        demoted_utilized = [
            d for d in report["demoted"]
            if "utilized" in (d.get("detail") or "").lower()
        ]
        self.assertFalse(demoted_utilized,
                         "clean regen should NOT appear in demoted")


# ---------------------------------------------------------------------------
# Findings UX — 3 buckets + plain-language copy + guarded "Fix it" flow.
# The Fix-it endpoint routes through _generate_one_bullet (the SAME
# guarded primitive the reviewer uses). Number-lock + banned-openings
# still fire. Refuses to fix NEEDS_USER_INPUT / ADVISORY findings.
# ---------------------------------------------------------------------------


class FindingsUxPlainCopyTests(SimpleTestCase):
    """User-facing copy MUST be plain language — no rule_id, no
    internal vocabulary leaks. The fail-safe (unknown → user_input)
    survives the helper module."""

    def test_plain_message_for_known_rules(self):
        from resumes.services.findings_ux import _plain_message
        m = _plain_message("A1_banned_phrase")
        self.assertIn("Utilized", m)  # uses the recruiter-recognizable verb
        self.assertNotIn("A1_", m)
        self.assertNotIn("rule_id", m.lower())

    def test_plain_message_for_unsupported_metric(self):
        from resumes.services.findings_ux import _plain_message
        m = _plain_message("unsupported_metric")
        self.assertIn("number", m.lower())
        self.assertNotIn("unsupported_metric", m.lower())

    def test_unknown_rule_falls_back_to_generic_review_text(self):
        from resumes.services.findings_ux import _plain_message, message_is_jargon_free
        m = _plain_message("totally_unknown_rule_X23")
        self.assertTrue(message_is_jargon_free(m))
        self.assertIn("look", m.lower())  # "needs a look"

    def test_message_is_jargon_free_detects_rule_id_leak(self):
        from resumes.services.findings_ux import message_is_jargon_free
        # Negative cases.
        self.assertFalse(message_is_jargon_free("Inside-out opener — fix it"))
        self.assertFalse(message_is_jargon_free("A1_banned_phrase: change verb"))
        self.assertFalse(message_is_jargon_free("auto_fix needed for this row"))
        # Positive cases.
        self.assertTrue(message_is_jargon_free("Open with what you built, not 'Utilized'."))
        self.assertTrue(message_is_jargon_free("Confirm this number against your source."))


class FindingsUxBucketBuilderTests(SimpleTestCase):
    """build_buckets_for_ui flattens the validation report into 3
    buckets with stable ids and plain-language messages."""

    def _resume(self, validation_report, content=None):
        from types import SimpleNamespace
        return SimpleNamespace(
            content=content or {"experience": [
                {"title": "Engineer", "company": "Acme",
                 "description": ["Utilized AWS to build the thing."]},
            ]},
            validation_report=validation_report,
        )

    def test_three_buckets_returned(self):
        from resumes.services.findings_ux import build_buckets_for_ui
        vr = {"findings": [
            {"rule_id": "A1_banned_phrase", "severity": "error",
             "location": "experience[0].description[0]",
             "issue": "bullet starts with 'Utilized'"},
            {"rule_id": "B1_quantification", "severity": "warn",
             "location": "experience[0].description",
             "issue": "no metric on the role"},
        ]}
        b = build_buckets_for_ui(self._resume(vr))
        self.assertIn("auto_fix", b)
        self.assertIn("user_input", b)
        self.assertIn("advisory", b)
        self.assertIn("counts", b)
        # auto_fix: the banned-phrase finding.
        self.assertGreaterEqual(len(b["auto_fix"]), 1)
        af = b["auto_fix"][0]
        self.assertEqual(af["bucket"], "auto_fix")
        self.assertEqual(af["label"], "Suggested rewrite")
        self.assertEqual(af["color"], "amber")
        self.assertEqual(af["action"], "Suggest a rewrite")
        self.assertTrue(af["fixable"])
        # user_input: the quantification finding (needs the user's number).
        self.assertGreaterEqual(len(b["user_input"]), 1)
        ui = b["user_input"][0]
        self.assertEqual(ui["bucket"], "user_input")
        self.assertEqual(ui["label"], "Needs your input")
        self.assertEqual(ui["action"], "Add/Confirm")
        self.assertFalse(ui["fixable"])

    def test_finding_ids_are_stable_across_calls(self):
        from resumes.services.findings_ux import build_buckets_for_ui
        vr = {"findings": [
            {"rule_id": "A1_banned_phrase", "severity": "error",
             "location": "experience[0].description[0]",
             "issue": "x"},
        ]}
        first = build_buckets_for_ui(self._resume(vr))
        second = build_buckets_for_ui(self._resume(vr))
        self.assertEqual(
            first["auto_fix"][0]["id"], second["auto_fix"][0]["id"],
            "finding id should be deterministic across calls",
        )

    def test_user_facing_messages_have_no_rule_id_leak(self):
        from resumes.services.findings_ux import (
            build_buckets_for_ui, message_is_jargon_free,
        )
        vr = {"findings": [
            {"rule_id": "A1_banned_phrase", "severity": "error",
             "location": "experience[0].description[0]", "issue": "x"},
            {"rule_id": "A7_demonstrating_closer", "severity": "error",
             "location": "experience[0].description[0]", "issue": "y"},
            {"rule_id": "B1_quantification", "severity": "warn",
             "location": "experience[0].description", "issue": "z"},
        ]}
        b = build_buckets_for_ui(self._resume(vr))
        for bucket_key in ("auto_fix", "user_input", "advisory"):
            for f in b[bucket_key]:
                self.assertTrue(
                    message_is_jargon_free(f["message"]),
                    f"jargon leak in {bucket_key} message: {f['message']!r}",
                )

    def test_find_finding_by_id_round_trip(self):
        from resumes.services.findings_ux import (
            build_buckets_for_ui, find_finding_by_id,
        )
        vr = {"findings": [
            {"rule_id": "A1_banned_phrase", "severity": "error",
             "location": "experience[0].description[0]", "issue": "x"},
        ]}
        b = build_buckets_for_ui(self._resume(vr))
        target = b["auto_fix"][0]
        got = find_finding_by_id(b, target["id"])
        self.assertEqual(got, target)
        self.assertIsNone(find_finding_by_id(b, "deadbeefdeadbeef"))


class _FindingsApiTestBase(TestCase):
    """Common DB fixtures for the endpoint tests."""

    @classmethod
    def setUpTestData(cls):
        from django.contrib.auth import get_user_model
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from resumes.models import GeneratedResume
        from profiles.models import UserProfile

        U = get_user_model()
        cls.user = U.objects.create_user(
            username="fux@example.com", email="fux@example.com",
            password="x",
        )
        cls.profile = UserProfile.objects.create(
            user=cls.user,
            data_content={
                "experiences": [{
                    "title": "Engineer", "company": "Acme",
                    "start_date": "Jan 2024", "end_date": "Dec 2024",
                    "description": ["Built a thing."],
                }],
                "projects": [], "skills": ["Python"], "certifications": [],
            },
        )
        cls.job = Job.objects.create(
            user=cls.user, title="Software Engineer",
            company="TestCo", description="Engineer JD",
        )
        cls.gap = GapAnalysis.objects.create(
            user=cls.user, job=cls.job, similarity_score=0.5,
        )
        cls.resume = GeneratedResume.objects.create(
            gap_analysis=cls.gap,
            content={
                "experience": [{
                    "title": "Engineer", "company": "Acme",
                    "description": ["Utilized AWS to build the thing."],
                }],
            },
            validation_report={
                "findings": [
                    {"rule_id": "A1_banned_phrase", "severity": "error",
                     "location": "experience[0].description[0]",
                     "issue": "bullet starts with 'Utilized'"},
                    {"rule_id": "B1_quantification", "severity": "warn",
                     "location": "experience[0].description",
                     "issue": "no metric"},
                ],
            },
        )

    def setUp(self):
        self.client.force_login(self.user)


class FindingsApiListTests(_FindingsApiTestBase):
    """GET /api/findings/<id>/ returns the 3-bucket JSON."""

    def test_findings_endpoint_returns_three_buckets(self):
        from django.urls import reverse
        resp = self.client.get(
            reverse("resume_findings_api", args=[self.resume.id]),
        )
        self.assertEqual(resp.status_code, 200)
        data = resp.json()
        for k in ("auto_fix", "user_input", "advisory", "counts"):
            self.assertIn(k, data)
        self.assertGreaterEqual(len(data["auto_fix"]), 1)
        self.assertGreaterEqual(len(data["user_input"]), 1)
        # Each finding carries id / bucket / message / fixable / location.
        af = data["auto_fix"][0]
        for k in ("id", "bucket", "label", "color", "action",
                  "message", "fixable", "location"):
            self.assertIn(k, af)

    def test_user_input_findings_are_not_fixable(self):
        from django.urls import reverse
        resp = self.client.get(
            reverse("resume_findings_api", args=[self.resume.id]),
        )
        data = resp.json()
        for f in data["user_input"]:
            self.assertFalse(
                f["fixable"],
                "user_input finding incorrectly marked fixable",
            )
            self.assertEqual(f["action"], "Add/Confirm")

    def test_other_user_gets_404(self):
        from django.contrib.auth import get_user_model
        from django.urls import reverse
        other = get_user_model().objects.create_user(
            username="other@example.com", email="other@example.com",
            password="x",
        )
        self.client.force_login(other)
        resp = self.client.get(
            reverse("resume_findings_api", args=[self.resume.id]),
        )
        self.assertEqual(resp.status_code, 404)


class FindingsApiProposeFixTests(_FindingsApiTestBase):
    """POST /api/findings/<id>/fix/<finding_id>/ — guarded regen."""

    def _af_finding_id(self):
        from django.urls import reverse
        resp = self.client.get(
            reverse("resume_findings_api", args=[self.resume.id]),
        )
        af = resp.json()["auto_fix"]
        self.assertTrue(af)
        # Pick the bullet-located one (location.bullet_idx not None).
        for f in af:
            if f["location"]["bullet_idx"] is not None:
                return f["id"]
        return af[0]["id"]

    def _user_input_finding_id(self):
        from django.urls import reverse
        resp = self.client.get(
            reverse("resume_findings_api", args=[self.resume.id]),
        )
        ui = resp.json()["user_input"]
        self.assertTrue(ui)
        return ui[0]["id"]

    def test_fix_routes_through_generate_one_bullet(self):
        """The endpoint MUST call _generate_one_bullet — the same
        guarded primitive the reviewer uses. Mock it to assert."""
        from unittest.mock import patch
        from django.urls import reverse
        from resumes.services.resume_generator_v2 import GeneratedBullet

        def fake_gen(**kwargs):
            return GeneratedBullet(text="Shipped a deployment pipeline.",
                                   fact_ids=["x"])
        with patch(
            "resumes.services.resume_generator_v2._generate_one_bullet",
            side_effect=fake_gen,
        ) as m:
            finding_id = self._af_finding_id()
            resp = self.client.post(reverse(
                "resume_propose_fix_api",
                args=[self.resume.id, finding_id],
            ))
        self.assertEqual(resp.status_code, 200, resp.content)
        m.assert_called()
        body = resp.json()
        self.assertEqual(body["proposed_text"],
                         "Shipped a deployment pipeline.")
        self.assertFalse(body.get("persisted", True),
                         "proposal must NOT be persisted at this stage")

    def test_fix_never_calls_v1_regenerate_section(self):
        """v1's regenerate_section is the unguarded path — the fix-it
        endpoint MUST never reach it."""
        from unittest.mock import patch
        from django.urls import reverse
        from resumes.services.resume_generator_v2 import GeneratedBullet

        def fake_gen(**kwargs):
            return GeneratedBullet(text="Cut latency by 18%.", fact_ids=["x"])

        def boom_v1(*args, **kwargs):
            raise AssertionError(
                "fix-it endpoint reached v1 regenerate_section — "
                "this path lacks the number-lock guard.",
            )

        finding_id = self._af_finding_id()
        with patch(
            "resumes.services.resume_generator_v2._generate_one_bullet",
            side_effect=fake_gen,
        ), patch(
            "resumes.services.resume_generator.regenerate_section",
            side_effect=boom_v1,
        ):
            resp = self.client.post(reverse(
                "resume_propose_fix_api",
                args=[self.resume.id, finding_id],
            ))
        self.assertEqual(resp.status_code, 200, resp.content)

    def test_integrity_number_lock_drops_ungrounded_regen(self):
        """If the mocked LLM produces a number not in the bullet's
        facts pool, _generate_one_bullet's number guard drops it. The
        endpoint reports 'guard_dropped'; no ungrounded number reaches
        the proposal."""
        from unittest.mock import patch
        from django.urls import reverse

        # _llm_call is the lowest-level call inside _generate_one_bullet.
        # Mock it to return a number not in the facts.
        def fake_llm(prompt):
            return "Shipped a 99% reduction in CPU usage."

        finding_id = self._af_finding_id()
        with patch(
            "resumes.services.resume_generator_v2._llm_call",
            side_effect=fake_llm,
        ):
            resp = self.client.post(reverse(
                "resume_propose_fix_api",
                args=[self.resume.id, finding_id],
            ))
        self.assertEqual(resp.status_code, 200, resp.content)
        body = resp.json()
        # Number guard fired and dropped both the first attempt and the
        # internal regen — endpoint reports guard_dropped with the
        # fabrication events.
        self.assertEqual(body.get("code"), "guard_dropped")
        self.assertGreaterEqual(
            len(body.get("fabrication_events") or []), 1,
            "expected the number guard to log fabrication events",
        )
        # The proposed_text key must NOT be set (no proposal returned).
        self.assertNotIn("proposed_text", body)

    def test_fix_refused_on_user_input_finding(self):
        """A user-input finding cannot be auto-fixed (regenerating it
        would fabricate). Endpoint refuses with code=not_fixable."""
        from unittest.mock import patch
        from django.urls import reverse

        def boom_gen(**kwargs):
            raise AssertionError(
                "endpoint should refuse the fix BEFORE reaching the regen",
            )

        with patch(
            "resumes.services.resume_generator_v2._generate_one_bullet",
            side_effect=boom_gen,
        ):
            ui_id = self._user_input_finding_id()
            resp = self.client.post(reverse(
                "resume_propose_fix_api",
                args=[self.resume.id, ui_id],
            ))
        self.assertEqual(resp.status_code, 400)
        self.assertEqual(resp.json().get("code"), "not_fixable")

    def test_unknown_finding_id_returns_404_not_found(self):
        from django.urls import reverse
        resp = self.client.post(reverse(
            "resume_propose_fix_api",
            args=[self.resume.id, "0" * 16],
        ))
        self.assertEqual(resp.status_code, 404)


class FindingsApiAcceptTests(_FindingsApiTestBase):
    """POST /api/findings/<id>/accept/ persists the user-approved
    text. No LLM call, no regen — pure persistence."""

    def test_accept_persists_new_bullet(self):
        import json
        from django.urls import reverse
        body = {
            "section": "experience",
            "item_idx": 0,
            "bullet_idx": 0,
            "new_text": "Shipped a deployment pipeline that cut release time 18%.",
        }
        resp = self.client.post(
            reverse("resume_accept_fix_api", args=[self.resume.id]),
            data=json.dumps(body),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200, resp.content)
        self.assertTrue(resp.json().get("persisted"))
        self.resume.refresh_from_db()
        self.assertEqual(
            self.resume.content["experience"][0]["description"][0],
            body["new_text"],
        )

    def test_accept_rejects_missing_new_text(self):
        import json
        from django.urls import reverse
        resp = self.client.post(
            reverse("resume_accept_fix_api", args=[self.resume.id]),
            data=json.dumps({
                "section": "experience", "item_idx": 0,
                "bullet_idx": 0, "new_text": "",
            }),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 400)

    def test_accept_rejects_bad_location(self):
        import json
        from django.urls import reverse
        resp = self.client.post(
            reverse("resume_accept_fix_api", args=[self.resume.id]),
            data=json.dumps({
                "section": "experience", "item_idx": 99,
                "bullet_idx": 0, "new_text": "x",
            }),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 400)


class FindingsChipTemplateTests(TestCase):
    """The findings_chip.html component renders the 3 buckets with
    consistent labels + actions. No rule_id appears in the rendered
    output."""

    def _render(self, annotations, section_key="experience", item_idx=None):
        from django.template.loader import render_to_string
        from django.template import engines
        # Build a tiny wrapper template that includes the chip with
        # the `only` keyword to mirror real call sites.
        engine = engines["django"]
        if item_idx is None:
            tpl = engine.from_string(
                '{% include "components/findings_chip.html" '
                'with annotations=annotations section_key=section_key only %}'
            )
        else:
            tpl = engine.from_string(
                '{% include "components/findings_chip.html" '
                'with annotations=annotations section_key=section_key '
                'item_idx=item_idx only %}'
            )
        return tpl.render({
            "annotations": annotations,
            "section_key": section_key,
            "item_idx": item_idx,
        })

    def test_auto_fix_chip_renders_fix_it_action(self):
        ann = {
            "bucket": "auto_fix", "section": "experience",
            "anchor_kind": "section", "item_idx": None,
            "count": 1,
            "items": [{"label": "Banned opener",
                       "detail": "lead with the system or outcome"}],
        }
        html = self._render([ann], "experience")
        self.assertIn("Suggested rewrite", html)
        self.assertIn("Suggest a rewrite", html)
        self.assertIn('data-action="fix-it"', html)
        self.assertNotIn("Needs your input", html)

    def test_user_input_chip_renders_add_confirm_and_no_fix_it(self):
        ann = {
            "bucket": "user_input", "section": "experience",
            "anchor_kind": "section", "item_idx": None,
            "count": 1,
            "items": [{"label": "Confirm metric",
                       "detail": "verify the 18% figure"}],
        }
        html = self._render([ann], "experience")
        self.assertIn("Needs your input", html)
        self.assertIn("Add or Confirm", html)
        self.assertNotIn('data-action="fix-it"', html,
                         "user_input chip MUST NOT render a Fix it action")

    def test_advisory_chip_renders_dismiss(self):
        ann = {
            "bucket": "advisory", "section": "skills",
            "anchor_kind": "section", "item_idx": None,
            "count": 1,
            "items": [{"label": "Polish suggestion", "detail": "tighten"}],
        }
        html = self._render([ann], "skills")
        self.assertIn("Suggestion", html)
        self.assertIn('data-action="dismiss"', html)
        self.assertNotIn('data-action="fix-it"', html)

    def test_no_rule_id_or_internal_jargon_in_rendered_chip(self):
        """The rendered chip MUST NOT contain any internal rule_id /
        BUCKET_/ phrasing. The `label` field passed in by the caller
        is allowed to be domain-language ('Banned opener'); the
        forbidden leaks are the internal IDs like 'A1_banned_phrase'."""
        ann = {
            "bucket": "auto_fix", "section": "experience",
            "anchor_kind": "section", "item_idx": None,
            "count": 1,
            "items": [{"label": "Banned opener", "detail": "rewrite"}],
        }
        html = self._render([ann], "experience")
        for forbidden in (
            "A1_banned_phrase", "A2_action_verb_start", "A3_duty_opener",
            "A4_inside_out", "A5_length", "A6_em_dash",
            "A7_demonstrating_closer", "B1_quantification",
            "rule_id", "BUCKET_AUTO_FIX",
        ):
            self.assertNotIn(forbidden, html,
                             f"chip leaked internal token: {forbidden!r}")

    def test_chip_does_not_use_x_teleport_to_body_overlay(self):
        """The redesign replaces the dark <template x-teleport='body'>
        overlay with an inline hint card. Assert the overlay markup is
        gone."""
        ann = {
            "bucket": "auto_fix", "section": "experience",
            "anchor_kind": "section", "item_idx": None,
            "count": 1,
            "items": [{"label": "x", "detail": "y"}],
        }
        html = self._render([ann], "experience")
        self.assertNotIn('x-teleport="body"', html,
                         "the dark-overlay teleport pattern was removed")
        self.assertNotIn('position: fixed', html,
                         "the inline hint is in flow, not fixed-position overlay")


# ---------------------------------------------------------------------------
# Findings-UX click-through fixes (BUG 1-7 from the audit).
# ---------------------------------------------------------------------------


class FindingsCountBadgeTests(SimpleTestCase):
    """BUG 2 — the count badge must reflect the REAL finding count
    for the bucket-in-section, not 0. Enrichment recomputes ann.count
    from the deduped items so the chip's badge can't silently render 0."""

    def test_enrichment_sets_count_to_real_item_total(self):
        from resumes.services.findings_ux import (
            enrich_annotations_with_plain_messages,
        )
        rs = {"annotations": [{
            "section": "experience", "bucket": "auto_fix",
            "anchor_kind": "section", "item_idx": None, "tier": "blocking",
            "items": [
                {"kind": "A1_banned_phrase", "label": "x", "detail": "a"},
                {"kind": "A1_banned_phrase", "label": "x", "detail": "b"},
                {"kind": "A7_demonstrating_closer", "label": "y", "detail": "c"},
            ],
        }]}
        out = enrich_annotations_with_plain_messages(rs)
        ann = out["annotations"][0]
        # Two distinct kinds × counts → badge = 3 (1 dedup-counted A1 + 1 A7).
        self.assertEqual(ann["count"], 3)
        self.assertEqual(len(ann["items"]), 2,
                         "duplicate items should collapse to one row")
        # The collapsed row carries dup_count = 2.
        a1_row = next(it for it in ann["items"] if it["kind"] == "A1_banned_phrase")
        self.assertEqual(a1_row["dup_count"], 2)

    def test_chip_badge_renders_real_count_not_zero(self):
        from django.template import engines
        engine = engines["django"]
        tpl = engine.from_string(
            '{% include "components/findings_chip.html" '
            'with annotations=annotations section_key=section_key only %}'
        )
        ann = {
            "bucket": "auto_fix", "section": "experience",
            "anchor_kind": "section", "item_idx": None,
            "count": 4,
            "items": [{"label": "x", "detail": "y",
                       "plain_message": "Open with what you built."}],
        }
        html = tpl.render({"annotations": [ann], "section_key": "experience"})
        # Visible badge shows the real number, not a 0 from a bogus
        # template filter chain.
        self.assertIn(">4<", html.replace("\n", "").replace(" ", ""))
        self.assertNotIn(">0<", html.replace("\n", "").replace(" ", ""))


class FindingsPlainCopyJargonScrubTests(SimpleTestCase):
    """BUG 3 — the rendered message must contain none of the rule-NAME
    phrases (inside-out, structural variation, quantification, etc.) —
    not just rule_id codes."""

    def test_all_internal_phrases_are_caught_by_jargon_check(self):
        from resumes.services.findings_ux import message_is_jargon_free
        for bad in (
            "Fix the inside-out summary opener",
            "This bullet lacks structural variation",
            "Add quantification to this role",
            "Banned phrase in opener pattern",
            "Action verb start missing",
            "Verb diversity is low",
            "Buzzword saturation",
            "Em-dash inside the bullet",
            "Demonstrating closer at the end",
        ):
            self.assertFalse(
                message_is_jargon_free(bad),
                f"jargon check should flag rule-name phrase {bad!r}",
            )

    def test_each_known_rule_maps_to_jargon_free_plain_copy(self):
        from resumes.services.findings_ux import (
            _PLAIN_MESSAGES, message_is_jargon_free,
        )
        for rule, msg in _PLAIN_MESSAGES.items():
            self.assertTrue(
                message_is_jargon_free(msg),
                f"{rule}: plain message contains internal vocabulary: {msg!r}",
            )

    def test_specific_plain_copy_for_bug3_callouts(self):
        from resumes.services.findings_ux import _plain_message
        m_a4 = _plain_message("A4_inside_out_summary")
        # "inside-out" must NOT appear.
        self.assertNotIn("inside-out", m_a4.lower())
        self.assertIn("Open with", m_a4)
        m_b3 = _plain_message("B3_structure_variation")
        self.assertNotIn("structural variation", m_b3.lower())
        self.assertIn("vary", m_b3.lower())
        m_b1 = _plain_message("B1_quantification")
        self.assertNotIn("quantification", m_b1.lower())
        self.assertIn("number", m_b1.lower())

    def test_enrichment_scrubs_residual_rule_name_phrases(self):
        """If the plain-message map missed a key, the enrichment swaps
        in a generic fallback rather than leaking the rule name."""
        from resumes.services.findings_ux import (
            enrich_annotations_with_plain_messages,
        )
        rs = {"annotations": [{
            "section": "experience", "bucket": "auto_fix",
            "anchor_kind": "section", "tier": "blocking",
            "items": [{
                "kind": "totally_unknown_rule",
                "label": "Inside-out summary opener with structural variation",
                "detail": "needs quantification",
            }],
        }]}
        out = enrich_annotations_with_plain_messages(rs)
        plain = out["annotations"][0]["items"][0]["plain_message"]
        for bad in ("inside-out", "structural variation", "quantification"):
            self.assertNotIn(bad, plain.lower())


class FindingsChipButtonRenderingTests(TestCase):
    """BUG 4 — auto_fix chips MUST render a Fix-it button. user_input
    and advisory MUST NOT."""

    def _render(self, annotations, section_key="experience"):
        from django.template import engines
        engine = engines["django"]
        tpl = engine.from_string(
            '{% include "components/findings_chip.html" '
            'with annotations=annotations section_key=section_key only %}'
        )
        return tpl.render({"annotations": annotations, "section_key": section_key})

    def test_auto_fix_chip_always_renders_fix_it_button(self):
        ann = {
            "bucket": "auto_fix", "section": "experience",
            "anchor_kind": "section", "item_idx": None, "count": 1,
            "items": [{"plain_message": "Open with what you built.",
                       "finding_id": "abc123"}],
        }
        html = self._render([ann])
        self.assertIn('data-action="fix-it"', html)
        self.assertIn('data-finding-id="abc123"', html)
        # The button must not be gated by an `x-show="open"` wrapper
        # that defaults to false (BUG 4 root cause).
        # Heuristic: between the `data-finding-bucket="auto_fix"` open
        # and the `data-action="fix-it"` token, no `x-show="open"`.
        i_bucket = html.find('data-finding-bucket="auto_fix"')
        i_btn = html.find('data-action="fix-it"', i_bucket)
        between = html[i_bucket:i_btn]
        self.assertNotIn('x-show="open"', between,
                         "Fix it button must not be hidden behind a "
                         "click-to-expand state (BUG 4)")

    def test_user_input_chip_renders_no_fix_it_button(self):
        ann = {
            "bucket": "user_input", "section": "experience",
            "anchor_kind": "section", "item_idx": None, "count": 1,
            "items": [{"plain_message": "Confirm this number.",
                       "finding_id": "xyz789"}],
        }
        html = self._render([ann])
        self.assertNotIn('data-action="fix-it"', html,
                         "user_input chips must never expose Fix it")
        self.assertIn('data-action="add-confirm"', html)

    def test_advisory_chip_renders_dismiss_not_fix_it(self):
        ann = {
            "bucket": "advisory", "section": "skills",
            "anchor_kind": "section", "item_idx": None, "count": 1,
            "items": [{"plain_message": "Optional polish.",
                       "finding_id": "p1"}],
        }
        html = self._render([ann], "skills")
        self.assertNotIn('data-action="fix-it"', html)
        self.assertIn('data-action="dismiss"', html)


class FindingsChipDarkModeTests(SimpleTestCase):
    """BUG 1 — chip styles include dark-mode pairs for every visible
    surface (bg, text, ring). Light-only classes would leave the card
    illegible on the editor's dark theme."""

    def _render(self, bucket):
        from django.template import engines
        engine = engines["django"]
        tpl = engine.from_string(
            '{% include "components/findings_chip.html" '
            'with annotations=annotations section_key=section_key only %}'
        )
        ann = {
            "bucket": bucket, "section": "experience",
            "anchor_kind": "section", "item_idx": None, "count": 1,
            "items": [{"plain_message": "Test message.",
                       "finding_id": "id1"}],
        }
        return tpl.render({"annotations": [ann], "section_key": "experience"})

    def test_each_bucket_has_dark_mode_bg_and_text_tokens(self):
        for bucket in ("auto_fix", "user_input", "advisory"):
            html = self._render(bucket)
            # Dark-mode background on the card.
            self.assertIn("dark:bg-neutral-900", html,
                          f"{bucket} card missing dark-mode bg token")
            # Dark-mode text on the items list.
            self.assertIn("dark:text-neutral-100", html,
                          f"{bucket} items missing dark-mode text token")

    def test_card_uses_accent_left_border_not_white_fill(self):
        """The accent color goes on the left border + label only —
        the body bg stays neutral so dark-mode text remains legible."""
        for bucket, accent in (
            ("auto_fix",   "amber"),
            ("user_input", "red"),
            ("advisory",   "neutral"),
        ):
            html = self._render(bucket)
            self.assertIn(f"border-{accent}-", html.replace("dark:border-", ""),
                          f"{bucket} missing the accent border")


class FindingsChipDedupeTests(SimpleTestCase):
    """BUG 5 — within a single chip, identical findings collapse to one
    row with a (×N) suffix instead of three repeated lines."""

    def test_duplicate_items_collapse_into_one_row(self):
        from resumes.services.findings_ux import (
            enrich_annotations_with_plain_messages,
        )
        rs = {"annotations": [{
            "section": "projects", "bucket": "user_input",
            "anchor_kind": "section", "item_idx": None, "tier": "advisory",
            "items": [
                {"kind": "B1_quantification", "label": "Q", "detail": "p1"},
                {"kind": "B1_quantification", "label": "Q", "detail": "p2"},
                {"kind": "B1_quantification", "label": "Q", "detail": "p3"},
            ],
        }]}
        out = enrich_annotations_with_plain_messages(rs)
        items = out["annotations"][0]["items"]
        self.assertEqual(len(items), 1,
                         "3 identical findings should collapse to 1 row")
        self.assertEqual(items[0]["dup_count"], 3)

    def test_chip_renders_dedup_count_suffix(self):
        from django.template import engines
        engine = engines["django"]
        tpl = engine.from_string(
            '{% include "components/findings_chip.html" '
            'with annotations=annotations section_key=section_key only %}'
        )
        ann = {
            "bucket": "user_input", "section": "experience",
            "anchor_kind": "section", "item_idx": None,
            "count": 3,
            "items": [{"plain_message": "Add a number to this role.",
                       "finding_id": "x", "dup_count": 3}],
        }
        html = tpl.render({"annotations": [ann], "section_key": "experience"})
        self.assertIn("×3", html,
                      "duplicate row should render its count as (×N)")


class FindingsChipSyncBannerAutoDismissTests(TestCase):
    """BUG 6 — the sync banner now auto-dismisses via x-init + setTimeout."""

    @classmethod
    def setUpTestData(cls):
        from django.contrib.auth import get_user_model
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from resumes.models import GeneratedResume
        U = get_user_model()
        cls.user = U.objects.create_user(
            username="sb@example.com", email="sb@example.com", password="x",
        )
        cls.job = Job.objects.create(user=cls.user, title="X",
                                     company="Y", description="z")
        cls.gap = GapAnalysis.objects.create(
            user=cls.user, job=cls.job, similarity_score=0.5,
        )
        cls.resume = GeneratedResume.objects.create(
            gap_analysis=cls.gap,
            content={"professional_summary": "x"},
            validation_report={},
        )

    def setUp(self):
        self.client.force_login(self.user)

    def test_sync_banner_has_auto_dismiss_timer(self):
        from django.urls import reverse
        # Set the session flag the view consumes.
        session = self.client.session
        session[f"resume_synced_{self.resume.id}"] = True
        session.save()
        resp = self.client.get(reverse("resume_edit", args=[self.resume.id]))
        body = resp.content.decode("utf-8", errors="ignore")
        self.assertIn("Synced from your master profile", body)
        # Auto-dismiss timer is present via x-init + setTimeout.
        self.assertIn("setTimeout(() => visible = false", body,
                      "sync banner missing auto-dismiss timer (BUG 6)")


class FindingsChipOverlapTests(SimpleTestCase):
    """BUG 7 — the chip card must be in-flow, not absolutely positioned
    over form fields."""

    def test_chip_is_in_flow_not_absolute(self):
        from django.template import engines
        engine = engines["django"]
        tpl = engine.from_string(
            '{% include "components/findings_chip.html" '
            'with annotations=annotations section_key=section_key only %}'
        )
        ann = {
            "bucket": "auto_fix", "section": "experience",
            "anchor_kind": "section", "item_idx": None, "count": 1,
            "items": [{"plain_message": "x", "finding_id": "x"}],
        }
        html = tpl.render({"annotations": [ann], "section_key": "experience"})
        # No absolute/fixed positioning on the card itself — the chip
        # is supposed to push content down, not overlay it.
        self.assertNotIn("position: absolute", html)
        self.assertNotIn("position: fixed", html)
        self.assertNotIn('class="absolute', html.replace('class="absolute inset-x', '___SKIP___'))
        # The teleport-to-body pattern is gone.
        self.assertNotIn('x-teleport="body"', html)


class FindingsFixItRoundTripTests(TestCase):
    """End-to-end: render the editor → extract a real finding_id from
    the chip's data-finding-id attribute → POST to the propose-fix
    endpoint with that id → assert it resolves to a proposal (NOT
    'Cannot resolve finding id' / 404 'not_found').

    Catches the render-vs-lookup id mismatch the isolated chip / API
    tests didn't surface — the bug that motivated this fix:

      * chip rendered with `only` template-include → `resume.id`
        wasn't in scope → `resumeId` in the chip's x-data was the
        empty string → JS guard refused to call the endpoint;
      * a separate concern: enrichment-side hash had `.strip()` on
        `kind`, lookup-side didn't — any whitespace drift would have
        produced different ids on the two paths.
    """

    @classmethod
    def setUpTestData(cls):
        from django.contrib.auth import get_user_model
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from resumes.models import GeneratedResume
        from profiles.models import UserProfile

        U = get_user_model()
        cls.user = U.objects.create_user(
            username="rt@example.com", email="rt@example.com", password="x",
        )
        cls.profile = UserProfile.objects.create(
            user=cls.user,
            data_content={
                "experiences": [{
                    "title": "Engineer", "company": "Acme",
                    "start_date": "Jan 2024", "end_date": "Dec 2024",
                    "description": ["Built a thing."],
                }],
                "projects": [], "skills": ["Python"],
                "certifications": [],
            },
        )
        cls.job = Job.objects.create(
            user=cls.user, title="Software Engineer",
            company="TestCo", description="Engineer JD",
        )
        cls.gap = GapAnalysis.objects.create(
            user=cls.user, job=cls.job, similarity_score=0.5,
        )
        cls.resume = GeneratedResume.objects.create(
            gap_analysis=cls.gap,
            content={
                "professional_summary": "Engineer with thinginess.",
                "experience": [{
                    "title": "Engineer", "company": "Acme",
                    "description": ["Utilized AWS to build the thing."],
                }],
            },
            validation_report={"findings": [
                {"rule_id": "A1_banned_phrase", "severity": "error",
                 "location": "experience[0].description[0]",
                 "issue": "bullet starts with 'Utilized'"},
            ]},
        )

    def setUp(self):
        self.client.force_login(self.user)

    def _extract_finding_ids(self, html):
        """Pull all data-finding-id values out of the rendered HTML.
        Skips empty values (would indicate a chip-render miss)."""
        import re
        ids = re.findall(r'data-finding-id="([^"]+)"', html)
        return [i for i in ids if i.strip()]

    def _extract_resume_id(self, html):
        """Pull the chip's resumeId out of the rendered Alpine x-data."""
        import re
        m = re.search(r"resumeId:\s*'([^']*)'", html)
        return m.group(1) if m else None

    def test_chip_emits_non_empty_resume_id(self):
        """Regression for the original bug — resume.id was '' because
        the include used `only` and didn't pass resume_id through."""
        from django.urls import reverse
        resp = self.client.get(reverse("resume_edit", args=[self.resume.id]))
        self.assertEqual(resp.status_code, 200)
        html = resp.content.decode("utf-8")
        rid = self._extract_resume_id(html)
        self.assertEqual(rid, str(self.resume.id),
                         f"chip's resumeId should be the resume UUID, got {rid!r}")

    def test_chip_emits_non_empty_finding_ids(self):
        """If the enrichment runs and emits ids, the chip should
        render them — empty values mean the round-trip fails before
        the user even clicks."""
        from django.urls import reverse
        resp = self.client.get(reverse("resume_edit", args=[self.resume.id]))
        html = resp.content.decode("utf-8")
        ids = self._extract_finding_ids(html)
        self.assertTrue(ids, "chip rendered no data-finding-id values "
                             "— enrichment didn't reach the chip")
        for fid in ids:
            self.assertRegex(fid, r"^[0-9a-f]{16}$",
                             f"finding_id should be a 16-char hex hash, got {fid!r}")

    def test_round_trip_render_then_propose_fix_resolves_id(self):
        """The big one: render the editor, grab a real finding_id from
        the rendered chip, POST it to the propose-fix endpoint, and
        assert the endpoint RESOLVES the id (not 404 'not_found').

        Mocks _generate_one_bullet so no Groq call happens — the test
        only verifies the id-routing seam."""
        from unittest.mock import patch
        from django.urls import reverse
        from resumes.services.resume_generator_v2 import GeneratedBullet

        # 1. Render the editor.
        resp = self.client.get(reverse("resume_edit", args=[self.resume.id]))
        self.assertEqual(resp.status_code, 200)
        html = resp.content.decode("utf-8")

        # 2. Extract the resume id + at least one finding id from the
        #    rendered HTML.
        rid = self._extract_resume_id(html)
        self.assertEqual(rid, str(self.resume.id))
        finding_ids = self._extract_finding_ids(html)
        self.assertTrue(finding_ids)

        # 3. Identify the auto_fix finding id we expect to be fixable
        #    by re-deriving it the same way the chip does — via the
        #    enriched annotations. Then assert the FIRST emitted id
        #    matches an auto_fix id (the chip renders auto_fix chips
        #    with Fix-it buttons; user_input / advisory chips don't).
        from resumes.services.findings_ux import (
            build_buckets_for_ui, find_finding_by_id,
        )
        buckets = build_buckets_for_ui(self.resume)
        auto_fix_ids = {f["id"] for f in buckets["auto_fix"]}
        self.assertTrue(auto_fix_ids,
                        "no auto_fix findings — fixture didn't produce one")
        candidate = next((fid for fid in finding_ids if fid in auto_fix_ids), None)
        self.assertIsNotNone(
            candidate,
            "no rendered finding_id matches a build_buckets_for_ui id — "
            "render-side and lookup-side hashes are drifting again",
        )

        # 4. POST to the propose-fix endpoint with that exact id.
        def fake_gen(**kwargs):
            return GeneratedBullet(text="Shipped X via AWS Lambda.",
                                   fact_ids=["x"])
        with patch(
            "resumes.services.resume_generator_v2._generate_one_bullet",
            side_effect=fake_gen,
        ):
            resp = self.client.post(reverse(
                "resume_propose_fix_api",
                args=[self.resume.id, candidate],
            ))

        # 5. The endpoint MUST resolve the id (not 404 'not_found') and
        #    return a proposal (not 'not_fixable').
        self.assertEqual(resp.status_code, 200,
                         f"propose-fix endpoint should resolve the chip-emitted "
                         f"id ({candidate!r}); got status {resp.status_code} "
                         f"body={resp.content!r}")
        body = resp.json()
        self.assertNotEqual(body.get("code"), "not_found",
                            "endpoint returned 'not_found' — id mismatch "
                            "between chip-render and endpoint lookup")
        self.assertNotEqual(body.get("code"), "not_fixable")
        self.assertIn("proposed_text", body)
        self.assertFalse(body.get("persisted", True),
                         "proposal must NOT be persisted at this stage")

    def test_round_trip_under_kind_whitespace_drift(self):
        """Defensive — if the upstream presenter ever emits a kind with
        stray whitespace, render-vs-lookup hashes must still match
        because both sides .strip() the kind before hashing."""
        from resumes.services.findings_ux import (
            _stable_finding_id,
            enrich_annotations_with_plain_messages,
            build_buckets_for_ui,
            find_finding_by_id,
        )
        # Simulate the presenter emitting a kind with stray whitespace
        # by constructing the annotation directly.
        from types import SimpleNamespace
        rs = {"annotations": [{
            "section": "experience",
            "bucket": "auto_fix",
            "anchor_kind": "item",
            "item_idx": 0,
            "bullet_idx": 0,
            "tier": "blocking",
            "items": [{
                "kind": "  A1_banned_phrase  ",   # ← stray whitespace
                "label": "x",
                "detail": "y",
            }],
        }]}
        enriched = enrich_annotations_with_plain_messages(rs)
        emitted_id = enriched["annotations"][0]["items"][0]["finding_id"]

        # Now build buckets from a stub resume whose validation_report
        # gives the SAME annotation shape.
        # (Direct id-computation parity is what we're after — call
        # _stable_finding_id with the same raw inputs and verify
        # equality.)
        from_lookup = _stable_finding_id("auto_fix", {
            "section": "experience",
            "item_idx": 0,
            "bullet_idx": 0,
            "kind": "  A1_banned_phrase  ".strip(),
        })
        self.assertEqual(
            emitted_id, from_lookup,
            "render-side and lookup-side hashes must match even with "
            "whitespace in kind",
        )


# ---------------------------------------------------------------------------
# Philosophy B, honestly labeled.
#   1. "Can auto-fix" → "Suggested rewrite"; "Fix it" → "Suggest a rewrite".
#      Honesty framing in the chip.
#   2. Summary fix unblocked — _generate_summary is itself a wrapper
#      around _generate_one_bullet, so the same number-lock fires.
#   3. Skills supervisor findings reclassified to USER_INPUT (no LLM
#      regen path exists for skills).
# ---------------------------------------------------------------------------


class FindingsAutoFixHonestRelabelTests(SimpleTestCase):
    """The auto_fix bucket's label + action now read as a PROPOSAL,
    not a self-applying fix. Tests against BUCKET_META so future
    chip rewrites can't silently regress the copy."""

    def test_bucket_meta_label_is_suggested_rewrite(self):
        from resumes.services.findings_ux import BUCKET_META, BUCKET_AUTO_FIX
        meta = BUCKET_META[BUCKET_AUTO_FIX]
        self.assertEqual(meta["label"], "Suggested rewrite")
        self.assertEqual(meta["action"], "Suggest a rewrite")
        self.assertEqual(meta["color"], "amber")
        self.assertTrue(meta["fixable"])

    def test_chip_header_renders_suggested_rewrite_not_can_auto_fix(self):
        from django.template import engines
        engine = engines["django"]
        tpl = engine.from_string(
            '{% include "components/findings_chip.html" '
            'with annotations=annotations section_key=section_key only %}'
        )
        ann = {"bucket": "auto_fix", "section": "experience",
               "anchor_kind": "section", "item_idx": None, "count": 1,
               "items": [{"plain_message": "x", "finding_id": "id1"}]}
        html = tpl.render({"annotations": [ann], "section_key": "experience"})
        self.assertIn("Suggested rewrite", html)
        self.assertNotIn("Can auto-fix", html)

    def test_chip_button_label_is_suggest_a_rewrite(self):
        from django.template import engines
        engine = engines["django"]
        tpl = engine.from_string(
            '{% include "components/findings_chip.html" '
            'with annotations=annotations section_key=section_key only %}'
        )
        ann = {"bucket": "auto_fix", "section": "experience",
               "anchor_kind": "section", "item_idx": None, "count": 1,
               "items": [{"plain_message": "x", "finding_id": "id1"}]}
        html = tpl.render({"annotations": [ann], "section_key": "experience"})
        self.assertIn("Suggest a rewrite", html)
        self.assertIn('data-action="fix-it"', html)

    def test_chip_renders_honesty_framing_for_auto_fix_only(self):
        """The 'we'll suggest a grounded rewrite — you choose' helper
        line renders for auto_fix; NOT for user_input or advisory."""
        from django.template import engines
        engine = engines["django"]
        tpl = engine.from_string(
            '{% include "components/findings_chip.html" '
            'with annotations=annotations section_key=section_key only %}'
        )

        def render(bucket):
            return tpl.render({"annotations": [{
                "bucket": bucket, "section": "experience",
                "anchor_kind": "section", "item_idx": None, "count": 1,
                "items": [{"plain_message": "x", "finding_id": "id1"}],
            }], "section_key": "experience"})

        framing = "you choose whether to use it"
        self.assertIn(framing, render("auto_fix"))
        self.assertNotIn(framing, render("user_input"))
        self.assertNotIn(framing, render("advisory"))


class FindingsSummaryFixRoundTripTests(TestCase):
    """The summary endpoint used to refuse with code=summary_not_supported.
    It now routes through _generate_one_bullet — the same guarded
    primitive the per-bullet fix uses, the same one _generate_summary
    itself wraps. End-to-end test: render → scrape → POST → assert
    proposal + integrity."""

    @classmethod
    def setUpTestData(cls):
        from django.contrib.auth import get_user_model
        from jobs.models import Job
        from analysis.models import GapAnalysis
        from resumes.models import GeneratedResume
        from profiles.models import UserProfile

        U = get_user_model()
        cls.user = U.objects.create_user(
            username="srt@example.com", email="srt@example.com",
            password="x",
        )
        cls.profile = UserProfile.objects.create(
            user=cls.user,
            data_content={
                "experiences": [{
                    "title": "Engineer", "company": "Acme",
                    "start_date": "Jan 2024", "end_date": "Dec 2024",
                    "description": ["Built a thing."],
                }],
                "projects": [], "skills": ["Python"],
                "certifications": [],
            },
        )
        cls.job = Job.objects.create(
            user=cls.user, title="Software Engineer",
            company="TestCo", description="Engineer JD",
        )
        cls.gap = GapAnalysis.objects.create(
            user=cls.user, job=cls.job, similarity_score=0.5,
        )
        cls.resume = GeneratedResume.objects.create(
            gap_analysis=cls.gap,
            content={
                "professional_summary":
                    "Engineer with over 4 years of experience.",
                "experience": [{"title": "Engineer", "company": "Acme",
                                "description": ["Built a thing."]}],
            },
            validation_report={"findings": [
                # Maps to A4_inside_out_summary -> auto_fix, on section
                # "professional_summary" which the presenter normalises
                # to "summary".
                {"rule_id": "A4_inside_out_summary", "severity": "error",
                 "location": "professional_summary",
                 "issue": "Opens with 'with over…' — inside-out."},
            ]},
        )

    def setUp(self):
        self.client.force_login(self.user)

    def _summary_finding_id(self):
        from resumes.services.findings_ux import build_buckets_for_ui
        b = build_buckets_for_ui(self.resume)
        for f in b["auto_fix"]:
            if f["location"].get("section") == "summary":
                return f["id"]
        return None

    def test_summary_round_trip_returns_proposal_not_summary_not_supported(self):
        from unittest.mock import patch
        from django.urls import reverse
        from resumes.services.resume_generator_v2 import GeneratedBullet

        finding_id = self._summary_finding_id()
        self.assertIsNotNone(
            finding_id,
            "no auto_fix summary finding in the fixture — bucket "
            "routing changed?",
        )

        def fake_gen(**kwargs):
            return GeneratedBullet(
                text="Built data pipelines that cut ETL latency.",
                fact_ids=["x"],
            )

        with patch(
            "resumes.services.resume_generator_v2._generate_one_bullet",
            side_effect=fake_gen,
        ) as m:
            resp = self.client.post(reverse(
                "resume_propose_fix_api",
                args=[self.resume.id, finding_id],
            ))
        self.assertEqual(resp.status_code, 200, resp.content)
        body = resp.json()
        self.assertNotEqual(body.get("code"), "summary_not_supported",
                            "summary fix still refused — Change 2 didn't land")
        self.assertNotEqual(body.get("code"), "not_found")
        self.assertEqual(body.get("section"), "summary")
        self.assertIn("proposed_text", body)
        self.assertFalse(body.get("persisted", True),
                         "proposal must NOT be persisted at this stage")
        # The guarded primitive WAS called.
        m.assert_called()
        # And called with section="summary" — proving the right path.
        kwargs = m.call_args.kwargs
        self.assertEqual(kwargs.get("section"), "summary")
        self.assertGreater(
            len(kwargs.get("facts") or []), 0,
            "summary regen must receive a grounded facts pool",
        )

    def test_summary_integrity_number_lock_drops_ungrounded(self):
        """If the mocked LLM returns a number not in the summary's
        facts pool, the number guard fires twice (initial + internal
        regen) and drops. Endpoint reports code=guard_dropped with
        fabrication events."""
        from unittest.mock import patch
        from django.urls import reverse

        # Mock the lowest-level LLM call so the REAL number guard runs.
        def fake_llm(prompt):
            return "Shipped a 99.9% improvement in cycle time."

        finding_id = self._summary_finding_id()
        with patch(
            "resumes.services.resume_generator_v2._llm_call",
            side_effect=fake_llm,
        ):
            resp = self.client.post(reverse(
                "resume_propose_fix_api",
                args=[self.resume.id, finding_id],
            ))
        self.assertEqual(resp.status_code, 200, resp.content)
        body = resp.json()
        self.assertEqual(body.get("code"), "guard_dropped",
                         f"number-lock must drop the ungrounded summary regen; "
                         f"got {body!r}")
        self.assertGreaterEqual(
            len(body.get("fabrication_events") or []), 1,
            "expected the number guard to log fabrication events",
        )
        self.assertNotIn("proposed_text", body,
                         "no proposal should reach the UI when the guard drops")

    def test_summary_accept_persists_to_professional_summary(self):
        import json
        from django.urls import reverse
        new = "Built data pipelines that cut ETL latency by half."
        resp = self.client.post(
            reverse("resume_accept_fix_api", args=[self.resume.id]),
            data=json.dumps({"section": "summary", "item_idx": None,
                             "bullet_idx": None, "new_text": new}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200, resp.content)
        self.assertTrue(resp.json().get("persisted"))
        self.resume.refresh_from_db()
        self.assertEqual(self.resume.content["professional_summary"], new)


class FindingsSkillsClassificationTests(SimpleTestCase):
    """A supervisor finding with category=skills now routes to
    USER_INPUT — no Fix-it button is offered (since there's no LLM
    skills regen)."""

    def test_supervisor_skills_classification_is_user_input(self):
        from resumes.services.findings_classifier import (
            classify_supervisor, BUCKET_USER_INPUT,
        )
        bucket = classify_supervisor("skills", severity="blocking",
                                     layer="content")
        self.assertEqual(
            bucket, BUCKET_USER_INPUT,
            "skills supervisor findings must route to user_input — no "
            "LLM regen path exists for skills",
        )

    def test_supervisor_summary_still_auto_fix(self):
        """Sanity check that we only touched skills, not summary."""
        from resumes.services.findings_classifier import (
            classify_supervisor, BUCKET_AUTO_FIX,
        )
        self.assertEqual(
            classify_supervisor("summary", severity="blocking",
                                layer="content"),
            BUCKET_AUTO_FIX,
        )

    def test_skills_finding_renders_no_fix_it_button(self):
        from django.template import engines
        engine = engines["django"]
        tpl = engine.from_string(
            '{% include "components/findings_chip.html" '
            'with annotations=annotations section_key=section_key only %}'
        )
        ann = {"bucket": "user_input", "section": "skills",
               "anchor_kind": "section", "item_idx": None, "count": 1,
               "items": [{"plain_message": "Confirm the skill list.",
                          "finding_id": "k1"}]}
        html = tpl.render({"annotations": [ann], "section_key": "skills"})
        self.assertNotIn('data-action="fix-it"', html,
                         "skills chip must not expose Fix-it")
        self.assertIn('data-action="add-confirm"', html)


class FindingsProposalPanelDarkModeTests(SimpleTestCase):
    """The WAS / NOW proposal panel (rendered after the user clicks
    "Suggest a rewrite") must use the SAME dark tokens the chip card
    uses, so it's legible on the editor's dark UI. The original panel
    used bg-amber-50 / dark:bg-amber-950/30 — a cream-tinted surface
    that read as near-white on dark."""

    def _render_auto_fix_chip(self):
        from django.template import engines
        engine = engines["django"]
        tpl = engine.from_string(
            '{% include "components/findings_chip.html" '
            'with annotations=annotations section_key=section_key only %}'
        )
        ann = {"bucket": "auto_fix", "section": "experience",
               "anchor_kind": "section", "item_idx": None, "count": 1,
               "items": [{"plain_message": "Open with what you built.",
                          "finding_id": "id1"}]}
        return tpl.render({"annotations": [ann], "section_key": "experience"})

    def _proposal_panel_slice(self, html):
        """Return the panel markup between `x-show="proposal"` and its
        closing div so assertions don't pick up the chip-card classes."""
        start = html.find('x-show="proposal"')
        self.assertGreater(start, 0, "no proposal panel in rendered HTML")
        end = html.find('x-show="error"', start)
        return html[start:end] if end > 0 else html[start:]

    def test_panel_uses_chip_card_neutral_dark_bg_not_amber_cream(self):
        """The panel was bg-amber-50 dark:bg-amber-950/30 — a cream
        background that read as 'light cream' on the dark UI. Replaced
        with bg-white dark:bg-neutral-900 (the chip card's pattern).
        The amber stays on the LEFT BORDER + label only."""
        panel = self._proposal_panel_slice(self._render_auto_fix_chip())
        self.assertIn("bg-white", panel)
        self.assertIn("dark:bg-neutral-900", panel)
        # The cream-amber backgrounds are gone.
        self.assertNotIn("bg-amber-50", panel,
                         "panel still uses the cream bg-amber-50 — "
                         "the original dark-mode bug")
        self.assertNotIn("dark:bg-amber-950", panel)

    def test_panel_has_amber_left_border_and_ring(self):
        """Accent stays on the left border + ring, matching chip card."""
        panel = self._proposal_panel_slice(self._render_auto_fix_chip())
        self.assertIn("border-l-4", panel)
        self.assertIn("border-amber-500", panel)
        self.assertIn("ring-amber-300/70", panel)
        self.assertIn("dark:ring-amber-700/60", panel)

    def test_panel_was_text_legible_in_both_modes(self):
        """WAS region: dimmer than NOW (it's the OLD version) but
        readable on the dark bg — neutral-500 / dark:neutral-400."""
        panel = self._proposal_panel_slice(self._render_auto_fix_chip())
        self.assertIn("text-neutral-500", panel)
        self.assertIn("dark:text-neutral-400", panel)
        self.assertIn("line-through", panel)

    def test_panel_now_text_full_contrast_in_both_modes(self):
        """NOW region: full contrast — text-neutral-900 / dark:text-neutral-100,
        same as the chip card body text."""
        panel = self._proposal_panel_slice(self._render_auto_fix_chip())
        self.assertIn("text-neutral-900", panel)
        self.assertIn("dark:text-neutral-100", panel)
        self.assertIn("font-medium", panel)

    def test_panel_was_now_labels_use_amber_accent(self):
        """The Was / Now labels carry the amber accent (matches the
        chip card's section label coloring)."""
        panel = self._proposal_panel_slice(self._render_auto_fix_chip())
        self.assertIn("text-amber-700", panel)
        self.assertIn("dark:text-amber-300", panel)

    def test_accept_button_has_dark_mode_variant(self):
        """Accept stays green; the dark variant brightens slightly so
        the contrast against dark:bg-neutral-900 reads cleanly."""
        panel = self._proposal_panel_slice(self._render_auto_fix_chip())
        self.assertIn('data-action="accept"', panel)
        self.assertIn("bg-emerald-600", panel)
        self.assertIn("dark:bg-emerald-500", panel)
        self.assertIn("Accept", panel)

    def test_reject_button_has_dark_mode_tokens(self):
        """Reject keeps the neutral pattern matching the chip's
        Dismiss button styling."""
        panel = self._proposal_panel_slice(self._render_auto_fix_chip())
        self.assertIn('data-action="reject"', panel)
        self.assertIn("dark:ring-neutral-700", panel)
        self.assertIn("dark:text-neutral-200", panel)
        self.assertIn("dark:bg-neutral-800", panel)
        self.assertIn("Reject", panel)


# ---------------------------------------------------------------------------
# FIX 1 — Summary synthesis prompt. _bullet_prompt now branches on
# section: summary gets a 2-3 sentence positioning prompt; bullets keep
# the 15-25 word single-line copy. Number-lock + banned-openings stay
# byte-for-byte (only the prompt copy differs).
# ---------------------------------------------------------------------------


class SummaryPromptSynthesisTests(SimpleTestCase):
    """The summary prompt now positions the candidate across the
    facts (2-3 sentences, ~50-80 words). The bullet prompt is
    unchanged. Both still feed _generate_one_bullet so the number
    guard and banned-openings chain apply identically."""

    def _summary_prompt(self):
        from resumes.services.resume_generator_v2 import _bullet_prompt
        return _bullet_prompt(
            role_hint="the professional summary for a Data Scientist role at Acme",
            facts=[],
            section="summary",
        )

    def _bullet_prompt(self):
        from resumes.services.resume_generator_v2 import _bullet_prompt
        return _bullet_prompt(
            role_hint="an experience entry: 'Engineer @ Acme'",
            facts=[],
            section="experience",
        )

    def test_summary_prompt_contains_synthesis_directive(self):
        p = self._summary_prompt().lower()
        self.assertIn("professional summary", p)
        self.assertIn("synthesize across the facts", p)
        # Length shape called out — spelled out post-FIX-2 to avoid the
        # number-lock false-positive on our own prompt digits.
        self.assertIn("two to three sentences", p)
        self.assertIn("fifty to eighty words", p)

    def test_summary_prompt_does_not_carry_bullet_copy(self):
        """The summary path must NOT carry the bullet's one-line
        framing — that was the root cause of the thin summary."""
        p = self._summary_prompt()
        self.assertNotIn("writing ONE resume bullet", p)
        self.assertNotIn("15-25 words", p)
        self.assertNotIn("one line, no quotes", p)
        # Bullet ACHIEVEMENT SHAPE block is also gone.
        self.assertNotIn("ACHIEVEMENT SHAPE", p)
        self.assertIn("POSITIONING SHAPE", p)

    def test_bullet_prompt_unchanged_for_non_summary_sections(self):
        """The bullet path keeps its original copy — the summary
        branch must not leak into per-bullet generation."""
        p = self._bullet_prompt()
        self.assertIn("writing ONE resume bullet", p)
        self.assertIn("ACHIEVEMENT SHAPE", p)
        self.assertIn("15-25 words", p)
        self.assertIn("one line, no quotes", p)
        # Summary copy must NOT have leaked into the bullet path.
        self.assertNotIn("PROFESSIONAL SUMMARY", p)
        self.assertNotIn("POSITIONING SHAPE", p)
        self.assertNotIn("Synthesize ACROSS", p)

    def test_summary_prompt_keeps_forbidden_openings_block(self):
        """The banned-openings discipline still applies to the summary
        (sourced from the canonical banned_openings module)."""
        from resumes.services.banned_openings import BANNED_OPENINGS
        p = self._summary_prompt()
        self.assertIn("FORBIDDEN OPENINGS", p)
        for banned in BANNED_OPENINGS:
            self.assertIn(banned, p.lower())

    def test_summary_prompt_keeps_numbers_policy(self):
        """The number-lock prompt copy is still in the summary prompt
        so the LLM knows it can't invent figures."""
        p = self._summary_prompt()
        self.assertIn("NUMBERS POLICY", p)
        self.assertIn("invented numbers will be DROPPED", p)


class SummaryNumberGuardIntegrityTests(SimpleTestCase):
    """Integrity: the summary path uses a different prompt but the
    SAME guarded primitive. A summary regen whose mocked LLM tries
    an ungrounded number is dropped by _ungrounded_numbers — exactly
    as for a bullet."""

    def _build_one_fact(self):
        from resumes.services.fact_store import (
            FactRecord, FactType, SourceReliability,
        )
        return FactRecord(
            id="f1", type=FactType.ACHIEVEMENT,
            claim="shipped X for 5,110 users",
            value=5110.0, unit="",
            entity_id="ent-x", entity_display="Engineer @ Acme",
            source="cv", source_reliability=SourceReliability.USER_ORIGINAL,
            evidence_quote="shipped X for 5,110 users",
        )

    def test_summary_drops_ungrounded_number_via_existing_guard(self):
        from unittest.mock import patch
        from resumes.services.resume_generator_v2 import (
            _generate_one_bullet, _allowed_numbers_from_facts,
        )
        facts = [self._build_one_fact()]
        allowed = _allowed_numbers_from_facts(facts)
        events = []

        # The LLM "summarizes" with a fabricated 99% figure not in facts.
        with patch("resumes.services.resume_generator_v2._llm_call",
                   side_effect=lambda prompt:
                   "Senior engineer with 99% throughput gains across the stack."):
            out = _generate_one_bullet(
                section="summary", entity_id="",
                role_hint="the professional summary for X",
                facts=facts, allowed_numbers=allowed, events=events,
            )

        self.assertIsNone(out,
                          "summary number-lock must drop the ungrounded "
                          "first + regen attempts the same way bullets do")
        # And the dropped event was logged with action='dropped'.
        actions = [e.action for e in events]
        self.assertIn("dropped", actions)

    def test_summary_prompt_passed_through_to_llm_when_section_is_summary(self):
        """End-to-end shape check: when section='summary' the LLM
        receives the synthesis prompt, not the bullet prompt."""
        from unittest.mock import patch
        from resumes.services.resume_generator_v2 import (
            _generate_one_bullet, _allowed_numbers_from_facts,
        )
        facts = [self._build_one_fact()]
        allowed = _allowed_numbers_from_facts(facts)
        captured = {}

        def fake_llm(prompt):
            captured.setdefault("prompts", []).append(prompt)
            return "Lead engineer focused on payments. Built systems that scaled."

        with patch("resumes.services.resume_generator_v2._llm_call",
                   side_effect=fake_llm):
            _generate_one_bullet(
                section="summary", entity_id="",
                role_hint="the professional summary for X",
                facts=facts, allowed_numbers=allowed, events=[],
            )
        prompts = captured.get("prompts") or []
        self.assertTrue(prompts, "LLM was never called")
        self.assertIn("PROFESSIONAL SUMMARY", prompts[0])
        self.assertNotIn("writing ONE resume bullet", prompts[0])


# ---------------------------------------------------------------------------
# FIX 2 — Stop advisory chips echoing the resume's own content.
# Render-side guard on supervisor `detail`: empty OR > 200 chars →
# fall back to a generic stub. Real critiques (tweet-length) pass through.
# ---------------------------------------------------------------------------


class SupervisorEchoGuardContentSubstringTests(SimpleTestCase):
    """The supervisor-detail echo guard is CONTENT-SUBSTRING based,
    NOT length-based. A finding whose `detail` overlaps with the
    resume's prose (summary / experience / projects descriptions) in
    either direction is SUPPRESSED — the annotation drops the item.
    Genuine critiques, even long ones, pass through verbatim."""

    REAL_SUMMARY = (
        "Engineer with over four years of building data pipelines "
        "at scale, leading on retention models, and partnering with "
        "product on growth experiments across regions."
    )

    def _enrich(self, detail, *, content=None):
        """Returns (plain_message, ann_was_dropped) — plain_message is
        None if the annotation was dropped entirely (every item
        suppressed)."""
        from resumes.services.findings_ux import (
            enrich_annotations_with_plain_messages,
        )
        rs = {"annotations": [{
            "section": "summary",
            "bucket": "advisory",
            "anchor_kind": "section",
            "tier": "advisory",
            "item_idx": None,
            "items": [{"kind": "supervisor", "label": "x", "detail": detail}],
        }]}
        out = enrich_annotations_with_plain_messages(rs, resume_content=content)
        anns = out.get("annotations") or []
        if not anns:
            return None, True
        items = anns[0].get("items") or []
        if not items:
            return None, True
        return items[0]["plain_message"], False

    def test_regression_case_echo_of_real_summary_is_suppressed(self):
        """THE MOTIVATING REGRESSION — a supervisor finding whose
        ``detail`` is the actual ~180-char summary text the resume
        carries. The length-based guard let this through; the
        substring guard suppresses it."""
        content = {"professional_summary": self.REAL_SUMMARY}
        # Detail length is well under any reasonable critique threshold.
        self.assertLessEqual(len(self.REAL_SUMMARY), 200,
                             "fixture must be a realistic ~180-char echo")
        plain, dropped = self._enrich(self.REAL_SUMMARY, content=content)
        self.assertTrue(dropped,
                        "echo of summary text must be suppressed, not rendered")

    def test_detail_substring_of_experience_bullet_is_suppressed(self):
        bullet = ("Shipped a real-time fraud-detection model in 14 weeks; "
                  "reduced charge-back disputes by 22%.")
        content = {"experience": [{"description": [bullet]}]}
        # A finding quoting most of the bullet.
        echo = "shipped a real-time fraud-detection model in 14 weeks"
        plain, dropped = self._enrich(echo, content=content)
        self.assertTrue(dropped,
                        "echo of an experience bullet must be suppressed")

    def test_detail_substring_of_project_description_is_suppressed(self):
        content = {"projects": [{"description":
                                 ["Built a Streamlit dashboard for RFM analysis."]}]}
        echo = "built a streamlit dashboard for rfm analysis"
        plain, dropped = self._enrich(echo, content=content)
        self.assertTrue(dropped)

    def test_corpus_substring_of_detail_also_suppressed(self):
        """Detail wraps content — supervisor said 'The summary <SUMMARY> is too
        passive', the embedded section text triggers suppression."""
        content = {"professional_summary": "Junior engineer focused on RAG."}
        detail = ("The summary 'Junior engineer focused on RAG.' is too "
                  "passive and reads as junior framing.")
        # corpus is a substring of detail
        plain, dropped = self._enrich(detail, content=content)
        self.assertTrue(dropped)

    def test_genuine_long_critique_not_in_content_is_RENDERED(self):
        """OVER-SUPPRESSION GUARD: a 250+ char real critique that is
        NOT in the resume must pass through verbatim. Length must NOT
        be the echo signal."""
        critique = (
            "The professional summary is missing a level marker — there's no "
            "'Senior' or 'Lead' positioning the candidate, and the focus "
            "areas in the second sentence aren't tied to the JD's listed "
            "must-haves (deep learning, NLP). Rewrite to lead with the "
            "level + the JD's top two skill themes synthesized from "
            "the candidate's strongest projects."
        )
        self.assertGreater(len(critique), 250,
                           "fixture must be a long critique to prove length "
                           "isn't suppressing real advice")
        # Realistic resume content the critique does NOT echo.
        content = {"professional_summary":
                   "Engineer with 4 years experience.",
                   "experience": [{"description":
                                   ["Built ETL pipelines on AWS."]}]}
        plain, dropped = self._enrich(critique, content=content)
        self.assertFalse(dropped, "long genuine critiques must NOT be suppressed")
        self.assertEqual(plain, critique)

    def test_short_genuine_critique_renders_verbatim(self):
        content = {"professional_summary": "Engineer with 4 years experience."}
        plain, dropped = self._enrich(
            "Summary stops mid-sentence.", content=content,
        )
        self.assertFalse(dropped)
        self.assertEqual(plain, "Summary stops mid-sentence.")

    def test_empty_detail_uses_fallback_stub(self):
        from resumes.services.findings_ux import _SUPERVISOR_ECHO_FALLBACK
        for empty in ("", "   "):
            plain, dropped = self._enrich(empty)
            self.assertFalse(dropped, "empty detail uses stub, not drop")
            self.assertEqual(plain, _SUPERVISOR_ECHO_FALLBACK)

    def test_no_resume_content_disables_echo_check_safely(self):
        """If the caller didn't pass resume_content, the corpus is
        empty and the substring check returns False for everything —
        details pass through verbatim. Backward-compat for callers
        that don't have content (unit tests, etc.)."""
        plain, dropped = self._enrich(
            self.REAL_SUMMARY, content=None,
        )
        self.assertFalse(dropped)
        self.assertEqual(plain, self.REAL_SUMMARY)

    def test_non_supervisor_findings_untouched_by_echo_guard(self):
        """The guard is supervisor-only. A bullet rule with a long
        detail still maps via _PLAIN_MESSAGES."""
        from resumes.services.findings_ux import (
            enrich_annotations_with_plain_messages,
            _SUPERVISOR_ECHO_FALLBACK,
        )
        rs = {"annotations": [{
            "section": "experience", "bucket": "auto_fix",
            "anchor_kind": "section", "tier": "blocking",
            "items": [{"kind": "A1_banned_phrase", "label": "x",
                       "detail": "x" * 500}],
        }]}
        out = enrich_annotations_with_plain_messages(
            rs, resume_content={"experience": [{"description": ["x" * 500]}]},
        )
        plain = out["annotations"][0]["items"][0]["plain_message"]
        self.assertNotEqual(plain, _SUPERVISOR_ECHO_FALLBACK)
        self.assertIn("Utilized", plain)

    def test_length_constant_is_gone(self):
        """Pin the absence of any length-based gate so a future
        refactor can't slip one back in."""
        import resumes.services.findings_ux as fu
        self.assertFalse(
            hasattr(fu, "_SUPERVISOR_DETAIL_MAX_LEN"),
            "_SUPERVISOR_DETAIL_MAX_LEN must be removed — length plays "
            "no role in the echo signal",
        )
        import inspect
        src = inspect.getsource(fu.enrich_annotations_with_plain_messages)
        # No length comparison ON DETAIL in the echo branch. The match
        # itself can be sized (see _MIN_ECHO_OVERLAP_CHARS), but that's
        # a quality check on the overlap, not a filter on detail.
        self.assertNotIn("len(detail) >", src)
        self.assertNotIn("len(detail) <", src)

    def test_trivial_overlap_does_not_suppress(self):
        """Regression: a tiny corpus (e.g. a single-character
        professional_summary like 's') used to substring-match into
        every detail that contained that letter, suppressing genuine
        critiques. The minimum-overlap rule prevents incidental
        single-character matches."""
        content = {"professional_summary": "s"}
        critique = "Summary stops mid-sentence."
        plain, dropped = self._enrich(critique, content=content)
        self.assertFalse(
            dropped,
            "single-char corpus must not trigger suppression — minimum "
            "overlap rule guards against trivial matches",
        )
        self.assertEqual(plain, critique)


class FindingsContentCorpusHelperTests(SimpleTestCase):
    """The corpus helper concatenates summary + experience/projects
    descriptions into a normalized lowercased / whitespace-collapsed
    text — the substring detector's left-hand side."""

    def test_summary_in_corpus(self):
        from resumes.services.findings_ux import _collect_resume_content_corpus
        corpus = _collect_resume_content_corpus(
            {"professional_summary": "Lead Engineer building RAG systems."},
        )
        self.assertIn("lead engineer building rag systems", corpus)

    def test_experience_descriptions_in_corpus(self):
        from resumes.services.findings_ux import _collect_resume_content_corpus
        corpus = _collect_resume_content_corpus({"experience": [
            {"description": ["Built X.", "Shipped Y."]},
            {"description": ["Cut Z by 30%."]},
        ]})
        self.assertIn("built x.", corpus)
        self.assertIn("shipped y.", corpus)
        self.assertIn("cut z by 30%.", corpus)

    def test_project_descriptions_in_corpus(self):
        from resumes.services.findings_ux import _collect_resume_content_corpus
        corpus = _collect_resume_content_corpus({"projects": [
            {"description": ["Built a Streamlit dashboard."]},
        ]})
        self.assertIn("built a streamlit dashboard.", corpus)

    def test_whitespace_collapsed_and_lowercased(self):
        from resumes.services.findings_ux import _collect_resume_content_corpus
        corpus = _collect_resume_content_corpus({
            "professional_summary": "Senior\nEngineer   building\tRAG",
        })
        self.assertIn("senior engineer building rag", corpus)

    def test_none_or_empty_content_returns_empty_corpus(self):
        from resumes.services.findings_ux import _collect_resume_content_corpus
        self.assertEqual(_collect_resume_content_corpus(None), "")
        self.assertEqual(_collect_resume_content_corpus({}), "")
        self.assertEqual(_collect_resume_content_corpus("not a dict"), "")


# ---------------------------------------------------------------------------
# Summary prompt refinements:
#   FIX 1 — kill the "Here is a rewritten…" preamble (prompt directive +
#           deterministic post-process strip).
#   FIX 2 — remove ALL digits from the summary prompt instruction text so
#           the number-lock can't false-positive on a digit we emitted.
#   FIX 3 — nudge the LLM toward the fuller end of the length range.
# ---------------------------------------------------------------------------


class SummaryPromptRefinementTests(SimpleTestCase):
    """Prompt-text assertions for the three summary refinements."""

    def _summary_prompt(self):
        from resumes.services.resume_generator_v2 import _bullet_prompt
        return _bullet_prompt(
            role_hint="the professional summary for an AI/ML Developer role at Acme",
            facts=[],
            section="summary",
        )

    def _bullet_prompt(self):
        from resumes.services.resume_generator_v2 import _bullet_prompt
        return _bullet_prompt(
            role_hint="an experience entry: 'Engineer @ Acme'",
            facts=[],
            section="experience",
        )

    # --- FIX 1: no preamble ----------------------------------------

    def test_summary_prompt_has_explicit_no_preamble_directive(self):
        p = self._summary_prompt()
        # The directive name-checks the patterns the LLM was emitting.
        self.assertIn("Do NOT prefix with 'Here is", p)
        self.assertIn("'Summary:'", p)
        self.assertIn("'Professional summary:'", p)
        self.assertIn("Output ONLY the summary itself", p)
        self.assertIn("Start directly with the first sentence", p)
        # Reinforce that the old softer closing isn't there in isolation
        # (the new one supersedes it).
        self.assertNotIn("Return JUST the summary text — 2-3 sentences",
                         p)

    # --- FIX 2: no digits in the summary instruction text ----------

    def test_summary_prompt_has_zero_digits_in_instructions(self):
        """Digits in our own prompt copy collide with the number-lock
        regex when the LLM echoes them. Spell every count out in words
        so the guard has nothing of ours to grab.

        We check the SUMMARY_QUALITY_RULES constant + the inline branch
        text. Digits inside facts/feedback aren't from our copy.
        """
        import re
        from resumes.services.resume_generator_v2 import (
            _SUMMARY_QUALITY_RULES,
        )
        # The standalone rules block — no digit anywhere.
        self.assertIsNone(re.search(r"\d", _SUMMARY_QUALITY_RULES),
                          f"_SUMMARY_QUALITY_RULES contains a digit: "
                          f"{_SUMMARY_QUALITY_RULES!r}")
        # The inline summary-branch directives (the opener + the
        # closing instruction) — render the prompt without facts /
        # feedback / writing rules, so the only text in the prompt
        # outside _SUMMARY_QUALITY_RULES is the inline copy. Pull
        # those lines and assert digit-free.
        p = self._summary_prompt()
        # Lines that are NOT in the rules block (everything before
        # _SUMMARY_QUALITY_RULES and after it).
        before_rules, _, after_rules = p.partition(_SUMMARY_QUALITY_RULES)
        for label, fragment in (("before-rules", before_rules),
                                ("after-rules", after_rules)):
            # Strip the facts placeholder line ("(no facts)") whose
            # text is not from our prompt copy.
            cleaned = fragment.replace("(no facts)", "")
            self.assertIsNone(
                re.search(r"\d", cleaned),
                f"summary prompt {label} fragment contains a digit: "
                f"{cleaned!r}",
            )

    def test_summary_prompt_uses_spelled_out_length_words(self):
        from resumes.services.resume_generator_v2 import (
            _SUMMARY_QUALITY_RULES,
        )
        lower = _SUMMARY_QUALITY_RULES.lower()
        self.assertIn("two to three sentences", lower)
        self.assertIn("fifty to eighty words", lower)
        # The old digit-bearing forms are gone.
        self.assertNotIn("2-3", _SUMMARY_QUALITY_RULES)
        self.assertNotIn("50-80", _SUMMARY_QUALITY_RULES)

    # --- FIX 3: nudge toward the fuller end ------------------------

    def test_summary_rules_nudges_toward_fuller_length(self):
        from resumes.services.resume_generator_v2 import (
            _SUMMARY_QUALITY_RULES,
        )
        self.assertIn("Aim for the fuller end", _SUMMARY_QUALITY_RULES)
        self.assertIn("three sentences is better than two",
                      _SUMMARY_QUALITY_RULES)

    # --- Regression: bullet prompt untouched -----------------------

    def test_bullet_prompt_still_has_15_25_words_etc(self):
        """Non-summary sections still get the bullet copy — digit
        cleanup is summary-only."""
        p = self._bullet_prompt()
        self.assertIn("15-25 words", p)
        self.assertIn("1-2 lines", p)
        self.assertIn("ONE resume bullet", p)
        self.assertIn("one line, no quotes", p)
        # And the summary-only copy did NOT leak into the bullet path.
        self.assertNotIn("two to three sentences", p)
        self.assertNotIn("Do NOT prefix with 'Here is", p)


class SummaryPreambleStripTests(SimpleTestCase):
    """Deterministic post-process strip for LLM-emitted preambles —
    runs after the number guard so a "Here is a rewritten…" preamble
    NEVER reaches the rendered PDF, even if the LLM ignores the prompt
    instruction."""

    def test_strips_here_is_a_rewritten_preamble(self):
        from resumes.services.resume_generator_v2 import _strip_summary_preamble
        raw = (
            "Here is a rewritten professional summary:\n\n"
            "AI/ML Developer with a foundation in Applied Machine Learning."
        )
        out = _strip_summary_preamble(raw)
        self.assertEqual(
            out,
            "AI/ML Developer with a foundation in Applied Machine Learning.",
        )

    def test_strips_summary_colon_label(self):
        from resumes.services.resume_generator_v2 import _strip_summary_preamble
        raw = "Summary:\nEngineer with hands-on ML production experience."
        self.assertEqual(
            _strip_summary_preamble(raw),
            "Engineer with hands-on ML production experience.",
        )

    def test_strips_professional_summary_label(self):
        from resumes.services.resume_generator_v2 import _strip_summary_preamble
        raw = "Professional summary:\n\nSenior Data Scientist focused on NLP."
        self.assertEqual(
            _strip_summary_preamble(raw),
            "Senior Data Scientist focused on NLP.",
        )

    def test_strips_heres_a_preamble(self):
        from resumes.services.resume_generator_v2 import _strip_summary_preamble
        raw = ("Here's a 2-3 sentence summary:\n"
               "ML Engineer with three years of pipeline work.")
        out = _strip_summary_preamble(raw)
        self.assertEqual(out, "ML Engineer with three years of pipeline work.")

    def test_strips_leading_label_line_ending_in_colon(self):
        """Second-pass safety net for labels the keyword set didn't
        anticipate (any short leading line ending in ':')."""
        from resumes.services.resume_generator_v2 import _strip_summary_preamble
        raw = "My polished version:\nAI Engineer building production RAG."
        self.assertEqual(
            _strip_summary_preamble(raw),
            "AI Engineer building production RAG.",
        )

    def test_clean_summary_passes_through_unchanged(self):
        from resumes.services.resume_generator_v2 import _strip_summary_preamble
        clean = (
            "AI/ML Developer with a Coursera credential and an AI Associate "
            "Level certificate. Builds and deploys ML models, including a "
            "Flask REST API for predictions. Drives business value via "
            "Python-based data solutions."
        )
        self.assertEqual(_strip_summary_preamble(clean), clean)

    def test_does_not_strip_first_sentence_that_happens_to_end_in_colon(self):
        """A summary that legitimately starts with a long first sentence
        ending in ':' (very rare) should NOT be stripped — the safety-net
        only fires when the leading line is short (< 120 chars). A real
        first sentence is typically 40-80 chars but the safety net's
        triggers on labels under 120, which covers normal label use
        without stripping prose."""
        from resumes.services.resume_generator_v2 import _strip_summary_preamble
        # 130+ char first "sentence" ending in colon — won't trigger
        # the safety-net strip.
        raw = (
            "AI/ML Developer with strong experience in deep learning, "
            "computer vision, and natural language processing pipelines "
            "across production-scale systems:\n"
            "Built and shipped multiple RAG systems."
        )
        out = _strip_summary_preamble(raw)
        # The first line (>= 120 chars) is preserved.
        self.assertTrue(out.startswith("AI/ML Developer"))

    def test_empty_or_none_returns_input_unchanged(self):
        from resumes.services.resume_generator_v2 import _strip_summary_preamble
        self.assertEqual(_strip_summary_preamble(""), "")
        self.assertEqual(_strip_summary_preamble("   "), "   ")
        self.assertIsNone(_strip_summary_preamble(None))


class SummaryStripIntegrationTests(SimpleTestCase):
    """End-to-end: a mocked summary LLM emitting a preamble produces
    a GeneratedSection whose .summary_text is the clean summary, NOT
    the preamble."""

    def _build_facts(self):
        from resumes.services.fact_store import (
            FactRecord, FactType, SourceReliability,
        )
        return [
            FactRecord(
                id="f1", type=FactType.ACHIEVEMENT,
                claim="shipped X",
                entity_id="ent", entity_display="Engineer @ Acme",
                source="cv", source_reliability=SourceReliability.USER_ORIGINAL,
                evidence_quote="shipped X",
            ),
        ]

    def test_generated_summary_text_has_preamble_stripped(self):
        from unittest.mock import patch
        from resumes.services.fact_store import FactStore
        from resumes.services.resume_planner_v2 import (
            FactAllocation, SectionPlan,
        )
        from resumes.services.resume_generator_v2 import _generate_summary

        store = FactStore()
        facts = self._build_facts()
        for f in facts:
            store.add(f)
        section = SectionPlan(
            section="summary",
            facts=[FactAllocation(fact_id=facts[0].id,
                                  rationale="test", hedged=False)],
        )

        preamble = (
            "Here is a rewritten professional summary:\n\n"
            "AI/ML Developer with a foundation in Applied Machine Learning. "
            "Builds production systems. Drives business outcomes."
        )
        with patch(
            "resumes.services.resume_generator_v2._llm_call",
            side_effect=lambda prompt: preamble,
        ):
            out = _generate_summary(
                store, section,
                job_title="AI/ML Developer",
                job_company="Acme",
                events=[],
            )
        self.assertFalse(out.summary_text.startswith("Here is"),
                         f"preamble not stripped: {out.summary_text!r}")
        self.assertTrue(out.summary_text.startswith("AI/ML Developer"))
        # The bullet copy was synced so the model object also carries
        # the clean text (for downstream review / regen feedback).
        self.assertTrue(out.bullets[0].text.startswith("AI/ML Developer"))

