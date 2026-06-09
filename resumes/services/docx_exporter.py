"""DOCX exporter for tailored resumes.

Mirrors the same `section_order`, the same fields, and the same content
ordering as the PDF exporter — but produces a single ATS-friendly DOCX
style rather than six visual variants. DOCX is overwhelmingly consumed
by ATS pipelines (LinkedIn Easy Apply, Workday, Greenhouse) where
visual flourish is at best ignored and at worst stripped/mangled, so
trying to replicate `pdf_template_zeyad`'s bold sans-serif character in
DOCX would just add code without helping the user.

Builds the document programmatically via python-docx (already in
requirements.txt for CV-parsing reads). No new dependency.
"""
from __future__ import annotations

import io
import logging
from typing import Iterable, Optional

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

logger = logging.getLogger(__name__)


# Mirror the PDF exporter's whitelist + default order. Importing from
# resumes.views would create a tighter coupling; redefining here keeps
# the docx exporter standalone-runnable from a Django shell or script.
RESUME_SECTION_KEYS = (
    'summary', 'skills', 'experience', 'education',
    'projects', 'certifications', 'awards', 'languages',
)
DEFAULT_SECTION_ORDER = list(RESUME_SECTION_KEYS)


# --- Style helpers ----------------------------------------------------------

# All visual choices live here so a future tweak (e.g. brand color) is
# one-line. Sizes intentionally restrained — ATS parsers do better with
# plain text than with stylized blocks.
# Type scale mirrors the PDF reference's hierarchy (pdf_base.html): a clear
# three-tier name > section-heading > body. The PDF runs name 22 / section 14 /
# body 10.5 (≈2.1x / 1.33x); these reproduce those PROPORTIONS on a 10pt body so
# the DOCX is scannable instead of flat.
NAME_PT = 20            # ≈2.0x body — a clear title (PDF: 22pt on 10.5)
TITLE_PT = 12
CONTACT_PT = 9.5
SECTION_HEADING_PT = 13  # ≈1.3x body — the middle tier (was 11 ≈ 1.1x: too flat)
ITEM_TITLE_PT = 10.5
ITEM_SUB_PT = 10
BODY_PT = 10
BULLET_PT = 10

ACCENT_RGB = RGBColor(0x1E, 0x3A, 0x8A)  # brand-900 — used sparingly
RULE_RGB = '334155'  # slate-700 — section + header underline; the PDF uses a
                     # dark rule (#000). Darker than the old slate-400 so the
                     # section separators actually read as separators.


def _add_hyperlink(paragraph, url: str, text: str, color: Optional[RGBColor] = None) -> None:
    """Add a clickable hyperlink to a paragraph.

    python-docx doesn't expose hyperlinks directly; we have to build the
    underlying XML. This helper keeps the inline call clean.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if color is not None:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), '%02X%02X%02X' % (color[0], color[1], color[2]))
        rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _set_run_font(run, *, size_pt: float = BODY_PT, bold: bool = False,
                  italic: bool = False, color: Optional[RGBColor] = None) -> None:
    run.font.name = 'Calibri'
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _add_section_heading(doc: Document, text: str) -> None:
    """An uppercase, bordered, brand-color heading. Restrained — DOCX
    section headings get aggressively re-styled by ATS pipelines anyway.

    Adds ``w:keepNext`` so Word never paginates the heading away from
    its first child paragraph — fixes the "underlined heading floating
    alone at the bottom of page 1 with the body on page 2" pattern that
    looks like an empty block above the body content.
    """
    p = doc.add_paragraph()
    # Breathing room that mirrors the PDF's rhythm: a clear gap ABOVE each
    # heading separates sections; the gap below sits the rule off the body so
    # the heading reads as a header, not another body line.
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    _set_run_font(run, size_pt=SECTION_HEADING_PT, bold=True, color=ACCENT_RGB)
    # Bottom border on the heading paragraph for the underlined look,
    # plus keepNext so the heading and its first body paragraph stay
    # on the same page.
    pPr = p._p.get_or_add_pPr()
    keep = OxmlElement('w:keepNext')
    pPr.append(keep)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), RULE_RGB)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _bullet(doc: Document, text: str) -> None:
    """Bulleted body paragraph. python-docx's built-in 'List Bullet' style
    is consistent across Word versions and doesn't depend on a numbering
    definition we'd have to ship.

    Empty / whitespace-only text is dropped — an empty bullet renders as
    a stray "•" with nothing after it, which reads as a layout glitch.
    """
    text = (text or '').strip()
    if not text:
        return
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.runs[0] if p.runs else p.add_run()
    if not p.runs[0].text:
        # python-docx creates an empty run with the style; use it.
        p.runs[0].text = text
    else:
        p.runs[0].text = text
    _set_run_font(p.runs[0], size_pt=BULLET_PT)


# PR1 Fix 4 — date-range normaliser. The LLM emits date strings in
# whatever month format it picks (full word, 3-letter abbrev, numeric);
# the docx then renders an inconsistent mix per resume. This helper
# coerces any pair of date-ish strings into "MMM YYYY – MMM YYYY".
_PRESENT_TOKENS = frozenset({'present', 'current', 'now', 'today'})


def _format_date_range(start: str | None, end: str | None,
                       is_current: bool | None = None) -> str:
    """Render a date range as ``"MMM YYYY – MMM YYYY"`` with en-dash.

    Rules:
      - Both parsed and in the same month/year → ``"MMM YYYY"`` once.
      - End in {Present, Current, Now, Today} AND ``is_current=True``
        (the source data explicitly marked the role ongoing)
        → ``"MMM YYYY – Present"``.
      - End empty/None OR a present-family token without
        ``is_current=True`` → render the start alone (honest: no
        invented "Present" for unknown ends; heals legacy LLM
        fabrications on re-render).
      - Year only (no month detected) → ``"YYYY"``.
      - Unparseable field → returned verbatim (never raises).

    Separator is ``" – "`` (space, en-dash, space).
    """
    from dateutil import parser as _dateutil_parser

    def _try_parse(value):
        """Parse a date-ish string. Returns (dt, month_present_bool) or
        None if unparseable. ``month_present_bool`` is False when only a
        year was extracted, so the caller can emit ``"YYYY"`` instead
        of forcing a month into the output."""
        if value is None:
            return None
        s = str(value).strip()
        if not s:
            return None
        try:
            dt = _dateutil_parser.parse(s, default=None, fuzzy=False)
        except (ValueError, TypeError, OverflowError):
            return None
        # dateutil with no default fills missing fields with TODAY's
        # values — that obscures "year-only" inputs. Re-parse with two
        # different defaults; a year-only string will give the same
        # year but different months, so we can detect "no month".
        from datetime import datetime
        d1 = datetime(1, 1, 1)
        d2 = datetime(2, 6, 15)
        try:
            r1 = _dateutil_parser.parse(s, default=d1, fuzzy=False)
            r2 = _dateutil_parser.parse(s, default=d2, fuzzy=False)
        except (ValueError, TypeError, OverflowError):
            return (dt, True)  # fall back, treat as full date
        month_present = (r1.month == r2.month)
        return (dt, month_present)

    def _render_single(value, parsed):
        if parsed is None:
            return str(value).strip() if value is not None else ''
        dt, month_present = parsed
        if month_present:
            return dt.strftime('%b %Y')
        return dt.strftime('%Y')

    start_str = (start or '').strip()
    end_str = (end or '').strip()

    # Honor "Present"-family tokens ONLY when is_current is True (the
    # source data explicitly says ongoing). An empty end_str or a
    # present-family token on a non-current record is treated as
    # unknown end → render the start alone (honest; legacy-heal).
    end_token_is_present_family = end_str.lower() in _PRESENT_TOKENS
    end_is_present = end_token_is_present_family and is_current is True

    parsed_start = _try_parse(start_str)
    parsed_end = (None if end_is_present or end_token_is_present_family or not end_str
                  else _try_parse(end_str))

    rendered_start = _render_single(start_str, parsed_start)
    if end_is_present:
        if parsed_start is None:
            return rendered_start
        return f"{rendered_start} – Present"
    # Empty end OR a present-family token without is_current=True →
    # render the start alone. Never invent "Present" for unknown ends.
    if not end_str or end_token_is_present_family:
        return rendered_start
    rendered_end = _render_single(end_str, parsed_end)

    if not rendered_start and not rendered_end:
        return ''
    if not rendered_start:
        return rendered_end
    if not rendered_end:
        return rendered_start
    # Same-output collapse — fires for "Aug 2025"/"Aug 2025" (both parsed
    # to same month) AND "2024"/"2024" (year-only on both sides). One
    # rendered string out, no range separator.
    if rendered_start == rendered_end:
        return rendered_start
    return f"{rendered_start} – {rendered_end}"


def _ensure_list(value) -> list:
    """description fields can be either a list or a multi-line string —
    normalize so the writer doesn't have to branch."""
    if value is None:
        return []
    if isinstance(value, str):
        return [line.strip() for line in value.split('\n') if line.strip()]
    if isinstance(value, list):
        return [str(v).strip() for v in value if str(v).strip()]
    return []


def _resolve_section_order(content: dict) -> list:
    saved = (content or {}).get('section_order') or []
    valid = [s for s in saved if s in RESUME_SECTION_KEYS]
    return valid + [s for s in DEFAULT_SECTION_ORDER if s not in valid]


# --- Header ---------------------------------------------------------------

def _write_header(doc: Document, profile, content: dict) -> None:
    """Name + professional title + contact line. Renders the full set of
    URLs the parser extracts (LinkedIn, GitHub, Portfolio, Kaggle, Scholar,
    other_urls) — same surface area the PDF templates show, so the DOCX
    doesn't silently drop links the candidate carefully filled in."""
    name = (getattr(profile, 'full_name', None) or 'Your Name').strip()
    title = (content or {}).get('professional_title') or ''

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(2)
    # Mixed case, like the PDF — an UPPERCASED name reads as a heavy block, not
    # a title. The size + the header rule below carry the prominence instead.
    run = p.add_run(name)
    _set_run_font(run, size_pt=NAME_PT, bold=True)

    # Professional title (skip when empty so we don't render a blank line)
    if title:
        pt = doc.add_paragraph()
        pt.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        pt.paragraph_format.space_after = Pt(2)
        run = pt.add_run(title)
        _set_run_font(run, size_pt=TITLE_PT, color=ACCENT_RGB)

    # Contact line — assemble plain-text bits + clickable hyperlinks.
    # We mix plain runs and hyperlink runs so emails/phones stay
    # readable while LinkedIn/GitHub/etc. stay clickable.
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_p.paragraph_format.space_after = Pt(10)
    # Header rule — a thin underline beneath the whole header block, mirroring
    # the PDF's .resume-header border-bottom; it separates the header from the
    # first section so the name/contact read as one unit.
    _hdr_pPr = contact_p._p.get_or_add_pPr()
    _hdr_pBdr = OxmlElement('w:pBdr')
    _hdr_bottom = OxmlElement('w:bottom')
    _hdr_bottom.set(qn('w:val'), 'single')
    _hdr_bottom.set(qn('w:sz'), '6')
    _hdr_bottom.set(qn('w:space'), '4')
    _hdr_bottom.set(qn('w:color'), RULE_RGB)
    _hdr_pBdr.append(_hdr_bottom)
    _hdr_pPr.append(_hdr_pBdr)

    SEP = ' · '
    first = True

    def _add_text(text: str):
        nonlocal first
        if not text:
            return
        if not first:
            sep_run = contact_p.add_run(SEP)
            _set_run_font(sep_run, size_pt=CONTACT_PT)
        run = contact_p.add_run(text)
        _set_run_font(run, size_pt=CONTACT_PT)
        first = False

    def _add_link(url: str, label: str):
        nonlocal first
        if not url:
            return
        if not first:
            sep_run = contact_p.add_run(SEP)
            _set_run_font(sep_run, size_pt=CONTACT_PT)
        _add_hyperlink(contact_p, url, label, color=ACCENT_RGB)
        first = False

    _add_text(getattr(profile, 'email', '') or '')
    _add_text(getattr(profile, 'phone', '') or '')
    _add_text(getattr(profile, 'location', '') or '')
    _add_link(getattr(profile, 'linkedin_url', '') or '', 'LinkedIn')
    _add_link(getattr(profile, 'github_url', '') or '', 'GitHub')
    _add_link(getattr(profile, 'portfolio_url', '') or '', 'Portfolio')
    _add_link(getattr(profile, 'kaggle_url', '') or '', 'Kaggle')
    _add_link(getattr(profile, 'scholar_url', '') or '', 'Scholar')
    for u in (getattr(profile, 'other_urls', None) or []):
        # Strip protocol/www. so the visible label stays compact, same
        # transformation the PDF templates do.
        label = u
        for prefix in ('https://', 'http://', 'www.'):
            if label.startswith(prefix):
                label = label[len(prefix):]
        _add_link(u, label[:24])


# --- Section writers ------------------------------------------------------

def _write_summary(doc: Document, content: dict) -> None:
    text = (content or {}).get('professional_summary') or ''
    objective = (content or {}).get('objective') or ''
    if not text and not objective:
        return
    _add_section_heading(doc, 'Professional Summary')
    if text:
        p = doc.add_paragraph()
        run = p.add_run(text)
        _set_run_font(run, size_pt=BODY_PT)
    if objective:
        p = doc.add_paragraph()
        label = p.add_run('Objective: ')
        _set_run_font(label, size_pt=BODY_PT, bold=True)
        run = p.add_run(objective)
        _set_run_font(run, size_pt=BODY_PT)


def _write_skills(doc: Document, content: dict) -> None:
    """Render the skills line as plain comma-separated text under the
    SKILLS heading. The previous "Core Competencies:" prefix was dead
    weight — the section heading already says "SKILLS"."""
    skills = (content or {}).get('skills') or []
    if not skills:
        return
    _add_section_heading(doc, 'Skills')
    p = doc.add_paragraph()
    body = p.add_run(', '.join(str(s) for s in skills))
    _set_run_font(body, size_pt=BODY_PT)


def _write_experience(doc: Document, content: dict) -> None:
    rows = (content or {}).get('experience') or []
    if not rows:
        return
    _add_section_heading(doc, 'Professional Experience')
    for exp in rows:
        if not isinstance(exp, dict):
            continue
        # Title + duration on the same paragraph: title left-aligned,
        # duration right-aligned via a tab stop.
        head = doc.add_paragraph()
        head.paragraph_format.space_before = Pt(8)
        head.paragraph_format.space_after = Pt(0)
        # Round 1.5.2: keep_with_next pins the job title to its
        # sub-header (company line). The audit caught an orphan widow
        # where "Information Technology Intern" sat alone at the bottom
        # of page 1 and its bullets started page 2.
        head.paragraph_format.keep_with_next = True
        # Set a right tab stop at the page width minus margins (~6.5 in).
        head.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_PARAGRAPH_ALIGNMENT.RIGHT)
        title_run = head.add_run(exp.get('title', '') or '')
        _set_run_font(title_run, size_pt=ITEM_TITLE_PT, bold=True)
        # PR1 Fix 4 — prefer the schema's start_date/end_date pair so we
        # can normalise the format; fall back to the legacy `duration`
        # string when those aren't present (e.g. offline-fallback path).
        date_text = ''
        start_date = exp.get('start_date') or ''
        end_date = exp.get('end_date') or ''
        is_current = exp.get('is_current') is True
        if start_date or end_date:
            date_text = _format_date_range(start_date, end_date, is_current=is_current)
        if not date_text:
            # Legacy fallback: respect a stored duration only when the
            # current is_current rule wouldn't itself heal it. If the
            # stored duration says "Present" but is_current is not True,
            # treat as unknown end (fall back to start_date alone).
            stored_dur = exp.get('duration') or ''
            if stored_dur and not is_current:
                import re as _re
                if _re.search(r'\b(present|current|currently|ongoing|now|today|to\s+date|till\s+now)\b',
                              stored_dur, _re.IGNORECASE):
                    stored_dur = start_date
            date_text = stored_dur
        if date_text:
            head.add_run('\t')
            date_run = head.add_run(date_text)
            _set_run_font(date_run, size_pt=ITEM_TITLE_PT)
        # Company · location · industry on the next line, italic accent
        sub_bits = [b for b in (exp.get('company'), exp.get('location'), exp.get('industry')) if b]
        if sub_bits:
            sub = doc.add_paragraph()
            sub.paragraph_format.space_after = Pt(2)
            # Same pin: sub-header sticks to the first bullet that follows.
            sub.paragraph_format.keep_with_next = True
            run = sub.add_run(' · '.join(sub_bits))
            _set_run_font(run, size_pt=ITEM_SUB_PT, italic=True, color=ACCENT_RGB)
        for bullet in _ensure_list(exp.get('description')):
            _bullet(doc, bullet)


def _write_education(doc: Document, content: dict) -> None:
    rows = (content or {}).get('education') or []
    if not rows:
        return
    _add_section_heading(doc, 'Education')
    for edu in rows:
        if not isinstance(edu, dict):
            continue
        degree = edu.get('degree') or ''
        field = edu.get('field') or ''
        institution = edu.get('institution') or ''
        location = edu.get('location') or ''
        gpa = edu.get('gpa') or ''
        honors = edu.get('honors') or []
        if isinstance(honors, str):
            honors = [line.strip() for line in honors.split('\n') if line.strip()]
        year = edu.get('graduation_year') or edu.get('year') or ''
        if field:
            degree_text = f"{degree} in {field}".strip()
        else:
            degree_text = degree
        head = doc.add_paragraph()
        head.paragraph_format.space_before = Pt(6)
        head.paragraph_format.space_after = Pt(0)
        head.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_PARAGRAPH_ALIGNMENT.RIGHT)
        title_run = head.add_run(degree_text)
        _set_run_font(title_run, size_pt=ITEM_TITLE_PT, bold=True)
        if year:
            head.add_run('\t')
            # PR1 Fix 4 — same normaliser as experience. For a single-year
            # field (the usual education shape), passing year as both
            # start and end triggers the "same month/year collapse" rule
            # and emits "MMM YYYY" or "YYYY" cleanly.
            year_text = _format_date_range(str(year), str(year))
            date_run = head.add_run(year_text or str(year))
            _set_run_font(date_run, size_pt=ITEM_TITLE_PT)
        sub_bits = [b for b in (institution, location, (f"GPA {gpa}" if gpa else '')) if b]
        if sub_bits:
            sub = doc.add_paragraph()
            sub.paragraph_format.space_after = Pt(2)
            run = sub.add_run(' · '.join(sub_bits))
            _set_run_font(run, size_pt=ITEM_SUB_PT, italic=True, color=ACCENT_RGB)
        if honors:
            h = doc.add_paragraph()
            h.paragraph_format.space_after = Pt(2)
            run = h.add_run(' · '.join(honors))
            _set_run_font(run, size_pt=BODY_PT)


def _write_projects(doc: Document, content: dict) -> None:
    rows = (content or {}).get('projects') or []
    if not rows:
        return
    _add_section_heading(doc, 'Projects')
    for proj in rows:
        if not isinstance(proj, dict):
            continue
        name = proj.get('name', '')
        url = proj.get('url', '')
        head = doc.add_paragraph()
        head.paragraph_format.space_before = Pt(8)
        head.paragraph_format.space_after = Pt(2)
        # Round 1.5: explicitly disable keep_with_next on project title
        # paragraphs so Word never sticks them to the following bullets
        # in a way that creates an empty band on the previous page when
        # the project doesn't fit. The audit caught half-page gaps from
        # this exact behaviour.
        head.paragraph_format.keep_with_next = False
        # Render the project NAME as a plain bold run (not wrapped in a
        # hyperlink). The audit tool was reading paragraph.runs and
        # finding bold=None/size=None because the hyperlink runs are
        # nested inside <w:hyperlink>, hidden from paragraph.runs. A
        # plain run is visible to every inspector and renders the same
        # visual weight as the experience-title row.
        name_run = head.add_run(name)
        _set_run_font(name_run, size_pt=ITEM_TITLE_PT, bold=True)
        # Tuck the URL behind a discreet " ↗" link suffix so the project
        # name itself stays a normal bold word recruiters can scan, and
        # the click-through is still one tap away. Same pattern as
        # _write_certifications.
        if url:
            sep = head.add_run(' ')
            _set_run_font(sep, size_pt=ITEM_TITLE_PT, bold=False)
            _add_hyperlink(head, url, '↗', color=ACCENT_RGB)
            # Round 1.5.2: shrink the arrow ~2pt below the title size so
            # it reads as a small affordance icon, not a glyph in the
            # title. Walks the last hyperlink in the paragraph (the one
            # we just added) and overrides w:sz on each run.
            last_links = head._p.findall(qn('w:hyperlink'))
            if last_links:
                arrow_size = str(int((ITEM_TITLE_PT - 2) * 2))
                for r in last_links[-1].iter(qn('w:r')):
                    rPr = r.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        r.insert(0, rPr)
                    for existing in rPr.findall(qn('w:sz')):
                        rPr.remove(existing)
                    sz = OxmlElement('w:sz')
                    sz.set(qn('w:val'), arrow_size)
                    rPr.append(sz)
        # Tech stack on its own italic line under the project name (ATS keywords)
        techs = proj.get('technologies') or []
        if isinstance(techs, str):
            techs = [t.strip() for t in techs.split(',') if t.strip()]
        if techs:
            sub = doc.add_paragraph()
            sub.paragraph_format.space_after = Pt(2)
            # No keep_with_next here either — let Word break the project
            # block naturally between tech stack and bullets.
            sub.paragraph_format.keep_with_next = False
            run = sub.add_run(' · '.join(techs))
            _set_run_font(run, size_pt=ITEM_SUB_PT, italic=True, color=ACCENT_RGB)
        for bullet in _ensure_list(proj.get('description')):
            _bullet(doc, bullet)


def _write_certifications(doc: Document, content: dict) -> None:
    """Render certifications with a softer visual weight than v1.

    v1 rendered the cert NAME as a bold blue underlined hyperlink, which
    turned an 8-cert section into a wall of heavy navy text that drowned
    out the body of the resume. v2 puts the name in plain bold body
    color (regular reading weight), and tucks the URL behind a small
    subdued " · verify" affordance — recruiters can still click through,
    but the section reads as supporting evidence, not a banner.
    """
    rows = (content or {}).get('certifications') or []
    if not rows:
        return
    _add_section_heading(doc, 'Certifications')
    for cert in rows:
        if not isinstance(cert, dict):
            continue
        name = (cert.get('name') or '').strip()
        issuer = (cert.get('issuer') or '').strip()
        date = (cert.get('date') or '').strip()
        url = (cert.get('url') or '').strip()
        duration = (cert.get('duration') or '').strip()
        if not name:
            continue
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(1)
        # Name: plain bold body color — same weight as a project name in
        # the Projects section, no underline, no accent color.
        name_run = p.add_run(name)
        _set_run_font(name_run, size_pt=BODY_PT, bold=True)
        # Issuer / duration / date suffix in regular body weight.
        suffix_bits = []
        if issuer:
            suffix_bits.append(f' - {issuer}')
        if duration:
            suffix_bits.append(f' · {duration}')
        if date:
            suffix_bits.append(f' ({date})')
        if suffix_bits:
            run = p.add_run(''.join(suffix_bits))
            # Explicit bold=False so the run doesn't inherit the
            # preceding name run's bold attribute (Word does inherit
            # in some style chains; the audit caught the whole cert
            # line rendering bold).
            _set_run_font(run, size_pt=BODY_PT, bold=False)
        # Verify link — discreet, smaller, no bold. Color is the accent
        # so a recruiter scanning for "is this real?" can spot it, but
        # the hyperlink no longer dominates the line.
        if url:
            sep = p.add_run(' · ')
            _set_run_font(sep, size_pt=BODY_PT, bold=False)
            _add_hyperlink(p, url, 'verify', color=ACCENT_RGB)
            # Shrink the verify link by a point so it reads as metadata
            # AND explicitly write <w:b w:val="0"/> so the hyperlink runs
            # don't inherit bold from the List Bullet style chain.
            # _add_hyperlink writes only <w:color> and <w:u> — no <w:b>
            # either way — so without this override the whole verify
            # link renders bold in any paragraph style that cascades
            # bold (which is what the multi-user audit caught).
            last_link = p._p.findall(qn('w:hyperlink'))
            if last_link:
                for r in last_link[-1].iter(qn('w:r')):
                    rPr = r.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        r.insert(0, rPr)
                    sz = OxmlElement('w:sz')
                    # w:sz is in half-points → (BODY_PT - 1) * 2.
                    sz.set(qn('w:val'), str(int((BODY_PT - 1) * 2)))
                    # Replace any prior size if present (defensive).
                    for existing in rPr.findall(qn('w:sz')):
                        rPr.remove(existing)
                    rPr.append(sz)
                    # Explicit bold-off marker.
                    for existing in rPr.findall(qn('w:b')):
                        rPr.remove(existing)
                    b = OxmlElement('w:b')
                    b.set(qn('w:val'), '0')
                    rPr.append(b)


def _write_awards(doc: Document, content: dict) -> None:
    """Render the Honors & Awards section. Each entry is rendered as a
    List Bullet so the visual weight matches the certifications
    section. Bold the first segment up to the first em-dash / colon
    (typically the award name) and leave the rest regular."""
    items = (content or {}).get('awards') or []
    if not items:
        return
    _add_section_heading(doc, 'Honors & Awards')
    for item in items:
        if not isinstance(item, str) or not item.strip():
            continue
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(1)
        # Split on the first em-dash / hyphen-spaces / colon to bold
        # the name and leave the issuer / date in regular weight.
        head, sep, tail = '', '', item.strip()
        for delim in (' — ', ' – ', ' - ', ': '):
            if delim in item:
                head, sep, tail = item.partition(delim)
                head = head.strip()
                tail = tail.strip()
                break
        if head and tail:
            name_run = p.add_run(head)
            _set_run_font(name_run, size_pt=BODY_PT, bold=True)
            rest = p.add_run(f"{sep}{tail}")
            _set_run_font(rest, size_pt=BODY_PT, bold=False)
        else:
            run = p.add_run(item.strip())
            _set_run_font(run, size_pt=BODY_PT, bold=True)


def _write_languages(doc: Document, content: dict) -> None:
    # Guard against the LLM misrouting technical skills into the
    # `languages` field (observed when data_content['languages'] is null
    # or empty — the model improvises). sanitize_languages_field drops
    # anything that doesn't look like a spoken human language and logs
    # a WARNING with what was filtered. If nothing real remains, skip
    # the section entirely — printing an empty "LANGUAGES" heading is
    # worse than no heading.
    from profiles.services.profile_sanitizer import sanitize_languages_field

    raw = (content or {}).get('languages') or []
    langs = sanitize_languages_field(raw)
    if not langs:
        return
    _add_section_heading(doc, 'Languages')
    p = doc.add_paragraph()
    run = p.add_run(', '.join(str(l) for l in langs))
    _set_run_font(run, size_pt=BODY_PT)


_SECTION_WRITERS = {
    'summary': _write_summary,
    'skills': _write_skills,
    'experience': _write_experience,
    'education': _write_education,
    'projects': _write_projects,
    'certifications': _write_certifications,
    'awards': _write_awards,
    'languages': _write_languages,
}


# --- Entry point ----------------------------------------------------------

def generate_docx(resume_obj, output: io.BytesIO | str | None = None) -> io.BytesIO:
    """Generate a DOCX from a GeneratedResume row.

    `output` can be:
      - a path string → docx is written there, BytesIO returned anyway
        for symmetry with the PDF exporter's API
      - a BytesIO → docx written into it
      - None → a fresh BytesIO is created and returned

    Honors the user's saved section_order from resume.content with the
    same defensive resolution as the PDF exporter.
    """
    user = resume_obj.gap_analysis.job.user
    profile = user.profile
    content = resume_obj.content or {}
    section_order = _resolve_section_order(content)

    doc = Document()
    # Tighten the default margins so a one-page resume actually fits on
    # one page in Word's default view. Match the PDF templates' visual
    # density.
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    _write_header(doc, profile, content)
    for key in section_order:
        writer = _SECTION_WRITERS.get(key)
        if writer is None:
            continue
        try:
            writer(doc, content)
        except Exception:
            # One section failing must not torpedo the whole DOCX export.
            # Log and continue — the user gets a slightly-shorter doc
            # rather than a 500.
            logger.exception("DOCX export: section '%s' failed; skipping", key)

    # Persist
    if isinstance(output, str):
        doc.save(output)
        # Read back into BytesIO so callers that want bytes don't have to
        # re-open the file.
        with open(output, 'rb') as f:
            return io.BytesIO(f.read())
    buf = output if isinstance(output, io.BytesIO) else io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
