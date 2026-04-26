"""
CV/Resume Data Extractor
Extracts and structures data from PDF and DOCX resume files.
"""

import re
import json
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
from difflib import SequenceMatcher
import sys

USE_PYMUPDF = False
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
    USE_PYMUPDF = True
except ImportError:
    try:
        import pdfplumber
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .llm_engine import get_llm_client, get_llm

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class CVExtractor:
    """Main class for extracting and processing CV/resume data."""
    
    def __init__(self, use_llm: bool = True):
        """
        Initialize the CV extractor.
        """
        self.use_spacy = False
        self.nlp = None
        
        # Initialize Cloud LLM
        self.use_llm = use_llm
        self.llm_client = get_llm_client(task="parser")
        self.llm_available = self.llm_client is not None
        
        # Common section headers with variations
        self.section_patterns = {
            'personal_info': [
                r'^(personal\s+information|contact\s+information|contact\s+details|about)$',
                r'^(name|email|phone|address|linkedin|github|portfolio)$'
            ],
            'summary': [
                r'^(summary|professional\s+summary|objective|profile|executive\s+summary|career\s+objective)$',
                r'^(about\s+me|overview)$'
            ],
            'experience': [
                r'^(work\s+experience|employment|professional\s+experience|experience|work\s+history|career\s+history)$',
                r'^(employment\s+history|professional\s+background)$'
            ],
            'education': [
                r'^(education|academic\s+background|qualifications|academic\s+qualifications)$',
                r'^(degrees?|university|college)$'
            ],
            'skills': [
                r'^(skills?|technical\s+skills?|core\s+skills?|competencies)$',
                r'^(proficiencies|expertise|technologies?)$'
            ],
            'certifications': [
                r'^(certifications?|licenses?|certificates?|professional\s+certifications?)$',
                r'^(licenses?\s+and\s+certifications?)$'
            ],
            'projects': [
                r'^(projects?|personal\s+projects?|key\s+projects?|notable\s+projects?)$',
                r'^(portfolio\s+projects?)$'
            ],
            'awards': [
                r'^(awards?|honors?|achievements?|recognition)$',
                r'^(awards?\s+and\s+achievements?)$'
            ],
            'publications': [
                r'^(publications?|research|papers?)$',
                r'^(published\s+works?)$'
            ],
            'volunteer': [
                r'^(volunteer\s+experience|volunteer\s+work|volunteering)$',
                r'^(community\s+service|volunteer)$'
            ],
            'languages': [
                r'^(languages?|language\s+skills?)$',
                r'^(spoken\s+languages?)$'
            ]
        }
        
        # Compiled regex patterns
        self.email_pattern = re.compile(
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        )
        self.phone_pattern = re.compile(
            r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}|\+\d{10,15}|\d{8,12}'
        )
        self.url_pattern = re.compile(
            r'https?://(?:[-\w.])+(?:[:\d]+)?(?:/(?:[\w/_.])*(?:\?(?:[\w&=%.])*)?(?:#(?:\w*))?)?'
        )
        self.date_pattern = re.compile(
            r'(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\b|\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}\s*[-–—]\s*(?:present|current|now)|\d{4}\s*[-–—]\s*\d{4})',
            re.IGNORECASE
        )
    
    def extract_from_pdf(self, file_path: str) -> str:
        """Stage 1: Robust PDF Ingestion with Noise Reduction"""
        if not PDF_AVAILABLE:
            raise ImportError("PDF extraction libraries not available. Install pdfplumber or PyMuPDF.")
        
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        text_parts = []
        
        try:
            if not USE_PYMUPDF:
                with pdfplumber.open(str(file_path)) as pdf:
                    for i, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            # Append embedded links so the LLM and regexes can find them
                            # Try hyperlinks attribute first, then fall back to annots
                            links_found = []
                            if hasattr(page, 'hyperlinks') and page.hyperlinks:
                                for link in page.hyperlinks:
                                    uri = link.get('uri')
                                    if uri and isinstance(uri, str) and 'http' in uri:
                                        links_found.append(uri)
                            elif hasattr(page, 'annots') and page.annots:
                                for annot in page.annots:
                                    uri = annot.get('uri')
                                    if uri and isinstance(uri, str) and 'http' in uri:
                                        links_found.append(uri)
                            for uri in links_found:
                                page_text += f"\n[Embedded Link: {uri}]"
                            text_parts.append(page_text)
            else:
                doc = fitz.open(str(file_path))
                for i, page_num in enumerate(range(len(doc)), 1):
                    page = doc[page_num]
                    page_text = page.get_text()
                    if page_text:
                        # Append embedded links
                        for link in page.get_links():
                            if link.get('kind') == fitz.LINK_URI:
                                uri = link.get('uri')
                                if uri and isinstance(uri, str) and 'http' in uri:
                                    # Fix malformed URIs like "github:%20https://..." -> "https://..."
                                    if ':%20http' in uri:
                                        uri = uri.split(':%20', 1)[1]
                                    elif ':http' in uri and not uri.startswith('http'):
                                        uri = 'http' + uri.split(':http', 1)[1]
                                    # Try to get the text of the link for better context
                                    rect = link.get('from')
                                    link_text = page.get_textbox(rect).strip() if rect else ""
                                    # Clean up letter-spaced text in link labels
                                    if link_text:
                                        link_text = re.sub(r'[ \t]+', ' ', link_text)
                                        # Collapse letter-spaced fragments (e.g. "K a g g l e" -> "Kaggle")
                                        link_text = re.sub(
                                            r'\b([A-Za-z]) (?=[A-Za-z] |[A-Za-z]\b)',
                                            r'\1',
                                            link_text
                                        )
                                        page_text += f"\n[Embedded Link: '{link_text}' -> {uri}]"
                                    else:
                                        page_text += f"\n[Embedded Link: {uri}]"
                        text_parts.append(page_text)
                doc.close()
            
            raw_text = '\n\n'.join(text_parts)
            
            # Stage 1: Sanitization
            return self._sanitize_text(raw_text)
            
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise
    
    def extract_from_docx(self, file_path: str) -> str:
        """Stage 1: Robust DOCX Ingestion with Table Support"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library not available. Install python-docx.")
        
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        try:
            doc = Document(str(file_path))
            text_parts = []
            
            for i, paragraph in enumerate(doc.paragraphs, 1):
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables (candidates often put skills in tables)
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
            
            # Extract URLs from hyperlinks in the document relationships
            try:
                for rel in doc.part.rels.values():
                    if "hyperlink" in rel.reltype:
                        if rel.target_ref and isinstance(rel.target_ref, str) and 'http' in rel.target_ref:
                            if rel.target_ref not in "\n".join(text_parts):
                                text_parts.append(f"\n[Embedded Link: {rel.target_ref}]")
            except Exception as e:
                logger.warning(f"Could not extract docx hyperlinks: {e}")
            
            raw_text = '\n\n'.join(text_parts)
            
            # Stage 1: Sanitization
            return self._sanitize_text(raw_text)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise
    
    def extract_text(self, file_path: str) -> str:
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return self.extract_from_pdf(str(file_path))
        elif suffix in ['.docx', '.doc']:
            return self.extract_from_docx(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {suffix}. Supported formats: .pdf, .docx")
    
    def _sanitize_text(self, text: str) -> str:
        """
        Stage 1: Noise Reduction
        - Repair letter-spaced words from PDF kerning artifacts
        - Remove header/footer artifacts (e.g., "Page 1 of 3", "confidential")
        - Collapse multiple newlines to save token costs
        """
        # --- Fix letter-spaced words (PDF kerning artifact) ---
        # PDFs with wide letter-spacing produce "B ACH ELOR" or "IN FR OM ATION".
        # Use a known-word list. Preserve original casing: detect whether the
        # match was all-caps, title case, or lowercase, and apply the same casing.
        letter_spaced_words = [
            r'B\s*ACH\s*ELOR',
            r'M\s*AST\s*ER',
            r'IN\s*FR\s*OM\s*ATION',
            r'IN\s*FORM\s*ATION',
            r'TECH\s*N\s*OL\s*O\s*G\s*Y',
            r'COM\s*PUTER',
            r'SCIEN\s*CE',
            r'DIG\s*ITAL',
            r'PION\s*E\s*E\s*R\s*S',
            r'IN\s*ITIATIVE',
            r'COUR\s*SER\s*A',
            r'DATA\s*CAM\s*P',
            r'ENG\s*IN\s*EER',
            r'MAN\s*AGE\s*MENT',
            r'CERT\s*IF\s*IC\s*AT',
            r'PROF\s*ESS\s*ION\s*AL',
            r'EXP\s*ER\s*IENCE',
            r'ED\s*UC\s*ATION',
            r'FOR\s*ENS\s*ICS',
            r'FOR\s*ENS\s*IC',
        ]

        def _case_preserving_collapse(match):
            """Collapse the matched letter-spaced text while preserving its original casing."""
            original = match.group(0)
            collapsed = re.sub(r'\s+', '', original)
            # Detect casing from first two non-space chars of the original
            non_space = [c for c in original if not c.isspace()]
            if not non_space:
                return collapsed
            if all(c.isupper() for c in non_space if c.isalpha()):
                return collapsed.upper()
            if non_space[0].isupper() and all(c.islower() for c in non_space[1:] if c.isalpha()):
                return collapsed[0].upper() + collapsed[1:].lower()
            if all(c.islower() for c in non_space if c.isalpha()):
                return collapsed.lower()
            # Mixed — fall back to title case
            return collapsed[0].upper() + collapsed[1:].lower()

        for pattern in letter_spaced_words:
            text = re.sub(pattern, _case_preserving_collapse, text, flags=re.IGNORECASE)

        # Remove common header/footer patterns
        noise_patterns = [
            r'Page \d+ of \d+',
            r'Confidential',
            r'Resume of .+',
            r'Curriculum Vitae',
            r'\d+/\d+/\d{4}',  # Standalone dates in headers
        ]

        for pattern in noise_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)

        # Collapse multiple newlines (3+ -> 2)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Remove excessive whitespace
        text = re.sub(r'[ \t]+', ' ', text)

        return text.strip()
    
    def fuzzy_match(self, text: str, patterns: List[str], threshold: float = 0.6) -> bool:
        text_lower = text.lower().strip()
        for pattern in patterns:
            if re.search(pattern, text_lower, re.IGNORECASE):
                return True
        for pattern_text in patterns:
            pattern_clean = re.sub(r'[^\w\s]', '', pattern_text).lower().strip()
            if pattern_clean and len(pattern_clean) > 2:
                similarity = SequenceMatcher(None, text_lower, pattern_clean).ratio()
                if similarity >= threshold:
                    return True
        return False
    
    def find_section_headers(self, text: str) -> Dict[str, List[int]]:
        lines = text.split('\n')
        section_positions = {key: [] for key in self.section_patterns.keys()}
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped or len(line_stripped) < 2:
                continue
            
            is_likely_header = (
                len(line_stripped) < 50 and
                (line_stripped.isupper() or 
                 (line_stripped[0].isupper() and line_stripped.replace(' ', '').replace('&', '').isalpha()) or
                 ':' in line_stripped or
                 line_stripped.startswith('##'))
            )
            
            if is_likely_header:
                for section, patterns in self.section_patterns.items():
                    if self.fuzzy_match(line_stripped, patterns):
                        section_positions[section].append(i)
                        break
        
        return section_positions
    
    def extract_personal_info(self, text: str) -> Dict[str, Any]:
        info = {
            'name': None, 'email': None, 'phone': None,
            'address': None, 'linkedin': None, 'github': None, 'portfolio': None
        }
        
        emails = self.email_pattern.findall(text)
        if emails:
            info['email'] = emails[0]
        
        phone_pattern_improved = re.compile(r'\+?\d{1,3}[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}|\+?\d{10,15}')
        phones = phone_pattern_improved.findall(text)
        if phones:
            phone = phones[0].strip()
            if phone.startswith('+'):
                digits = '+' + ''.join(filter(str.isdigit, phone[1:]))
            else:
                digits = ''.join(filter(str.isdigit, phone))
            if len(digits.replace('+', '')) >= 10:
                info['phone'] = phone
        
        # Extract URLs from [Embedded Link: ...] tags (most reliable for PDFs)
        embedded_link_pattern = re.compile(r"\[Embedded Link: '([^']*?)' -> (https?://[^\]]+)\]")
        embedded_links = embedded_link_pattern.findall(text)

        for label, url in embedded_links:
            url_lower = url.lower()
            label_lower = label.lower().strip('| ').strip()
            if 'linkedin.com' in url_lower:
                info['linkedin'] = url
            elif 'github.com' in url_lower:
                info['github'] = url
            elif 'kaggle' in url_lower or 'kaggle' in label_lower:
                info['kaggle'] = url
            elif 'portfolio' in label_lower or 'portfolio' in url_lower:
                info['portfolio'] = url

        # Fallback: scan raw text URLs (for cases without embedded link tags)
        if not info.get('linkedin') or not info.get('github'):
            urls = self.url_pattern.findall(text)
            for url in urls:
                url_lower = url.lower()
                if 'linkedin.com' in url_lower and not info.get('linkedin'):
                    info['linkedin'] = url
                elif 'github.com' in url_lower and not info.get('github'):
                    info['github'] = url
                elif 'kaggle.com' in url_lower and not info.get('kaggle'):
                    info['kaggle'] = url
        
        # Name extraction — conservative. Only accept if the line looks
        # clearly like a person name (not a placeholder, section header,
        # job title, or academic field). Leave null on uncertainty; the
        # LLM validator will extract from raw text.
        NAME_BLOCKLIST = {
            'first last', 'john doe', 'jane doe', 'firstname lastname',
            'computer science', 'software engineering', 'data science',
            'curriculum vitae', 'resume', 'cv', 'contact', 'summary',
            'profile', 'about', 'objective', 'experience', 'education',
            'skills', 'projects', 'certifications', 'work experience',
            'professional experience', 'relevant experience',
            'work history', 'employment history', 'career history',
            'technical skills', 'core skills', 'key skills',
            'cornell university', 'stanford university', 'mit',
        }
        # Section header words — any line containing these in a short
        # 2-4 word context is almost certainly not a person name.
        SECTION_WORDS = {
            'experience', 'education', 'skills', 'contact', 'summary',
            'profile', 'objective', 'projects', 'certifications',
            'references', 'awards', 'publications', 'employment',
            'qualifications', 'achievements', 'history', 'university',
            'college', 'institute', 'school',
        }
        ROLE_WORDS = {
            'developer', 'engineer', 'analyst', 'manager', 'lead',
            'scientist', 'designer', 'architect', 'consultant',
            'specialist', 'intern', 'associate', 'director', 'officer',
            'administrator', 'coordinator', 'assistant', 'technician',
            'researcher', 'programmer',
        }
        lines = text.split('\n')
        for i, line in enumerate(lines[:15]):
            line_s = line.strip()
            if not line_s or len(line_s) > 60:
                continue
            words = line_s.split()
            if not (2 <= len(words) <= 4):
                continue
            lower = line_s.lower()
            if lower in NAME_BLOCKLIST:
                continue
            word_lowers = {w.lower() for w in words}
            if word_lowers & SECTION_WORDS:
                continue
            if word_lowers & ROLE_WORDS:
                continue
            if self.email_pattern.search(line_s) or phone_pattern_improved.search(line_s):
                continue
            # ALL CAPS (e.g., "JOHANN BACH") — each word 2+ chars, no digits
            if line_s.isupper() and all(len(w) >= 2 and w.isalpha() for w in words):
                info['name'] = line_s.title()
                break
            # Title case (e.g., "Karen Santos") — each word starts uppercase, rest lowercase
            if all(w[0].isupper() and w[1:].islower() and w.isalpha() for w in words):
                info['name'] = line_s
                break

        # Location extraction — strict. Only accept "City, State" or
        # "City, Country" patterns, or a short known-city list.
        # Don't guess on single-word capitalized lines (too risky — catches
        # names, headers, companies).
        US_STATES = {
            'AL','AK','AZ','AR','CA','CO','CT','DE','FL','GA','HI','ID','IL',
            'IN','IA','KS','KY','LA','ME','MD','MA','MI','MN','MS','MO','MT',
            'NE','NV','NH','NJ','NM','NY','NC','ND','OH','OK','OR','PA','RI',
            'SC','SD','TN','TX','UT','VT','VA','WA','WV','WI','WY','DC',
        }
        KNOWN_CITIES = {
            'cairo', 'giza', 'alexandria', 'london', 'new york', 'dubai',
            'berlin', 'paris', 'milan', 'rome', 'madrid', 'barcelona',
            'amsterdam', 'toronto', 'vancouver', 'sydney', 'melbourne',
            'singapore', 'tokyo', 'seoul', 'beijing', 'shanghai', 'mumbai',
            'bangalore', 'delhi', 'riyadh', 'doha', 'istanbul',
        }
        COUNTRIES = {
            'usa','us','uk','uae','egypt','germany','france','italy','spain',
            'netherlands','canada','australia','india','china','japan',
            'singapore','saudi arabia','qatar','brazil','mexico', 'ireland',
        }
        for line in lines[:20]:
            line_s = line.strip()
            if not line_s or len(line_s) > 50 or len(line_s) < 4:
                continue
            if self.email_pattern.search(line_s) or self.phone_pattern.search(line_s):
                continue
            lower = line_s.lower()
            # Skip section headers and non-location lines
            if lower in NAME_BLOCKLIST or lower in ('contact', 'contact info', 'contact information'):
                continue
            # Strict pattern: "City, Region/Country" with comma
            if ',' in line_s:
                parts = [p.strip() for p in line_s.split(',')]
                if len(parts) == 2:
                    city, region = parts
                    # Each part must be short (<=3 words), alphabetic-ish
                    if (1 <= len(city.split()) <= 3 and 1 <= len(region.split()) <= 3
                        and all(c.isalpha() or c.isspace() for c in city)
                        and all(c.isalpha() or c.isspace() for c in region)):
                        # Region must be a US state, country, or short region name
                        region_lower = region.lower()
                        region_upper = region.upper()
                        if (region_upper in US_STATES
                            or region_lower in COUNTRIES
                            or region_lower in KNOWN_CITIES
                            or (len(region.split()) <= 2 and city.lower() not in NAME_BLOCKLIST)):
                            info['address'] = line_s
                            break
            # Single known city as a standalone line
            elif lower in KNOWN_CITIES:
                info['address'] = line_s
                break
             
        # LinkedIn handle fallback
        if not info.get('linkedin'):
            # Look for lines with "LinkedIn" and extract handle nearby or next line
            for i, line in enumerate(lines[:20]):
                if 'linkedin' in line.lower():
                    # Check if handle is on the next line (often the case in styled CVs)
                    if i + 1 < len(lines):
                        next_line = lines[i+1].strip()
                        # Handles are usually short, no spaces (or only one), often capitalized
                        if 3 < len(next_line) < 30 and '@' not in next_line:
                             info['linkedin'] = f"https://www.linkedin.com/in/{next_line.replace(' ', '')}"
                             break

        return info
    
    def extract_summary(self, text: str, section_positions: Dict[str, List[int]]) -> Optional[str]:
        lines = text.split('\n')
        if 'summary' in section_positions and section_positions['summary']:
            start_idx = section_positions['summary'][0]
            end_idx = len(lines)
            for i in range(start_idx + 1, len(lines)):
                if lines[i].strip() and any(
                    self.fuzzy_match(lines[i].strip(), self.section_patterns[section])
                    for section in ['experience', 'education', 'skills']
                ):
                    end_idx = i
                    break
            summary_lines = [line.strip() for line in lines[start_idx+1:end_idx] if line.strip()]
            return ' '.join(summary_lines) if summary_lines else None
        return None
    
    def extract_experience(self, text: str, section_positions: Dict[str, List[int]]) -> List[Dict[str, Any]]:
        experiences = []
        lines = text.split('\n')
        if 'experience' not in section_positions or not section_positions['experience']:
            for i, line in enumerate(lines):
                if 'PROFESSIONAL EXPERIENCE' in line.upper() or 'WORK EXPERIENCE' in line.upper():
                    section_positions['experience'] = [i]
                    break
        
        if 'experience' not in section_positions or not section_positions['experience']:
            return experiences
            
        start_idx = section_positions['experience'][0]
        end_idx = len(lines)
        for section in ['education', 'skills', 'certifications', 'projects']:
            if section in section_positions and section_positions[section]:
                for pos in section_positions[section]:
                    if pos > start_idx:
                        end_idx = min(end_idx, pos)
        
        experience_text = '\n'.join(lines[start_idx:end_idx])
        entries = re.split(r'\n(?=(?:[A-Z][A-Z\s&]+(?:INTERN|ENGINEER|DEVELOPER|ANALYST|MANAGER|LEAD|SENIOR)\s+[–—]|\d{4}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}))', experience_text, flags=re.IGNORECASE)
        
        if len(entries) <= 1:
            entries = re.split(r'\n(?=[A-Z][A-Z\s&]{10,}\s+[–—])', experience_text)
            
        for entry in entries[1:]:
            entry_dict = self._parse_experience_entry(entry)
            if entry_dict:
                experiences.append(entry_dict)
        return experiences

    def _parse_experience_entry(self, entry):
        entry = entry.strip()
        if not entry or len(entry) < 20:
            return None
            
        exp = {'company': None, 'position': None, 'location': None, 'start_date': None, 'end_date': None, 'responsibilities': []}
        
        dates = self.date_pattern.findall(entry)
        if dates:
            exp['start_date'] = dates[-1] if len(dates) > 0 else None
            exp['end_date'] = dates[0] if len(dates) > 1 else ('Present' if 'present' in entry.lower() else None)
        
        entry_lines = [line.strip() for line in entry.split('\n') if line.strip()]
        if entry_lines:
            first_line = entry_lines[0]
            parts = re.split(r'\s+[–—]\s+', first_line)
            if len(parts) >= 2:
                exp['position'] = parts[0].strip()
                exp['company'] = parts[1].strip()
            # Flexible separator check
            elif ' at ' in first_line:
                parts = first_line.split(' at ')
                exp['position'] = parts[0].strip()
                exp['company'] = parts[1].strip()
                parts = first_line.split(' at ')
                exp['position'] = parts[0].strip()
                exp['company'] = parts[1].strip()
            elif ',' in first_line:
                # Heuristic: Position, Company
                parts = first_line.split(',')
                if len(parts) == 2:
                     exp['position'] = parts[0].strip()
                     exp['company'] = parts[1].strip()
            elif '-' in first_line: # ASCII hyphen already checked, try shorter hyphen
                parts = first_line.split('-')
                if len(parts) >= 2:
                     exp['position'] = parts[0].strip()
                     exp['company'] = parts[1].strip()
            else:
                # Fallback: check if line ends with a known company entity or looks like one?
                # For DEPI case: "Digital Egypt Pioneers Initiative - DEPI" might be the line
                # If only one string, assume it's the Company if previous line was date? 
                # Or Position? Usually Position is pivotal. 
                # Let's leave as Position for single line, LLM refines it better.
                exp['position'] = first_line
        
        for line in entry_lines[1:]:
            # Cleanup Artifacts
            line_clean = re.sub(r'^(Remote|Hybrid|On-site|Full-time|Part-time)\s*[•\-\*●]*\s*', '', line, flags=re.IGNORECASE)
            
            if line_clean.startswith(('•', '-', '*', '●')):
                content = line_clean[1:].strip()
                if content: exp['responsibilities'].append(content)
            elif len(line_clean) > 30:
                exp['responsibilities'].append(line_clean)
        return exp

    def extract_education(self, text: str, section_positions: Dict[str, List[int]]) -> List[Dict[str, Any]]:
        educations = []
        lines = text.split('\n')
        if 'education' not in section_positions or not section_positions['education']:
            return educations
            
        start_idx = section_positions['education'][0]
        end_idx = len(lines)
        for section in ['experience', 'skills', 'certifications', 'projects']:
            if section in section_positions and section_positions[section]:
                for pos in section_positions[section]:
                    if pos > start_idx:
                        end_idx = min(end_idx, pos)
        
        education_text = '\n'.join(lines[start_idx:end_idx])
        entries = re.split(r'\n(?=\d{4}|\b(?:Bachelor|Master|PhD|Doctorate|B\.?S|M\.?S))', education_text, flags=re.IGNORECASE)
        
        for entry in entries[1:]:
            edu = self._parse_education_entry(entry)
            if edu:
                educations.append(edu)
        return educations

    def _parse_education_entry(self, entry):
        entry = entry.strip()
        if not entry: return None
        edu = {'institution': None, 'degree': None, 'field_of_study': None, 'graduation_date': None}
        
        # Improved Date pattern for "[2024]" or "2024"
        # Use the LAST date found — date ranges list start first, graduation last
        dates = self.date_pattern.findall(entry)
        if dates:
            edu['graduation_date'] = dates[-1]
        else:
            # Fallback for simple Year
            year_match = re.search(r'\b(?:20|19)\d{2}\b', entry)
            if year_match:
                edu['graduation_date'] = year_match.group(0)
        
        degree_match = re.search(r'\b(BACHELOR|MASTER|PHD|B\.?S\.?|B\.?A\.?|M\.?S\.?)\b.*?(?:in|of)?\s+([A-Za-z\s]+)', entry, re.IGNORECASE)
        if degree_match:
            edu['degree'] = degree_match.group(1).title()
            edu['field_of_study'] = degree_match.group(2).strip()
        
        # Simple institution extraction
        entry_lines = entry.split('\n')
        for line in entry_lines:
            if 'University' in line or 'College' in line or 'Institute' in line:
                edu['institution'] = line.strip()
                break
        if not edu['institution'] and entry_lines:
            edu['institution'] = entry_lines[0].strip()
            
        return edu

    def extract_skills(self, text: str, section_positions: Dict[str, List[int]]) -> Dict[str, List[str]]:
        skills = {'technical_skills': [], 'soft_skills': [], 'tools': [], 'frameworks': []}
        lines = text.split('\n')
        if 'skills' not in section_positions or not section_positions['skills']:
            return skills
            
        start_idx = section_positions['skills'][0]
        end_idx = len(lines)
        for section in ['experience', 'education', 'certifications', 'projects']:
            if section in section_positions and section_positions[section]:
                for pos in section_positions[section]:
                    if pos > start_idx:
                        end_idx = min(end_idx, pos)
        
        skills_text = '\n'.join(lines[start_idx:end_idx])
        items = re.split(r'[,•\-\*\|]|\n', skills_text)
        
        expanded_items = []
        for item in items:
            item = item.strip()
            # Split by parentheses: "Data Analysis (Pandas, ...)" -> "Data Analysis", "Pandas", ...
            if '(' in item and ')' in item:
                base = item.split('(')[0].strip()
                if base: expanded_items.append(base)
                content = item[item.find('(')+1:item.find(')')]
                # Split content by comma
                sub_items = [s.strip() for s in content.split(',')]
                expanded_items.extend(sub_items)
            else:
                 expanded_items.append(item)

        for item in expanded_items:
            item = item.strip()
            if not item or len(item) < 2 or 'skill' in item.lower(): continue
            
            # Basic categorization
            item_lower = item.lower()
            if any(x in item_lower for x in ['communication', 'team', 'leadership']):
                skills['soft_skills'].append(item)
            elif any(x in item_lower for x in ['react', 'django', 'spring', 'flask']):
                skills['frameworks'].append(item)
            elif any(x in item_lower for x in ['docker', 'git', 'aws', 'jira', 'github', 'linux']):
                skills['tools'].append(item)
            else:
                skills['technical_skills'].append(item)
                
        # Deduplicate
        for k in skills:
            skills[k] = list(set(skills[k]))
        return skills

    def extract_projects(self, text: str, section_positions: Dict[str, List[int]]) -> List[Dict[str, Any]]:
        projects = []
        lines = text.split('\n')
        if 'projects' not in section_positions or not section_positions['projects']:
            return projects
            
        start_idx = section_positions['projects'][0]
        end_idx = len(lines)
        for section in ['experience', 'education', 'skills', 'certifications']:
            if section in section_positions and section_positions[section]:
                for pos in section_positions[section]:
                    if pos > start_idx:
                        end_idx = min(end_idx, pos)
        
        proj_text = '\n'.join(lines[start_idx:end_idx])
        entries = re.split(r'\n(?=[A-Z][A-Z\s]+(?:PROJECT|APP|SYSTEM))', proj_text)
        if len(entries) <= 1:
             entries = re.split(r'\n(?=[A-Z][A-Z\s]{10,})', proj_text)
             
        for entry in entries[1:]:
            entry_lines = entry.strip().split('\n')
            if entry_lines:
                proj = {'title': entry_lines[0].strip(), 'description': ' '.join(entry_lines[1:]).strip()}
                if len(proj['title']) > 3:
                     projects.append(proj)
        return projects

    def extract_certifications(self, text: str, section_positions: Dict[str, List[int]]) -> List[Dict[str, Any]]:
        certs = []
        lines = text.split('\n')
        if 'certifications' not in section_positions or not section_positions['certifications']:
            return certs
            
        start_idx = section_positions['certifications'][0]
        end_idx = len(lines)
        for section in ['experience', 'education', 'skills', 'projects']:
            if section in section_positions and section_positions[section]:
                for pos in section_positions[section]:
                    if pos > start_idx:
                        end_idx = min(end_idx, pos)
                        
        cert_text = '\n'.join(lines[start_idx:end_idx])
        for line in cert_text.split('\n'):
            line = line.strip()
            if len(line) > 5 and not line.upper().startswith('CERTIFI'):
                certs.append({'name': line})
        return certs

    def parse(self, file_path: str, use_llm_refinement: bool = False) -> Dict[str, Any]:
        text = self.extract_text(file_path)
        section_positions = self.find_section_headers(text)
        
        result = {
            'personal_information': self.extract_personal_info(text),
            'work_experience': self.extract_experience(text, section_positions),
            'education': self.extract_education(text, section_positions),
            'skills': self.extract_skills(text, section_positions),
            'projects': self.extract_projects(text, section_positions),
            'certifications': self.extract_certifications(text, section_positions),
            'raw_text': text
        }
        
        if use_llm_refinement and self.use_llm and self.llm_available:
            try:
                result = self._refine_data_with_llm(text, result)
            except Exception as e:
                logger.error(f"LLM refinement failed: {e}")
                
        return result

    def _refine_data_with_llm(self, text: str, initial_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Use Cerebras Cloud LLM to refine the extracted data, specifically for skills and experience.
        """
        # Truncate text if too long for context window (rough estimate)
        MAX_CHARS = 12000 
        context_text = text[:MAX_CHARS]
        
        prompt = f"""
        You are an expert CV parser. I have extracted some data using regex, but I need you to correct and structure it better.
        
        Raw CV Text:
        {context_text}
        
        Current Extracted Data (JSON):
        {json.dumps(initial_data, default=str)}
        
        Task:
        1. Fix any extracted fields that look wrong based on the Raw CV Text.
        2. Categorize 'skills' into 'technical_skills', 'tools', 'frameworks'. 
           - CRITICAL: Remove ALL soft skills. Keep ONLY hard/technical skills explicitly listed.
           - IMPORTANT: Split combined skills, e.g. "Python, Java" -> ["Python", "Java"]. Separate skills inside parentheses.
        3. Ensure 'work_experience' has correct 'company', 'position', 'duration' (standardize dates to YYYY-MM if possible) and 'responsibilities'.
           - Remove artifacts like "Remote", "Hybrid" from the start of responsibility bullets.
           - Correctly identify Company Name vs Position if they were swapped or merged.
        4. Standardize 'education' degrees.
        5. CRITICAL: Extract missing arrays from the Raw text and add these keys if present:
           - "languages": list of strings
           - "volunteer_experience": list of objects
           - "awards": list of objects
           - "publications": list of objects
           - "patents": list of objects
        6. **URL EXTRACTION**: Pay close attention to `[Embedded Link: ...]` tags in the text. 
           - Links are often embedded in icons (empty text), titles, or "View Project" buttons. 
           - You MUST map these URLs into the `url` fields of the corresponding `projects` and `certifications` objects!
        
        === STRICT ANTI-HALLUCINATION RULE (CRITICAL) ===
        - Never invent, add, or imply skills, keywords, achievements, metrics, job titles, or any other content not present in the original resume.
        - Only rewrite and restructure what already exists.
        
        === REMOVE FROM EXTRACTED DATA ===
        - Street/home address (city and country are fine)
        - Objective statements
        - Graduation year if the degree is more than 10 years old
        - Work experience older than 15 years (20 years max for executive roles)
        - High school experience
        - GPA or university grades
        - Headshot or photo references
        - Salary expectations
        
        7. Return the COMPLETE corrected data structure as valid JSON.
        
        IMPORTANT: Return ONLY the JSON object. No markdown formatting, no explanations.
        """
        
        try:
            print(f"[INFO] Starting LLM Refinement for CV (Length: {len(context_text)} chars)...")
            response = self.llm_client.chat.completions.create(
                model=LLM_MODEL, 
                messages=[{'role': 'user', 'content': prompt}],
                temperature=0.0,
                response_format={"type": "json_object"}
            )
            print("[INFO] LLM Refinement Completed.")

            
            refined_json = response.choices[0].message.content
            
            # Clean up markdown formatting
            refined_json = refined_json.strip()
            if refined_json.startswith("```json"):
                refined_json = refined_json[7:]
            if refined_json.startswith("```"):
                refined_json = refined_json[3:]
            if refined_json.endswith("```"):
                refined_json = refined_json[:-3]
            
            refined_data = json.loads(refined_json.strip())
            
            # Merge/Validate: Ensure we don't lose key fields if LLM hallucinates emptiness
            # For now, trust LLM if valid JSON is returned, but fallback to initial for empty critical methods?
            # Let's simple return refined_data but ensure keys exist
            for key in ['personal_information', 'work_experience', 'education', 'skills', 'projects', 'certifications']:
                if key not in refined_data:
                    refined_data[key] = initial_data.get(key)
            
            # Keep raw text
            refined_data['raw_text'] = text
            
            return refined_data
            
        except Exception as e:
            logger.error(f"Cerebras chat error: {e}")
            return initial_data

# Wrapper for existing app integration
# Patterns that mark a "skill" string as PDF noise rather than a real skill.
# Hit list comes from benchmarks/results/2026-04-25 — strings like
# "increased sales by 40%.", "[Embedded Link: '", "www.enhancv.com",
# "Developing a High", "Powered by'" leaked through regex splitting on PDFs
# whose skills text was glued onto bullet content.
_SKILL_NOISE_SUBSTRINGS = (
    '[embedded',     # "[Embedded Link: ..." artifacts
    'www.', 'http',  # URL fragments
    '\\u',           # PUA glyphs picked up from icon fonts
)
_SKILL_PERCENT_RE = re.compile(r'\d+\s*%')


def _is_plausible_skill_name(name: str) -> bool:
    """Conservative filter that keeps real skills and drops PDF/bullet noise.

    Tuned against the parser_eval fixtures: each rejected pattern was an
    actual hit in the 2026-04-25 results, not a hypothetical.
    """
    if not name:
        return False
    s = name.strip()
    if len(s) < 2 or len(s) > 40:
        return False
    if s.endswith('.'):                      # sentence fragment
        return False
    if not s[0].isalpha():                   # "(React)" / digit-leading noise
        return False
    if len(s.split()) > 4:                   # bullet body, not a skill
        return False
    low = s.lower()
    if any(token in low for token in _SKILL_NOISE_SUBSTRINGS):
        return False
    if _SKILL_PERCENT_RE.search(s):          # "increased sales by 40%"
        return False
    return True


def parse_cv(file_path):
    """
    Wrapper function compatible with the existing view logic.
    LLM refinement is handled downstream by llm_validator —
    skipping it here avoids redundant API calls and token waste.
    """
    extractor = CVExtractor(use_llm=True)
    data = extractor.parse(file_path, use_llm_refinement=False)

    # Flatten skills for view compatibility
    flat_skills = []
    if 'skills' in data and isinstance(data['skills'], dict):
        for category, skills in data['skills'].items():
            if isinstance(skills, list):
                for skill in skills:
                    if not _is_plausible_skill_name(skill):
                        continue
                    flat_skills.append({"name": skill.strip(), "proficiency": None, "category": category})
            
    personal_info = data.get('personal_information', {})

    # Collect other profile URLs (Kaggle, etc.)
    other_urls = []
    if personal_info.get('kaggle'):
        other_urls.append(personal_info['kaggle'])

    return {
        'full_name': personal_info.get('name', ''),
        'email': personal_info.get('email', ''),
        'phone': personal_info.get('phone', ''),
        'location': personal_info.get('address', ''),
        'linkedin_url': personal_info.get('linkedin', ''),
        'github_url': personal_info.get('github', ''),
        'portfolio_url': personal_info.get('portfolio', ''),
        'other_urls': other_urls,
        'skills': flat_skills,
        'experiences': data.get('work_experience', []),
        'education': data.get('education', []),
        'projects': data.get('projects', []),
        'certifications': data.get('certifications', []),
        'languages': data.get('languages', []),
        'volunteer_experience': data.get('volunteer_experience', []),
        'awards': data.get('awards', []),
        'publications': data.get('publications', []),
        'patents': data.get('patents', []),
        'raw_text': data.get('raw_text', '')
    }

if __name__ == "__main__":
    # Test block
    pass
