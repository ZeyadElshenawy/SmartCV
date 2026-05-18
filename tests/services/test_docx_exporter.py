"""Renderer correctness tests for resumes/services/docx_exporter.py.

PR 1 — Fix 1: certifications-bold isolation.
                  ↑ Fix 2 + Fix 4 tests appended below.
"""
from __future__ import annotations

import io

import pytest
from docx import Document
from docx.oxml.ns import qn


# --------------------------------------------------------------------------
# Fix 1 — un-bold cert issuer / date / verify link
# --------------------------------------------------------------------------

def _find_run_with_text(paragraph, needle: str):
    """Return the python-docx Run whose text contains `needle`, or None."""
    for r in paragraph.runs:
        if needle in (r.text or ''):
            return r
    return None


def _find_hyperlink_run_with_text(paragraph, needle: str):
    """Return the raw <w:r> element nested inside a <w:hyperlink>
    whose <w:t> text contains `needle`. Hyperlink runs are NOT
    surfaced via paragraph.runs, so we walk the XML directly."""
    for hl in paragraph._p.findall(qn('w:hyperlink')):
        for r in hl.iter(qn('w:r')):
            t = r.find(qn('w:t'))
            if t is not None and needle in (t.text or ''):
                return r
    return None


def _has_bold_off(r_element) -> bool:
    """True iff the raw <w:r> element has an explicit <w:b w:val="0"/>."""
    rPr = r_element.find(qn('w:rPr'))
    if rPr is None:
        return False
    b = rPr.find(qn('w:b'))
    if b is None:
        return False
    val = b.get(qn('w:val'))
    return val in ('0', 'false')


def _render_one_cert(cert_dict: dict):
    """Render a single-cert resume via _write_certifications + return
    the certifications paragraph (the first paragraph after the section
    heading, which itself is paragraph[0])."""
    from resumes.services.docx_exporter import _write_certifications
    doc = Document()
    _write_certifications(doc, {'certifications': [cert_dict]})
    # Heading is paragraph[0], cert bullet is paragraph[1].
    return doc.paragraphs[1]


class TestCertificationBoldIsolation:
    """Fix 1 — only the cert NAME is bold. Issuer, date, separator, and
    the verify hyperlink are explicitly non-bold so they don't inherit
    bold from the List Bullet style chain."""

    def setup_method(self):
        self.cert_p = _render_one_cert({
            'name': 'Test Cert',
            'issuer': 'Test Issuer',
            'date': 'Jan 2025',
            'url': 'https://example.com/verify',
        })

    def test_cert_name_is_bold(self):
        run = _find_run_with_text(self.cert_p, 'Test Cert')
        assert run is not None, 'cert name run not found'
        assert run.bold is True, f'expected name bold=True, got {run.bold!r}'

    def test_issuer_is_explicitly_not_bold(self):
        # The suffix run contains issuer + date together: " - Test Issuer (Jan 2025)"
        run = _find_run_with_text(self.cert_p, 'Test Issuer')
        assert run is not None, 'issuer run not found'
        # Explicit False (NOT None) — python-docx writes <w:b w:val="0"/>
        # which surfaces as run.bold == False.
        assert run.bold is False, f'expected issuer bold=False (explicit), got {run.bold!r}'

    def test_date_is_explicitly_not_bold(self):
        # Same suffix run contains the date.
        run = _find_run_with_text(self.cert_p, 'Jan 2025')
        assert run is not None, 'date run not found'
        assert run.bold is False, f'expected date bold=False (explicit), got {run.bold!r}'

    def test_separator_is_explicitly_not_bold(self):
        run = _find_run_with_text(self.cert_p, '·')
        assert run is not None, 'separator run not found'
        assert run.bold is False, f'expected separator bold=False, got {run.bold!r}'

    def test_verify_hyperlink_run_has_explicit_bold_off(self):
        # Hyperlink runs are nested inside <w:hyperlink> — not surfaced
        # by paragraph.runs. The fix writes <w:b w:val="0"/> directly
        # onto each run inside the verify hyperlink.
        r_element = _find_hyperlink_run_with_text(self.cert_p, 'verify')
        assert r_element is not None, 'verify hyperlink run not found'
        assert _has_bold_off(r_element), (
            'verify hyperlink run should have explicit <w:b w:val="0"/> '
            "so it doesn't inherit bold from the List Bullet style chain"
        )


# --------------------------------------------------------------------------
# Fix 2 — LANGUAGES section guard
# --------------------------------------------------------------------------

def _render_languages(languages):
    """Render only the languages section + return the resulting Document
    so the test can introspect whether the section was emitted at all."""
    from resumes.services.docx_exporter import _write_languages
    doc = Document()
    _write_languages(doc, {'languages': languages} if languages is not None else {})
    return doc


def _has_languages_heading(doc) -> bool:
    """True iff a paragraph contains the all-caps 'LANGUAGES' heading
    text (case-insensitive). _add_section_heading uppercases its input."""
    for p in doc.paragraphs:
        if (p.text or '').strip().upper() == 'LANGUAGES':
            return True
    return False


class TestLanguagesGuard:
    """Fix 2 — _write_languages routes its input through
    sanitize_languages_field. Tech skills get dropped; if nothing real
    is left, the section header itself is skipped."""

    def test_all_tech_skills_drops_the_whole_section(self):
        doc = _render_languages(['Python', 'TensorFlow', 'Machine Learning'])
        assert not _has_languages_heading(doc), (
            'tech-only "languages" list should produce no LANGUAGES heading'
        )

    def test_real_languages_render_with_heading(self):
        doc = _render_languages(['English (Fluent)', 'Arabic (Native)'])
        assert _has_languages_heading(doc), 'real languages should produce a LANGUAGES heading'
        body = ' '.join(p.text for p in doc.paragraphs)
        assert 'English (Fluent)' in body
        assert 'Arabic (Native)' in body

    def test_mixed_keeps_only_real_languages(self):
        doc = _render_languages(['English', 'Python', 'Arabic'])
        assert _has_languages_heading(doc)
        body = ' '.join(p.text for p in doc.paragraphs)
        assert 'English' in body
        assert 'Arabic' in body
        # The renderer joins by ", " — so the surviving string is
        # exactly "English, Arabic". Python must not appear standalone.
        rendered_line = next(
            (p.text for p in doc.paragraphs
             if 'English' in (p.text or '') or 'Arabic' in (p.text or '')),
            '',
        )
        assert 'Python' not in rendered_line, (
            f'tech skill leaked into rendered languages line: {rendered_line!r}'
        )

    def test_empty_list_skips_section(self):
        doc = _render_languages([])
        assert not _has_languages_heading(doc)

    def test_none_skips_section(self):
        doc = _render_languages(None)
        assert not _has_languages_heading(doc)


# --------------------------------------------------------------------------
# Fix 4 — Date range normaliser
# --------------------------------------------------------------------------

class TestFormatDateRange:
    """Fix 4 — _format_date_range coerces any pair of date-ish strings
    into 'MMM YYYY – MMM YYYY' with en-dash, with Present handling,
    same-month collapse, year-only fallback, and verbatim preservation
    on parse failure."""

    @pytest.mark.parametrize('start,end,expected', [
        # Spec examples.
        ('August 2025', 'Sep 2025', 'Aug 2025 – Sep 2025'),
        ('July 2024', 'Sep 2024', 'Jul 2024 – Sep 2024'),
        ('Jun 2025', 'Dec 2025', 'Jun 2025 – Dec 2025'),
        ('2025-06', '2025-12', 'Jun 2025 – Dec 2025'),
        ('Jun 2025', 'Present', 'Jun 2025 – Present'),
        ('Jun 2025', None, 'Jun 2025 – Present'),
        # Same-month collapse.
        ('Aug 2025', 'Aug 2025', 'Aug 2025'),
        # Unparseable field → returned verbatim, real one normalised.
        ('garbage', 'Sep 2024', 'garbage – Sep 2024'),
    ])
    def test_spec_cases(self, start, end, expected):
        from resumes.services.docx_exporter import _format_date_range
        assert _format_date_range(start, end) == expected

    def test_present_variants(self):
        from resumes.services.docx_exporter import _format_date_range
        assert _format_date_range('Jun 2025', 'Current') == 'Jun 2025 – Present'
        assert _format_date_range('Jun 2025', 'now') == 'Jun 2025 – Present'
        assert _format_date_range('Jun 2025', '') == 'Jun 2025 – Present'

    def test_year_only_input(self):
        from resumes.services.docx_exporter import _format_date_range
        # Bare year on both sides → bare year out.
        assert _format_date_range('2024', '2024') == '2024'
        # Year range across two years (no month info).
        assert _format_date_range('2023', '2024') == '2023 – 2024'

    def test_both_empty_returns_empty(self):
        from resumes.services.docx_exporter import _format_date_range
        assert _format_date_range('', '') == ''
        assert _format_date_range(None, None) == ''
